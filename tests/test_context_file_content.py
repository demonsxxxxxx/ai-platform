import hashlib
import io
from zipfile import ZIP_DEFLATED, ZIP_STORED, ZipFile

import pytest
from docx import Document
from openpyxl import Workbook
from pypdf import PdfWriter

from app.context.api import ContextFileContentError
from app.context.file_content import (
    MAX_CONTEXT_FILE_STAGE_BYTES,
    validate_context_file_for_stage,
)
from app.file_parser_contracts import XLSX_CONTENT_TYPE


DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PDF_CONTENT_TYPE = "application/pdf"
PPTX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.presentationml.presentation"
_CFB_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"


def _row(name: str, content_type: str, raw: bytes) -> dict[str, object]:
    return {
        "id": "file-a",
        "original_name": name,
        "content_type": content_type,
        "size_bytes": len(raw),
        "sha256": hashlib.sha256(raw).hexdigest(),
    }


def _docx_bytes() -> bytes:
    stream = io.BytesIO()
    document = Document()
    document.add_paragraph("First paragraph")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    document.save(stream)
    return stream.getvalue()


def _zip_with_added_part(
    raw: bytes,
    name: str,
    payload: bytes,
    *,
    compression: int = ZIP_DEFLATED,
) -> bytes:
    stream = io.BytesIO(raw)
    with ZipFile(stream, "a", compression=compression) as archive:
        archive.writestr(name, payload)
    return stream.getvalue()


def _opaque_pdf_bytes(*, encrypted: bool = False, page_count: int = 1) -> bytes:
    stream = io.BytesIO()
    writer = PdfWriter()
    for _ in range(page_count):
        writer.add_blank_page(width=20, height=20)
    if encrypted:
        writer.encrypt("secret")
    writer.write(stream)
    return stream.getvalue()


def _xlsx_bytes_with_cells(cell_count: int) -> bytes:
    stream = io.BytesIO()
    workbook = Workbook()
    sheet = workbook.active
    for row in range(1, cell_count + 1):
        sheet.cell(row=row, column=1, value=row)
    workbook.save(stream)
    return stream.getvalue()


def _xlsx_with_embedded_object(
    relationship_type: str,
    payload: bytes,
    *,
    part_name: str = "xl/embeddings/object.bin",
) -> bytes:
    source = ZipFile(io.BytesIO(_xlsx_bytes_with_cells(2)))
    stream = io.BytesIO()
    with source, ZipFile(stream, "w", compression=ZIP_DEFLATED) as archive:
        for entry in source.infolist():
            entry_payload = source.read(entry)
            if entry.filename == "xl/_rels/workbook.xml.rels":
                entry_payload = entry_payload.replace(
                    b"</Relationships>",
                    (
                        f'<Relationship Id="rId900" Type="{relationship_type}" '
                        f'Target="{part_name.removeprefix("xl/")}" /></Relationships>'
                    ).encode(),
                )
            archive.writestr(entry, entry_payload)
        archive.writestr(part_name, payload)
    return stream.getvalue()


def test_validate_context_file_for_stage_accepts_xlsx_above_legacy_cell_limit():
    raw = _xlsx_bytes_with_cells(2_049)

    validate_context_file_for_stage(_row("large.xlsx", XLSX_CONTENT_TYPE, raw), raw)


@pytest.mark.parametrize(
    ("name", "content_type", "raw"),
    [
        ("notes.txt", "text/plain", b"bounded text"),
        ("source.doc", "application/msword", _CFB_MAGIC + b"legacy Word"),
        ("source.docx", DOCX_CONTENT_TYPE, _docx_bytes()),
        ("source.pdf", PDF_CONTENT_TYPE, _opaque_pdf_bytes()),
        ("book.xls", "application/vnd.ms-excel", _CFB_MAGIC + b"legacy Excel"),
        ("book.xlsx", XLSX_CONTENT_TYPE, _xlsx_bytes_with_cells(2)),
        ("slides.ppt", "application/vnd.ms-powerpoint", _CFB_MAGIC + b"legacy PowerPoint"),
        ("slides.pptx", PPTX_CONTENT_TYPE, b"PK\x03\x04opaque PowerPoint"),
        ("source.bin", "application/octet-stream", b"opaque binary"),
    ],
    ids=("text", "doc", "docx", "pdf", "xls", "xlsx", "ppt", "pptx", "binary"),
)
def test_validate_context_file_for_stage_accepts_opaque_authorized_types(name, content_type, raw):
    validate_context_file_for_stage(_row(name, content_type, raw), raw)


def test_validate_context_file_for_stage_rejects_identity_mismatch_without_parsing_type():
    raw = b"hello"
    with pytest.raises(ContextFileContentError, match="context_file_identity_mismatch"):
        validate_context_file_for_stage(
            {**_row("notes.txt", "text/plain", raw), "sha256": "0" * 64}, raw
        )

    validate_context_file_for_stage(_row("notes.pdf", PDF_CONTENT_TYPE, raw), raw)


def test_validate_context_file_for_stage_rejects_oversize_file():
    raw = b"a" * (MAX_CONTEXT_FILE_STAGE_BYTES + 1)

    with pytest.raises(ContextFileContentError, match="context_file_too_large"):
        validate_context_file_for_stage(_row("notes.txt", "text/plain", raw), raw)


def test_validate_context_file_for_stage_accepts_docx_above_legacy_16_mib_limit():
    stream = io.BytesIO(_docx_bytes())
    with ZipFile(stream, "a", compression=ZIP_STORED) as archive:
        archive.writestr("customXml/bounded-padding.bin", b"x" * (17 * 1024 * 1024))
    raw = stream.getvalue()

    assert 16 * 1024 * 1024 < len(raw) <= MAX_CONTEXT_FILE_STAGE_BYTES
    validate_context_file_for_stage(_row("source.docx", DOCX_CONTENT_TYPE, raw), raw)


@pytest.mark.parametrize(
    "raw",
    [
        b"%PDF-intentionally-malformed",
        _opaque_pdf_bytes(encrypted=True),
        _opaque_pdf_bytes(page_count=1_001),
    ],
    ids=("malformed", "encrypted", "above-former-page-limit"),
)
def test_validate_context_file_for_stage_accepts_opaque_pdf_without_parsing(raw):
    validate_context_file_for_stage(_row("source.pdf", PDF_CONTENT_TYPE, raw), raw)


@pytest.mark.parametrize(
    ("name", "content_type"),
    [
        ("source.xlsx", XLSX_CONTENT_TYPE),
        ("source.xlsx", "application/octet-stream"),
        ("source.bin", XLSX_CONTENT_TYPE),
    ],
)
def test_validate_context_file_for_stage_rejects_invalid_declared_xlsx_archive(
    name, content_type
):
    raw = b"PK-not-a-zip"

    with pytest.raises(ContextFileContentError, match="context_file_xlsx_archive_invalid"):
        validate_context_file_for_stage(_row(name, content_type, raw), raw)


def test_validate_context_file_for_stage_rejects_xlsx_ole_content():
    ole_relationship = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/oleObject"
    )
    package_relationship = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/package"
    )
    vba_relationship = (
        "http://schemas.microsoft.com/office/2006/relationships/vbaProject"
    )
    cases = (
        _xlsx_with_embedded_object(ole_relationship, b"opaque"),
        _xlsx_with_embedded_object(package_relationship, b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1payload"),
        _xlsx_with_embedded_object(
            package_relationship,
            b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1payload",
            part_name="xl/data/payload.bin",
        ),
        _xlsx_with_embedded_object(
            vba_relationship,
            b"opaque",
            part_name="xl/data/payload.bin",
        ),
    )

    for raw in cases:
        with pytest.raises(ContextFileContentError, match="context_file_xlsx_archive_invalid"):
            validate_context_file_for_stage(_row("source.xlsx", XLSX_CONTENT_TYPE, raw), raw)


def test_validate_context_file_for_stage_accepts_docx_cfb_payload_without_inspection():
    raw = _zip_with_added_part(
        _docx_bytes(),
        "word/embeddings/process-flow.bin",
        b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1opaque payload",
    )

    validate_context_file_for_stage(_row("source.docx", DOCX_CONTENT_TYPE, raw), raw)


@pytest.mark.parametrize(
    "unsafe_name",
    ["../xl/workbook.xml", "XL/workbook.xml", "xl//workbook.xml", "xl/./workbook.xml"],
)
def test_validate_context_file_for_stage_rejects_unsafe_xlsx_paths(unsafe_name):
    raw = _zip_with_added_part(
        _xlsx_bytes_with_cells(2),
        unsafe_name,
        b'<workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main" />',
    )

    with pytest.raises(ContextFileContentError, match="context_file_xlsx_archive_invalid"):
        validate_context_file_for_stage(_row("source.xlsx", XLSX_CONTENT_TYPE, raw), raw)


def test_validate_context_file_for_stage_rejects_xlsx_missing_required_part():
    source = ZipFile(io.BytesIO(_xlsx_bytes_with_cells(2)))
    stream = io.BytesIO()
    with source, ZipFile(stream, "w", compression=ZIP_DEFLATED) as archive:
        for entry in source.infolist():
            if entry.filename != "xl/workbook.xml":
                archive.writestr(entry, source.read(entry))
    raw = stream.getvalue()

    with pytest.raises(ContextFileContentError, match="context_file_xlsx_archive_invalid"):
        validate_context_file_for_stage(_row("source.xlsx", XLSX_CONTENT_TYPE, raw), raw)

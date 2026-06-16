#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Document Reviewer - Unified pipeline runner

This script provides a deterministic execution path for the local qa-file-reviewer skill:
1) parse DOCX
2) execute review branches in parallel
3) build review-schema JSON
4) validate hard gates
5) generate the default user-facing commented Word copy and keep txt as diagnostics
"""

from __future__ import annotations

import argparse
import json
import os
import re
import subprocess
import sys
import time
from collections import Counter, defaultdict
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

from docx import Document
from docx.oxml.ns import qn

from add_word_comments_v3 import add_comments_improved
from qa_comment_adjudicator import adjudicate_issues

try:
    from validate_agent_context_package import validate_package as validate_agent_context_package
except ModuleNotFoundError:  # Support package-style imports in tests/tools.
    from .validate_agent_context_package import validate_package as validate_agent_context_package


SCRIPT_DIR = Path(__file__).resolve().parent
SKILL_DIR = SCRIPT_DIR.parent
VALIDATOR_PATH = SCRIPT_DIR / "validate_pipeline_gate.py"
MINIMAX_DOCX_ENGINE_PATH = SKILL_DIR.parent / "minimax-docx" / "docx_engine.py"

SEVERITY_ORDER = {"关键": 0, "主要": 1, "次要": 2}
TYPE_ORDER = {"内容一致性": 0, "大模型全文审核": 1, "双语一致性": 2, "英文语言": 3, "中文语言": 4, "项目号": 5, "格式": 6}
CONTENT_CONSISTENCY_BRANCH = "content_consistency"
LLM_REVIEW_BRANCH = "llm_full_review"
MAX_DOCUMENT_END_FALLBACK_COMMENTS = 3
MAX_DOCUMENT_END_FALLBACK_RATIO = 0.5
MAX_COMMENT_FIELD_CHARS = 180
COMMENT_FIELD_LIMITS = {
    "issue": 130,
    "original": 115,
    "evidence": 115,
    "suggestion": 130,
}

BRANCH_META = {
    "format": {"agent_role": "format-agent", "type": "格式"},
    "project_number": {"agent_role": "project-number-agent", "type": "项目号"},
    CONTENT_CONSISTENCY_BRANCH: {"agent_role": "content-consistency-agent", "type": "内容一致性"},
    LLM_REVIEW_BRANCH: {"agent_role": "llm-full-review-agent", "type": "大模型全文审核"},
}
LLM_CATEGORY_RULES = {
    "zh_language": "LLM-ZH-001",
    "en_language": "LLM-EN-001",
    "bilingual_consistency": "LLM-BI-001",
    "semantic_consistency": "LLM-SEM-001",
}
LLM_ALLOWED_CATEGORIES = set(LLM_CATEGORY_RULES)
LLM_ALLOWED_SEVERITIES = {"关键", "主要", "次要"}
LLM_ALLOWED_CONFIDENCES = {"high", "medium", "low"}
ISSUE_STATUS_CONFIRMED = "confirmed"
ISSUE_STATUS_NEEDS_USER_CHECK = "needs_user_check"
CHECK_STATUSES = {ISSUE_STATUS_NEEDS_USER_CHECK}
WORD_VISIBLE_STATUSES = {ISSUE_STATUS_CONFIRMED, ISSUE_STATUS_NEEDS_USER_CHECK}
ALLOWED_COVERAGE_DOMAINS = {
    "format",
    "structure",
    "zh_language",
    "en_language",
    "bilingual",
    "data_consistency",
    "terminology",
    "external_check",
}
ALLOWED_REVIEW_BASIS = {"company_standard", "single_doc_internal", "agent_semantic", "external_required"}
ALLOWED_EXTERNAL_EVIDENCE_TYPES = {"none", "record", "protocol", "sample_info", "lims", "other"}
ALLOWED_COMMENT_INTENTS = {"suggest_change", "request_check", "global_summary"}
AGENT_ROLE_CATEGORY_DEFAULTS = {
    "zh_language": "zh_language",
    "en_language": "en_language",
    "bilingual_consistency": "bilingual_consistency",
    "semantic_consistency": "semantic_consistency",
    "risk_classifier": "semantic_consistency",
    "structure": "semantic_consistency",
}
AGENT_CATEGORY_SYNONYMS = {
    "bilingual": "bilingual_consistency",
    "bilingualconsistency": "bilingual_consistency",
    "semantic": "semantic_consistency",
    "semanticconsistency": "semantic_consistency",
    "dataconsistency": "semantic_consistency",
    "riskclassification": "semantic_consistency",
    "structure": "semantic_consistency",
    "format": "semantic_consistency",
    "zhlanguage": "zh_language",
    "enlanguage": "en_language",
}
AGENT_COVERAGE_DOMAIN_SYNONYMS = {
    "bilingualconsistency": "bilingual",
    "structureandformat": "structure",
    "zhlanguage": "zh_language",
    "enlanguage": "en_language",
    "dataconsistency": "data_consistency",
    "terminology": "terminology",
    "structure": "structure",
    "format": "format",
}
AGENT_COMMENT_INTENT_SYNONYMS = {
    "suggestedit": "suggest_change",
    "suggestchange": "suggest_change",
    "fixsuggestion": "suggest_change",
    "suggestion": "suggest_change",
    "requestcheck": "request_check",
    "needcheck": "request_check",
    "globalsummary": "global_summary",
}
PAGE_STANDARD = {"top_cm": 2.5, "bottom_cm": 2.5, "left_cm": 3.0, "right_cm": 2.5}
A4_STANDARD = {"width_cm": 21.0, "height_cm": 29.7}
SECTION_DISTANCE_STANDARD = {"gutter_cm": 0.5, "header_cm": 1.5, "footer_cm": 1.75}
BODY_FONT_SIZE_PT = 12.0
TABLE_FONT_SIZE_PT = 10.5
FOOTER_FONT_SIZE_PT = 8.0
FOOTER_REQUIRED_FRAGMENTS = ("公司机密-仅供内部使用", "Company Confidential-Internal Use Only")
APPROVAL_TABLE_KEYWORDS = ("审批", "姓名", "部门", "职务", "签名", "日期", "approval")
PROJECT_PATTERN = re.compile(r"(?<![A-Za-z0-9])[A-Z]{2,4}\d{3,5}[A-Z]?(?![A-Za-z0-9])")
ZH_CHAR_PATTERN = re.compile(r"[\u4e00-\u9fff]")
EN_CHAR_PATTERN = re.compile(r"[A-Za-z]")
HANGUL_PATTERN = re.compile(r"[\uac00-\ud7af]")
ASCII_WORD_PATTERN = re.compile(r"[A-Za-z]{2,}")
TABLE_EXTRA_SPACE_PATTERN = re.compile(r"(?:\n\s{2,}|[A-Za-z]\s{2,}[A-Za-z])")
TABLE_NOTE_PREFIXES = ("注：", "注:", "Note:", "NOTE:", "备注：", "说明：")
ZH_DOMINANT_TERM_ENDINGS = {"肽"}
ZH_NEAR_VARIANT_ENDINGS = {"肽", "太", "胎"}
EN_DOMINANT_TERM_SUFFIXES = ("glutide",)
MONTH_YEAR_PATTERN = re.compile(
    r"\b(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|"
    r"Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+20\d{2}\b",
    flags=re.IGNORECASE,
)
YEAR_DOT_MONTH_PATTERN = re.compile(r"\b20\d{2}\.(?:0?[1-9]|1[0-2])\b")
PASSIVE_FILTER_PATTERN = re.compile(r"\b(was|were|is|are|be|been|being)\s+filter\b", flags=re.IGNORECASE)
RELEASE_DATA_NOTE_PATTERN = re.compile(r"\bdata\s+reference\s+to\s+the\s+releasing\s+data\b", flags=re.IGNORECASE)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Run the document review local pipeline.")
    parser.add_argument("input_docx", help="Input Word document path.")
    parser.add_argument(
        "output_dir",
        nargs="?",
        default=str(Path.cwd() / "qa_review_output"),
        help="Output directory. Defaults to ./qa_review_output",
    )
    comment_group = parser.add_mutually_exclusive_group()
    comment_group.add_argument(
        "--with-comments",
        dest="with_comments",
        action="store_true",
        help="Generate a commented Word copy after validation passes.",
    )
    comment_group.add_argument(
        "--no-comments",
        dest="with_comments",
        action="store_false",
        help="Skip generating a commented Word copy after validation passes.",
    )
    parser.set_defaults(with_comments=True)
    parser.add_argument("--author", default="文件审核系统", help="Author name for new Word comments.")
    parser.add_argument(
        "--original-filename",
        default="",
        help="Original uploaded filename for diagnostic text report and output artifact names.",
    )
    parser.add_argument(
        "--report-file",
        default="",
        help="Optional path for internal diagnostic text report. Default: <output_dir>/<doc_stem>_审核详细报告.txt",
    )
    parser.add_argument(
        "--keep-json-artifacts",
        action="store_true",
        help="Keep all internal JSON artifacts, including validation/document_map/comment_plan/docx_audit.",
    )
    parser.add_argument(
        "--agent-review-json",
        action="append",
        default=[],
        help=(
            "External semantic review JSON produced by the calling agent. "
            "May be passed multiple times. The script validates and merges it, but never calls an LLM itself."
        ),
    )
    parser.add_argument(
        "--agent-context-manifest",
        default="",
        help=(
            "Optional agent_context_manifest.json from export_agent_review_context.py --context-version v2. "
            "When omitted, <output_dir>/agent_context_manifest.json is used if present."
        ),
    )
    return parser.parse_args()


def ensure_utf8_stdio() -> None:
    for name in ("stdout", "stderr"):
        stream = getattr(sys, name)
        if hasattr(stream, "reconfigure"):
            stream.reconfigure(encoding="utf-8")


def load_document(input_docx: Path) -> Document:
    if not input_docx.exists():
        raise FileNotFoundError(f"input_docx not found: {input_docx}")
    return Document(str(input_docx))


def paragraph_text(paragraph: Any) -> str:
    return paragraph.text or ""


ANCHOR_CHAR_TRANSLATION = str.maketrans(
    {
        "“": '"',
        "”": '"',
        "„": '"',
        "＂": '"',
        "‘": "'",
        "’": "'",
        "‚": "'",
        "＇": "'",
        "µ": "μ",
    }
)


def normalize_anchor_char(char: str) -> str:
    return char.translate(ANCHOR_CHAR_TRANSLATION)


def normalized_text(text: str) -> str:
    return re.sub(r"\s+", "", (text or "").translate(ANCHOR_CHAR_TRANSLATION))


def normalized_text_with_map(text: str) -> Tuple[str, List[int]]:
    normalized_chars: List[str] = []
    mapping: List[int] = []
    for index, char in enumerate(text or ""):
        if char.isspace():
            continue
        normalized_chars.append(normalize_anchor_char(char))
        mapping.append(index)
    return "".join(normalized_chars), mapping


def split_non_empty_lines(text: str) -> List[str]:
    return [line.strip() for line in text.splitlines() if line.strip()]


def extract_document_map(input_docx: Path, output_path: Path) -> Dict[str, Any]:
    command = [
        sys.executable,
        str(MINIMAX_DOCX_ENGINE_PATH),
        "review-map",
        str(input_docx),
        str(output_path),
    ]
    process = subprocess.run(command, capture_output=True, text=True, encoding="utf-8", errors="replace")
    if process.returncode != 0:
        raise RuntimeError(
            "minimax-docx review-map failed: "
            + (process.stderr.strip() or process.stdout.strip() or f"exit={process.returncode}")
        )
    if not output_path.exists():
        raise RuntimeError(f"document_map not generated: {output_path}")
    return json.loads(output_path.read_text(encoding="utf-8"))


def get_review_paragraphs(document_map: Dict[str, Any]) -> List[Dict[str, Any]]:
    paragraphs = document_map.get("paragraphs", [])
    if not isinstance(paragraphs, list):
        return []
    review_paragraphs: List[Dict[str, Any]] = []
    in_generated_review_appendix = False
    for item in paragraphs:
        text = str(item.get("text", "")).strip()
        if not text:
            continue
        if text == "未定位审核意见":
            in_generated_review_appendix = True
            continue
        if text.startswith("审核意见："):
            continue
        if in_generated_review_appendix:
            continue
        review_paragraphs.append(item)
    return review_paragraphs


def ordered_unique(values: Iterable[str]) -> List[str]:
    seen: set[str] = set()
    result: List[str] = []
    for value in values:
        if value in seen:
            continue
        seen.add(value)
        result.append(value)
    return result


def user_filename_basename(filename: str) -> str:
    cleaned = str(filename or "").strip().replace("\\", "/")
    return cleaned.rsplit("/", 1)[-1].strip()


def output_stem_from_filename(filename: str, fallback_stem: str) -> str:
    name = user_filename_basename(filename)
    stem = Path(name).stem if name else ""
    if not stem:
        stem = fallback_stem
    stem = re.sub(r"[\x00-\x1f<>:\"/\\|?*]+", "_", stem).strip(" ._")
    return stem or fallback_stem or "document"


def display_name_from_filename(filename: str, fallback_name: str) -> str:
    name = user_filename_basename(filename)
    return name or fallback_name


def project_aliases(current_project: str) -> set[str]:
    return set(PROJECT_PATTERN.findall(current_project or ""))


def project_alias_group_from_text(text: str) -> List[str]:
    for line in split_non_empty_lines(text or ""):
        line_matches = [
            (match.group(0), match.start(), match.end())
            for match in PROJECT_PATTERN.finditer(line)
        ]
        if len(line_matches) < 2:
            continue
        for parenthetical in re.finditer(r"[\(（][^\)）]*[\)）]", line):
            inner_matches = [
                item
                for item in line_matches
                if parenthetical.start() <= item[1] and item[2] <= parenthetical.end()
            ]
            if not inner_matches:
                continue
            before_matches = [item for item in line_matches if item[2] <= parenthetical.start()]
            after_matches = [item for item in line_matches if item[1] >= parenthetical.end()]
            anchor = before_matches[-1] if before_matches else (after_matches[0] if after_matches else None)
            if anchor is None:
                continue
            alias_matches = sorted([anchor] + inner_matches, key=lambda item: item[1])
            return ordered_unique(item[0] for item in alias_matches)
    return []


def detect_current_project_aliases(input_docx: Path, review_paragraphs: Sequence[Dict[str, Any]]) -> List[str]:
    file_name = user_filename_basename(str(input_docx))
    file_stem = Path(file_name).stem if file_name else input_docx.stem
    file_matches = ordered_unique(PROJECT_PATTERN.findall(file_stem))
    if len(file_matches) > 1:
        return file_matches
    if len(file_matches) == 1:
        file_project = file_matches[0]
        for para in review_paragraphs[:80]:
            aliases = project_alias_group_from_text(str(para.get("text", "")))
            if file_project in aliases:
                return aliases
        return file_matches

    for para in review_paragraphs[:80]:
        aliases = project_alias_group_from_text(str(para.get("text", "")))
        if aliases:
            return aliases

    counts: Counter[str] = Counter()
    for para in review_paragraphs[:80]:
        for match in PROJECT_PATTERN.findall(str(para.get("text", ""))):
            counts[match] += 1

    if len(counts) == 1:
        return [next(iter(counts))]

    return []


def detect_current_project(input_docx: Path, review_paragraphs: Sequence[Dict[str, Any]]) -> str:
    aliases = detect_current_project_aliases(input_docx, review_paragraphs)
    return " / ".join(aliases) if aliases else "UNKNOWN"


def paragraph_location(record: Dict[str, Any]) -> str:
    logical_index = record.get("logical_index") or record.get("xml_index") or "?"
    xml_index = record.get("xml_index") or "?"
    return f"第{logical_index}段（XML:{xml_index}）"


def paragraph_anchor_locator(record: Dict[str, Any]) -> str:
    return f"paragraph={record.get('xml_index', '')}"


def paragraph_document_zone(record: Dict[str, Any]) -> str:
    zone = str(record.get("document_zone") or "body")
    return zone if zone in {"body", "table", "metadata"} else "body"


def paragraph_location_kind(record: Dict[str, Any]) -> str:
    return "table" if paragraph_document_zone(record) == "table" else "paragraph"


def has_chinese(text: str) -> bool:
    return bool(ZH_CHAR_PATTERN.search(text or ""))


def has_english(text: str) -> bool:
    return bool(EN_CHAR_PATTERN.search(text or ""))


def find_exact_span(text: str, needle: str) -> Optional[Dict[str, Any]]:
    if not text or not needle:
        return None
    indices: List[int] = []
    index = text.find(needle)
    while index >= 0:
        indices.append(index)
        index = text.find(needle, index + 1)
    if len(indices) == 1:
        start = indices[0]
        return {"start": start, "end": start + len(needle), "unit": "char"}

    normalized_source, mapping = normalized_text_with_map(text)
    normalized_needle, _ = normalized_text_with_map(needle)
    if len(normalized_needle) < 2 or not normalized_source:
        return None
    normalized_indices: List[int] = []
    index = normalized_source.find(normalized_needle)
    while index >= 0:
        normalized_indices.append(index)
        index = normalized_source.find(normalized_needle, index + 1)
    if len(normalized_indices) == 1:
        start = normalized_indices[0]
        return {
            "start": mapping[start],
            "end": mapping[start + len(normalized_needle) - 1] + 1,
            "unit": "char",
        }
    return None


def find_first_exact_span(text: str, needle: str) -> Optional[Dict[str, Any]]:
    if not text or not needle:
        return None
    index = text.find(needle)
    if index >= 0:
        return {"start": index, "end": index + len(needle), "unit": "char"}

    normalized_source, mapping = normalized_text_with_map(text)
    normalized_needle, _ = normalized_text_with_map(needle)
    if len(normalized_needle) < 2 or not normalized_source:
        return None
    normalized_index = normalized_source.find(normalized_needle)
    if normalized_index >= 0:
        return {
            "start": mapping[normalized_index],
            "end": mapping[normalized_index + len(normalized_needle) - 1] + 1,
            "unit": "char",
        }
    return None


def find_paragraph_record_for_text(
    paragraphs: Sequence[Dict[str, Any]],
    text: str,
    *,
    document_zone: str = "",
    table_index: Optional[int] = None,
    row_index: Optional[int] = None,
    cell_index: Optional[int] = None,
) -> Optional[Dict[str, Any]]:
    needle = str(text or "").strip()
    if not needle:
        return None
    needle_norm = normalized_text(needle)
    if len(needle_norm) < 2:
        return None

    for record in paragraphs:
        if document_zone and paragraph_document_zone(record) != document_zone:
            continue
        if table_index is not None and record.get("table_index") != table_index:
            continue
        if row_index is not None and record.get("row_index") != row_index:
            continue
        if cell_index is not None and record.get("cell_index") != cell_index:
            continue

        source_text = str(record.get("text") or "")
        source_norm = normalized_text(source_text)
        if not source_norm:
            continue
        if needle_norm in source_norm or source_norm in needle_norm:
            return record

    return None


def anchor_fields_from_record(record: Optional[Dict[str, Any]], text: str = "") -> Dict[str, Any]:
    if not record:
        return {
            "paragraph_index": None,
            "anchor_locator": "",
            "anchor_span": None,
            "anchor_text": "",
            "comments_added": 0,
            "match_method": "inference",
        }

    source_text = str(record.get("text") or "")
    anchor_text = str(text or source_text).strip() or source_text
    span = find_first_exact_span(source_text, anchor_text) or (
        {"start": 0, "end": len(source_text), "unit": "char"} if source_text else None
    )
    return {
        "paragraph_index": None,
        "anchor_locator": paragraph_anchor_locator(record),
        "anchor_span": span,
        "anchor_text": anchor_text,
        "comments_added": 1 if span else 0,
        "match_method": "span" if span else "inference",
    }


def first_table_header_anchor_record(
    paragraphs: Sequence[Dict[str, Any]],
    table_index: int,
) -> Optional[Dict[str, Any]]:
    table_records = [
        record
        for record in paragraphs
        if paragraph_document_zone(record) == "table"
        and record.get("table_index") == table_index
        and normalized_text(str(record.get("text") or ""))
    ]
    if not table_records:
        return None

    first_row_records = [record for record in table_records if record.get("row_index") == 1]
    candidates = first_row_records or table_records
    return sorted(
        candidates,
        key=lambda record: (
            int(record.get("row_index") or 999999),
            int(record.get("cell_index") or 999999),
            int(record.get("xml_index") or 999999),
        ),
    )[0]


def env_flag(name: str, default: bool) -> bool:
    value = os.environ.get(name)
    if value is None:
        return default
    return value.strip().lower() not in {"0", "false", "no", "off", "disabled"}


def env_int(name: str, default: int) -> int:
    try:
        return int(os.environ.get(name, "").strip())
    except ValueError:
        return default


def approx_equal(left: float, right: float, tolerance: float = 0.05) -> bool:
    return abs(left - right) <= tolerance


def compact_text(text: str) -> str:
    return re.sub(r"\s+", "", text or "")


def has_unstable_table_extra_space(text: str) -> bool:
    if not text:
        return False

    normalized = text.replace("\u00a0", " ").replace("\u3000", " ")
    if TABLE_EXTRA_SPACE_PATTERN.search(normalized):
        return True

    for line in normalized.splitlines():
        if not line.strip():
            continue
        if line[:1].isspace():
            return True

    return False


def style_name(paragraph: Any) -> str:
    style = getattr(paragraph, "style", None)
    if style is None:
        return ""
    return str(getattr(style, "name", "") or "")


def resolve_style_font_size(style: Any) -> Optional[float]:
    seen: set[int] = set()
    current = style
    while current is not None and id(current) not in seen:
        seen.add(id(current))
        size = getattr(getattr(current, "font", None), "size", None)
        if size is not None:
            return round(size.pt, 1)
        current = getattr(current, "base_style", None)
    return None


def paragraph_effective_font_sizes(paragraph: Any) -> List[float]:
    sizes = sorted(
        {
            round(run.font.size.pt, 1)
            for run in getattr(paragraph, "runs", [])
            if getattr(run, "text", "").strip() and getattr(getattr(run, "font", None), "size", None) is not None
        }
    )
    if sizes:
        return sizes

    fallback = resolve_style_font_size(getattr(paragraph, "style", None))
    return [fallback] if fallback is not None else []


def is_body_text_paragraph(paragraph: Any) -> bool:
    text = getattr(paragraph, "text", "") or ""
    name = style_name(paragraph)
    return bool(text.strip()) and (name == "Normal (Web)" or "正文" in name)


def length_pt(value: Any) -> Optional[float]:
    if value is None:
        return None
    pt = getattr(value, "pt", None)
    if pt is None:
        return None
    return round(float(pt), 2)


def paragraph_has_zero_before_after_spacing(paragraph: Any) -> bool:
    before = paragraph_spacing_pt(paragraph, "space_before")
    after = paragraph_spacing_pt(paragraph, "space_after")
    return before is not None and after is not None and approx_equal(before, 0.0) and approx_equal(after, 0.0)


def paragraph_spacing_pt(paragraph: Any, attr_name: str) -> Optional[float]:
    paragraph_format = getattr(paragraph, "paragraph_format", None)
    if paragraph_format is not None:
        direct = length_pt(getattr(paragraph_format, attr_name, None))
        if direct is not None:
            return direct

    seen: set[int] = set()
    current = getattr(paragraph, "style", None)
    while current is not None and id(current) not in seen:
        seen.add(id(current))
        style_format = getattr(current, "paragraph_format", None)
        if style_format is not None:
            inherited = length_pt(getattr(style_format, attr_name, None))
            if inherited is not None:
                return inherited
        current = getattr(current, "base_style", None)
    return None


def paragraph_is_adjacent_to_table(paragraph: Any) -> bool:
    element = getattr(paragraph, "_p", None)
    if element is None:
        return False
    previous_element = element.getprevious()
    next_element = element.getnext()
    table_tag = qn("w:tbl")
    return (
        (previous_element is not None and previous_element.tag == table_tag)
        or (next_element is not None and next_element.tag == table_tag)
    )


def is_table_note_paragraph(paragraph: Any) -> bool:
    text = (getattr(paragraph, "text", "") or "").strip()
    if not text:
        return False
    sizes = paragraph_effective_font_sizes(paragraph)
    has_table_note_font = bool(sizes) and all(approx_equal(size, TABLE_FONT_SIZE_PT) for size in sizes)
    if not has_table_note_font:
        return False
    if not paragraph_has_zero_before_after_spacing(paragraph):
        return False
    if not paragraph_is_adjacent_to_table(paragraph):
        return False
    return text.startswith(TABLE_NOTE_PREFIXES) or len(text) <= 160


def footer_compact_text(section: Any) -> str:
    text = "".join(paragraph.text for paragraph in section.footer.paragraphs if paragraph.text.strip())
    return compact_text(text)


def section_has_footer_text(section: Any) -> bool:
    text = footer_compact_text(section)
    return all(compact_text(fragment) in text for fragment in FOOTER_REQUIRED_FRAGMENTS)


def footer_font_sizes(section: Any) -> List[float]:
    sizes: set[float] = set()
    for paragraph in section.footer.paragraphs:
        sizes.update(paragraph_effective_font_sizes(paragraph))
    return sorted(sizes)


def table_has_repeat_header(table: Any) -> bool:
    first_row = table.rows[0]._tr if table.rows else None
    if first_row is None or first_row.trPr is None:
        return False
    for child in first_row.trPr.iterchildren():
        if child.tag == qn("w:tblHeader"):
            return True
    return False


def table_preview_text(table: Any, limit: int = 8) -> str:
    chunks: List[str] = []
    for row in table.rows[:limit]:
        for cell in row.cells[: min(len(row.cells), 4)]:
            for paragraph in cell.paragraphs:
                text = (paragraph.text or "").strip()
                if text:
                    chunks.append(text)
                if len(chunks) >= limit:
                    return " ".join(chunks)
    return " ".join(chunks)


def is_approval_table(table: Any) -> bool:
    preview = table_preview_text(table).lower()
    return len(table.rows) <= 10 and any(keyword in preview for keyword in APPROVAL_TABLE_KEYWORDS)


def format_examples(examples: Sequence[str], limit: int = 6) -> str:
    unique_examples = list(dict.fromkeys(examples))
    if len(unique_examples) <= limit:
        return "；".join(unique_examples)
    remaining = len(unique_examples) - limit
    return "；".join(list(unique_examples[:limit]) + [f"其余{remaining}处省略"])


def format_grouped_examples(examples: Sequence[Tuple[str, str]], limit: int = 6) -> str:
    grouped: Dict[str, Dict[str, Any]] = {}
    for location, text in examples:
        key = compact_text(text)
        if key not in grouped:
            grouped[key] = {"location": location, "text": text, "count": 0}
        grouped[key]["count"] += 1

    rendered: List[str] = []
    for item in grouped.values():
        suffix = f"（共{item['count']}处）" if item["count"] > 1 else ""
        rendered.append(f"{item['location']} `{item['text']}`{suffix}")

    return format_examples(rendered, limit=limit)


def iter_table_paragraphs(doc: Document) -> Iterable[Tuple[int, int, int, Any, Any]]:
    for table_index, table in enumerate(doc.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            for cell_index, cell in enumerate(row.cells, start=1):
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        yield table_index, row_index, cell_index, table, paragraph


def build_approval_table_index_set(doc: Document) -> set[int]:
    indexes: set[int] = set()
    for table_index, table in enumerate(doc.tables, start=1):
        if is_approval_table(table):
            indexes.add(table_index)
    return indexes


def is_approval_record(record: Dict[str, Any], approval_table_indexes: set[int]) -> bool:
    table_index = record.get("table_index")
    if table_index is None:
        return False
    try:
        return int(table_index) in approval_table_indexes
    except (TypeError, ValueError):
        return False


def table_record_location(record: Dict[str, Any]) -> str:
    table_index = record.get("table_index") or "?"
    row_index = record.get("row_index") or "?"
    cell_index = record.get("cell_index") or "?"
    logical_index = record.get("logical_index") or record.get("xml_index") or "?"
    xml_index = record.get("xml_index") or "?"
    return f"表{table_index} R{row_index}C{cell_index}（第{logical_index}段 / XML:{xml_index}）"


def short_table_record_location(record: Dict[str, Any]) -> str:
    table_index = record.get("table_index") or "?"
    row_index = record.get("row_index") or "?"
    cell_index = record.get("cell_index") or "?"
    return f"表{table_index} R{row_index}C{cell_index}"


def short_heading_location(index: int, paragraph: Any) -> str:
    return f"第{index}段（{style_name(paragraph) or 'unknown-style'}）"


def is_bilingual_heading_paragraph(paragraph: Any) -> bool:
    name = style_name(paragraph)
    if name in {"Heading 1", "Heading 2", "三级标题", "表格名称", "toc 1", "toc 2"}:
        text = (getattr(paragraph, "text", "") or "").strip()
        if text.startswith(("Note:", "注：", "注:")):
            return False
        return True
    text = (getattr(paragraph, "text", "") or "").strip()
    if not text:
        return False
    return bool(re.match(r"^\d+(?:\.\d+)*[\.、]?", text)) and (has_chinese(text) or has_english(text))


def infer_document_zone(location_kind: str) -> str:
    if location_kind in {"table"}:
        return "table"
    if location_kind in {"property", "section"}:
        return "metadata"
    return "body"


def parse_bool(value: Any) -> bool:
    if isinstance(value, bool):
        return value
    if value is None:
        return False
    return str(value).strip().lower() in {"1", "true", "yes", "y", "是", "需要", "required"}


def enum_token(value: Any) -> str:
    return re.sub(r"[\s_\-]+", "", str(value or "")).strip().casefold()


def normalize_status(status: str, *, requires_external_evidence: bool = False) -> str:
    value = str(status or "").strip()
    if requires_external_evidence:
        return ISSUE_STATUS_NEEDS_USER_CHECK
    if value in {ISSUE_STATUS_CONFIRMED, ISSUE_STATUS_NEEDS_USER_CHECK}:
        return value
    return ISSUE_STATUS_CONFIRMED


def infer_requires_external_evidence(
    *,
    requires_external_evidence: Any = False,
    external_evidence_type: Any = "",
    review_basis: Any = "",
    coverage_domain: Any = "",
    comment_intent: Any = "",
) -> bool:
    if parse_bool(requires_external_evidence):
        return True
    external_type = str(external_evidence_type or "").strip()
    if external_type and external_type != "none":
        return True
    if str(review_basis or "").strip() == "external_required":
        return True
    if str(coverage_domain or "").strip() == "external_check":
        return True
    return False


def default_coverage_domain(branch: str) -> str:
    if branch == "format":
        return "format"
    if branch == "project_number":
        return "terminology"
    if branch == CONTENT_CONSISTENCY_BRANCH:
        return "data_consistency"
    if branch == LLM_REVIEW_BRANCH:
        return "bilingual"
    return "structure"


def default_review_basis(branch: str) -> str:
    if branch in {"format", "project_number"}:
        return "company_standard"
    if branch == CONTENT_CONSISTENCY_BRANCH:
        return "single_doc_internal"
    if branch == LLM_REVIEW_BRANCH:
        return "agent_semantic"
    return "single_doc_internal"


def normalize_enum(value: Any, allowed: set[str], default: str) -> str:
    text = str(value or "").strip()
    return text if text in allowed else default


def default_coverage_domain_for_category(category: str, agent_role: str = "") -> str:
    if category == "zh_language":
        return "zh_language"
    if category == "en_language":
        return "en_language"
    if category == "bilingual_consistency":
        return "bilingual"
    if agent_role == "structure":
        return "structure"
    if agent_role == "risk_classifier":
        return "data_consistency"
    return "data_consistency"


def normalize_agent_category(value: Any, *, agent_role: str = "") -> str:
    text = str(value or "").strip()
    if text in LLM_ALLOWED_CATEGORIES:
        return text
    token = enum_token(text)
    if token in AGENT_CATEGORY_SYNONYMS:
        return AGENT_CATEGORY_SYNONYMS[token]
    if agent_role in AGENT_ROLE_CATEGORY_DEFAULTS:
        return AGENT_ROLE_CATEGORY_DEFAULTS[agent_role]
    if "中文" in text:
        return "zh_language"
    if "英文" in text:
        return "en_language"
    if "双语" in text:
        return "bilingual_consistency"
    return text


def is_risk_classifier_agent_role(agent_role: str) -> bool:
    return "riskclassifier" in enum_token(agent_role)


def normalize_agent_comment_intent(value: Any, *, requires_external: bool = False) -> str:
    text = str(value or "").strip()
    if text in ALLOWED_COMMENT_INTENTS:
        return text
    token = enum_token(text)
    if token in AGENT_COMMENT_INTENT_SYNONYMS:
        return AGENT_COMMENT_INTENT_SYNONYMS[token]
    if requires_external:
        return "request_check"
    if any(marker in text for marker in ("核对", "确认")) or "check" in token:
        return "request_check"
    if any(marker in text for marker in ("修改", "修订")) or any(token_part in token for token_part in ("suggest", "change", "edit", "fix")):
        return "suggest_change"
    return text


def normalize_agent_external_evidence_type(value: Any, *, requires_external: bool = False) -> str:
    text = str(value or "").strip()
    if text in ALLOWED_EXTERNAL_EVIDENCE_TYPES:
        return text
    if not text:
        return "other" if requires_external else "none"
    lowered = text.casefold()
    if "record" in lowered or "记录" in text:
        return "record"
    if "protocol" in lowered or "方案" in text:
        return "protocol"
    if "sample" in lowered or "样品" in text:
        return "sample_info"
    if "lims" in lowered:
        return "lims"
    return "other"


def normalize_agent_review_basis(value: Any, *, requires_external: bool = False, agent_role: str = "", category: str = "") -> str:
    text = str(value or "").strip()
    if text in ALLOWED_REVIEW_BASIS:
        return text
    if requires_external:
        return "external_required"
    token = enum_token(text)
    if token in {"agentsemantic", "semantic", "llm"}:
        return "agent_semantic"
    if any(marker in text for marker in ("中文", "英文", "双语", "语言", "语法", "表达", "用词", "语义")):
        return "agent_semantic"
    if any(marker in text for marker in ("一致性", "编号", "交叉引用", "完整性", "术语", "结构", "逻辑连贯")):
        return "single_doc_internal"
    if any(marker in text for marker in ("模板", "标准", "规范", "官方名称")):
        return "company_standard"
    if category in {"zh_language", "en_language", "bilingual_consistency", "semantic_consistency"}:
        return "agent_semantic"
    if agent_role == "structure":
        return "single_doc_internal"
    return text


def normalize_agent_coverage_domain(
    value: Any,
    *,
    requires_external: bool = False,
    agent_role: str = "",
    category: str = "",
    comment_intent: str = "",
) -> str:
    text = str(value or "").strip()
    if text in ALLOWED_COVERAGE_DOMAINS:
        return text
    if requires_external:
        return "external_check"
    token = enum_token(text)
    if not token:
        return default_coverage_domain_for_category(category, agent_role)
    if token == "riskclassification":
        return default_coverage_domain_for_category(category, agent_role)
    if token in AGENT_COVERAGE_DOMAIN_SYNONYMS:
        return AGENT_COVERAGE_DOMAIN_SYNONYMS[token]
    if "中文" in text:
        return "zh_language"
    if "英文" in text:
        return "en_language"
    if "双语" in text:
        return "bilingual"
    if any(marker in text for marker in ("一致性", "编号", "引用")):
        return "data_consistency"
    if any(marker in text for marker in ("结构", "格式")):
        return "structure"
    return text


def normalize_agent_issue_payload(raw_issue: Dict[str, Any], *, agent_role: str = "") -> Dict[str, Any]:
    normalized = dict(raw_issue)
    normalized["category"] = normalize_agent_category(normalized.get("category"), agent_role=agent_role)
    normalized["severity"] = str(normalized.get("severity") or "").strip()
    normalized["confidence"] = str(normalized.get("confidence") or "").strip()
    initial_requires_external = infer_requires_external_evidence(
        requires_external_evidence=normalized.get("requires_external_evidence"),
        external_evidence_type=normalized.get("external_evidence_type"),
        review_basis=normalized.get("review_basis"),
        coverage_domain=normalized.get("coverage_domain"),
        comment_intent=normalized.get("comment_intent"),
    )
    normalized["comment_intent"] = normalize_agent_comment_intent(
        normalized.get("comment_intent"),
        requires_external=initial_requires_external,
    )
    normalized["review_basis"] = normalize_agent_review_basis(
        normalized.get("review_basis"),
        requires_external=initial_requires_external,
        agent_role=agent_role,
        category=normalized["category"],
    )
    normalized["external_evidence_type"] = normalize_agent_external_evidence_type(
        normalized.get("external_evidence_type"),
        requires_external=initial_requires_external,
    )
    normalized["coverage_domain"] = normalize_agent_coverage_domain(
        normalized.get("coverage_domain"),
        requires_external=initial_requires_external,
        agent_role=agent_role,
        category=normalized["category"],
        comment_intent=normalized["comment_intent"],
    )
    normalized["requires_external_evidence"] = infer_requires_external_evidence(
        requires_external_evidence=normalized.get("requires_external_evidence"),
        external_evidence_type=normalized.get("external_evidence_type"),
        review_basis=normalized.get("review_basis"),
        coverage_domain=normalized.get("coverage_domain"),
        comment_intent=normalized.get("comment_intent"),
    )
    return normalized


def trim_comment_field(value: str, limit: int = MAX_COMMENT_FIELD_CHARS) -> str:
    text = re.sub(r"\s+", " ", str(value or "")).strip()
    if len(text) <= limit:
        return text
    return text[: max(0, limit - 1)].rstrip() + "…"


def strip_internal_comment_locators(value: str) -> str:
    """Remove document-map locators from Word-visible comment text.

    P342 / T7R5C3 style coordinates are useful in JSON diagnostics, but they
    read like system logs in Word comments. Strip only locator-shaped prefixes
    and keep the source quote or user-facing sentence that follows.
    """

    text = str(value or "")
    text = re.sub(
        r"(?<![A-Za-z0-9])P\d{1,6}\s*(?:原文|正文|记录章节|表格|章节)\s*(?:为|写为|列出)?\s*[:：]?",
        "",
        text,
    )
    text = re.sub(
        r"(?<![A-Za-z0-9])P\d{1,6}\s*(?:（[^）]{0,40}）|\([^)]{0,40}\))\s*(?:为|写为|列出)?\s*[:：]?",
        "",
        text,
    )
    text = re.sub(r"(?<![A-Za-z0-9])P\d{1,6}\s*[:：]", "", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text.lstrip("；;，,。:： ")


def visible_comment_field(value: str, field_name: str) -> str:
    return trim_comment_field(strip_internal_comment_locators(value), comment_field_limit(field_name))


def localize_visible_issue_text(value: str) -> str:
    text = str(value or "").strip()
    if not text or has_chinese(text):
        return text

    compact = re.sub(r"\s+", " ", text).strip()
    match = re.search(r"missing space between\s+(.+?)\s+and\s+(.+?)[.。]?$", compact, flags=re.IGNORECASE)
    if match:
        return f"{match.group(1).strip()} 与 {match.group(2).strip()} 之间缺少空格。"

    match = re.search(r"duplicate adjacent token[:：]?\s*[`\"'“”‘’]?([^`\"'“”‘’]+)[`\"'“”‘’]?", compact, flags=re.IGNORECASE)
    if match:
        return f"英文存在相邻重复词：`{match.group(1).strip()}`。"

    match = re.search(
        r"(?:typographical error|typo|spelling error)[:：]?\s*[`\"'“”‘’]?([^`\"'“”‘’]+)[`\"'“”‘’]?\s+(?:should be|is misspelled and should be)\s+[`\"'“”‘’]?([^`\"'“”‘’]+)[`\"'“”‘’]?",
        compact,
        flags=re.IGNORECASE,
    )
    if match:
        return f"英文拼写错误：`{match.group(1).strip()}` 应为 `{match.group(2).strip()}`。"

    match = re.search(r"[`\"'“”‘’]([^`\"'“”‘’]{1,80})[`\"'“”‘’]\s+should be\s+[`\"'“”‘’]([^`\"'“”‘’]{1,80})[`\"'“”‘’]", compact, flags=re.IGNORECASE)
    if match:
        return f"英文表述错误：`{match.group(1).strip()}` 应为 `{match.group(2).strip()}`。"

    if any(token in compact.casefold() for token in ("grammar", "grammatical", "passive voice", "verb form")):
        return "英文语法或动词形式存在需修订的问题。"
    if any(token in compact.casefold() for token in ("mismatch", "inconsistent", "differs", "different")):
        return "原文存在前后或中英文不一致的问题。"
    return "英文原文存在需修订的问题。"


def localize_visible_suggestion_text(value: str) -> str:
    text = str(value or "").strip()
    if not text or has_chinese(text):
        return text

    compact = re.sub(r"\s+", " ", text).strip()
    match = re.search(r"insert a space between\s+(.+?)\s+and\s+(.+?)[.。]?$", compact, flags=re.IGNORECASE)
    if match:
        return f"在 {match.group(1).strip()} 和 {match.group(2).strip()} 之间补充空格。"

    replacement_patterns = (
        r"(?:replace|change|revise|correct)\s+[`\"'“”‘’]?([^`\"'“”‘’]+?)[`\"'“”‘’]?\s+(?:with|to)\s+[`\"'“”‘’]?([^`\"'“”‘’]+?)[`\"'“”‘’]?[.。]?$",
        r"(?:replace|change|revise|correct)\s+.+?\s+(?:with|to)\s+[`\"'“”‘’]([^`\"'“”‘’]+)[`\"'“”‘’][.。]?$",
    )
    for pattern in replacement_patterns:
        match = re.search(pattern, compact, flags=re.IGNORECASE)
        if match:
            source = match.group(1).strip()
            target = match.group(2).strip()
            if source and target:
                return f"将 `{source}` 改为 `{target}`。"

    if compact.casefold().startswith("add "):
        return f"按批注位置补充缺失内容：{compact[4:].strip().rstrip('.')}。"
    return "按批注位置核对并修订该处英文表述。"


def comment_field_limit(field_name: str) -> int:
    return COMMENT_FIELD_LIMITS.get(field_name, MAX_COMMENT_FIELD_CHARS)


def is_global_summary_issue(issue: Dict[str, Any]) -> bool:
    if str(issue.get("comment_intent") or "").strip() == "global_summary":
        return True
    if str(issue.get("location_kind") or "").strip() in {"global", "property", "section", "footer"}:
        return True
    if str(issue.get("anchor_locator") or "").strip() in {"section_properties", "footer_properties"}:
        return True
    location = str(issue.get("location") or "").strip()
    if location in {"全文审核意见", "全文人工复核"}:
        return True
    if str(issue.get("paragraph_index") or "").strip():
        return False
    combined = "\n".join(str(issue.get(key) or "") for key in ("original", "issue", "evidence"))
    return any(marker in combined for marker in ("全文", "全篇", "整体", "全局"))


def is_explicit_agent_global_summary_issue(raw_issue: Dict[str, Any], comment_intent: str = "") -> bool:
    return str(comment_intent or raw_issue.get("comment_intent") or "").strip() == "global_summary"


def agent_issue_requires_unit_locator(raw_issue: Dict[str, Any], comment_intent: str = "") -> bool:
    return True


def has_explicit_source_evidence(raw_issue: Dict[str, Any]) -> bool:
    evidence = str(raw_issue.get("evidence") or "").strip()
    if not evidence:
        return False
    normalized = re.sub(r"[\s,，。.;；:：]+", "", evidence).casefold()
    generic_values = {
        "见原文",
        "原文可见",
        "根据原文",
        "同上",
        "同原文",
        "需核对",
        "待确认",
        "seesource",
        "source",
        "originaltext",
        "asabove",
        "checkneeded",
    }
    if normalized in generic_values:
        return False
    if len(normalized) < 4:
        return False
    return True


def is_actionable_concise_suggestion(suggestion: str) -> bool:
    raw_text = re.sub(r"\s+", " ", str(suggestion or "")).strip()
    text = re.sub(r"\s+", "", raw_text)
    lowered = raw_text.casefold()
    if not text or len(raw_text) > 220:
        return False
    if has_multi_option_suggestion(raw_text):
        return False
    if re.search(r"\b(no change needed|no changes? required|无需修改|不需要修改|无需变更)\b", lowered):
        return False
    vague_exact = {
        "请核对",
        "请检查",
        "请确认",
        "请核对是否一致",
        "请检查是否一致",
        "建议优化",
        "建议润色",
        "请进一步核对",
        "pleasecheck",
        "pleaseconfirm",
        "checkconsistency",
        "reviewmanually",
        "considerrephrasing",
        "improvereadability",
    }
    stripped = text.rstrip("。.;；")
    if stripped in vague_exact:
        return False
    if stripped.startswith(("请核对是否", "请检查是否", "请确认是否")) and len(stripped) <= 14:
        return False
    if re.fullmatch(r"(please\s+)?(check|confirm|review)(\s+(it|this|consistency|accuracy))?\.?", lowered):
        return False
    actionable_markers = (
        "改为",
        "修改为",
        "统一为",
        "补充",
        "删除",
        "替换",
        "更正",
        "调整为",
        "保持一致",
        "明确",
        "统一",
    )
    english_action = re.search(
        r"\b(change|replace|revise|correct|insert|add|delete|remove|standardize|align|unify|use|make)\b",
        lowered,
    )
    english_specific = re.search(r"\b(add|insert)\s+(a\s+)?space\b|\bchange\s+.+\s+to\s+.+|\breplace\s+.+\s+with\s+.+", lowered)
    return any(marker in text for marker in actionable_markers) or bool(english_action or english_specific)


def bare_replacement_suggestion(raw_issue: Dict[str, Any]) -> str:
    suggestion = re.sub(r"\s+", " ", str(raw_issue.get("suggestion") or "")).strip()
    if not suggestion or len(suggestion) > 180:
        return ""
    if is_actionable_concise_suggestion(suggestion):
        return ""
    if has_multi_option_suggestion(suggestion):
        return ""
    lowered_suggestion = suggestion.casefold()
    if re.search(r"\b(or|and/or|check|confirm|review|verify|consider|ensure)\b", lowered_suggestion):
        return ""
    if any(token in suggestion for token in ("或", "或者", "核对", "确认", "检查", "请", "建议")):
        return ""

    original = strip_agent_text_label(str(raw_issue.get("original") or "")).strip(" \t\r\n:：;；,，。.-")
    original_norm = normalized_text(original).casefold()
    suggestion_norm = normalized_text(suggestion).casefold()
    if not original_norm or len(suggestion_norm) < 2 or suggestion_norm == original_norm:
        return ""

    combined = "\n".join(str(raw_issue.get(key) or "") for key in ("issue", "evidence", "original"))
    lowered = combined.casefold()
    direct_markers = (
        "should be",
        "should use",
        "incorrect",
        "error",
        "wrong",
        "duplicate",
        "misspell",
        "spelling",
        "typographical",
        "grammar",
        "grammatical",
        "word form",
        "contradict",
        "mismatch",
        "does not match",
        "错误",
        "应为",
        "应改为",
        "重复",
        "错译",
        "不一致",
        "矛盾",
        "拼写",
        "语法",
        "词性",
    )
    if not any(marker in lowered for marker in direct_markers):
        return ""
    return suggestion


def has_actionable_agent_suggestion(raw_issue: Dict[str, Any]) -> bool:
    suggestion = str(raw_issue.get("suggestion") or "")
    return is_actionable_concise_suggestion(suggestion) or bool(bare_replacement_suggestion(raw_issue))


def has_multi_option_suggestion(suggestion: str) -> bool:
    text = re.sub(r"\s+", " ", str(suggestion or "")).strip()
    if not text:
        return False
    lowered = text.casefold()
    if re.search(r"\beither\b.{0,120}\bor\b", lowered):
        return True
    if re.search(r"[`\"'“”][^`\"'“”]{2,80}[`\"'“”]\s+or\s+[`\"'“”][^`\"'“”]{2,80}[`\"'“”]", text, flags=re.IGNORECASE):
        return True
    if re.search(r"\b(change|replace|revise|correct|align|standardize|use|process|treat)\b.{0,220}\bor\b", lowered):
        return True
    if re.search(r"\be\.g\.?,?\s+.{0,180}\bor\s+(add|use|ensure|revise|change|replace|insert)\b", lowered):
        return True
    if re.search(r"(?:改为|修改为|更正为|统一为|替换为).{0,80}(?:或|或者).{0,80}", text):
        return True
    return False


def has_weak_agent_anchor(anchor_text: str) -> bool:
    text = str(anchor_text or "").strip()
    compact = normalized_text(text).casefold()
    if len(compact) < 3:
        return True
    if re.fullmatch(r"[Nn]\s*=\s*\d+", text):
        return True
    if re.fullmatch(r"[A-Za-z]+", text):
        lowered = text.casefold()
        if lowered in {
            "a",
            "an",
            "and",
            "or",
            "the",
            "to",
            "of",
            "in",
            "on",
            "for",
            "with",
            "by",
            "as",
            "is",
            "are",
            "was",
            "were",
            "than",
            "less",
            "from",
            "into",
            "that",
            "this",
            "these",
            "those",
            "not",
            "no",
            "high",
            "low",
            "data",
            "table",
            "report",
            "solution",
            "method",
            "methods",
            "protocol",
            "protocols",
            "experiment",
            "experiments",
            "polymer",
            "aggregate",
            "aggregates",
            "peak",
            "peaks",
        }:
            return True
    if compact in {
        "条件",
        "条件下",
        "结果",
        "结果表明",
        "显示",
        "详见",
        "见表",
        "原文",
        "报告",
        "表格",
        "数据",
        "方法",
        "方法确认",
        "试验",
        "实验",
        "聚合体",
        "单体",
    }:
        return True
    return False


def is_unsafe_embedded_numeric_anchor(
    paragraph_text: str,
    span: Optional[Dict[str, Any]],
    anchor_text: str,
    raw_issue: Dict[str, Any],
) -> bool:
    """Reject short numeric anchors that only matched inside a larger value.

    Weak reviewers often cite a bare value such as "109" while pointing to the
    neighboring "109%" cell. For Word comments this is worse than filtering:
    the comment appears on the value that already has the requested symbol.
    """

    if not paragraph_text or not isinstance(span, dict):
        return False
    anchor = str(anchor_text or "").strip()
    if not re.fullmatch(r"\d+(?:\.\d+)?", anchor):
        return False
    combined = "\n".join(str(raw_issue.get(key) or "") for key in ("issue", "suggestion", "evidence")).casefold()
    if "%" not in combined and "percent" not in combined and "百分" not in combined:
        return False
    try:
        start = int(span.get("start"))
        end = int(span.get("end"))
    except (TypeError, ValueError):
        return False
    before = paragraph_text[start - 1] if start > 0 else ""
    after = paragraph_text[end] if 0 <= end < len(paragraph_text) else ""
    return before in {"%", "％"} or after in {"%", "％"}


def has_speculative_language(raw_issue: Dict[str, Any]) -> bool:
    combined = "\n".join(
        str(raw_issue.get(key) or "")
        for key in ("issue", "suggestion", "evidence")
    ).casefold()
    return any(
        marker in combined
        for marker in (
            "可能",
            "疑似",
            "似乎",
            "不确定",
            "可能被误解",
            "可能会",
            "建议确认",
            "建议核对",
            "may be",
            "might be",
            "possibly",
            "appears to",
            "seems to",
            "could be interpreted",
            "may be interpreted",
            "can be interpreted",
            "could be mistaken",
            "may cause confusion",
            "need to confirm",
        )
    )


def is_external_brand_or_product_name_claim(raw_issue: Dict[str, Any]) -> bool:
    combined = "\n".join(str(raw_issue.get(key) or "") for key in ("issue", "suggestion", "evidence"))
    lowered = combined.casefold()
    if str(raw_issue.get("category") or "") not in {"en_language", "semantic_consistency"}:
        return False
    if not any(
        marker in lowered
        for marker in (
            "thermo fisher",
            "sciex",
            "manufacturer",
            "vendor",
            "official",
            "brand",
            "product is named",
            "system name uses",
            "instrument model",
            "model name",
            "设备厂家",
            "厂家",
            "制造商",
            "供应商",
            "品牌",
            "设备型号",
            "型号",
            "产品名",
            "官方",
        )
    ):
        return False
    if not any(marker in lowered for marker in ("spelling", "typographical", "misspelled", "should be", "拼写", "应为", "正确写法")):
        shape_or_direct_edit_claim = any(
            marker in lowered
            for marker in (
                "letter",
                "letters",
                "mixed up",
                "confused",
                "confusion",
                "字母",
                "混淆",
                "改为",
                "建议改为",
                "正确型号",
            )
        )
        if not shape_or_direct_edit_claim:
            return False
    return True


def has_approved_term_or_template_basis(raw_issue: Dict[str, Any]) -> bool:
    """Return true only for runner-verified approved rule sources.

    Agent metadata or prose such as `company_standard` is not itself proof of a
    controlled template, terminology source, or company rule. Deterministic
    runner branches encode the currently verified company-standard rules.
    """

    return False


def has_explicit_controlled_term_source(raw_issue: Dict[str, Any]) -> bool:
    """Return true only for runner-verified controlled sources.

    Agent prose or metadata such as `company_standard` is not proof. Until a
    glossary/template source is loaded and verified by the runner, model-only
    official terminology rewrites stay external.
    """

    return False


def is_external_official_terminology_rewrite(raw_issue: Dict[str, Any]) -> bool:
    """Filter official-term rewrites that need a controlled terminology source."""

    if has_explicit_controlled_term_source(raw_issue):
        return False
    combined = "\n".join(str(raw_issue.get(key) or "") for key in ("original", "issue", "suggestion", "evidence"))
    lowered = combined.casefold()
    acronym_authority_claim = (
        any(
            token in combined
            for token in (
                "缩写",
                "缩略语",
                "英文全称",
                "正确全称",
                "标准全称",
                "官方全称",
                "标准英文名称",
                "官方中文名称",
                "官方英文名称",
            )
        )
        or any(
            token in lowered
            for token in (
                "acronym",
                "abbreviation",
                "correct full name",
                "standard full name",
                "official full name",
                "official english name",
                "official chinese name",
                "regulatory abbreviation",
            )
        )
    )
    domain_claim = any(
        token in combined
        for token in (
            "包材",
            "包装材料",
            "包装系统",
            "容器密封",
            "物料名称",
            "物料描述",
            "试剂名称",
            "耗材名称",
            "辅料名称",
            "注册名称",
            "药典名称",
            "官方名称",
            "官方术语",
            "标准名称",
            "官方",
        )
    ) or any(
        token in lowered
        for token in (
            "packaging material",
            "container closure",
            "official term",
            "official name",
            "registered name",
            "controlled term",
            "standard name",
            "material name",
            "official",
        )
    )
    rewrite_claim = any(
        token in combined
        for token in (
            "统一表述为",
            "统一为",
            "应表述为",
            "应改为",
            "改为",
            "标准表述",
            "规范表述",
            "正确全称",
            "正确名称",
            "更正为",
        )
    ) or any(
        token in lowered
        for token in (
            "rewrite as",
            "rename to",
            "standardize as",
            "should be named",
            "should be written as",
            "correct full name",
            "correct name",
            "revise to",
        )
    )
    if acronym_authority_claim and rewrite_claim:
        return True
    if not domain_claim or not rewrite_claim:
        return False
    if any(token in combined for token in ("数值", "温度", "条件", "方法", "DS", "DP", "原液", "制剂不一致")):
        return False
    return True


def is_unapproved_terminology_standardization(raw_issue: Dict[str, Any]) -> bool:
    if has_approved_term_or_template_basis(raw_issue):
        return False
    combined = "\n".join(str(raw_issue.get(key) or "") for key in ("original", "issue", "suggestion", "evidence"))
    lowered = combined.casefold()
    coverage = str(raw_issue.get("coverage_domain") or "").strip()
    terminology_signal = (
        coverage == "terminology"
        or any(
            token in lowered
            for token in (
                "terminology inconsistency",
                "term inconsistency",
                "terminology standardization",
                "terminology preference",
                "established terminology",
                "abbreviation table",
                "document elsewhere",
                "rest of the document",
                "for consistency with",
            )
        )
        or any(token in combined for token in ("术语不一致", "术语统一", "统一术语", "缩略语表", "全文其他位置"))
    )
    if not terminology_signal:
        return False
    direct_mistranslation = any(
        token in lowered
        for token in (
            "mistranslation",
            "wrong translation",
            "incorrect translation",
            "means",
            "does not mean",
        )
    ) or any(token in combined for token in ("错译", "误译", "意思是", "并非"))
    if direct_mistranslation and not any(
        token in lowered for token in ("terminology inconsistency", "term inconsistency", "abbreviation table", "rest of the document")
    ):
        return False
    return True


def is_high_risk_formula_semantic_issue(raw_issue: Dict[str, Any]) -> bool:
    original = str(raw_issue.get("original") or "")
    combined = "\n".join(
        str(raw_issue.get(key) or "")
        for key in ("issue", "suggestion", "evidence")
    )
    basis = str(raw_issue.get("review_basis") or "").strip()
    if basis == "company_standard" and has_approved_term_or_template_basis(raw_issue):
        return False
    formula_like = bool(re.search(r"[=<>|±*/]", original)) or (
        "%" in original and bool(re.search(r"[A-Za-z]", original))
    )
    if not formula_like:
        return False
    return any(
        token in combined.casefold()
        for token in (
            "公式",
            "绝对值",
            "absolute",
            "百分号",
            "%符号",
            "单位",
            "symbol",
            "formula",
        )
    )


def is_objective_formula_case_issue(raw_issue: Dict[str, Any]) -> bool:
    """Allow anchored chemical formula case fixes such as Nacl -> NaCl.

    These are not stylistic capitalization preferences. They are compact,
    source-local spelling/case defects that weak models should be allowed to
    surface when the original token and corrected token differ only by case.
    """

    original = str(raw_issue.get("original") or "")
    combined = "\n".join(
        str(raw_issue.get(key) or "")
        for key in ("issue", "suggestion", "evidence")
    )
    if not any(marker in combined.casefold() for marker in ("化学式", "chemical formula")):
        return False
    token_pattern = r"(?<![A-Za-z0-9])([A-Za-z][A-Za-z0-9]{1,9})(?![A-Za-z0-9])"
    original_tokens = re.findall(token_pattern, original)
    suggested_tokens = re.findall(token_pattern, combined)
    for source_token in original_tokens:
        for target_token in suggested_tokens:
            if source_token != target_token and source_token.casefold() == target_token.casefold():
                return True
    return False


def has_repeated_whitespace_span(value: str) -> bool:
    normalized_original = str(value or "").replace("\u00a0", " ").replace("\u3000", " ").replace("\t", "  ")
    has_repeated_inline_space = bool(re.search(r"\S {2,}\S", normalized_original))
    has_repeated_leading_space = bool(re.search(r"(?m)^ {2,}\S", normalized_original))
    return has_repeated_inline_space or has_repeated_leading_space


def is_objective_repeated_whitespace_issue(raw_issue: Dict[str, Any]) -> bool:
    """Allow exact repeated-whitespace defects while still filtering style-only spacing claims."""

    category = str(raw_issue.get("category") or "").strip()
    if category not in {"en_language", "zh_language", "bilingual_consistency", "semantic_consistency"}:
        return False
    original = str(raw_issue.get("original") or "")
    if not original:
        return False
    if not has_repeated_whitespace_span(original):
        return False
    combined = "\n".join(
        str(raw_issue.get(key) or "")
        for key in ("issue", "suggestion", "evidence", "original")
    )
    lowered = combined.casefold()
    has_space_marker = any(
        marker in lowered
        for marker in (
            "extra space",
            "extra spaces",
            "multiple spaces",
            "consecutive spaces",
            "repeated spaces",
            "redundant spaces",
            "重复空格",
            "连续空格",
            "多个空格",
            "多余空格",
            "空格过多",
            "前导空格",
            "缩进",
        )
    )
    if not has_space_marker:
        return False
    return has_actionable_agent_suggestion(raw_issue)


def is_objective_source_local_typo_issue(raw_issue: Dict[str, Any]) -> bool:
    """Allow exact, source-local typo findings while still filtering style noise.

    Weak reviewers sometimes wrap an objective typo in hedging language such as
    "appears to be". If the finding is otherwise anchored to a concrete source
    typo with a direct correction, keep it eligible for the Word gate. This does
    not make external brand/model spelling claims acceptable.
    """

    category = str(raw_issue.get("category") or "").strip()
    if category not in {"en_language", "zh_language", "bilingual_consistency", "semantic_consistency"}:
        return False
    if is_external_brand_or_product_name_claim(raw_issue):
        return False
    if is_unapproved_terminology_standardization(raw_issue):
        return False
    original = strip_agent_text_label(str(raw_issue.get("original") or "")).strip(" \t\r\n:：;；,，。.-")
    if not original or has_weak_agent_anchor(original):
        return False
    combined = "\n".join(
        str(raw_issue.get(key) or "")
        for key in ("issue", "suggestion", "evidence", "original")
    )
    lowered = combined.casefold()
    typo_markers = (
        "typo",
        "typographical",
        "misspell",
        "spelling error",
        "spelled",
        "missing letter",
        "missing word",
        "缺字",
        "漏字",
        "缺少",
        "拼写",
        "错拼",
    )
    if not any(marker in lowered for marker in typo_markers):
        return False
    if not has_actionable_agent_suggestion(raw_issue):
        return False
    return True


def is_objective_repeated_whitespace_anchor(raw_issue: Dict[str, Any], anchor_text: str) -> bool:
    return is_objective_repeated_whitespace_issue(raw_issue) and has_repeated_whitespace_span(anchor_text)


def is_direct_internal_source_contradiction(issue: Dict[str, Any]) -> bool:
    """Return true for high-confidence issues proven by the document itself.

    The risk classifier is intentionally conservative, but it must not hide
    exact, anchored contradictions where one reviewer already cites the local
    Chinese/English or method/condition mismatch and proposes a direct edit.
    """

    if str(issue.get("branch") or "") != LLM_REVIEW_BRANCH:
        return False
    if str(issue.get("confidence") or "") != "high":
        return False
    if issue.get("requires_external_evidence"):
        return False
    if str(issue.get("comment_intent") or "") == "request_check":
        return False
    if str(issue.get("review_basis") or "") not in {"single_doc_internal", "agent_semantic", ""}:
        return False
    if not issue.get("anchor_span") or not str(issue.get("anchor_text") or "").strip():
        return False

    combined = "\n".join(
        str(issue.get(key) or "")
        for key in ("original", "issue", "suggestion", "evidence")
    )
    lowered = combined.casefold()
    has_internal_pair = (
        ("中文" in combined and "英文" in combined)
        or ("chinese" in lowered and "english" in lowered)
        or ("nrce" in lowered and "reduced ce-sds" in lowered)
    )
    has_contradiction_marker = any(
        marker in lowered
        for marker in (
            "不一致",
            "不符",
            "不同",
            "矛盾",
            "错误复制",
            "误粘贴",
            "copy",
            "copied",
            "instead of",
            "mismatch",
            "inconsistent",
            "differs",
        )
    )
    has_direct_suggestion = any(
        marker in lowered
        for marker in (
            "应改为",
            "改为",
            "更正为",
            "replace",
            "change",
            "revise",
            "should be",
            "一致",
        )
    )
    soft_check_only = any(
        marker in lowered
        for marker in (
            "需核对",
            "核对原始",
            "核对试验",
            "confirm against",
            "after verifying",
            "source record",
        )
    ) and not any(marker in lowered for marker in ("错误复制", "误粘贴", "copy", "copied"))
    return has_internal_pair and has_contradiction_marker and has_direct_suggestion and not soft_check_only


def normalized_agent_visible_severity(severity: str, raw_issue: Dict[str, Any]) -> str:
    value = str(severity or "").strip()
    if value not in {"关键", "主要"}:
        return value
    combined = "\n".join(str(raw_issue.get(key) or "") for key in ("original", "issue", "suggestion", "evidence"))
    lowered = combined.casefold()
    if value == "关键" and ("浊度" in combined or "turbidity" in lowered or "urbidity" in lowered):
        if any(token in lowered for token in ("less urbidity", "lessurbidity", "standard solution")):
            return "主要"
    if value == "关键" and "nrce" in lowered and "reduced ce-sds" in lowered:
        return "主要"
    if any(token in combined for token in ("搭配不当", "结构冗余", "成分冗余", "风格欠佳")):
        return "次要"
    low_scope = any(
        token in lowered
        for token in (
            "missing period",
            "extra period",
            "sentence-ending",
            "missing space",
            "extra space",
            "spacing",
            "punctuation",
            "标点",
            "句号",
            "分号",
            "空格",
            "多余",
            "冗余",
            "赘余",
        )
    )
    high_impact = bool(
        re.search(
            r"\b(data|value|temperature|condition|method|reduced|non-reduced|ds|dp)\b",
            lowered,
        )
    ) or any(
        token in combined
        for token in ("原液", "制剂", "成品", "数据", "温度", "条件", "方法")
    )
    if low_scope and not high_impact:
        return "次要"
    return value


def passes_user_visible_agent_quality_gate(
    *,
    raw_issue: Dict[str, Any],
    has_stable_anchor: bool,
    confidence: str,
    requires_external: bool,
    comment_intent: str,
) -> bool:
    """Gate agent findings before they can become user-visible review items."""

    if confidence == "low":
        return False
    objective_typo = is_objective_source_local_typo_issue(raw_issue)
    if has_speculative_language(raw_issue) and not objective_typo:
        return False
    if is_external_brand_or_product_name_claim(raw_issue):
        return False
    if is_external_official_terminology_rewrite(raw_issue):
        return False
    if is_unapproved_terminology_standardization(raw_issue):
        return False
    if is_high_risk_formula_semantic_issue(raw_issue):
        return False
    if not has_explicit_source_evidence(raw_issue):
        return False
    if not has_actionable_agent_suggestion(raw_issue):
        return False
    if requires_external or comment_intent == "request_check":
        return False
    if has_stable_anchor:
        return True
    return False


def is_current_date_regulatory_status_noise(issue: Dict[str, Any]) -> bool:
    """Filter model-only checks about public regulatory version effective status.

    These findings are usually generated from the model/runtime current date,
    not from a document-internal contradiction. They are not useful as Word
    comments unless a separate configured rule or document evidence proves a
    conflict, which the agent JSON contract does not provide.
    """

    if not infer_requires_external_evidence(
        requires_external_evidence=issue.get("requires_external_evidence"),
        external_evidence_type=issue.get("external_evidence_type"),
        review_basis=issue.get("review_basis"),
        coverage_domain=issue.get("coverage_domain"),
        comment_intent=issue.get("comment_intent"),
    ):
        return False

    combined = "\n".join(
        str(issue.get(key) or "")
        for key in ("original", "issue", "suggestion", "evidence")
    )
    if "当前日期" not in combined:
        return False
    if not any(token in combined for token in ("正式实施", "正式发布", "生效", "施行")):
        return False
    if not any(token in combined for token in ("药典", "法规", "指导原则", "标准")):
        return False
    return True


def build_comment_text(
    issue: str,
    suggestion: str,
    severity: str,
    *,
    issue_type: str = "",
    original: str = "",
    evidence: str = "",
    status: str = ISSUE_STATUS_CONFIRMED,
    comment_intent: str = "suggest_change",
    review_basis: str = "",
    external_evidence_type: str = "none",
) -> str:
    check_required = status in CHECK_STATUSES or comment_intent == "request_check"
    issue_label = "核对项" if check_required else "发现"
    suggestion_label = "核对建议" if check_required else "建议"

    parts = []
    if check_required and "需核对" not in issue and "人工核对" not in issue:
        issue = f"{issue}（需核对后确认）"
    parts.append(f"{issue_label}：{visible_comment_field(localize_visible_issue_text(issue), 'issue')}")
    if original:
        parts.append(f"原文：{visible_comment_field(original, 'original')}")
    if suggestion:
        if check_required and "核对" not in suggestion:
            suggestion = f"先核对相关依据；{suggestion}"
        parts.append(f"{suggestion_label}：{visible_comment_field(localize_visible_suggestion_text(suggestion), 'suggestion')}")
    return "\n".join(parts)


def make_issue(
    issue_id: int,
    *,
    rule_id: str,
    branch: str,
    paragraph_index: Optional[int],
    location: str,
    original: str,
    issue: str,
    suggestion: str,
    severity: str,
    evidence: str,
    document_zone: str = "body",
    location_kind: str = "paragraph",
    anchor_locator: str = "",
    anchor_span: Optional[Dict[str, Any]] = None,
    anchor_text: str = "",
    status: str = "confirmed",
    comments_added: int = 1,
    confidence: str = "high",
    match_method: str = "exact",
    notes: str = "",
    requires_external_evidence: bool = False,
    external_evidence_type: str = "none",
    coverage_domain: str = "",
    review_basis: str = "",
    comment_intent: str = "",
    comment_visibility: str = "word_comment",
) -> Dict[str, Any]:
    meta = BRANCH_META[branch]
    requires_external = infer_requires_external_evidence(
        requires_external_evidence=requires_external_evidence,
        external_evidence_type=external_evidence_type,
        review_basis=review_basis,
        coverage_domain=coverage_domain,
        comment_intent=comment_intent,
    )
    external_type = normalize_enum(
        external_evidence_type,
        ALLOWED_EXTERNAL_EVIDENCE_TYPES,
        "other" if requires_external else "none",
    )
    if requires_external and external_type == "none":
        external_type = "other"
    normalized_status = normalize_status(status, requires_external_evidence=requires_external)
    normalized_review_basis = normalize_enum(
        review_basis,
        ALLOWED_REVIEW_BASIS,
        "external_required" if requires_external else default_review_basis(branch),
    )
    normalized_coverage_domain = normalize_enum(
        coverage_domain,
        ALLOWED_COVERAGE_DOMAINS,
        "external_check" if requires_external else default_coverage_domain(branch),
    )
    normalized_comment_intent = normalize_enum(
        comment_intent,
        ALLOWED_COMMENT_INTENTS,
        "request_check" if normalized_status in CHECK_STATUSES else "suggest_change",
    )
    if requires_external:
        normalized_comment_intent = "request_check"
        normalized_status = ISSUE_STATUS_NEEDS_USER_CHECK
        normalized_review_basis = "external_required"
        normalized_coverage_domain = "external_check"
    normalized_comment_visibility = comment_visibility or "word_comment"
    normalized_comments_added = int(comments_added or 0)
    if normalized_comments_added <= 0 and normalized_comment_visibility == "word_comment":
        normalized_comment_visibility = "internal"
    if normalized_comment_visibility != "word_comment":
        normalized_comments_added = 0
    issue_data: Dict[str, Any] = {
        "id": f"issue-{issue_id:04d}",
        "rule_id": rule_id,
        "type": meta["type"],
        "branch": branch,
        "agent_role": meta["agent_role"],
        "document_zone": document_zone,
        "location_kind": location_kind,
        "location": location,
        "anchor_locator": anchor_locator,
        "anchor_span": anchor_span,
        "anchor_text": anchor_text or original,
        "original": original,
        "issue": issue,
        "severity": severity,
        "suggestion": suggestion,
        "comment_text": build_comment_text(
            issue,
            suggestion,
            severity,
            issue_type=meta["type"],
            original=original,
            evidence=evidence,
            status=normalized_status,
            comment_intent=normalized_comment_intent,
            review_basis=normalized_review_basis,
            external_evidence_type=external_type,
        ),
        "evidence": evidence,
        "match_method": match_method,
        "preexisting_comment_count": 0,
        "comments_added": normalized_comments_added,
        "confidence": confidence,
        "status": normalized_status,
        "comment_visibility": normalized_comment_visibility,
        "requires_external_evidence": requires_external,
        "external_evidence_type": external_type,
        "coverage_domain": normalized_coverage_domain,
        "review_basis": normalized_review_basis,
        "comment_intent": normalized_comment_intent,
        "source": "qa-file-reviewer",
        "source_agent": meta["agent_role"],
        "notes": notes,
    }
    if paragraph_index is not None and not anchor_locator:
        issue_data["anchor_locator"] = f"paragraph={paragraph_index}"
    if anchor_span is None:
        issue_data["anchor_span"] = None
    return issue_data


def issue_anchor_length(issue: Dict[str, Any]) -> int:
    span = issue.get("anchor_span")
    if isinstance(span, dict):
        try:
            start = int(span.get("start"))
            end = int(span.get("end"))
        except (TypeError, ValueError):
            start = end = -1
        if start >= 0 and end > start:
            return end - start
    original = str(issue.get("original") or issue.get("anchor_text") or "")
    return max(len(original), 9999)


def issue_source_family(issue: Dict[str, Any]) -> str:
    source_agent = str(issue.get("source_agent") or "").strip()
    if source_agent == "structure":
        return "structure"
    if source_agent == "bilingual_consistency":
        return "bilingual"
    return "textual"


def normalized_issue_target(issue: Dict[str, Any]) -> str:
    suggestion = str(issue.get("suggestion") or "")
    original_norm = normalized_text(str(issue.get("original") or issue.get("anchor_text") or "")).casefold()
    replacement_match = re.search(r"(?:改为|修改为|统一为|建议改为|建议修改为)[:：]?\s*[\"'“”‘’]?([^，。,；;\"'“”‘’]+)", suggestion)
    if replacement_match:
        replacement = normalized_text(replacement_match.group(1)).casefold()
        if replacement and replacement != original_norm:
            return replacement
    quoted = re.findall(r"[\"'“”‘’]([^\"'“”‘’]{2,})[\"'“”‘’]", suggestion)
    candidates = [
        normalized_text(candidate).casefold()
        for candidate in quoted
        if normalized_text(candidate).casefold() and normalized_text(candidate).casefold() != original_norm
    ]
    if candidates:
        candidates.sort(key=len)
        return candidates[0]
    match = re.search(r"(?:改为|修改为|统一为|建议改为|建议修改为)[:：]?\s*(.+)", suggestion)
    if match:
        return normalized_text(match.group(1)).casefold()
    return normalized_text(suggestion).casefold()


def issue_replacement_signatures(issue: Dict[str, Any]) -> set[Tuple[str, str]]:
    suggestion = str(issue.get("suggestion") or "")
    signatures: set[Tuple[str, str]] = set()
    quote = r"[\"'“”‘’]"
    patterns = [
        rf"{quote}([^\"'“”‘’]{{1,40}}){quote}\s*(?:改为|修改为|更正为|统一为)\s*{quote}([^\"'“”‘’]{{1,40}}){quote}",
        rf"(?:将|把)\s*{quote}([^\"'“”‘’]{{1,40}}){quote}\s*(?:改为|修改为|更正为|统一为)\s*{quote}([^\"'“”‘’]{{1,40}}){quote}",
    ]
    for pattern in patterns:
        for before, after in re.findall(pattern, suggestion):
            before_norm = normalized_text(before).casefold()
            after_norm = normalized_text(after).casefold()
            if before_norm and after_norm and before_norm != after_norm:
                signatures.add((before_norm, after_norm))
    return signatures


def issue_location_key(issue: Dict[str, Any]) -> str:
    anchor_locator = normalized_text(str(issue.get("anchor_locator") or "")).casefold()
    if anchor_locator:
        return anchor_locator
    paragraph_index = str(issue.get("paragraph_index") or "").strip()
    if paragraph_index:
        return f"paragraph={paragraph_index}"
    return normalized_text(str(issue.get("location") or "")).casefold()


def issue_semantic_dedupe_key(issue: Dict[str, Any]) -> Optional[Tuple[str, str, str]]:
    if str(issue.get("branch") or "") != LLM_REVIEW_BRANCH:
        return None
    location = issue_location_key(issue)
    original = normalized_text(str(issue.get("original") or issue.get("anchor_text") or "")).casefold()
    target = normalized_issue_target(issue)
    if len(location) < 2 or len(original) < 2 or len(target) < 2:
        return None
    return (location, original, target)


def issue_origin_dedupe_key(issue: Dict[str, Any]) -> Optional[Tuple[str, str, str]]:
    if str(issue.get("branch") or "") != LLM_REVIEW_BRANCH:
        return None
    location = issue_location_key(issue)
    original = normalized_text(str(issue.get("original") or issue.get("anchor_text") or "")).casefold()
    family = issue_source_family(issue)
    if family != "textual" or len(location) < 2 or len(original) < 2:
        return None
    return (family, location, original)


def is_secondary_textual_agent(issue: Dict[str, Any]) -> bool:
    return str(issue.get("source_agent") or "").strip() in {"risk_classifier", "semantic_consistency"}


def issues_have_contained_original_with_same_target(first: Dict[str, Any], second: Dict[str, Any]) -> bool:
    if str(first.get("branch") or "") != LLM_REVIEW_BRANCH or str(second.get("branch") or "") != LLM_REVIEW_BRANCH:
        return False
    if issue_location_key(first) != issue_location_key(second):
        return False
    first_target = normalized_issue_target(first)
    second_target = normalized_issue_target(second)
    same_target = len(first_target) >= 2 and first_target == second_target
    same_replacement = bool(issue_replacement_signatures(first) & issue_replacement_signatures(second))
    if not same_target and not same_replacement:
        return False
    first_original = normalized_text(str(first.get("original") or first.get("anchor_text") or "")).casefold()
    second_original = normalized_text(str(second.get("original") or second.get("anchor_text") or "")).casefold()
    if len(first_original) < 4 or len(second_original) < 4 or first_original == second_original:
        return False
    return first_original in second_original or second_original in first_original


def prefers_issue(candidate: Dict[str, Any], current: Dict[str, Any]) -> bool:
    confidence_rank = {"high": 3, "medium": 2, "low": 1}
    source_rank = {"risk_classifier": 0, "semantic_consistency": 1}

    def score(issue: Dict[str, Any]) -> Tuple[int, int, int, int, int, int]:
        source_agent = str(issue.get("source_agent") or "").strip()
        return (
            1 if str(issue.get("status") or "") == ISSUE_STATUS_CONFIRMED else 0,
            1 if not issue.get("requires_external_evidence") else 0,
            1 if str(issue.get("comments_added") or 0) else 0,
            confidence_rank.get(str(issue.get("confidence") or "").strip(), 0),
            source_rank.get(source_agent, 2),
            -issue_anchor_length(issue),
        )

    return score(candidate) > score(current)


def prefers_contained_duplicate_issue(candidate: Dict[str, Any], current: Dict[str, Any]) -> bool:
    confidence_rank = {"high": 3, "medium": 2, "low": 1}

    def score(issue: Dict[str, Any]) -> Tuple[int, int, int, int, int]:
        return (
            1 if str(issue.get("status") or "") == ISSUE_STATUS_CONFIRMED else 0,
            1 if not issue.get("requires_external_evidence") else 0,
            1 if str(issue.get("comments_added") or 0) else 0,
            confidence_rank.get(str(issue.get("confidence") or "").strip(), 0),
            -issue_anchor_length(issue),
        )

    return score(candidate) > score(current)


def unique_issues(issues: Iterable[Dict[str, Any]]) -> List[Dict[str, Any]]:
    seen = set()
    semantic_seen: Dict[Tuple[str, str, str], int] = {}
    origin_seen: Dict[Tuple[str, str, str], int] = {}
    result: List[Dict[str, Any]] = []
    for item in issues:
        key = (
            item.get("rule_id"),
            item.get("type"),
            item.get("document_zone"),
            item.get("location_kind"),
            item.get("location"),
            item.get("anchor_locator"),
            item.get("original"),
            item.get("severity"),
        )
        if key in seen:
            continue
        semantic_key = issue_semantic_dedupe_key(item)
        if semantic_key is not None and semantic_key in semantic_seen:
            index = semantic_seen[semantic_key]
            if prefers_issue(item, result[index]):
                result[index] = item
            continue
        origin_key = issue_origin_dedupe_key(item)
        if origin_key is not None and origin_key in origin_seen:
            index = origin_seen[origin_key]
            existing = result[index]
            if is_secondary_textual_agent(item) or is_secondary_textual_agent(existing):
                if prefers_issue(item, existing):
                    result[index] = item
                continue
        contained_duplicate_index = next(
            (
                index
                for index, existing in enumerate(result)
                if issues_have_contained_original_with_same_target(item, existing)
            ),
            None,
        )
        if contained_duplicate_index is not None:
            if prefers_contained_duplicate_issue(item, result[contained_duplicate_index]):
                result[contained_duplicate_index] = item
            continue
        seen.add(key)
        if semantic_key is not None:
            semantic_seen[semantic_key] = len(result)
        if origin_key is not None:
            origin_seen[origin_key] = len(result)
        result.append(item)
    return result


def sort_issues(issues: Iterable[Dict[str, Any]]) -> List[Dict[str, Any]]:
    return sorted(
        issues,
        key=lambda item: (
            SEVERITY_ORDER.get(str(item.get("severity")), 99),
            TYPE_ORDER.get(str(item.get("type")), 99),
            str(item.get("location")),
        ),
    )


def run_format_branch(
    doc: Document,
    paragraphs: Sequence[Dict[str, Any]],
    next_issue_id: int,
) -> List[Dict[str, Any]]:
    issues: List[Dict[str, Any]] = []
    issue_id = next_issue_id
    approval_table_indexes = build_approval_table_index_set(doc)
    mismatched_sections: List[str] = []
    geometry_mismatched_sections: List[str] = []
    footer_text_sections: List[str] = []
    footer_size_sections: List[str] = []
    body_font_examples: List[Dict[str, Any]] = []
    table_font_examples: List[Dict[str, Any]] = []
    missing_repeat_header_tables: List[Tuple[int, int]] = []
    table_spacing_examples: List[Tuple[str, str, Dict[str, Any]]] = []

    for index, section in enumerate(doc.sections, start=1):
        current = {
            "top_cm": round(section.top_margin.cm, 2),
            "bottom_cm": round(section.bottom_margin.cm, 2),
            "left_cm": round(section.left_margin.cm, 2),
            "right_cm": round(section.right_margin.cm, 2),
        }
        if current != PAGE_STANDARD:
            mismatched_sections.append(
                f"第{index}节(top={current['top_cm']}, bottom={current['bottom_cm']}, left={current['left_cm']}, right={current['right_cm']})"
            )
        page_width = round(section.page_width.cm, 2)
        page_height = round(section.page_height.cm, 2)
        header_distance = round(section.header_distance.cm, 2)
        footer_distance = round(section.footer_distance.cm, 2)
        gutter = round(section.gutter.cm, 2)
        geometry_checks = {
            "width_cm": page_width,
            "height_cm": page_height,
            "header_cm": header_distance,
            "footer_cm": footer_distance,
            "gutter_cm": gutter,
        }
        if (
            not approx_equal(page_width, A4_STANDARD["width_cm"])
            and not approx_equal(page_height, A4_STANDARD["width_cm"])
        ) or (
            not approx_equal(page_width, A4_STANDARD["height_cm"])
            and not approx_equal(page_height, A4_STANDARD["height_cm"])
        ) or not approx_equal(header_distance, SECTION_DISTANCE_STANDARD["header_cm"]) or not approx_equal(
            footer_distance, SECTION_DISTANCE_STANDARD["footer_cm"]
        ) or not approx_equal(gutter, SECTION_DISTANCE_STANDARD["gutter_cm"]):
            geometry_mismatched_sections.append(
                f"第{index}节(page={page_width}x{page_height}, header={header_distance}, footer={footer_distance}, gutter={gutter})"
            )

        if not section_has_footer_text(section):
            footer_text_sections.append(f"第{index}节")

        sizes = footer_font_sizes(section)
        if sizes and any(not approx_equal(size, FOOTER_FONT_SIZE_PT) for size in sizes):
            footer_size_sections.append(f"第{index}节({', '.join(str(size) for size in sizes)}pt)")

    if mismatched_sections:
        evidence = "；".join(mismatched_sections)
        issues.append(
            make_issue(
                issue_id,
                rule_id="FMT-PAGE-001",
                branch="format",
                paragraph_index=None,
                location="；".join(f"第{part.split('(')[0][1:]}" for part in mismatched_sections),
                original=evidence,
                issue="页面设置的页边距不符合公司标准。",
                suggestion="将相关节页边距统一为：上2.5cm、下2.5cm、左3.0cm、右2.5cm。",
                severity="主要",
                evidence=evidence,
                document_zone="metadata",
                location_kind="property",
                anchor_locator="section_properties",
                comments_added=0,
                confidence="high",
                match_method="inference",
                notes="页面设置问题不写入正文批注。",
            )
        )
        issue_id += 1

    if geometry_mismatched_sections:
        evidence = "；".join(geometry_mismatched_sections)
        issues.append(
            make_issue(
                issue_id,
                rule_id="FMT-PAGE-002",
                branch="format",
                paragraph_index=None,
                location="；".join(item.split("(")[0] for item in geometry_mismatched_sections),
                original=evidence,
                issue="页面设置中的纸张尺寸、装订线或页眉页脚距离不符合公司标准。",
                suggestion="将纸张统一为 A4，装订线设为左侧 0.5cm，页眉 1.5cm，页脚 1.75cm。",
                severity="主要",
                evidence=evidence,
                document_zone="metadata",
                location_kind="property",
                anchor_locator="section_properties",
                comments_added=0,
                confidence="high",
                match_method="inference",
                notes="页面几何设置问题不写入正文批注。",
            )
        )
        issue_id += 1

    if footer_text_sections:
        evidence = "；".join(footer_text_sections)
        issues.append(
            make_issue(
                issue_id,
                rule_id="FMT-FTR-001",
                branch="format",
                paragraph_index=None,
                location="；".join(footer_text_sections),
                original=evidence,
                issue="页脚未完整包含公司固定文本。",
                suggestion="补齐页脚固定文本：公司机密-仅供内部使用 Company Confidential-Internal Use Only。",
                severity="主要",
                evidence=evidence,
                document_zone="footer",
                location_kind="footer",
                anchor_locator="footer_properties",
                comments_added=0,
                confidence="high",
                match_method="inference",
                notes="页脚固定文本问题不写入正文批注。",
            )
        )
        issue_id += 1

    if footer_size_sections:
        evidence = "；".join(footer_size_sections)
        issues.append(
            make_issue(
                issue_id,
                rule_id="FMT-FTR-002",
                branch="format",
                paragraph_index=None,
                location="；".join(item.split("(")[0] for item in footer_size_sections),
                original=evidence,
                issue="页脚字号不符合公司标准的 8pt。",
                suggestion="将页脚文字字号统一调整为 8pt。",
                severity="主要",
                evidence=evidence,
                document_zone="footer",
                location_kind="footer",
                anchor_locator="footer_properties",
                comments_added=0,
                confidence="high",
                match_method="inference",
                notes="页脚字号问题不写入正文批注。",
            )
        )
        issue_id += 1

    for paragraph in doc.paragraphs:
        if not is_body_text_paragraph(paragraph):
            continue
        if is_table_note_paragraph(paragraph):
            continue
        sizes = paragraph_effective_font_sizes(paragraph)
        if sizes and any(not approx_equal(size, BODY_FONT_SIZE_PT) for size in sizes):
            text = paragraph.text.strip()
            body_font_examples.append(
                {
                    "display": f"`{text[:30]}`({', '.join(str(size) for size in sizes)}pt)",
                    "text": text,
                    "record": find_paragraph_record_for_text(paragraphs, text, document_zone="body"),
                }
            )

    if body_font_examples:
        first_example = body_font_examples[0]
        anchor_fields = anchor_fields_from_record(first_example.get("record"), str(first_example.get("text") or ""))
        evidence = format_examples(str(item["display"]) for item in body_font_examples)
        issues.append(
            make_issue(
                issue_id,
                rule_id="FMT-BODY-001",
                branch="format",
                paragraph_index=anchor_fields["paragraph_index"],
                location="正文段落",
                original=evidence,
                issue="正文段落字号不符合公司标准的 12pt。",
                suggestion="将正文段落字号统一调整为 12pt（小四）。",
                severity="主要",
                evidence=evidence,
                document_zone="body",
                location_kind="paragraph" if anchor_fields["anchor_locator"] else "property",
                anchor_locator=anchor_fields["anchor_locator"] or "body_font_styles",
                anchor_span=anchor_fields["anchor_span"],
                anchor_text=anchor_fields["anchor_text"] or evidence,
                comments_added=anchor_fields["comments_added"],
                confidence="medium",
                match_method=anchor_fields["match_method"],
                notes="按正文样式识别，建议人工复核极少数特殊段落。",
            )
        )
        issue_id += 1

    for table_index, row_index, cell_index, table, paragraph in iter_table_paragraphs(doc):
        if is_approval_table(table):
            continue
        sizes = paragraph_effective_font_sizes(paragraph)
        if sizes and any(not approx_equal(size, TABLE_FONT_SIZE_PT) for size in sizes):
            text = paragraph.text.strip()
            table_font_examples.append(
                {
                    "display": f"表{table_index} R{row_index}C{cell_index} `{text[:24]}`({', '.join(str(size) for size in sizes)}pt)",
                    "text": text,
                    "record": find_paragraph_record_for_text(
                        paragraphs,
                        text,
                        document_zone="table",
                        table_index=table_index,
                        row_index=row_index,
                        cell_index=cell_index,
                    ),
                }
            )

    if table_font_examples:
        first_example = table_font_examples[0]
        anchor_fields = anchor_fields_from_record(first_example.get("record"), str(first_example.get("text") or ""))
        evidence = format_examples(str(item["display"]) for item in table_font_examples)
        issues.append(
            make_issue(
                issue_id,
                rule_id="FMT-TBL-001",
                branch="format",
                paragraph_index=anchor_fields["paragraph_index"],
                location="表格内容",
                original=evidence,
                issue="表格内容字号不符合公司标准的 10.5pt。",
                suggestion="将表格内容字号统一调整为 10.5pt（五号）。",
                severity="主要",
                evidence=evidence,
                document_zone="table",
                location_kind="table" if anchor_fields["anchor_locator"] else "property",
                anchor_locator=anchor_fields["anchor_locator"] or "table_font_styles",
                anchor_span=anchor_fields["anchor_span"],
                anchor_text=anchor_fields["anchor_text"] or evidence,
                comments_added=anchor_fields["comments_added"],
                confidence="medium",
                match_method=anchor_fields["match_method"],
                notes="按表格段落与 run/style 字号综合判断。",
            )
        )
        issue_id += 1

    for table_index, table in enumerate(doc.tables, start=1):
        if is_approval_table(table):
            continue
        if len(table.rows) >= 10 and not table_has_repeat_header(table):
            missing_repeat_header_tables.append((table_index, len(table.rows)))

    if missing_repeat_header_tables:
        table_labels = [f"表{table_index}(rows={row_count})" for table_index, row_count in missing_repeat_header_tables]
        evidence = "；".join(table_labels)
        first_table_index = missing_repeat_header_tables[0][0]
        anchor_record = first_table_header_anchor_record(paragraphs, first_table_index)
        anchor_fields = anchor_fields_from_record(anchor_record, str(anchor_record.get("text") or "") if anchor_record else "")
        issues.append(
            make_issue(
                issue_id,
                rule_id="FMT-TBL-002",
                branch="format",
                paragraph_index=anchor_fields["paragraph_index"],
                location="；".join(table_labels),
                original=evidence,
                issue="大表未启用跨页重复表头。",
                suggestion="为跨页或长表的首行启用“重复标题行”。",
                severity="主要",
                evidence=evidence,
                document_zone="table",
                location_kind="table",
                anchor_locator=anchor_fields["anchor_locator"] or "table_header_properties",
                anchor_span=anchor_fields["anchor_span"],
                anchor_text=anchor_fields["anchor_text"] or evidence,
                comments_added=anchor_fields["comments_added"],
                confidence="high",
                match_method=anchor_fields["match_method"],
                comment_visibility="internal",
                notes=(
                    "按行数推断的跨页重复表头风险仅保留为内部诊断，不写入 Word 批注。"
                    if anchor_fields["comments_added"]
                    else "表格结构问题不写入正文批注。"
                ),
            )
        )
        issue_id += 1

    for record in paragraphs:
        if paragraph_document_zone(record) != "table" or is_approval_record(record, approval_table_indexes):
            continue
        text = str(record.get("text", ""))
        if not has_unstable_table_extra_space(text):
            continue
        table_spacing_examples.append((short_table_record_location(record), text.strip()[:36], record))

    if table_spacing_examples:
        first_record = table_spacing_examples[0][2]
        anchor_text = str(first_record.get("text") or "").strip()
        anchor_fields = anchor_fields_from_record(first_record, anchor_text)
        evidence = format_grouped_examples((item[0], item[1]) for item in table_spacing_examples)
        has_stable_anchor = bool(anchor_fields["comments_added"])
        issues.append(
            make_issue(
                issue_id,
                rule_id="FMT-TBL-003",
                branch="format",
                paragraph_index=anchor_fields["paragraph_index"],
                location="表格双语说明行",
                original=evidence,
                issue="表格单元格中存在多余空格或缩进，影响表内双语说明的版式一致性。",
                suggestion="清理表格单元格中的多余空格、前导缩进和重复空白，保持中英文说明行对齐一致。",
                severity="次要",
                evidence=evidence,
                document_zone="table",
                location_kind="table",
                anchor_locator=anchor_fields["anchor_locator"] or "table_whitespace_styles",
                anchor_span=anchor_fields["anchor_span"],
                anchor_text=anchor_fields["anchor_text"] or evidence,
                comments_added=anchor_fields["comments_added"],
                confidence="medium",
                match_method=anchor_fields["match_method"],
                comment_visibility="word_comment" if has_stable_anchor else "internal",
                notes=(
                    "表格连续空格/前导缩进为客观格式错误，已有稳定锚点时写入代表性批注。"
                    if has_stable_anchor
                    else "表格空白/缩进格式项缺少稳定锚点，仅保留为内部诊断。"
                ),
            )
        )

    return issues


def run_project_branch(paragraphs: Sequence[Dict[str, Any]], current_project: str, next_issue_id: int) -> List[Dict[str, Any]]:
    issues: List[Dict[str, Any]] = []
    issue_id = next_issue_id
    allow_words = ("参考", "根据", "引用", "参照", "来源", "见", "reference")
    aliases = project_aliases(current_project)
    if not aliases:
        return []

    for record in paragraphs:
        text = str(record.get("text", ""))
        if paragraph_document_zone(record) == "table":
            continue
        candidate_matches: List[Tuple[str, re.Match[str]]] = []
        for match_obj in PROJECT_PATTERN.finditer(text):
            item = match_obj.group(0)
            if item in aliases:
                continue
            context = text[max(0, match_obj.start() - 30) : min(len(text), match_obj.end() + 30)].lower()
            if any(word.lower() in context for word in allow_words):
                continue
            if is_project_code_biological_context(item, context):
                continue
            if is_current_project_subentity_context(item, aliases, context):
                continue
            candidate_matches.append((item, match_obj))
        if not candidate_matches:
            continue
        location = paragraph_location(record)
        anchor_locator = paragraph_anchor_locator(record)
        document_zone = paragraph_document_zone(record)
        location_kind = paragraph_location_kind(record)

        other, match_obj = candidate_matches[0]
        span = find_exact_span(text, other)
        status = "confirmed" if span else ISSUE_STATUS_NEEDS_USER_CHECK
        issues.append(
            make_issue(
                issue_id,
                rule_id="PROJ-ID-001",
                branch="project_number",
                paragraph_index=None,
                location=location,
                original=other,
                issue=f"正文出现当前项目号 `{current_project}` 以外的项目号 `{other}`。",
                suggestion="确认是否为模板残留；如无明确引用背景，请统一改为当前项目号。",
                severity="主要",
                evidence=text.strip(),
                document_zone=document_zone,
                location_kind=location_kind,
                anchor_locator=anchor_locator,
                anchor_span=span,
                anchor_text=other,
                status=status,
                comments_added=1 if status == "confirmed" else 0,
                match_method="span" if span else "inference",
                notes="" if status == "confirmed" else "项目号出现多处或锚点不唯一，需人工确认后批注。",
            )
        )
        issue_id += 1

    return unique_issues(issues)


def is_project_code_biological_context(code: str, context: str) -> bool:
    value = str(code or "").upper()
    text = str(context or "").lower()
    if not value.startswith(("HEK", "CHO")):
        return False
    return any(term in text for term in ("cell", "cells", "cell line", "细胞", "细胞系"))


def is_current_project_subentity_context(code: str, current_aliases: set[str], context: str) -> bool:
    value = str(code or "").upper()
    aliases = {str(alias or "").upper() for alias in current_aliases if alias}
    if not any(len(value) == len(alias) + 1 and value.startswith(alias) and value[-1:].isalpha() for alias in aliases):
        return False
    text = str(context or "").lower()
    if re.search(rf"{re.escape(value.lower())}\s*-", text):
        return True
    subentity_terms = (
        "分子",
        "氨基酸",
        "序列",
        "目的基因",
        "基因",
        "信号肽",
        "表达载体",
        "载体",
        "质粒",
        "片段",
        "构建体",
        "构建",
        "克隆",
        "酶切",
        "转化",
        "菌落",
        "重组",
        "snts",
        "sp1",
        "sp2",
        "ptt5",
        "molecule",
        "sequence",
        "gene",
        "signal peptide",
        "expression vector",
        "vector",
        "plasmid",
        "fragment",
        "construct",
        "construction",
        "clone",
        "cloning",
        "digestion",
        "transformation",
        "recombinant",
    )
    return any(term in text for term in subentity_terms)


def content_issue(
    issue_id: int,
    record: Dict[str, Any],
    *,
    rule_id: str,
    original: str,
    issue: str,
    suggestion: str,
    severity: str = "主要",
    evidence: str = "",
    confidence: str = "high",
) -> Dict[str, Any]:
    text = str(record.get("text", ""))
    span = find_exact_span(text, original)
    if span is None and original in text:
        span = find_first_exact_span(text, original)
    return make_issue(
        issue_id,
        rule_id=rule_id,
        branch=CONTENT_CONSISTENCY_BRANCH,
        paragraph_index=None,
        location=paragraph_location(record),
        original=original,
        issue=issue,
        suggestion=suggestion,
        severity=severity,
        evidence=evidence or text,
        document_zone=paragraph_document_zone(record),
        location_kind=paragraph_location_kind(record),
        anchor_locator=paragraph_anchor_locator(record),
        anchor_span=span,
        anchor_text=original,
        status="confirmed" if span else ISSUE_STATUS_NEEDS_USER_CHECK,
        comments_added=1 if span else 0,
        confidence=confidence,
        match_method="span" if span else "inference",
        notes="" if span else "内容一致性问题定位不稳定，需人工复核。",
    )


def same_table_cell(left: Dict[str, Any], right: Dict[str, Any]) -> bool:
    return all(
        left.get(key) == right.get(key) and left.get(key) is not None
        for key in ("table_index", "row_index", "cell_index")
    )


TEMP_TOLERANCE_PATTERN = re.compile(
    r"(?P<base>\d+(?:\.\d+)?)\s*°\s*C\s*±\s*(?P<tolerance>\d+(?:\.\d+)?)\s*°\s*C",
    flags=re.IGNORECASE,
)


def format_numeric_token(value: str) -> str:
    text = str(value or "").strip()
    if "." not in text:
        return text
    return text.rstrip("0").rstrip(".")


def format_temperature_tolerance(base: str, tolerance: str) -> str:
    return f"{format_numeric_token(base)}°C±{format_numeric_token(tolerance)}°C"


def iter_temperature_tolerances(text: str) -> Iterable[Dict[str, Any]]:
    for match in TEMP_TOLERANCE_PATTERN.finditer(str(text or "")):
        yield {
            "base": format_numeric_token(match.group("base")),
            "tolerance": format_numeric_token(match.group("tolerance")),
            "text": match.group(0),
            "start": match.start(),
            "end": match.end(),
        }


def detect_table_temperature_tolerance_mismatches(
    paragraphs: Sequence[Dict[str, Any]]
) -> List[Tuple[Dict[str, Any], str, str, str, str]]:
    """Find same-cell bilingual temperature tolerance mismatches such as 25°C±2°C vs 25°C±5°C."""

    global_counts: Dict[str, Counter[str]] = defaultdict(Counter)
    for record in paragraphs:
        for item in iter_temperature_tolerances(str(record.get("text", ""))):
            global_counts[item["base"]][item["tolerance"]] += 1

    findings: List[Tuple[Dict[str, Any], str, str, str, str]] = []
    for record in paragraphs:
        if paragraph_document_zone(record) != "table":
            continue
        text = str(record.get("text", ""))
        if not (has_chinese(text) and has_english(text)):
            continue
        by_base: Dict[str, List[Dict[str, Any]]] = defaultdict(list)
        for item in iter_temperature_tolerances(text):
            by_base[item["base"]].append(item)
        for base, items in by_base.items():
            tolerances = {item["tolerance"] for item in items}
            if len(tolerances) < 2:
                continue
            counts = global_counts.get(base) or Counter()
            expected_tolerance = counts.most_common(1)[0][0] if counts else items[0]["tolerance"]
            if expected_tolerance not in tolerances:
                expected_tolerance = items[0]["tolerance"]
            mismatch = next((item for item in items if item["tolerance"] != expected_tolerance), items[-1])
            expected = format_temperature_tolerance(base, expected_tolerance)
            actual = format_temperature_tolerance(base, mismatch["tolerance"])
            findings.append((record, mismatch["text"], actual, expected, text))
    return findings


def collect_protocol_group_numbers(paragraphs: Sequence[Dict[str, Any]]) -> set[int]:
    numbers: set[int] = set()
    for record in paragraphs:
        if paragraph_document_zone(record) != "table":
            continue
        text = str(record.get("text", "")).strip()
        match = re.match(r"^(?:Group|组)\s*(\d+)\b", text, flags=re.IGNORECASE)
        if match:
            numbers.add(int(match.group(1)))
    return numbers


def method_tokens_from_chinese_text(text: str) -> set[str]:
    """Extract method abbreviations written as '<token>法' from Chinese text."""
    return {
        match.group(1)
        for match in re.finditer(r"(?<![A-Za-z0-9])([A-Za-z][A-Za-z0-9-]{1,})\s*法", text or "")
    }


def has_substantive_chinese_text_outside_method_marker(text: str) -> bool:
    """Return true when a Chinese source line contains more than a bare method marker."""

    stripped = str(text or "")
    for term in sorted(method_tokens_from_chinese_text(stripped), key=len, reverse=True):
        stripped = re.sub(
            rf"[（(]?\s*{re.escape(term)}\s*法\s*[)）]?",
            "",
            stripped,
            flags=re.IGNORECASE,
        )
    remaining_zh = "".join(ZH_CHAR_PATTERN.findall(stripped))
    return len(remaining_zh) >= 2


def bounded_edit_distance(left: str, right: str, limit: int = 1) -> int:
    if left == right:
        return 0
    if abs(len(left) - len(right)) > limit:
        return limit + 1
    previous = list(range(len(right) + 1))
    for i, left_char in enumerate(left, start=1):
        current = [i]
        row_min = i
        for j, right_char in enumerate(right, start=1):
            cost = 0 if left_char == right_char else 1
            value = min(previous[j] + 1, current[j - 1] + 1, previous[j - 1] + cost)
            current.append(value)
            row_min = min(row_min, value)
        if row_min > limit:
            return limit + 1
        previous = current
    return previous[-1]


def iter_zh_name_candidates(text: str, endings: set[str]) -> Iterable[str]:
    for match in re.finditer(r"[\u4e00-\u9fff]{4,32}", text or ""):
        run = match.group(0)
        for end_index, char in enumerate(run):
            if char not in endings:
                continue
            max_len = min(12, end_index + 1)
            for length in range(4, max_len + 1):
                start = end_index - length + 1
                if start >= 0:
                    yield run[start : end_index + 1]


def iter_en_name_candidates(text: str) -> Iterable[str]:
    for match in re.finditer(r"\b[A-Z][A-Za-z]{6,}\b", text or ""):
        token = match.group(0)
        lower = token.lower()
        if any(lower.endswith(suffix) for suffix in EN_DOMINANT_TERM_SUFFIXES):
            yield token


def iter_formula_like_tokens(text: str) -> Iterable[str]:
    for match in re.finditer(r"(?<![A-Za-z0-9])([A-Z][A-Za-z0-9]{1,11})(?![A-Za-z0-9])", text or ""):
        token = match.group(1)
        if len(token) < 3:
            continue
        if not any(char.islower() for char in token):
            continue
        if not any(char in "Il1" for char in token):
            continue
        yield token


def confusable_formula_key(value: str) -> str:
    return re.sub(r"[il1]", "l", str(value or "").casefold())


def find_confusable_dominant_formula_variant(
    candidate: str,
    dominant_counts: Counter[str],
    all_counts: Counter[str],
    *,
    min_dominant_count: int = 2,
) -> Optional[str]:
    candidate_key = confusable_formula_key(candidate)
    candidate_count = all_counts.get(candidate, 0)
    best: Optional[Tuple[int, int, str]] = None
    for dominant, count in dominant_counts.items():
        if dominant == candidate or count < min_dominant_count or candidate_count >= count:
            continue
        if len(dominant) != len(candidate):
            continue
        if confusable_formula_key(dominant) != candidate_key:
            continue
        diff_positions = [index for index, pair in enumerate(zip(candidate, dominant)) if pair[0] != pair[1]]
        if not diff_positions or any(candidate[index] not in "Il1" and dominant[index] not in "Il1" for index in diff_positions):
            continue
        score = (-count, -len(dominant), dominant)
        if best is None or score < best:
            best = score
    return best[2] if best else None


def find_near_dominant_variant(
    candidate: str,
    dominant_counts: Counter[str],
    all_counts: Counter[str],
    *,
    min_dominant_count: int = 2,
) -> Optional[str]:
    best: Optional[Tuple[int, int, str]] = None
    candidate_count = all_counts.get(candidate, 0)
    for dominant, count in dominant_counts.items():
        if dominant == candidate or count < min_dominant_count or len(dominant) != len(candidate):
            continue
        if candidate_count >= count:
            continue
        distance = bounded_edit_distance(candidate.casefold(), dominant.casefold(), 1)
        if distance != 1:
            continue
        score = (-count, -len(dominant), dominant)
        if best is None or score < best:
            best = score
    return best[2] if best else None


def detect_dominant_term_variant_issues(
    paragraphs: Sequence[Dict[str, Any]]
) -> List[Tuple[Dict[str, Any], str, str, str, str]]:
    zh_dominants: Counter[str] = Counter()
    zh_all: Counter[str] = Counter()
    en_dominants: Counter[str] = Counter()
    en_all: Counter[str] = Counter()

    for record in paragraphs:
        text = str(record.get("text", ""))
        zh_candidates = list(iter_zh_name_candidates(text, ZH_NEAR_VARIANT_ENDINGS))
        zh_all.update(zh_candidates)
        zh_dominants.update(candidate for candidate in zh_candidates if candidate[-1] in ZH_DOMINANT_TERM_ENDINGS)
        en_candidates = list(iter_en_name_candidates(text))
        en_all.update(en_candidates)
        en_dominants.update(en_candidates)

    findings: List[Tuple[Dict[str, Any], str, str, str, str]] = []
    for record in paragraphs:
        text = str(record.get("text", ""))
        selected_zh: List[str] = []
        for candidate in sorted(set(iter_zh_name_candidates(text, ZH_NEAR_VARIANT_ENDINGS)), key=len, reverse=True):
            if candidate[-1] in ZH_DOMINANT_TERM_ENDINGS or any(candidate in existing for existing in selected_zh):
                continue
            expected = find_near_dominant_variant(candidate, zh_dominants, zh_all)
            if expected:
                selected_zh.append(candidate)
                findings.append((record, "CONTENT-TERM-ZH-001", candidate, expected, "zh"))

        for candidate in sorted(set(iter_en_name_candidates(text)), key=len, reverse=True):
            expected = find_near_dominant_variant(candidate, en_dominants, en_all)
            if expected:
                findings.append((record, "CONTENT-TERM-EN-001", candidate, expected, "en"))

    return findings


def detect_confusable_formula_variant_issues(
    paragraphs: Sequence[Dict[str, Any]]
) -> List[Tuple[Dict[str, Any], str, str, str, str]]:
    formula_counts: Counter[str] = Counter()
    for record in paragraphs:
        formula_counts.update(iter_formula_like_tokens(str(record.get("text", ""))))

    findings: List[Tuple[Dict[str, Any], str, str, str, str]] = []
    for record in paragraphs:
        text = str(record.get("text", ""))
        seen: set[str] = set()
        for candidate in iter_formula_like_tokens(text):
            if candidate in seen:
                continue
            seen.add(candidate)
            expected = find_confusable_dominant_formula_variant(candidate, formula_counts, formula_counts)
            if expected:
                findings.append((record, "CONTENT-EN-FORMULA-CONFUSABLE-001", candidate, expected, "en"))
    return findings


SEC_MONOMER_ZH_PATTERN = re.compile(
    r"SEC\s*的?\s*单体\s*%\s*下降了?\s*(?P<value>\d+(?:\.\d+)?%)",
    flags=re.IGNORECASE,
)
SEC_MONOMER_EN_PATTERN = re.compile(
    r"\bSEC\s+monomer\s*%\s+decreased\s+by\s+(?P<value>\d+(?:\.\d+)?)%",
    flags=re.IGNORECASE,
)
ABSOLUTE_FORMULA_PATTERN = re.compile(
    r"(?P<left>[A-Za-z][A-Za-z0-9%_./()]{0,16})\s*=\s*\|\s*(?P<body>[^|=\r\n]{1,60}?[-+−][^|=\r\n]{1,60}?)\s*\|",
)


def formula_token_pattern(token: str) -> str:
    escaped = re.escape(token)
    return r"\s*".join(re.escape(char) for char in token) if len(token) <= 6 else escaped.replace(r"\ ", r"\s*")


def normalize_formula_token(value: str) -> str:
    return normalized_text(str(value or "").replace("−", "-")).casefold()


def split_binary_formula_body(value: str) -> Optional[Tuple[str, str, str]]:
    match = re.fullmatch(
        r"\s*(?P<left>[A-Za-z0-9%_./()]+)\s*(?P<op>[-+−])\s*(?P<right>[A-Za-z0-9%_./()]+)\s*",
        value or "",
    )
    if not match:
        return None
    return match.group("left"), "-" if match.group("op") == "−" else match.group("op"), match.group("right")


def has_absolute_formula_marker_near(text: str, start: int, end: int) -> bool:
    before = text[max(0, start - 100) : start]
    after = text[end : min(len(text), end + 80)]
    window = f"{before} {after}"
    return bool(re.search(r"\b(?:absolute|formula|calculation|difference)\b", window, flags=re.IGNORECASE))


def detect_absolute_value_formula_mismatches(
    paragraphs: Sequence[Dict[str, Any]]
) -> List[Tuple[Dict[str, Any], str, str, str, str]]:
    """Find bilingual formula mismatches where an absolute-value bar is lost in English."""

    findings: List[Tuple[Dict[str, Any], str, str, str, str]] = []
    for record in paragraphs:
        text = str(record.get("text", ""))
        if not text or not has_chinese(text) or not has_english(text):
            continue
        if "|" not in text or "=" not in text or not re.search(r"绝对值|绝对差值|absolute", text, flags=re.IGNORECASE):
            continue
        for absolute_match in ABSOLUTE_FORMULA_PATTERN.finditer(text):
            left = absolute_match.group("left")
            body = absolute_match.group("body")
            parts = split_binary_formula_body(body)
            if not parts:
                continue
            first, operator, second = parts
            left_pattern = formula_token_pattern(left)
            first_pattern = formula_token_pattern(first)
            second_pattern = formula_token_pattern(second)
            missing_bar_pattern = re.compile(
                rf"{left_pattern}\s*=\s*(?!\|)\s*{first_pattern}\s*{re.escape(operator)}\s*{second_pattern}(?!\s*\|)",
                flags=re.IGNORECASE,
            )
            expected_norm = normalize_formula_token(f"{left}=|{first}{operator}{second}|")
            for missing_match in missing_bar_pattern.finditer(text):
                if absolute_match.start() <= missing_match.start() and missing_match.end() <= absolute_match.end():
                    continue
                nearby = text[max(0, missing_match.start() - 8) : min(len(text), missing_match.end() + 8)]
                if "|" in nearby or re.search(r"\babs\s*\($", text[max(0, missing_match.start() - 8) : missing_match.start()], flags=re.IGNORECASE):
                    continue
                if not has_absolute_formula_marker_near(text, missing_match.start(), missing_match.end()):
                    continue
                original = missing_match.group(0).strip()
                if normalize_formula_token(original) == expected_norm:
                    continue
                expected = f"{left} = |{first} {operator} {second}|"
                evidence = f"中文公式：{absolute_match.group(0).strip()}；英文公式：{original}"
                findings.append(
                    (
                        record,
                        original,
                        "中英文公式绝对值符号不一致：中文公式包含绝对值符号，英文对应公式未保留绝对值。",
                        f"将英文公式统一为 `{expected}`，或按源文件标准保留等价的绝对值表达。",
                        evidence,
                    )
                )
                break
    return findings


def detect_sec_monomer_bilingual_number_mismatches(
    paragraphs: Sequence[Dict[str, Any]]
) -> List[Tuple[Dict[str, Any], str, str, str, str]]:
    findings: List[Tuple[Dict[str, Any], str, str, str, str]] = []
    for record in paragraphs:
        text = str(record.get("text", ""))
        if not text or not has_chinese(text) or not has_english(text):
            continue
        zh_match = SEC_MONOMER_ZH_PATTERN.search(text)
        en_match = SEC_MONOMER_EN_PATTERN.search(text)
        if not zh_match or not en_match:
            continue
        zh_value = zh_match.group("value")
        en_value = f"{en_match.group('value')}%"
        if zh_value == en_value:
            continue
        original = en_match.group(0).strip()
        issue = f"中英文 SEC 单体%下降数据不一致：中文为 {zh_value}，英文为 {en_value}。"
        suggestion = f"核对后统一中英文 SEC 单体%下降数据；如中文为准，将英文改为 `SEC monomer % decreased by {zh_value}`。"
        evidence = str(record.get("text", "")).strip()
        findings.append((record, original, issue, suggestion, evidence))
    return findings


REPEATED_TURBIDITY_ZH_PATTERN = re.compile(
    r"((?:深于|浅于|大于|小于)\s*\d+(?:\.\d+)?\s*号浊)\s*\1\s*度标准液"
)


def detect_source_local_text_defects(
    paragraphs: Sequence[Dict[str, Any]]
) -> List[Tuple[Dict[str, Any], str, str, str, str, str]]:
    findings: List[Tuple[Dict[str, Any], str, str, str, str, str]] = []
    for record in paragraphs:
        text = str(record.get("text", ""))
        if not text:
            continue
        for marker, issue, suggestion in (
            ("；；", "中文正文出现连续两个分号，属于多余标点。", "删除多余分号，保留一个分号。"),
            ("；。", "中文正文出现分号后紧跟句号的异常标点组合。", "删除分号或改为单个句号，保持句读清晰。"),
        ):
            if marker in text:
                findings.append((record, "CONTENT-PUNCT-001", marker, issue, suggestion, "次要"))
        for match in REPEATED_TURBIDITY_ZH_PATTERN.finditer(text):
            repeated_prefix = match.group(1)
            original = match.group(0)
            corrected = f"{repeated_prefix}度标准液"
            findings.append(
                (
                    record,
                    "CONTENT-ZH-DUP-001",
                    original,
                    "中文浊度结果中同一短语重复，影响结果表达清晰性。",
                    f"删除重复短语，改为 `{corrected}`。",
                    "次要",
                )
            )
    return findings


ZH_DURATION_PATTERN = re.compile(r"(?P<minute>\d{1,2})\s*分\s*(?P<second>\d{1,2})\s*秒")
EN_DURATION_PATTERN = re.compile(
    r"(?P<minute>\d{1,2})\s+minute[s]?\s+(?P<second>\d{1,2})\s+second[s]?",
    flags=re.IGNORECASE,
)


def format_english_duration(minute: str, second: str) -> str:
    suffix = "second" if int(second) == 1 else "seconds"
    return f"{minute} minute {second} {suffix}"


def detect_adjacent_time_translation_mismatches(
    paragraphs: Sequence[Dict[str, Any]]
) -> List[Tuple[Dict[str, Any], str, str, str, str]]:
    findings: List[Tuple[Dict[str, Any], str, str, str, str]] = []
    for index, record in enumerate(paragraphs[:-1]):
        next_record = paragraphs[index + 1]
        if not same_table_cell(record, next_record):
            continue
        zh_text = str(record.get("text", ""))
        en_text = str(next_record.get("text", ""))
        zh_match = ZH_DURATION_PATTERN.search(zh_text)
        en_match = EN_DURATION_PATTERN.search(en_text)
        if not zh_match or not en_match:
            continue
        zh_pair = (int(zh_match.group("minute")), int(zh_match.group("second")))
        en_pair = (int(en_match.group("minute")), int(en_match.group("second")))
        if zh_pair == en_pair:
            continue
        original = en_match.group(0).strip()
        expected = format_english_duration(zh_match.group("minute"), zh_match.group("second"))
        evidence = f"中文：{zh_match.group(0)}；英文：{original}"
        findings.append(
            (
                next_record,
                original,
                "同一表格单元格内中英文时间不一致。",
                f"将英文时间改为 `{expected}`，并保持 minute/second 单复数一致。",
                evidence,
            )
        )
    return findings


def run_content_consistency_branch(paragraphs: Sequence[Dict[str, Any]], next_issue_id: int) -> List[Dict[str, Any]]:
    issues: List[Dict[str, Any]] = []
    issue_id = next_issue_id
    protocol_group_numbers = collect_protocol_group_numbers(paragraphs)
    release_data_note_records: List[Dict[str, Any]] = []

    def add(record: Dict[str, Any], **kwargs: Any) -> None:
        nonlocal issue_id
        issues.append(content_issue(issue_id, record, **kwargs))
        issue_id += 1

    for record, rule_id, original, expected, language in detect_dominant_term_variant_issues(paragraphs):
        if language == "zh":
            add(
                record,
                rule_id=rule_id,
                original=original,
                issue=f"全文主写法为 `{expected}`，此处出现近似写法 `{original}`，疑似关键产品名称错别字。",
                suggestion=f"将 `{original}` 统一为全文主写法 `{expected}`；如确为不同术语，请补充依据。",
                severity="关键",
                evidence=str(record.get("text", "")).strip(),
            )
        else:
            add(
                record,
                rule_id=rule_id,
                original=original,
                issue=f"The dominant product-name spelling is `{expected}`, but this paragraph uses `{original}`.",
                suggestion=f"Change `{original}` to `{expected}` unless this is an intentional different term.",
                severity="关键",
                evidence=str(record.get("text", "")).strip(),
            )

    for record, rule_id, original, expected, _language in detect_confusable_formula_variant_issues(paragraphs):
        add(
            record,
            rule_id=rule_id,
            original=original,
            issue=f"同一文档主流写法为 `{expected}`，此处 `{original}` 疑似存在 I/l/1 混淆。",
            suggestion=f"将 `{original}` 改为 `{expected}`。",
            severity="主要",
            evidence=str(record.get("text", "")).strip(),
        )

    for record, original, actual, expected, evidence in detect_table_temperature_tolerance_mismatches(paragraphs):
        add(
            record,
            rule_id="CONTENT-BI-TEMP-001",
            original=original,
            issue=f"同一表格单元格内温度允差中英文不一致：同一温度条件同时出现 `{expected}` 和 `{actual}`。",
            suggestion=f"将该处温度允差统一为 `{expected}`；如 `{actual}` 才是正确条件，请同步修订中文、英文及相关表格依据。",
            severity="主要",
            evidence=evidence,
        )

    for record, original, issue, suggestion, evidence in detect_sec_monomer_bilingual_number_mismatches(paragraphs):
        add(
            record,
            rule_id="CONTENT-BI-SEC-MONOMER-001",
            original=original,
            issue=issue,
            suggestion=suggestion,
            severity="主要",
            evidence=evidence,
        )

    for record, original, issue, suggestion, evidence in detect_absolute_value_formula_mismatches(paragraphs):
        add(
            record,
            rule_id="CONTENT-BI-FORMULA-001",
            original=original,
            issue=issue,
            suggestion=suggestion,
            severity="关键",
            evidence=evidence,
        )

    for record, rule_id, original, issue, suggestion, severity in detect_source_local_text_defects(paragraphs):
        add(
            record,
            rule_id=rule_id,
            original=original,
            issue=issue,
            suggestion=suggestion,
            severity=severity,
            evidence=str(record.get("text", "")).strip(),
        )

    for record, original, issue, suggestion, evidence in detect_adjacent_time_translation_mismatches(paragraphs):
        add(
            record,
            rule_id="CONTENT-BI-TIME-001",
            original=original,
            issue=issue,
            suggestion=suggestion,
            severity="主要",
            evidence=evidence,
        )

    for record in paragraphs:
        text = str(record.get("text", "")).strip()
        lower = text.lower()
        if not text:
            continue

        for grammar_match in PASSIVE_FILTER_PATTERN.finditer(text):
            original = grammar_match.group(0)
            aux = grammar_match.group(1)
            add(
                record,
                rule_id="CONTENT-EN-GRAMMAR-001",
                original=original,
                issue=f"英文被动语态动词形式错误：`{original}` 应为 `{aux} filtered`。",
                suggestion=f"将 `{original}` 改为 `{aux} filtered`。",
                severity="次要",
                evidence=text,
            )

        release_note_match = RELEASE_DATA_NOTE_PATTERN.search(text)
        if release_note_match:
            release_data_note_records.append(record)

        broken_zh_ref = re.search(r"见\s*(0|[。．.])", text)
        if broken_zh_ref:
            add(
                record,
                rule_id="CONTENT-REF-001",
                original=broken_zh_ref.group(0),
                issue="引用位置不是有效表号或章节号。",
                suggestion="补充实际表号、章节号或附件编号，并同步中英文引用。",
            )
        broken_en_ref = re.search(r"\b(section|table|appendix)\s+0\b", lower)
        if broken_en_ref:
            add(
                record,
                rule_id="CONTENT-REF-001",
                original=broken_en_ref.group(0),
                issue="英文引用位置不是有效章节号、表号或附件号。",
                suggestion="改为实际章节、表格或附件编号。",
            )
        vague_en_ref = re.search(r"\brefer to the (details|following|below)\b", lower)
        if vague_en_ref:
            add(
                record,
                rule_id="CONTENT-REF-001",
                original=vague_en_ref.group(0),
                issue="英文引用缺少明确目标，无法追溯具体方案。",
                suggestion="补充实际引用对象，例如“Table X”“Section X”或“Appendix X”。",
            )

        if HANGUL_PATTERN.search(text):
            add(
                record,
                rule_id="CONTENT-LANG-001",
                original=text,
                issue="中英文审批信息中混入第三语言模板文本。",
                suggestion="删除非中英文模板残留，并按公司模板统一为中英文并列。",
            )

        group_match = re.search(r"(?:组|Group)\s*(\d+)", text, flags=re.IGNORECASE)
        if paragraph_document_zone(record) == "body" and group_match and protocol_group_numbers:
            group_number = int(group_match.group(1))
            if group_number not in protocol_group_numbers:
                add(
                    record,
                    rule_id="CONTENT-GROUP-001",
                    original=group_match.group(0),
                    issue=f"正文组别编号 {group_match.group(0)} 与表格方案组别不一致。",
                    suggestion="核对实验方案组别；如对应成品组，应统一为正确组别编号。",
                    severity="关键",
                    evidence=f"表格方案组别: {', '.join('Group ' + str(item) for item in sorted(protocol_group_numbers))}; 原文: {text}",
                )

    for index, record in enumerate(paragraphs[:-1]):
        text = str(record.get("text", "")).strip()
        next_record = paragraphs[index + 1]
        next_text = str(next_record.get("text", "")).strip()
        if not same_table_cell(record, next_record):
            continue

        if has_chinese(text) and has_english(next_text):
            if has_substantive_chinese_text_outside_method_marker(text):
                missing_terms = sorted(
                    term for term in method_tokens_from_chinese_text(text) if term.lower() not in next_text.lower()
                )
            else:
                missing_terms = []
            if missing_terms:
                add(
                    next_record,
                    rule_id="CONTENT-BI-002",
                    original=next_text,
                    issue="检测项英文遗漏中文 X 项中的关键方法或指标。",
                    suggestion=f"英文补齐遗漏项目：{', '.join(missing_terms)}。",
                    severity="关键",
                        evidence=f"中文: {text}; 英文: {next_text}",
                    )

    if release_data_note_records:
        first_record = release_data_note_records[0]
        locations = [paragraph_location(record).split("（", 1)[0] for record in release_data_note_records]
        visible_locations = "、".join(locations[:2])
        if len(locations) > 2:
            visible_locations = f"{visible_locations}，其余{len(locations) - 2}处"
        evidence = (
            f"同类英文脚注出现于{visible_locations}；"
            f"示例：{str(first_record.get('text', '')).strip()}"
        )
        add(
            first_record,
            rule_id="CONTENT-EN-GRAMMAR-002",
            original="data reference to the releasing data",
            issue="英文脚注表达不清：`data reference to the releasing data` 不是规范英语表达。",
            suggestion="改为 `data references the release data` 或 `data refers to the release data`。",
            severity="次要",
            evidence=evidence,
        )

    return unique_issues(issues)


def describe_exception(exc: Exception) -> str:
    message = str(exc).strip()
    if message:
        return f"{type(exc).__name__}: {message}"
    return type(exc).__name__


def repair_common_agent_json(text: str) -> str:
    repaired = str(text or "").strip()
    repaired = re.sub(r"^\s*```(?:json)?\s*", "", repaired, flags=re.IGNORECASE)
    repaired = re.sub(r"\s*```\s*$", "", repaired)
    fixed_lines: List[str] = []
    for line in repaired.splitlines():
        fixed_lines.append(repair_json_string_line(line))
    repaired = "\n".join(fixed_lines)
    repaired = re.sub(r",(\s*[}\]])", r"\1", repaired)
    return repaired


def repair_json_string_line(line: str) -> str:
    colon = line.find(":")
    if colon < 0:
        return line
    index = colon + 1
    while index < len(line) and line[index].isspace():
        index += 1
    if index >= len(line) or line[index] != '"':
        return line
    right = len(line) - 1
    while right > index and line[right].isspace():
        right -= 1
    if right <= index:
        return line
    if line[right] == ",":
        right -= 1
        while right > index and line[right].isspace():
            right -= 1
    if right <= index or line[right] != '"':
        return line
    prefix = line[: index + 1]
    middle = line[index + 1 : right]
    suffix = line[right:]
    middle = re.sub(r'(?<!\\)"', r'\\"', middle)
    return f"{prefix}{middle}{suffix}"


def load_json_with_repair(path: Path) -> Dict[str, Any]:
    raw_text = path.read_text(encoding="utf-8-sig")
    try:
        payload = json.loads(raw_text)
    except json.JSONDecodeError:
        repaired = repair_common_agent_json(raw_text)
        try:
            payload = json.loads(repaired)
        except json.JSONDecodeError:
            decoder = json.JSONDecoder()
            payload, end_index = decoder.raw_decode(repaired)
            trailing = repaired[end_index:].strip()
            if trailing and not re.fullmatch(r"[}\]\s`]*", trailing):
                raise
    if not isinstance(payload, dict):
        raise ValueError("top-level JSON payload must be an object")
    return payload


def load_agent_context_unit_ids(
    manifest_path: Path,
    *,
    require_manifest: bool = False,
) -> Tuple[Optional[set[str]], Dict[str, Any]]:
    """Load allowed unit IDs from a context v2 package if one is available.

    Missing manifests keep backward compatibility and return ``None``. Existing
    but invalid manifests fail closed with an empty set so stale or incomplete
    context packages cannot authorize agent findings outside the reviewed
    context.
    """

    metadata: Dict[str, Any] = {
        "context_manifest": str(manifest_path),
        "context_manifest_found": False,
        "context_manifest_valid": False,
        "context_unit_count": 0,
    }
    if not manifest_path.exists():
        if require_manifest:
            metadata["status"] = "failed"
            metadata["context_manifest_error"] = f"agent context manifest not found: {manifest_path}"
            metadata["context_manifest_mode"] = "required_missing"
            return set(), metadata
        metadata["context_manifest_mode"] = "legacy_no_context_manifest"
        metadata["agent_merge_unrestricted_for_legacy_context"] = True
        return None, metadata

    metadata["context_manifest_found"] = True
    try:
        validation_errors = validate_agent_context_package(manifest_path)
        metadata["context_manifest_validation_errors"] = validation_errors
        if validation_errors:
            raise ValueError("context package validation failed: " + "; ".join(validation_errors[:10]))

        manifest = load_json_with_repair(manifest_path)
        unit_index_name = str(manifest.get("unit_index") or "").strip()
        if not unit_index_name:
            raise ValueError("agent context manifest missing unit_index")
        unit_index_path = manifest_path.parent / unit_index_name
        if not unit_index_path.exists():
            raise FileNotFoundError(f"agent context unit index not found: {unit_index_path}")

        unit_ids: set[str] = set()
        with unit_index_path.open("r", encoding="utf-8-sig") as handle:
            for line_number, line in enumerate(handle, start=1):
                if not line.strip():
                    continue
                row = json.loads(line)
                if not isinstance(row, dict):
                    raise ValueError(f"{unit_index_path.name}:{line_number}: JSONL row must be an object")
                unit_id = str(row.get("unit_id") or "").strip()
                if unit_id:
                    unit_ids.add(unit_id)
        metadata["context_unit_index"] = str(unit_index_path)
        metadata["context_unit_count"] = len(unit_ids)
        metadata["context_manifest_valid"] = True
        metadata["context_manifest_mode"] = "v2_manifest"
        return unit_ids, metadata
    except Exception as exc:
        metadata["context_manifest_error"] = describe_exception(exc)
        metadata["status"] = "failed"
        return set(), metadata


def make_human_review_item(
    branch: str,
    reason: str,
    location: str,
    *,
    issue_id: str = "",
    user_visible: bool = True,
    skip_category: str = "",
) -> Dict[str, Any]:
    item = {
        "issue_id": issue_id,
        "branch": branch,
        "reason": reason,
        "location": location,
        "user_visible": user_visible,
    }
    if skip_category:
        item["skip_category"] = skip_category
    return item


def summarize_chunk_failures(chunk_failures: Sequence[str], limit: int = 3) -> str:
    visible = list(chunk_failures[:limit])
    if len(chunk_failures) > limit:
        visible.append(f"... and {len(chunk_failures) - limit} more")
    return "; ".join(visible)


def llm_paragraph_target(value: Any) -> str:
    match = re.search(r"\d+", str(value or ""))
    if not match:
        return ""
    return str(int(match.group(0)))


def locate_llm_record(paragraphs: Sequence[Dict[str, Any]], raw_issue: Dict[str, Any]) -> Optional[Dict[str, Any]]:
    target = llm_paragraph_target(raw_issue.get("paragraph_index", ""))
    exact_needles = [
        str(raw_issue.get(key, "")).strip()
        for key in ("original", "evidence")
        if str(raw_issue.get(key, "")).strip()
    ]
    search_needles = ordered_unique(exact_needles + llm_anchor_candidates(raw_issue))

    def exact_matches(records: Sequence[Dict[str, Any]]) -> List[Dict[str, Any]]:
        return [
            record
            for record in records
            if any(find_exact_span(str(record.get("text", "")), needle) for needle in search_needles)
        ]

    if target:
        target_groups = [
            [record for record in paragraphs if target == str(record.get("logical_index", ""))],
            [record for record in paragraphs if target == str(record.get("xml_index", ""))],
        ]
        for matches in target_groups:
            matches_with_exact_needles = exact_matches(matches)
            if len(matches_with_exact_needles) == 1:
                return matches_with_exact_needles[0]
        for matches in target_groups:
            if len(matches) == 1:
                return matches[0]

    for needle in search_needles:
        if not needle:
            continue
        matches = [record for record in paragraphs if find_exact_span(str(record.get("text", "")), needle)]
        if len(matches) == 1:
            return matches[0]
    for needle in search_needles:
        normalized_needle = normalized_text(needle)
        if len(normalized_needle) >= 2:
            matches = [
                record
                for record in paragraphs
                if normalized_needle in normalized_text(str(record.get("text", "")))
            ]
            if len(matches) == 1:
                return matches[0]
    return None


def int_record_field(record: Dict[str, Any], field_name: str) -> Optional[int]:
    try:
        value = record.get(field_name)
        if value is None or value == "":
            return None
        return int(value)
    except (TypeError, ValueError):
        return None


def agent_unit_id(raw_issue: Dict[str, Any]) -> str:
    return str(raw_issue.get("problem_unit_id") or raw_issue.get("unit_id") or "").strip()


def add_record_unit_key(index: Dict[str, Dict[str, Any]], key: Any, record: Dict[str, Any]) -> None:
    text = str(key or "").strip()
    if text:
        index.setdefault(text.casefold(), record)


def build_record_unit_index(paragraphs: Sequence[Dict[str, Any]]) -> Dict[str, Dict[str, Any]]:
    index: Dict[str, Dict[str, Any]] = {}
    for ordinal, record in enumerate(paragraphs, start=1):
        unit_id = str(record.get("unit_id") or "").strip()
        paragraph_id = str(record.get("paragraph_id") or "").strip()
        for key in (unit_id, paragraph_id):
            if not key:
                continue
            add_record_unit_key(index, key, record)
            add_record_unit_key(index, f"u:{key}", record)

        xml_index = int_record_field(record, "xml_index")
        if xml_index is not None:
            add_record_unit_key(index, f"xml-{xml_index:05d}", record)
            add_record_unit_key(index, f"xml:{xml_index}", record)

        add_record_unit_key(index, f"u-{ordinal:05d}", record)
    return index


def locate_record_by_unit_id(
    paragraphs: Sequence[Dict[str, Any]],
    raw_issue: Dict[str, Any],
) -> Optional[Dict[str, Any]]:
    unit_id = agent_unit_id(raw_issue)
    if not unit_id:
        return None
    return build_record_unit_index(paragraphs).get(unit_id.casefold())


def find_anchor_quote_span(record: Dict[str, Any], raw_issue: Dict[str, Any]) -> Tuple[Optional[Dict[str, Any]], str]:
    text = str(record.get("text") or "")
    candidate = strip_agent_text_label(str(raw_issue.get("anchor_quote") or "")).strip(" \t\r\n:：;；,，。.-")
    if not candidate or has_weak_agent_anchor(candidate):
        return None, ""
    span = find_exact_span(text, candidate)
    if span:
        return span, candidate
    return None, ""


def compact_source_anchor_text(value: Any) -> str:
    text = strip_agent_text_label(str(value or "")).strip(" \t\r\n:：;；,，。.-")
    return normalized_text(text).casefold()


def source_anchor_texts_compatible(left: Any, right: Any) -> bool:
    left_norm = compact_source_anchor_text(left)
    right_norm = compact_source_anchor_text(right)
    if not left_norm or not right_norm:
        return False
    return left_norm in right_norm or right_norm in left_norm


def nearby_llm_target_records(
    paragraphs: Sequence[Dict[str, Any]],
    raw_issue: Dict[str, Any],
    record: Optional[Dict[str, Any]],
    *,
    window: int = 2,
) -> List[Dict[str, Any]]:
    """Return conservative nearby records for paired bilingual lines.

    Some DOCX maps split a Chinese note and its English note into adjacent
    paragraph records. Weak agents often cite the Chinese paragraph index while
    reporting an English issue from the paired next line. Keep this retry local
    to the cited index to avoid turning broad searches into unstable anchors.
    """

    target = llm_paragraph_target(raw_issue.get("paragraph_index", ""))
    target_numbers = {int(target)} if target else set()
    if record:
        for field_name in ("logical_index", "xml_index"):
            value = int_record_field(record, field_name)
            if value is not None:
                target_numbers.add(value)

    if not target_numbers:
        return []

    nearby: List[Dict[str, Any]] = []
    for candidate in paragraphs:
        for field_name in ("logical_index", "xml_index"):
            value = int_record_field(candidate, field_name)
            if value is None:
                continue
            if any(abs(value - target_value) <= window for target_value in target_numbers):
                if candidate not in nearby:
                    nearby.append(candidate)
                break
    return nearby


AGENT_TEXT_LABEL = r"(?:中文|英文|CN|EN|Chinese|English|原文|证据|依据|source|original|evidence)"


def strip_agent_text_label(text: str) -> str:
    return re.sub(rf"^\s*{AGENT_TEXT_LABEL}\s*[:：]\s*", "", text.strip(), flags=re.IGNORECASE)


def split_agent_labeled_fragments(text: str) -> List[str]:
    value = str(text or "").strip()
    if not value:
        return []
    fragments: List[str] = []
    pattern = re.compile(
        rf"(?:^|[\s;；,，]){AGENT_TEXT_LABEL}\s*[:：]\s*(.*?)(?=(?:\s+{AGENT_TEXT_LABEL}\s*[:：])|[;；\r\n]|$)",
        flags=re.IGNORECASE,
    )
    for match in pattern.finditer(value):
        fragment = match.group(1).strip(" \t\r\n:：;；,，。")
        if fragment and fragment not in fragments:
            fragments.append(fragment)
    stripped = strip_agent_text_label(value).strip(" \t\r\n:：;；,，。")
    if stripped and stripped != value and stripped not in fragments:
        fragments.append(stripped)
    return fragments


def is_code_like_anchor_candidate(text: str) -> bool:
    raw_value = str(text or "").strip()
    if re.search(r"\s", raw_value) and len(re.findall(r"[A-Za-z]{2,}", raw_value)) >= 2:
        return False
    value = re.sub(r"\s+", "", raw_value)
    if not value:
        return True
    letters = re.findall(r"[A-Za-z]+", value)
    if len(letters) >= 2 and any(len(item) >= 4 for item in letters):
        return False
    return bool(re.fullmatch(r"[A-Za-z0-9_.()/+-]+", value))


def llm_anchor_candidates(
    raw_issue: Dict[str, Any],
    *,
    field_names: Sequence[str] = ("original", "evidence", "evidence_quote", "source_quote"),
) -> List[str]:
    candidates: List[str] = []

    def add(value: str) -> None:
        text = strip_agent_text_label(value)
        text = text.strip(" \t\r\n:：;；,，。.-")
        if 3 <= len(text) <= 180 and text not in candidates:
            candidates.append(text)

    fields = [
        str(raw_issue.get(key) or "")
        for key in field_names
    ]

    for source in fields:
        for fragment in split_agent_labeled_fragments(source):
            add(fragment)
        for pattern in (
            r"[`\"“”']([^`\"“”']{3,120})[`\"“”']",
            r"(\d+(?:\.\d+)?\s*°?C\s*±\s*\d+(?:\.\d+)?\s*°?C)",
            r"\b([A-Z][A-Za-z]+(?:\s+(?:of|and|for|the|to|in|by|with|[A-Z][A-Za-z]+)){2,})\b",
        ):
            for match in re.findall(pattern, source):
                add(match if isinstance(match, str) else " ".join(match))

    for source in fields:
        for line in re.split(r"[\r\n]+", source):
            for part in re.split(r"\.{3,}|…+|；|;|，|,|\(|\)|（|）", line):
                add(part)
                if " in " in part:
                    add(part.split(" in ", 1)[0])
                if " vs " in part:
                    for side in part.split(" vs "):
                        add(side)

    return candidates


def resolve_llm_anchor(
    paragraph_text: str,
    raw_issue: Dict[str, Any],
) -> Tuple[Optional[Dict[str, Any]], str]:
    raw_original = strip_agent_text_label(str(raw_issue.get("original", "")))
    if is_objective_repeated_whitespace_issue(raw_issue):
        original = raw_original.strip(" \t\r\n")
    else:
        original = raw_original.strip(" \t\r\n:：;；,，。.-")
    reject_replacement_target_anchor = not is_objective_formula_case_issue(raw_issue)
    preferred_error = preferred_error_anchor_from_issue(paragraph_text, raw_issue, original)
    if preferred_error:
        preferred_span = find_exact_span(paragraph_text, preferred_error) or find_first_exact_span(paragraph_text, preferred_error)
        if preferred_span and not (
            reject_replacement_target_anchor
            and should_reject_replacement_target_anchor(raw_issue, preferred_error)
        ):
            return preferred_span, preferred_error
    if original:
        original_span = find_exact_span(paragraph_text, original) or find_first_exact_span(paragraph_text, original)
        if original_span and not (
            reject_replacement_target_anchor
            and should_reject_replacement_target_anchor(raw_issue, original)
        ):
            return original_span, original

    candidates = sorted(
        llm_anchor_candidates(raw_issue),
        key=lambda item: (
            is_code_like_anchor_candidate(item),
            bool(ZH_CHAR_PATTERN.search(item)) and len(item) > 20,
            len(item) > 80,
            len(item),
        ),
    )
    for candidate in candidates:
        if (
            reject_replacement_target_anchor
            and should_reject_replacement_target_anchor(raw_issue, candidate)
        ):
            continue
        span = find_exact_span(paragraph_text, candidate) or find_first_exact_span(paragraph_text, candidate)
        if span:
            return span, candidate
    return None, str(raw_issue.get("original", "")).strip()


def find_unique_source_local_typo_anchor(
    paragraphs: Sequence[Dict[str, Any]],
    raw_issue: Dict[str, Any],
) -> Optional[Tuple[Dict[str, Any], Dict[str, Any], str]]:
    """Retry a wrong paragraph index only for exact, unique source-local typos."""

    if not is_objective_source_local_typo_issue(raw_issue):
        return None
    original = strip_agent_text_label(str(raw_issue.get("original") or "")).strip(" \t\r\n:：;；,，。.-")
    candidates = ordered_unique([original, *llm_anchor_candidates(raw_issue)])
    for candidate in candidates:
        anchor = str(candidate or "").strip(" \t\r\n:：;；,，。.-")
        if not anchor or has_weak_agent_anchor(anchor):
            continue
        if anchor_is_replacement_target(raw_issue, anchor):
            continue
        matches: List[Tuple[Dict[str, Any], Dict[str, Any], str]] = []
        for record in paragraphs:
            paragraph_text = str(record.get("text", ""))
            span = find_exact_span(paragraph_text, anchor) or find_first_exact_span(paragraph_text, anchor)
            if span and not is_unsafe_embedded_numeric_anchor(paragraph_text, span, anchor, raw_issue):
                matches.append((record, span, anchor))
                if len(matches) > 1:
                    break
        if len(matches) == 1:
            return matches[0]
    return None


def replacement_target_norms_from_issue(raw_issue: Dict[str, Any]) -> set[str]:
    suggestion = str(raw_issue.get("suggestion") or "")
    targets: set[str] = set()
    bare_target = bare_replacement_suggestion(raw_issue)
    if bare_target:
        targets.add(normalized_text(bare_target).casefold())
    patterns = (
        r"(?:改为|修改为|更正为|统一为|替换为)[:：]?\s*[`\"'“”‘’]?([^`\"'“”‘’，。,；;]{2,80})[`\"'“”‘’]?",
        r"[`\"'“”‘’][^`\"'“”‘’]{1,80}[`\"'“”‘’]\s*(?:改为|修改为|更正为|统一为|替换为)\s*[`\"'“”‘’]([^`\"'“”‘’]{2,80})[`\"'“”‘’]",
        r"\b(?:change|replace|revise|correct)\b\s+.+?\b(?:to|with)\b\s*[`\"'“”‘’]([^`\"'“”‘’]{2,80})[`\"'“”‘’]",
    )
    for pattern in patterns:
        for match in re.findall(pattern, suggestion, flags=re.IGNORECASE):
            value = str(match).strip(" \t\r\n:：;；,，。.`\"'“”‘’")
            norm = normalized_text(value).casefold()
            if norm:
                targets.add(norm)
    for marker in ("改为", "修改为", "更正为", "统一为", "替换为"):
        if marker not in suggestion:
            continue
        tail = suggestion.split(marker, 1)[1]
        quoted = re.search(r"[`\"'“”‘’]([^`\"'“”‘’]{2,80})[`\"'“”‘’]", tail)
        value = quoted.group(1) if quoted else re.split(r"[。；;，,]", tail, maxsplit=1)[0]
        value = str(value).strip(" \t\r\n:：;；,，。.`\"'“”‘’")
        norm = normalized_text(value).casefold()
        if norm:
            targets.add(norm)
    return targets


def anchor_is_replacement_target(raw_issue: Dict[str, Any], anchor_text: str) -> bool:
    anchor_norm = normalized_text(str(anchor_text or "")).casefold()
    return bool(anchor_norm and anchor_norm in replacement_target_norms_from_issue(raw_issue))


def should_reject_replacement_target_anchor(raw_issue: Dict[str, Any], anchor_text: str) -> bool:
    if not anchor_is_replacement_target(raw_issue, anchor_text):
        return False
    if is_objective_formula_case_issue(raw_issue):
        return False
    if is_objective_repeated_whitespace_anchor(raw_issue, anchor_text):
        return False
    combined = "\n".join(str(raw_issue.get(key) or "") for key in ("issue", "suggestion", "evidence", "original"))
    lowered = combined.casefold()
    if ("缺少空格" in combined or "missing space" in lowered) and str(raw_issue.get("original") or "").strip() == str(anchor_text or "").strip():
        return False
    return True


def preferred_error_anchor_from_issue(paragraph_text: str, raw_issue: Dict[str, Any], original: str) -> str:
    if not paragraph_text:
        return ""
    issue_text = str(raw_issue.get("issue") or "")
    evidence_text = str(raw_issue.get("evidence") or "")
    fields = "\n".join(str(raw_issue.get(key) or "") for key in ("issue", "suggestion", "evidence", "original"))
    candidates: List[str] = []
    combined_norm = normalized_text(fields).casefold()
    original_norm = normalized_text(original).casefold()
    target_norms = replacement_target_norms_from_issue(raw_issue)
    original_span_exists = bool(original and (find_exact_span(paragraph_text, original) or find_first_exact_span(paragraph_text, original)))
    missing_than_after_less = (
        original_norm == "less" or original_norm.startswith("lessurbidity")
    ) and (
        "missingthanafterless" in combined_norm
        or "缺少than" in combined_norm
        or "lessthanturbiditystandardsolution" in combined_norm
    )
    if missing_than_after_less and "turbiditystandardsolution" in combined_norm:
        match = re.search(
            r"\bless\s+[A-Za-z]*urbidity\s+standard\s+solution\s+\d+(?:\.\d+)?",
            paragraph_text,
            flags=re.IGNORECASE,
        )
        if match:
            candidates.append(match.group(0).strip(" \t\r\n:：;；,，。.-"))
    for pattern in (
        r"拼写为[\"'“”‘’]([^\"'“”‘’]{3,40})[\"'“”‘’]",
        r"misspell(?:ed|ing)?\s+(?:as\s+)?[\"'“”‘’]([^\"'“”‘’]{3,40})[\"'“”‘’]",
        r"spelled\s+as\s+[\"'“”‘’]([^\"'“”‘’]{3,40})[\"'“”‘’]",
        r"(?:replace|change|correct|revise)\s+[\"'“”‘’]([^\"'“”‘’]{2,80})[\"'“”‘’]\s+(?:to|with)\s+[\"'“”‘’][^\"'“”‘’]{1,80}[\"'“”‘’]",
    ):
        for match in re.findall(pattern, issue_text, flags=re.IGNORECASE):
            candidates.append(str(match).strip())
    if (not original_span_exists) or anchor_is_replacement_target(raw_issue, original) or len(original) > 80:
        for quoted in re.findall(r"[`\"'“”‘’]([^`\"'“”‘’]{2,80})[`\"'“”‘’]", f"{issue_text}\n{evidence_text}"):
            value = str(quoted).strip(" \t\r\n:：;；,，。.-")
            norm = normalized_text(value).casefold()
            if not norm or norm == original_norm or norm in target_norms:
                continue
            if has_weak_agent_anchor(value):
                continue
            if find_exact_span(paragraph_text, value) or find_first_exact_span(paragraph_text, value):
                candidates.insert(0, value)
    for pattern in (
        r"原文为[\"'“”‘’]([^\"'“”‘’]{3,80})[\"'“”‘’]",
        r"source text (?:contains|is)\s+[\"'“”‘’]([^\"'“”‘’]{3,80})[\"'“”‘’]",
    ):
        for match in re.findall(pattern, fields, flags=re.IGNORECASE):
            candidates.append(str(match).strip())
    original_norm = normalized_text(original).casefold()
    for candidate in candidates:
        candidate = candidate.strip(" \t\r\n:：;；,，。.-")
        if not candidate or normalized_text(candidate).casefold() == original_norm:
            continue
        if has_weak_agent_anchor(candidate):
            continue
        span = find_exact_span(paragraph_text, candidate) or find_first_exact_span(paragraph_text, candidate)
        if span:
            return candidate
        for part in re.split(r"\s+", candidate):
            part = part.strip(" \t\r\n:：;；,，。.-")
            if part and normalized_text(part).casefold() != original_norm and not has_weak_agent_anchor(part):
                if find_exact_span(paragraph_text, part) or find_first_exact_span(paragraph_text, part):
                    return part
    return ""


def validate_llm_issue_payload(raw_issue: Dict[str, Any]) -> List[str]:
    errors: List[str] = []
    category = str(raw_issue.get("category", "")).strip()
    severity = str(raw_issue.get("severity", "")).strip()
    confidence = str(raw_issue.get("confidence", "")).strip()
    coverage_domain = str(raw_issue.get("coverage_domain", "")).strip()
    review_basis = str(raw_issue.get("review_basis", "")).strip()
    external_evidence_type = str(raw_issue.get("external_evidence_type", "")).strip()
    comment_intent = str(raw_issue.get("comment_intent", "")).strip()
    if category not in LLM_ALLOWED_CATEGORIES:
        errors.append(f"category invalid: {category or '<empty>'}")
    if severity not in LLM_ALLOWED_SEVERITIES:
        errors.append(f"severity invalid: {severity or '<empty>'}")
    if confidence not in LLM_ALLOWED_CONFIDENCES:
        errors.append(f"confidence invalid: {confidence or '<empty>'}")
    if coverage_domain and coverage_domain not in ALLOWED_COVERAGE_DOMAINS:
        errors.append(f"coverage_domain invalid: {coverage_domain}")
    if review_basis and review_basis not in ALLOWED_REVIEW_BASIS:
        errors.append(f"review_basis invalid: {review_basis}")
    if external_evidence_type and external_evidence_type not in ALLOWED_EXTERNAL_EVIDENCE_TYPES:
        errors.append(f"external_evidence_type invalid: {external_evidence_type}")
    if comment_intent and comment_intent not in ALLOWED_COMMENT_INTENTS:
        errors.append(f"comment_intent invalid: {comment_intent}")
    for field in ("original", "issue", "suggestion"):
        if not str(raw_issue.get(field, "")).strip():
            errors.append(f"{field} is empty")
    return errors


def is_low_value_llm_style_issue(raw_issue: Dict[str, Any]) -> bool:
    original = str(raw_issue.get("original", "")).strip()
    issue = str(raw_issue.get("issue", "")).strip()
    suggestion = str(raw_issue.get("suggestion", "")).strip()
    evidence = str(raw_issue.get("evidence", "")).strip()
    combined = f"{issue}\n{suggestion}\n{evidence}"
    lowered = combined.casefold()

    if is_objective_formula_case_issue(raw_issue):
        return False
    if is_objective_repeated_whitespace_issue(raw_issue):
        return False
    if is_signature_date_placeholder_column_claim(raw_issue):
        return True

    issue_suggestion = f"{issue}\n{suggestion}"
    lowered_issue_suggestion = issue_suggestion.casefold()
    table_title_context = (
        bool(re.search(r"\bTable\s*\d+\s*[.:]\s*\S", original))
        or any(token in issue_suggestion for token in ("表格标题", "表题", "标题格式", "表编号", "表号"))
        or any(token in lowered_issue_suggestion for token in ("table title", "table caption", "caption", "heading", "table number"))
    )
    table_title_punctuation_or_format = (
        any(token in issue_suggestion for token in ("标点", "空格", "冒号", "句点", "句号", "标题格式", "表题格式"))
        or any(token in lowered_issue_suggestion for token in ("punctuation", "space", "spacing", "period", "colon", "title format", "caption format"))
        or (
            any(token in issue_suggestion for token in ("格式", "保持一致", "统一"))
            and any(token in issue_suggestion for token in ("其余表格", "文档其余", "同级标题", "其他表格"))
        )
        or (
            any(token in lowered_issue_suggestion for token in ("format", "consistent", "consistency", "standardize"))
            and any(token in lowered_issue_suggestion for token in ("other table", "rest of the document", "same-level heading"))
        )
    )
    semantic_title_content = (
        any(token in issue_suggestion for token in ("语义", "含义", "错译", "误译", "数值", "数据", "条件", "样品", "对象", "检测项", "方法名称"))
        or any(
            token in lowered_issue_suggestion
            for token in (
                "semantic",
                "meaning",
                "mistranslation",
                "wrong translation",
                "sample",
                "value",
                "condition",
                "object",
                "incorrect term",
                "wrong term",
            )
        )
    )
    if (
        table_title_context
        and table_title_punctuation_or_format
        and not semantic_title_content
        and not has_approved_term_or_template_basis(raw_issue)
    ):
        return True

    table_title_spacing = (
        re.search(r"\bTable\s*\d+\s*[.:]\s*[A-Za-z]", original)
        and any(token in lowered for token in ("space", "spacing", "punctuation"))
        and any(token in lowered for token in ("table number", "table title", "caption", "heading", "period", "colon"))
        and not any(token in lowered for token in ("data", "method", "condition", "mistranslation", "incorrect term"))
    )
    if table_title_spacing and not has_approved_term_or_template_basis(raw_issue):
        return True

    capitalization_claim = any(
        token in lowered
        for token in (
            "improper capitalization",
            "capitalized",
            "capitalize",
            "lowercase",
            "upper case",
            "lower case",
            "common noun",
            "mid-sentence",
        )
    ) or any(token in combined for token in ("大小写", "首字母大写", "首字母小写"))
    capitalization_has_business_basis = (
        has_approved_term_or_template_basis(raw_issue)
        or "abbreviation" in lowered
        or "proper noun" in lowered and "not a proper noun" not in lowered
        or any(token in lowered for token in ("misspell", "spelling", "brand", "registered", "official"))
        or any(token in combined for token in ("缩写", "专有名词", "注册", "官方"))
    )
    if capitalization_claim and not capitalization_has_business_basis:
        return True

    if any(token in combined for token in ("标点层次不统一", "顿号和逗号混用", "标点层次", "标点混用")) and not any(
        token in combined for token in ("连续", "重复", "缺少", "多余", "错误", "数值", "方法")
    ):
        return True

    if re.search(r"中英文均有.{0,12}(?:标点|空格).{0,12}问题", issue) and not any(
        token in combined for token in ("数值不一致", "方法不一致", "条件不一致", "错译", "含义")
    ):
        return True

    if any(token in lowered for token in ("percentage notation", "percent notation", "percentage values include", "suffix")) and any(
        token in lowered for token in ("standardize", "add '%' to all", "remove it from all", "统一")
    ):
        return True

    if (YEAR_DOT_MONTH_PATTERN.search(original) or MONTH_YEAR_PATTERN.search(original)) and (
        "日期格式" in combined
        or "中文通常" in combined
        or "不一致" in combined
        or "中文翻译" in combined
    ):
        return True

    if "语序" in combined and not any(
        token in combined for token in ("歧义", "错译", "语法", "含义", "合规", "误解", "相反")
    ):
        return True

    if any(token in combined for token in ("句子结构堆叠", "衔接不自然", "语意不清", "合理衔接")) and not any(
        token in combined for token in ("数据", "数值", "方法", "条件", "错译", "矛盾", "不一致", "合规")
    ):
        return True

    if (
        "lux" in lowered
        and "lumen per square meter" in lowered
        and any(token in combined for token in ("不够规范", "定义方式", "考虑改为"))
        and not any(token in combined for token in ("不一致", "矛盾", "错译", "数值", "条件"))
    ):
        return True

    if (
        "目录" in combined
        and "空格" in combined
        and any(token in combined for token in ("排版规范", "多余空格", "格式", "标题"))
        and not has_approved_term_or_template_basis(raw_issue)
    ):
        return True

    if (
        ("百分号" in combined or "%号" in combined or "% 号" in combined)
        and any(token in combined for token in ("格式不一致", "统一", "部分数值", "部分单元格", "部分带", "部分不带"))
        and not any(token in combined for token in ("数值不一致", "数据不一致", "中英文", "条件", "方法", "错译", "矛盾"))
    ):
        return True

    if (
        "空格" in combined
        and "同一表格" in combined
        and any(token in combined for token in ("格式不一致", "部分单元格", "部分写法", "前后空格"))
        and not any(token in combined for token in ("缺少空格", "漏空格", "数值", "方法", "条件", "错译", "矛盾"))
    ):
        return True

    if (
        "/" in original
        and ("空格" in combined or "space" in lowered or "spacing" in lowered)
        and (
            any(token in combined for token in ("斜杠", "前后空格", "空格格式", "格式统一"))
            or any(token in lowered for token in ("slash", "solidus"))
        )
        and not is_objective_repeated_whitespace_issue(raw_issue)
        and not any(token in combined for token in ("数值不一致", "方法不一致", "条件不一致", "错译", "矛盾", "含义改变"))
    ):
        return True

    if (
        any(token in combined for token in ("同一文档", "不同表格", "混用"))
        and any(token in combined for token in ("格式不一致", "格式统一", "应统一", "建议统一"))
        and not any(token in combined for token in ("缺少空格", "漏空格", "拼写错误", "数值", "方法", "条件", "错译", "矛盾"))
    ):
        return True

    if (
        not has_approved_term_or_template_basis(raw_issue)
        and (
            ("heading" in lowered and any(token in lowered for token in ("numbering prefix", "numeric prefix", "section 1")))
            or any(token in combined for token in ("标题编号", "编号前缀", "章节编号", "缺少编号"))
        )
        and any(token in lowered for token in ("missing", "prefix", "consistency", "add the prefix"))
    ):
        return True

    if (
        any(token in lowered for token in ("punctuation/capitalization", "sentence capitalization", "should be capitalized", "following word"))
        and any(token in lowered for token in ("minor", "after", "period"))
        and not any(token in lowered for token in ("proper noun", "abbreviation", "misspelled", "mistranslation", "data", "method", "condition"))
    ):
        return True

    if any(token in combined for token in ("更自然", "更通顺", "表达风格", "建议润色")) and not any(
        token in combined for token in ("歧义", "错译", "语法错误", "事实", "合规")
    ):
        return True

    capitalization_has_source_local_basis = any(
        token in lowered for token in ("terminology", "proper noun", "abbreviation", "misspelled")
    )
    if any(token in lowered for token in ("inconsistent capitalization", "capitalization style", "title case")) and not (
        has_approved_term_or_template_basis(raw_issue) or capitalization_has_source_local_basis
    ):
        return True

    if (
        any(token in combined for token in ("动词冗余", "语义重复", "赘余动词", "结构冗余", "成分冗余", "进行执行"))
        or any(token in lowered for token in ("redundant verb", "redundant wording", "word redundancy"))
    ) and not any(token in combined for token in ("数据", "数值", "方法不一致", "条件不一致", "错译", "合规", "影响含义")):
        return True

    if (
        ("空格" in combined or "space" in lowered or "spacing" in lowered)
        and any(token in combined for token in ("中文逗号", "逗号后", "， "))
        and not any(token in combined for token in ("数值", "方法", "条件", "错译", "矛盾", "合规"))
    ):
        return True

    if (
        "、" in original
        and any(token in lowered for token in ("english title", "english document title", "chinese enumeration comma"))
        and not has_approved_term_or_template_basis(raw_issue)
    ):
        return True

    if is_unapproved_terminology_standardization(raw_issue):
        return True
    if any(
        token in lowered
        for token in (
            "more natural",
            "more fluent",
            "awkward phrasing",
            "better style",
            "style preference",
            "wording preference",
            "consider rephrasing",
            "improve readability",
            "sounds better",
            "inconsistent capitalization",
            "capitalization style",
        )
    ) and not any(
        token in lowered
        for token in (
            "ambiguous",
            "mistranslation",
            "grammar error",
            "grammatical error",
            "factual",
            "compliance",
            "inconsistent",
            "missing",
            "incorrect",
            "misspelled",
        )
    ):
        return True

    if any(token in combined for token in ("大小写", "首字母大写", "首字母未大写", "标题大小写")) and not (
        has_approved_term_or_template_basis(raw_issue)
        or any(token in combined for token in ("术语", "专有名词", "缩写", "拼写"))
    ):
        return True

    return False


def is_nonlocal_na_percentage_table_claim(raw_issue: Dict[str, Any], record: Optional[Dict[str, Any]]) -> bool:
    """Filter table cross-column claims that are not proven at the anchor cell.

    A single table cell containing `106%` does not by itself prove it is the
    English counterpart of a neighboring `N/A` cell. Those claims require row
    schema or source-record review and should not become direct Word comments.
    """

    if record is None or paragraph_document_zone(record) != "table":
        return False
    paragraph_text = str(record.get("text") or "").strip()
    combined = "\n".join(
        str(raw_issue.get(key) or "")
        for key in ("original", "issue", "suggestion", "evidence")
    )
    lowered = combined.casefold()
    if "n/a" not in lowered or not re.search(r"\d+(?:\.\d+)?%", combined):
        return False
    if not (("中文" in combined and "英文" in combined) or ("chinese" in lowered and "english" in lowered)):
        return False
    if "n/a" in paragraph_text.casefold():
        return False
    if not re.fullmatch(r"\s*\d+(?:\.\d+)?%\s*", paragraph_text):
        return False
    if any(marker in lowered for marker in ("lims", "原始", "source record", "record")):
        return True
    return "核对" in combined and any(token in combined for token in ("数据", "记录", "检测"))


def is_signature_date_placeholder_column_claim(raw_issue: Dict[str, Any]) -> bool:
    """Filter electronic-form signature/date placeholder column complaints.

    SOP and execution-form drafts often expose print-signature placeholder
    headers while row content only carries business columns. Without a verified
    template rule, this is not a source-local document defect.
    """

    combined = "\n".join(
        str(raw_issue.get(key) or "")
        for key in ("original", "issue", "suggestion", "evidence", "anchor_quote")
    )
    lowered = combined.casefold()
    signature_headers = (
        ("签字/日期" in combined or "签名/日期" in combined or "signature/date" in lowered)
        and ("复核/日期" in combined or "审核/日期" in combined or "review/date" in lowered)
    )
    if not signature_headers:
        return False
    column_mismatch_claim = (
        (
            "表头" in combined
            and "数据行" in combined
            and any(token in combined for token in ("列数", "列结构", "定义了", "只有", "仅包含", "无第", "缺少"))
        )
        or (
            "header row" in lowered
            and "data row" in lowered
            and any(token in lowered for token in ("column", "absent", "missing", "only have"))
        )
    )
    if not column_mismatch_claim:
        return False
    correction_targets_signature_columns = any(
        token in combined
        for token in (
            "补充签字/日期",
            "补充签名/日期",
            "补充复核/日期",
            "缩减表头",
            "调整表头列数",
            "统一表头与数据行",
        )
    ) or any(token in lowered for token in ("signature/date", "review/date", "adjust the header", "align the column"))
    return correction_targets_signature_columns and not has_approved_term_or_template_basis(raw_issue)


def make_llm_issue(
    issue_id: int,
    paragraphs: Sequence[Dict[str, Any]],
    raw_issue: Dict[str, Any],
    *,
    require_unit_locator: bool = False,
    allowed_unit_ids: set[str] | None = None,
) -> Optional[Dict[str, Any]]:
    if validate_llm_issue_payload(raw_issue):
        return None
    if is_low_value_llm_style_issue(raw_issue):
        return None
    category = str(raw_issue.get("category", "")).strip()
    severity = str(raw_issue.get("severity", "")).strip()
    confidence = str(raw_issue.get("confidence", "")).strip()
    original = str(raw_issue.get("original", "")).strip()
    issue = str(raw_issue.get("issue", "")).strip()
    suggestion = str(raw_issue.get("suggestion", "")).strip()
    if not original or not issue or not suggestion:
        return None

    unit_id = agent_unit_id(raw_issue)
    if allowed_unit_ids is not None and unit_id:
        normalized_allowed_unit_ids = {str(allowed_unit_id).strip().casefold() for allowed_unit_id in allowed_unit_ids}
        if unit_id.casefold() not in normalized_allowed_unit_ids:
            return None
    early_comment_intent = normalize_enum(raw_issue.get("comment_intent"), ALLOWED_COMMENT_INTENTS, "")
    if is_explicit_agent_global_summary_issue(raw_issue, early_comment_intent) and not unit_id:
        return None
    if require_unit_locator and agent_issue_requires_unit_locator(raw_issue, early_comment_intent) and not unit_id:
        return None
    if unit_id:
        record = locate_record_by_unit_id(paragraphs, raw_issue)
        if record is None:
            return None
        span, anchor_text = find_anchor_quote_span(record, raw_issue)
        if not span:
            return None
        if not source_anchor_texts_compatible(original, anchor_text):
            return None
    else:
        record = locate_llm_record(paragraphs, raw_issue)
        paragraph_text_value = str(record.get("text", "")) if record else ""
        span, anchor_text = resolve_llm_anchor(paragraph_text_value, raw_issue) if record else (None, original)
    if is_nonlocal_na_percentage_table_claim(raw_issue, record):
        return None
    paragraph_text_value = str(record.get("text", "")) if record else ""
    if span and should_reject_replacement_target_anchor(raw_issue, anchor_text):
        span = None
    if is_unsafe_embedded_numeric_anchor(paragraph_text_value, span, anchor_text, raw_issue):
        span = None
    if unit_id and not span:
        return None
    if record and not span:
        for nearby_record in nearby_llm_target_records(paragraphs, raw_issue, record):
            if nearby_record is record:
                continue
            nearby_text = str(nearby_record.get("text", ""))
            nearby_span, nearby_anchor_text = resolve_llm_anchor(nearby_text, raw_issue)
            if nearby_span and should_reject_replacement_target_anchor(raw_issue, nearby_anchor_text):
                nearby_span = None
            if (
                nearby_span
                and not has_weak_agent_anchor(nearby_anchor_text)
                and not is_unsafe_embedded_numeric_anchor(nearby_text, nearby_span, nearby_anchor_text, raw_issue)
            ):
                record = nearby_record
                span = nearby_span
                anchor_text = nearby_anchor_text
                break
    if record and not span:
        unique_anchor = find_unique_source_local_typo_anchor(paragraphs, raw_issue)
        if unique_anchor:
            record, span, anchor_text = unique_anchor
    has_stable_anchor = bool(record and span) and not has_weak_agent_anchor(anchor_text)
    comment_intent = early_comment_intent
    external_type = str(raw_issue.get("external_evidence_type") or "").strip()
    coverage_domain = str(raw_issue.get("coverage_domain") or "").strip()
    review_basis = str(raw_issue.get("review_basis") or "").strip()
    requires_external = infer_requires_external_evidence(
        requires_external_evidence=raw_issue.get("requires_external_evidence"),
        external_evidence_type=external_type,
        review_basis=review_basis,
        coverage_domain=coverage_domain,
        comment_intent=comment_intent,
    )
    if not passes_user_visible_agent_quality_gate(
        raw_issue=raw_issue,
        has_stable_anchor=has_stable_anchor,
        confidence=confidence,
        requires_external=requires_external,
        comment_intent=comment_intent,
    ):
        return None
    status = ISSUE_STATUS_CONFIRMED
    is_global_summary = is_explicit_agent_global_summary_issue(raw_issue, comment_intent)
    severity = normalized_agent_visible_severity(severity, raw_issue)
    issue_data = make_issue(
        issue_id,
        rule_id=LLM_CATEGORY_RULES[category],
        branch=LLM_REVIEW_BRANCH,
        paragraph_index=None,
        location=paragraph_location(record) if record else "全文审核意见",
        original=original,
        issue=issue,
        suggestion=suggestion,
        severity=severity,
        evidence=str(raw_issue.get("evidence", "")).strip(),
        document_zone=paragraph_document_zone(record) if record else "body",
        location_kind=paragraph_location_kind(record) if record else "paragraph",
        anchor_locator=paragraph_anchor_locator(record) if record else "",
        anchor_span=span,
        anchor_text=anchor_text,
        status=status,
        comments_added=1 if has_stable_anchor else 0,
        confidence=confidence,
        match_method="span" if span else "inference",
        notes="全文级审核意见，追加到文末全局审核意见。" if is_global_summary and not has_stable_anchor else "",
        requires_external_evidence=requires_external,
        external_evidence_type=external_type or ("other" if requires_external else "none"),
        coverage_domain=coverage_domain,
        review_basis=review_basis,
        comment_intent=comment_intent,
    )
    if unit_id:
        issue_data["unit_id"] = unit_id
    for field_name in ("problem_unit_id", "pair_id", "anchor_quote", "evidence_unit_ids"):
        if field_name in raw_issue:
            issue_data[field_name] = raw_issue[field_name]
    return issue_data


def classify_skipped_agent_issue(
    raw_issue: Dict[str, Any],
    paragraphs: Sequence[Dict[str, Any]],
    *,
    allowed_unit_ids: set[str] | None = None,
) -> Tuple[str, str]:
    """Classify hidden semantic findings so audit logs distinguish noise from retriable misses."""

    schema_errors = validate_llm_issue_payload(raw_issue)
    if schema_errors:
        return "schema_invalid", "; ".join(schema_errors)
    if is_low_value_llm_style_issue(raw_issue):
        return "filtered_low_value", "low-value style or wording preference filtered"

    confidence = str(raw_issue.get("confidence") or "").strip()
    comment_intent = normalize_enum(raw_issue.get("comment_intent"), ALLOWED_COMMENT_INTENTS, "")
    external_type = str(raw_issue.get("external_evidence_type") or "").strip()
    coverage_domain = str(raw_issue.get("coverage_domain") or "").strip()
    review_basis = str(raw_issue.get("review_basis") or "").strip()
    requires_external = infer_requires_external_evidence(
        requires_external_evidence=raw_issue.get("requires_external_evidence"),
        external_evidence_type=external_type,
        review_basis=review_basis,
        coverage_domain=coverage_domain,
        comment_intent=comment_intent,
    )
    if requires_external or comment_intent == "request_check":
        return "filtered_external", "requires external evidence or manual source records"
    if agent_issue_requires_unit_locator(raw_issue, comment_intent) and not agent_unit_id(raw_issue):
        return "missing_unit_locator", "non-global agent finding missing required unit_id locator"
    unit_id = agent_unit_id(raw_issue)
    if allowed_unit_ids is not None and unit_id:
        normalized_allowed_unit_ids = {str(allowed_unit_id).strip().casefold() for allowed_unit_id in allowed_unit_ids}
        if unit_id.casefold() not in normalized_allowed_unit_ids:
            return "missing_unit_locator", "agent finding unit_id is not present in the exported context package"
    if confidence == "low":
        return "filtered_low_confidence", "low-confidence finding filtered"
    objective_typo = is_objective_source_local_typo_issue(raw_issue)
    if has_speculative_language(raw_issue) and not objective_typo:
        return "filtered_quality_gate", "speculative language filtered"
    if is_external_brand_or_product_name_claim(raw_issue):
        return "filtered_external", "external brand or product-name spelling claim requires approved terminology evidence"
    if is_external_official_terminology_rewrite(raw_issue):
        return "filtered_external", "official terminology or material-name rewrite requires approved terminology evidence"
    if is_unapproved_terminology_standardization(raw_issue):
        return "filtered_low_value", "terminology standardization requires approved glossary or direct local mistranslation"
    if is_high_risk_formula_semantic_issue(raw_issue):
        return "filtered_quality_gate", "high-risk formula or symbol claim needs stronger proof"
    if not has_explicit_source_evidence(raw_issue):
        return "filtered_quality_gate", "missing explicit source-backed evidence"
    if not has_actionable_agent_suggestion(raw_issue):
        return "filtered_quality_gate", "suggestion is vague or not actionable"

    record = locate_llm_record(paragraphs, raw_issue)
    if is_nonlocal_na_percentage_table_claim(raw_issue, record):
        return "filtered_external", "table cross-column N/A versus percentage claim requires row schema or source records"
    paragraph_text_value = str(record.get("text", "")) if record else ""
    span, anchor_text = resolve_llm_anchor(paragraph_text_value, raw_issue) if record else (None, "")
    if span and should_reject_replacement_target_anchor(raw_issue, anchor_text):
        span = None
    if is_unsafe_embedded_numeric_anchor(paragraph_text_value, span, anchor_text, raw_issue):
        span = None
    has_stable_anchor = bool(record and span) and not has_weak_agent_anchor(anchor_text)
    if agent_issue_requires_unit_locator(raw_issue, comment_intent):
        anchor_quote = strip_agent_text_label(str(raw_issue.get("anchor_quote") or "")).strip(" \t\r\n:：;；,，。.-")
        if not agent_unit_id(raw_issue) or not anchor_quote or has_weak_agent_anchor(anchor_quote):
            return "missing_unit_locator", "agent finding missing required unit_id or anchor_quote locator"
    if not has_stable_anchor and not is_explicit_agent_global_summary_issue(raw_issue, comment_intent):
        return "anchor_failure_should_retry", "high-value finding could not be mapped to a stable source span"

    return "filtered_quality_gate", "failed user-visible quality gate"


def semantic_overlap_tokens(item: Dict[str, Any]) -> set[str]:
    blob = "\n".join(
        str(item.get(key) or "")
        for key in ("original", "anchor_text", "issue", "suggestion", "evidence")
    )
    blob = blob.translate(ANCHOR_CHAR_TRANSLATION).casefold()
    tokens: set[str] = set()
    for match in re.findall(r"[a-z0-9][a-z0-9_.+-]{2,}", blob):
        tokens.add(match)
    for match in re.findall(r"[\u4e00-\u9fff]{2,}", blob):
        if len(match) <= 4:
            tokens.add(match)
            continue
        tokens.update(match[index : index + 2] for index in range(len(match) - 1))
    return tokens


def semantic_overlap_ratio(left: Dict[str, Any], right: Dict[str, Any]) -> float:
    left_tokens = semantic_overlap_tokens(left)
    right_tokens = semantic_overlap_tokens(right)
    if not left_tokens or not right_tokens:
        return 0.0
    return len(left_tokens & right_tokens) / max(1, min(len(left_tokens), len(right_tokens)))


def issue_location_matches_raw_agent_issue(issue: Dict[str, Any], raw_issue: Dict[str, Any]) -> bool:
    target = llm_paragraph_target(raw_issue.get("paragraph_index", ""))
    if not target:
        return False
    location_text = " ".join(
        str(issue.get(key) or "")
        for key in ("location", "anchor_locator", "paragraph_index")
    )
    return any(str(int(value)) == target for value in re.findall(r"\d+", location_text))


def risk_external_veto_matches_issue(veto: Dict[str, Any], issue: Dict[str, Any]) -> bool:
    if str(issue.get("branch") or "") != LLM_REVIEW_BRANCH:
        return False
    if str(issue.get("source_agent") or "") == "risk_classifier":
        return False
    if is_direct_internal_source_contradiction(issue):
        return False
    if not issue_location_matches_raw_agent_issue(issue, veto):
        return False

    veto_category = str(veto.get("category") or "")
    issue_rule = str(issue.get("rule_id") or "")
    if veto_category == "bilingual_consistency" and issue_rule not in {"LLM-BI-001", "LLM-SEM-001"}:
        return False
    if veto_category == "en_language" and issue_rule not in {"LLM-EN-001", "LLM-SEM-001"}:
        return False
    if veto_category == "zh_language" and issue_rule not in {"LLM-ZH-001", "LLM-SEM-001"}:
        return False

    veto_original = normalized_text(str(veto.get("original") or "")).casefold()
    issue_original = normalized_text(str(issue.get("original") or issue.get("anchor_text") or "")).casefold()
    if len(veto_original) >= 8 and len(issue_original) >= 8 and (
        veto_original in issue_original or issue_original in veto_original
    ):
        return True

    return semantic_overlap_ratio(veto, issue) >= 0.25


def load_agent_review_issues(
    review_json_paths: Sequence[str],
    paragraphs: Sequence[Dict[str, Any]],
    next_issue_id: int,
    *,
    allowed_unit_ids: set[str] | None = None,
) -> Tuple[List[Dict[str, Any]], Dict[str, Any]]:
    """Load semantic findings produced by an outer agent or SDK subagents.

    This keeps model-dependent review outside this deterministic runner. The
    runner only validates schema, maps stable anchors, and routes uncertain
    findings into the human review queue.
    """

    issues: List[Dict[str, Any]] = []
    human_review_items: List[Dict[str, Any]] = []
    source_files: List[str] = []
    errors: List[str] = []
    schema_invalid_count = 0
    skipped_agent_issue_count = 0
    skip_categories: Counter[str] = Counter()
    issue_id = next_issue_id
    external_vetoes: List[Dict[str, Any]] = []

    for raw_path in review_json_paths:
        path = Path(raw_path).expanduser().resolve()
        source_files.append(str(path))
        try:
            payload = load_json_with_repair(path)
        except Exception as exc:
            errors.append(f"{path.name}: {describe_exception(exc)}")
            human_review_items.append(
                make_human_review_item(
                    LLM_REVIEW_BRANCH,
                    f"agent review JSON could not be loaded: {describe_exception(exc)}",
                    path.name,
                    user_visible=False,
                    skip_category="load_error",
                )
            )
            skip_categories["load_error"] += 1
            continue

        raw_issues = payload.get("issues") if isinstance(payload, dict) else None
        if not isinstance(raw_issues, list):
            errors.append(f"{path.name}: field `issues` must be a list")
            human_review_items.append(
                make_human_review_item(
                    LLM_REVIEW_BRANCH,
                    "agent review JSON schema invalid: field `issues` must be a list",
                    path.name,
                    user_visible=False,
                    skip_category="schema_invalid",
                )
            )
            skip_categories["schema_invalid"] += 1
            continue

        for raw_index, raw_issue in enumerate(raw_issues, start=1):
            if not isinstance(raw_issue, dict):
                schema_invalid_count += 1
                skipped_agent_issue_count += 1
                skip_categories["schema_invalid"] += 1
                continue
            normalized_issue = normalize_agent_issue_payload(
                raw_issue,
                agent_role=str(payload.get("agent_role") or "").strip(),
            )
            normalized_agent_role = str(payload.get("agent_role") or "").strip()
            is_risk_classifier_agent = is_risk_classifier_agent_role(normalized_agent_role)
            is_risk_classifier_veto = bool(
                is_risk_classifier_agent
                and (
                    normalized_issue.get("requires_external_evidence")
                    or str(normalized_issue.get("comment_intent") or "") == "request_check"
                )
            )
            if is_risk_classifier_veto:
                external_vetoes.append(normalized_issue)
            if is_current_date_regulatory_status_noise(normalized_issue):
                skipped_agent_issue_count += 1
                skip_categories["filtered_external"] += 1
                continue
            schema_errors = validate_llm_issue_payload(normalized_issue)
            if schema_errors:
                schema_invalid_count += 1
                skipped_agent_issue_count += 1
                skip_categories["schema_invalid"] += 1
                continue
            if is_risk_classifier_agent and not is_risk_classifier_veto:
                skipped_agent_issue_count += 1
                skip_categories["filtered_quality_gate"] += 1
                human_review_items.append(
                    make_human_review_item(
                        LLM_REVIEW_BRANCH,
                        f"{path.name} issue #{raw_index} skipped: risk classifier outputs are gating signals, not direct findings",
                        str(normalized_issue.get("unit_id") or normalized_issue.get("paragraph_index") or path.name),
                        user_visible=False,
                        skip_category="filtered_quality_gate",
                    )
                )
                continue
            issue = make_llm_issue(
                issue_id,
                paragraphs,
                normalized_issue,
                require_unit_locator=True,
                allowed_unit_ids=allowed_unit_ids,
            )
            if issue:
                issue["source"] = "qa-file-reviewer-agent-review"
                issue["source_agent"] = str(payload.get("agent_role") or issue.get("source_agent") or "calling-agent")
                issues.append(issue)
                issue_id += 1
                continue
            location = str(normalized_issue.get("paragraph_index", "")).strip()
            skipped_agent_issue_count += 1
            skip_category, skip_reason = classify_skipped_agent_issue(
                normalized_issue,
                paragraphs,
                allowed_unit_ids=allowed_unit_ids,
            )
            skip_categories[skip_category] += 1
            human_review_items.append(
                make_human_review_item(
                    LLM_REVIEW_BRANCH,
                    f"{path.name} issue #{raw_index} skipped: {skip_reason}",
                    f"P{location.lstrip('Pp')}" if location else path.name,
                    user_visible=False,
                skip_category=skip_category,
                )
            )

    if external_vetoes and issues:
        kept_issues: List[Dict[str, Any]] = []
        for issue in issues:
            matched_veto = next((veto for veto in external_vetoes if risk_external_veto_matches_issue(veto, issue)), None)
            if matched_veto is None:
                kept_issues.append(issue)
                continue
            skipped_agent_issue_count += 1
            skip_categories["filtered_external_veto"] += 1
            human_review_items.append(
                make_human_review_item(
                    LLM_REVIEW_BRANCH,
                    "confirmed agent issue skipped: risk classifier marked the same source-backed concern as requiring external evidence",
                    f"P{llm_paragraph_target(matched_veto.get('paragraph_index', ''))}" if llm_paragraph_target(matched_veto.get("paragraph_index", "")) else LLM_REVIEW_BRANCH,
                    issue_id=str(issue.get("id") or ""),
                    user_visible=False,
                    skip_category="filtered_external_veto",
                )
            )
        issues = kept_issues

    metadata = {
        "status": "completed_with_errors" if errors else "",
        "error": summarize_chunk_failures(errors) if errors else "",
        "human_review_items": human_review_items,
        "source_files": source_files,
        "source_file_count": len(source_files),
        "loaded_issue_count": len(issues),
        "schema_invalid_count": schema_invalid_count,
        "skipped_agent_issue_count": skipped_agent_issue_count,
        "skip_categories": dict(skip_categories),
    }
    return unique_issues(issues), metadata


def build_manifest_entry(
    branch: str,
    issues: Sequence[Dict[str, Any]],
    duration_ms: int,
    error: str = "",
    status_override: str = "",
    *,
    human_review_items: Sequence[Dict[str, Any]] = (),
    branch_details: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    status = "completed_with_human_review" if any(item.get("status") in CHECK_STATUSES for item in issues) else "completed"
    visible_human_review_items = [item for item in human_review_items if item.get("user_visible") is not False]
    if visible_human_review_items and status == "completed":
        status = "completed_with_human_review"
    if error:
        status = "failed"
    if status_override:
        status = status_override
    entry = {
        "branch": branch,
        "agent_role": BRANCH_META[branch]["agent_role"],
        "status": status,
        "issue_count": len(issues),
        "duration_ms": duration_ms,
        "retry_count": 0,
        "error": error,
    }
    if human_review_items:
        entry["human_review_items"] = list(human_review_items)
    if branch_details:
        for key in (
            "chunk_count",
            "completed_chunks",
            "failed_chunks",
            "source",
            "source_files",
            "source_file_count",
            "loaded_issue_count",
            "schema_invalid_count",
            "skipped_agent_issue_count",
            "skip_categories",
            "agent_context",
        ):
            if key in branch_details:
                entry[key] = branch_details[key]
    return entry


def build_human_review_queue(
    issues: Sequence[Dict[str, Any]],
    manifest: Sequence[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    queue: List[Dict[str, Any]] = []
    for issue in issues:
        if issue.get("status") not in CHECK_STATUSES:
            continue
        queue.append(
            make_human_review_item(
                str(issue.get("branch", "")),
                str(issue.get("notes") or issue.get("comment_intent") or ISSUE_STATUS_NEEDS_USER_CHECK),
                str(issue.get("location", "")),
                issue_id=str(issue.get("id", "")),
            )
        )

    for item in manifest:
        if not isinstance(item, dict):
            continue
        branch = str(item.get("branch", ""))
        for extra in item.get("human_review_items", []):
            if not isinstance(extra, dict):
                continue
            if extra.get("user_visible") is False:
                continue
            queue.append(
                make_human_review_item(
                    str(extra.get("branch") or branch),
                    str(extra.get("reason") or ISSUE_STATUS_NEEDS_USER_CHECK),
                    str(extra.get("location") or branch or "unknown"),
                    issue_id=str(extra.get("issue_id", "")),
                )
            )

    unique_queue: List[Dict[str, Any]] = []
    seen = set()
    for item in queue:
        key = (
            item.get("issue_id", ""),
            item.get("branch", ""),
            item.get("reason", ""),
            item.get("location", ""),
        )
        if key in seen:
            continue
        seen.add(key)
        unique_queue.append(item)
    return unique_queue


def is_user_visible_issue(issue: Dict[str, Any]) -> bool:
    if issue.get("comment_visibility", "word_comment") != "word_comment":
        return False
    if str(issue.get("branch") or "") == LLM_REVIEW_BRANCH and (
        issue.get("requires_external_evidence")
        or str(issue.get("comment_intent") or "") == "request_check"
    ):
        return False
    if issue.get("status") not in WORD_VISIBLE_STATUSES:
        return False
    if issue.get("comments_added", 1) != 0:
        return True
    return bool(issue.get("append_to_document_end"))


def user_visible_issues(issues: Sequence[Dict[str, Any]]) -> List[Dict[str, Any]]:
    return [item for item in issues if is_user_visible_issue(item)]


def internal_diagnostic_issues(issues: Sequence[Dict[str, Any]]) -> List[Dict[str, Any]]:
    return [item for item in issues if not is_user_visible_issue(item)]


def build_summary(issues: Sequence[Dict[str, Any]]) -> Dict[str, Any]:
    visible_issues = user_visible_issues(issues)
    diagnostic_count = len(issues) - len(visible_issues)
    by_severity: Counter[str] = Counter()
    by_type: Counter[str] = Counter()
    for issue in visible_issues:
        by_severity[str(issue.get("severity", "次要"))] += 1
        by_type[str(issue.get("type", "其他"))] += 1
    return {
        "total_issues": len(visible_issues),
        "user_visible_issue_count": len(visible_issues),
        "diagnostic_issue_count": diagnostic_count,
        "by_severity": {
            "关键": by_severity.get("关键", 0),
            "主要": by_severity.get("主要", 0),
            "次要": by_severity.get("次要", 0),
        },
        "by_type": dict(sorted(by_type.items(), key=lambda item: item[0])),
    }


def build_quality_score(
    issues: Sequence[Dict[str, Any]],
    manifest: Sequence[Dict[str, Any]],
    commenting: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    penalties: List[Dict[str, Any]] = []
    by_severity = Counter(str(issue.get("severity", "次要")) for issue in issues)
    human_review_count = sum(1 for issue in issues if issue.get("status") in CHECK_STATUSES)

    def add_penalty(name: str, count: int, weight: int, cap: int) -> None:
        if count <= 0:
            return
        value = min(count * weight, cap)
        penalties.append({"name": name, "count": count, "penalty": value})

    add_penalty("critical_issues", by_severity.get("关键", 0), 25, 50)
    add_penalty("major_issues", by_severity.get("主要", 0), 10, 40)
    add_penalty("minor_issues", by_severity.get("次要", 0), 3, 15)
    add_penalty("human_review_items", human_review_count, 5, 20)

    failed_branches = [item for item in manifest if str(item.get("status")) == "failed"]
    partial_branches = [item for item in manifest if str(item.get("status")) == "completed_with_errors"]
    add_penalty("failed_branches", len(failed_branches), 20, 40)
    add_penalty("partial_branches", len(partial_branches), 10, 30)

    positioning = (commenting or {}).get("positioning_quality", {}) if isinstance(commenting, dict) else {}
    if isinstance(positioning, dict):
        add_penalty("comment_failures", int((commenting or {}).get("failed_total") or 0), 10, 30)
        document_end_fallback_count = (
            int(positioning.get("by_anchor_failure_document_end") or 0)
            if "by_anchor_failure_document_end" in positioning
            else int(positioning.get("by_document_end") or 0)
        )
        add_penalty("document_end_fallback_comments", document_end_fallback_count, 2, 20)
        add_penalty(
            "high_risk_ambiguous_anchor_comments",
            int(positioning.get("by_high_risk_ambiguous_anchor") or 0),
            5,
            20,
        )

    score = max(0, 100 - sum(int(item["penalty"]) for item in penalties))
    if score >= 90:
        grade = "A"
    elif score >= 75:
        grade = "B"
    elif score >= 60:
        grade = "C"
    else:
        grade = "D"
    def semantic_manifest_complete(item: Dict[str, Any]) -> bool:
        if str(item.get("branch", "")) != LLM_REVIEW_BRANCH:
            return False
        status = str(item.get("status", ""))
        if status == "completed":
            return True
        if status != "completed_with_human_review":
            return False
        return int(item.get("loaded_issue_count") or 0) > 0 and int(item.get("schema_invalid_count") or 0) == 0

    semantic_complete = any(semantic_manifest_complete(item) for item in manifest)
    return {
        "score": score,
        "grade": grade,
        "basis": "100 minus deterministic penalties for issue severity, incomplete branches, human review items, and comment positioning.",
        "penalties": penalties,
        "coverage": {
            "required_branches": [branch for branch in ("format", "project_number", CONTENT_CONSISTENCY_BRANCH)],
            "executed_branches": [str(item.get("branch", "")) for item in manifest],
            "semantic_agent_review": semantic_complete,
            "semantic_agent_review_attempted": any(str(item.get("branch", "")) == LLM_REVIEW_BRANCH for item in manifest),
        },
    }


def validate_comment_positioning_quality(commenting: Dict[str, Any]) -> List[str]:
    positioning = commenting.get("positioning_quality", {}) if isinstance(commenting, dict) else {}
    if not isinstance(positioning, dict):
        return ["commenting.positioning_quality must be an object after comment insertion"]

    failures: List[str] = []
    expected_total = int(commenting.get("expected_total") or 0)
    high_risk_ambiguous = int(positioning.get("by_high_risk_ambiguous_anchor") or 0)
    if high_risk_ambiguous > 0:
        failures.append(
            f"high-risk ambiguous comment anchors need adjudication: {high_risk_ambiguous}/{expected_total}"
        )

    document_end = (
        int(positioning.get("by_anchor_failure_document_end") or 0)
        if "by_anchor_failure_document_end" in positioning
        else int(positioning.get("by_document_end") or 0)
    )
    if document_end > 0:
        failures.append(f"anchor-failed comments must not fall back to document end: {document_end}/{expected_total}")
        return failures
    if expected_total <= 0 or document_end <= MAX_DOCUMENT_END_FALLBACK_COMMENTS:
        return failures

    ratio = document_end / expected_total
    if ratio > MAX_DOCUMENT_END_FALLBACK_RATIO:
        failures.append(
            "too many comments fell back to document end: "
            f"{document_end}/{expected_total} exceeds "
            f"{MAX_DOCUMENT_END_FALLBACK_COMMENTS} and {MAX_DOCUMENT_END_FALLBACK_RATIO:.0%}"
        )
    return failures


def write_json(path: Path, payload: Dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def write_text(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")


def build_comment_plan(
    input_docx: Path,
    issues: Sequence[Dict[str, Any]],
    document_map_path: Path,
    document_display_name: str = "",
    human_review_queue: Sequence[Dict[str, Any]] = (),
) -> Dict[str, Any]:
    issue_ids = {str(item.get("id", "")) for item in issues if isinstance(item, dict)}
    visible_issues = [
        comment_plan_issue(item)
        for item in issues
        if is_issue_visible_in_comment_plan(item)
    ]
    visible_issues.extend(
        human_review_queue_comment_issue(item, index)
        for index, item in enumerate(human_review_queue, start=1)
        if item.get("user_visible", True)
        and (
            not str(item.get("issue_id", "")).strip()
            or str(item.get("issue_id", "")).strip() not in issue_ids
        )
    )
    return {
        "schema_version": "qa-file-reviewer.comment-plan.v1",
        "document": str(input_docx),
        "document_display_name": document_display_name or input_docx.name,
        "generated_at": time.strftime("%Y-%m-%dT%H:%M:%S"),
        "document_map": str(document_map_path),
        "issues": visible_issues,
    }


def is_issue_visible_in_comment_plan(issue: Dict[str, Any]) -> bool:
    return is_user_visible_issue(issue)


def human_review_queue_comment_issue(item: Dict[str, Any], index: int) -> Dict[str, Any]:
    branch = str(item.get("branch") or "human_review")
    reason = str(item.get("reason") or "需人工核对").strip()
    location = str(item.get("location") or "全文人工复核").strip()
    issue_id = str(item.get("issue_id") or f"human-review-{index:04d}").strip()
    comment_text = build_comment_text(
        reason,
        "请核对该项对应的语义审核、外部依据或系统执行状态，确认后再修订文件。",
        "主要",
        issue_type="人工复核",
        original=location,
        evidence=reason,
        status=ISSUE_STATUS_NEEDS_USER_CHECK,
        comment_intent="request_check",
        review_basis="external_required",
        external_evidence_type="other",
    )
    return {
        "id": issue_id,
        "rule_id": "HUMAN-REVIEW-QUEUE",
        "type": "人工复核",
        "branch": branch,
        "agent_role": "human-review-queue",
        "document_zone": "body",
        "location_kind": "global",
        "location": location,
        "anchor_locator": "",
        "anchor_span": None,
        "anchor_text": location,
        "original": location,
        "issue": reason,
        "severity": "主要",
        "suggestion": "请核对该项对应的语义审核、外部依据或系统执行状态，确认后再修订文件。",
        "comment_text": comment_text,
        "evidence": reason,
        "match_method": "inference",
        "preexisting_comment_count": 0,
        "comments_added": 1,
        "confidence": "medium",
        "status": ISSUE_STATUS_NEEDS_USER_CHECK,
        "comment_visibility": "word_comment",
        "requires_external_evidence": True,
        "external_evidence_type": "other",
        "coverage_domain": "external_check",
        "review_basis": "external_required",
        "comment_intent": "request_check",
        "source": "qa-file-reviewer-human-review-queue",
        "source_agent": "human-review-queue",
        "notes": reason,
        "append_to_document_end": True,
        "appendix_reason": reason,
    }


def comment_plan_issue(issue: Dict[str, Any]) -> Dict[str, Any]:
    copied = dict(issue)
    if copied.get("comments_added", 1) == 0:
        copied["comments_added"] = 1
        copied["append_to_document_end"] = True
        copied["appendix_reason"] = copied.get("notes") or "原文位置无法稳定定位，已追加到文末全局审核意见。"
    return copied


def adjudicate_issues_before_comment_plan(
    issues: Sequence[Dict[str, Any]],
    review_paragraphs: Sequence[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    adjudicated = adjudicate_issues(issues, review_paragraphs)
    return renumber_issues(sort_issues(adjudicated))


def sync_manifest_issue_counts(manifest: Sequence[Dict[str, Any]], issues: Sequence[Dict[str, Any]]) -> None:
    """Keep branch manifest counts aligned with final post-adjudication issues."""

    counts = Counter(str(issue.get("branch") or "") for issue in issues)
    for entry in manifest:
        branch = str(entry.get("branch") or "")
        if branch:
            entry["issue_count"] = int(counts.get(branch, 0))


def audit_docx_with_minimax(document_path: Path, report_path: Path) -> Dict[str, Any]:
    command = [
        sys.executable,
        str(MINIMAX_DOCX_ENGINE_PATH),
        "audit",
        str(document_path),
    ]
    process = subprocess.run(command, capture_output=True, text=True, encoding="utf-8", errors="replace")
    report = {
        "success": process.returncode == 0,
        "stdout": process.stdout,
        "stderr": process.stderr,
        "exit_code": process.returncode,
        "document": str(document_path),
    }
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    return report


def validate_pipeline(review_json: Path, report_json: Path) -> Dict[str, Any]:
    command = [sys.executable, str(VALIDATOR_PATH), str(review_json), "--report-json", str(report_json)]
    process = subprocess.run(command, capture_output=True, text=True, encoding="utf-8", errors="replace")
    report = json.loads(report_json.read_text(encoding="utf-8"))
    report["stdout"] = process.stdout
    report["stderr"] = process.stderr
    report["exit_code"] = process.returncode
    return report


def renumber_issues(issues: Sequence[Dict[str, Any]]) -> List[Dict[str, Any]]:
    renumbered: List[Dict[str, Any]] = []
    for index, item in enumerate(issues, start=1):
        copied = dict(item)
        copied["id"] = f"issue-{index:04d}"
        renumbered.append(copied)
    return renumbered


def execute_branches(
    doc: Document,
    review_paragraphs: Sequence[Dict[str, Any]],
    current_project: str,
) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
    branch_results: List[Dict[str, Any]] = []
    issues: List[Dict[str, Any]] = []
    branch_order = [
        "format",
        "project_number",
        CONTENT_CONSISTENCY_BRANCH,
    ]

    def execute_one(branch: str, issue_base: int) -> Tuple[str, List[Dict[str, Any]], Dict[str, Any]]:
        start = time.perf_counter()
        error = ""
        status_override = ""
        branch_issues: List[Dict[str, Any]] = []
        human_review_items: List[Dict[str, Any]] = []
        branch_details: Optional[Dict[str, Any]] = None
        try:
            if branch == "format":
                branch_issues = run_format_branch(doc, review_paragraphs, issue_base)
            elif branch == "project_number":
                branch_issues = run_project_branch(review_paragraphs, current_project, issue_base)
            elif branch == CONTENT_CONSISTENCY_BRANCH:
                branch_issues = run_content_consistency_branch(review_paragraphs, issue_base)
            else:
                raise ValueError(f"unknown review branch: {branch}")
        except Exception as exc:  # pragma: no cover - defensive
            error = str(exc)
            branch_issues = []
        duration_ms = int((time.perf_counter() - start) * 1000)
        manifest_entry = build_manifest_entry(
            branch,
            branch_issues,
            duration_ms,
            error=error,
            status_override=status_override,
            human_review_items=human_review_items,
            branch_details=branch_details,
        )
        return branch, branch_issues, manifest_entry

    max_workers = min(len(branch_order), max(1, env_int("QA_BRANCH_MAX_WORKERS", len(branch_order))))
    branch_payloads: Dict[str, Tuple[List[Dict[str, Any]], Dict[str, Any]]] = {}
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = {
            executor.submit(execute_one, branch, 1 + index * 10000): branch
            for index, branch in enumerate(branch_order)
        }
        for future in as_completed(futures):
            branch, branch_issues, manifest_entry = future.result()
            branch_payloads[branch] = (branch_issues, manifest_entry)

    for branch in branch_order:
        branch_issues, manifest_entry = branch_payloads[branch]
        branch_results.append(manifest_entry)
        issues.extend(branch_issues)

    return branch_results, renumber_issues(sort_issues(unique_issues(issues)))


def build_review_payload(
    input_docx: Path,
    issues: Sequence[Dict[str, Any]],
    manifest: Sequence[Dict[str, Any]],
    *,
    document_map_path: Path,
    comment_plan_path: Path,
    document_display_name: str = "",
) -> Dict[str, Any]:
    human_review_queue = build_human_review_queue(issues, manifest)
    visible_issues = user_visible_issues(issues)
    diagnostics = internal_diagnostic_issues(issues)
    return {
        "document": str(input_docx),
        "document_display_name": document_display_name or input_docx.name,
        "review_time": time.strftime("%Y-%m-%dT%H:%M:%S"),
        "reviewer": "Codex Document Reviewer",
        "review_scope": "full",
        "branch": "merge",
        "branch_execution_manifest": list(manifest),
        "human_review_queue": human_review_queue,
        "commenting": {
            "policy": "preserve_human_strip_prior_automated",
            "existing_total": 0,
            "removed_existing_automated": 0,
            "added_total": 0,
            "failed_total": 0,
            "expected_total": 0,
            "actual_total": 0,
            "positioning_quality": {
                "by_span": 0,
                "by_exact_text": 0,
                "by_contains_text": 0,
                "by_inference": 0,
                "by_document_end": 0,
                "by_expected_document_end": 0,
                "by_anchor_failure_document_end": 0,
                "by_ambiguous_anchor": 0,
                "by_high_risk_ambiguous_anchor": 0,
                "failed": 0,
                "warnings": [],
            },
        },
        "summary": {
            **build_summary(issues),
            "quality_score": build_quality_score(visible_issues, manifest),
        },
        "internal_diagnostics": diagnostics,
        "artifacts": {
            "document_map": str(document_map_path),
            "comment_plan": str(comment_plan_path),
            "reviewed_docx": "",
            "validation_report": "",
        },
        "issues": list(issues),
    }


def print_summary(payload: Dict[str, Any], validation_report: Dict[str, Any]) -> None:
    summary = payload["summary"]
    queue_count = len(payload.get("human_review_queue", []))
    print("文件审核")
    print(f"总问题数: {summary['total_issues']}")
    print(f"关键/主要/次要: {summary['by_severity']['关键']}/{summary['by_severity']['主要']}/{summary['by_severity']['次要']}")
    print(f"人工复核: {queue_count}")
    if not validation_report.get("passed"):
        print("系统校验: 未通过")


def truncate_text(text: str, limit: int = 160) -> str:
    value = (text or "").strip().replace("\n", " ")
    if len(value) <= limit:
        return value
    return value[: limit - 3] + "..."


def friendly_positioning_reason(reason: str) -> str:
    text = str(reason or "").strip()
    lowered = text.lower()
    if "no stable anchor matched" in lowered:
        return "该问题的原文片段在目标段落中无法形成唯一、非重叠的批注锚点。"
    if "unsupported xml container" in lowered:
        return "该段文字位于当前批注器尚未安全支持的 Word XML 容器中。"
    if "paragraph has no text runs" in lowered:
        return "该位置没有可写入批注的正文文本。"
    if "anchor span could not be materialized" in lowered:
        return "已找到目标片段，但无法安全拆分为 Word 批注范围。"
    if "no stable anchor" in lowered:
        return "未找到稳定的批注锚点。"
    return text or "定位失败。"


def build_detailed_report(
    payload: Dict[str, Any],
    validation_report: Dict[str, Any],
    *,
    current_project: str,
    commented_docx: str,
) -> str:
    lines: List[str] = []
    summary = payload.get("summary", {})
    issues = user_visible_issues(payload.get("issues", []))
    review_time = payload.get("review_time", "")
    doc_path = payload.get("document_display_name") or payload.get("document", "")
    human_review_queue = payload.get("human_review_queue", [])
    commenting = payload.get("commenting", {})

    lines.append("文件审核报告")
    lines.append("=" * 60)
    lines.append(f"文档: {doc_path}")
    lines.append(f"审核时间: {review_time}")
    lines.append(f"识别项目号: {current_project}")
    lines.append("")

    lines.append("审核概览")
    lines.append("-" * 60)
    lines.append(f"总问题数: {summary.get('total_issues', 0)}")
    by_severity = summary.get("by_severity", {})
    lines.append(
        f"关键/主要/次要: {by_severity.get('关键', 0)}/{by_severity.get('主要', 0)}/{by_severity.get('次要', 0)}"
    )
    lines.append(f"需人工复核: {len(human_review_queue)}")
    if not validation_report.get("passed"):
        lines.append("系统校验: 未通过，建议联系管理员重新审核。")
    lines.append("")

    if commented_docx or any(int(commenting.get(key, 0) or 0) for key in ("expected_total", "added_total", "failed_total")):
        positioning_quality = commenting.get("positioning_quality", {}) or {}
        warnings = positioning_quality.get("warnings", []) or []
        lines.append("批注概览")
        lines.append("-" * 60)
        lines.append(
            "预计/新增/失败: "
            f"{int(commenting.get('expected_total', 0) or 0)}/"
            f"{int(commenting.get('added_total', 0) or 0)}/"
            f"{int(commenting.get('failed_total', 0) or 0)}"
        )
        removed_existing = int(commenting.get("removed_existing_automated", 0) or 0)
        if removed_existing:
            lines.append(f"已剥离旧自动批注: {removed_existing}")
        if warnings:
            lines.append("未写入批注的定位问题:")
            for item in warnings[:5]:
                reason = truncate_text(friendly_positioning_reason(str(item.get("reason", "定位失败"))), 160)
                location = truncate_text(str(item.get("location", "")), 120)
                lines.append(f"- {location}: {reason}")
        lines.append("")

    lines.append("问题清单")
    lines.append("-" * 60)
    for index, issue in enumerate(issues, start=1):
        lines.append(f"{index}. [{issue.get('severity', '')}] {issue.get('issue', '')}")
        lines.append(f"   位置: {issue.get('location', '')}")
        lines.append(f"   原文: {truncate_text(str(issue.get('original', '')), 180)}")
        lines.append(f"   建议: {truncate_text(str(issue.get('suggestion', '')), 220)}")
        evidence = truncate_text(str(issue.get("evidence", "")), 240)
        if evidence and evidence != str(issue.get("original", "")):
            lines.append(f"   依据: {evidence}")
    if not issues:
        lines.append("未发现需要提示的问题。")
    lines.append("")

    if human_review_queue:
        lines.append("人工复核项")
        lines.append("-" * 60)
        for index, item in enumerate(human_review_queue, start=1):
            lines.append(f"{index}. {truncate_text(str(item.get('reason', '需人工复核')), 220)}")
            lines.append(f"   位置: {truncate_text(str(item.get('location', '')), 180)}")
        lines.append("")

    if commented_docx:
        lines.append(f"批注版 Word: {commented_docx}")
    return "\n".join(lines) + "\n"


def remove_if_exists(path: Path) -> None:
    try:
        if path.exists():
            path.unlink()
    except OSError:
        pass


def main() -> int:
    ensure_utf8_stdio()
    args = parse_args()

    input_docx = Path(args.input_docx).expanduser().resolve()
    output_dir = Path(args.output_dir).expanduser().resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    document_display_name = display_name_from_filename(args.original_filename, input_docx.name)
    output_stem = output_stem_from_filename(args.original_filename, input_docx.stem)
    project_source = Path(document_display_name)
    document_map_json = output_dir / "document_map.json"
    comment_plan_json = output_dir / "comment_plan.json"
    report_path = (
        Path(args.report_file).expanduser().resolve()
        if args.report_file
        else output_dir / f"{output_stem}_审核详细报告.txt"
    )
    cleanup_json_artifacts: List[Path] = [document_map_json, comment_plan_json]

    doc = load_document(input_docx)
    document_map = extract_document_map(input_docx, document_map_json)
    review_paragraphs = get_review_paragraphs(document_map)
    current_project = detect_current_project(project_source, review_paragraphs)

    manifest, issues = execute_branches(doc, review_paragraphs, current_project)
    if args.agent_review_json:
        agent_review_started = time.perf_counter()
        context_manifest_path = (
            Path(args.agent_context_manifest).expanduser().resolve()
            if args.agent_context_manifest
            else output_dir / "agent_context_manifest.json"
        )
        allowed_agent_unit_ids, agent_context_metadata = load_agent_context_unit_ids(
            context_manifest_path,
            require_manifest=bool(args.agent_context_manifest),
        )
        if agent_context_metadata.get("status") == "failed":
            agent_issues = []
            agent_metadata = {
                "status": "failed",
                "error": str(agent_context_metadata.get("context_manifest_error") or "agent context manifest invalid"),
                "human_review_items": [],
                "source_files": list(args.agent_review_json),
                "source_file_count": len(args.agent_review_json),
                "loaded_issue_count": 0,
                "schema_invalid_count": 0,
                "skipped_agent_issue_count": 0,
                "skip_categories": {},
            }
        else:
            agent_issues, agent_metadata = load_agent_review_issues(
                args.agent_review_json,
                review_paragraphs,
                next_issue_id=len(issues) + 1,
                allowed_unit_ids=allowed_agent_unit_ids,
            )
        manifest.append(
            build_manifest_entry(
                LLM_REVIEW_BRANCH,
                agent_issues,
                int((time.perf_counter() - agent_review_started) * 1000),
                error=str(agent_metadata.get("error", "")),
                status_override=str(agent_metadata.get("status", "")),
                human_review_items=agent_metadata.get("human_review_items", []),
                branch_details={
                    "source": "external_agent_review_json",
                    "source_files": agent_metadata.get("source_files", []),
                    "source_file_count": agent_metadata.get("source_file_count", 0),
                    "loaded_issue_count": agent_metadata.get("loaded_issue_count", 0),
                    "schema_invalid_count": agent_metadata.get("schema_invalid_count", 0),
                    "skipped_agent_issue_count": agent_metadata.get("skipped_agent_issue_count", 0),
                    "skip_categories": agent_metadata.get("skip_categories", {}),
                    "agent_context": agent_context_metadata,
                },
            )
        )
        issues = renumber_issues(sort_issues(unique_issues([*issues, *agent_issues])))
    pre_adjudication_issue_count = len(issues)
    issues = adjudicate_issues_before_comment_plan(issues, review_paragraphs)
    deduped_issue_count = max(0, pre_adjudication_issue_count - len(issues))
    if deduped_issue_count:
        for entry in manifest:
            if entry.get("branch") == LLM_REVIEW_BRANCH:
                entry["deduped_issue_count"] = deduped_issue_count
                break
    sync_manifest_issue_counts(manifest, issues)
    human_review_queue = build_human_review_queue(issues, manifest)
    comment_plan = build_comment_plan(
        input_docx,
        issues,
        document_map_json,
        document_display_name,
        human_review_queue=human_review_queue,
    )
    write_json(comment_plan_json, comment_plan)
    payload = build_review_payload(
        input_docx,
        issues,
        manifest,
        document_map_path=document_map_json,
        comment_plan_path=comment_plan_json,
        document_display_name=document_display_name,
    )

    review_json = output_dir / "review_result.json"
    validation_report_json = output_dir / "validation_report.json"
    cleanup_json_artifacts.append(review_json)
    cleanup_json_artifacts.append(validation_report_json)
    write_json(review_json, payload)

    validation_report = validate_pipeline(review_json, validation_report_json)
    commented_docx = ""
    docx_audit_report_json = output_dir / "docx_audit_report.json"
    if not validation_report.get("passed"):
        report_text = build_detailed_report(
            payload,
            validation_report,
            current_project=current_project,
            commented_docx=commented_docx,
        )
        write_text(report_path, report_text)
        print_summary(payload, validation_report)
        print("门禁校验失败，未继续生成批注版 Word。")
        if validation_report.get("errors"):
            for item in validation_report["errors"]:
                print(f"- {item}")
        print(f"详细报告: {report_path}")
        if not args.keep_json_artifacts:
            for artifact in cleanup_json_artifacts:
                remove_if_exists(artifact)
        return 1

    if args.with_comments:
        commented_docx = str(output_dir / f"{output_stem}_reviewed.docx")
        comment_result = add_comments_improved(str(input_docx), str(comment_plan_json), commented_docx, author=args.author)
        if not comment_result.get("success"):
            report_text = build_detailed_report(
                payload,
                validation_report,
                current_project=current_project,
                commented_docx=commented_docx,
            )
            write_text(report_path, report_text)
            print_summary(payload, validation_report)
            print("批注写回失败。")
            print(comment_result.get("error", "unknown error"))
            print(f"详细报告: {report_path}")
            if not args.keep_json_artifacts:
                for artifact in cleanup_json_artifacts:
                    remove_if_exists(artifact)
            return 1

        payload["commenting"].update(
            {
                "existing_total": comment_result.get("existing_comments", 0),
                "removed_existing_automated": comment_result.get("removed_existing_comments", 0),
                "added_total": comment_result.get("comments_added", 0),
                "failed_total": comment_result.get("comments_failed", 0),
                "expected_total": comment_result.get("comments_expected", 0),
                "actual_total": comment_result.get("comments_actual", 0),
                "positioning_quality": comment_result.get("positioning_quality", payload["commenting"]["positioning_quality"]),
            }
        )
        payload["summary"]["quality_score"] = build_quality_score(user_visible_issues(issues), manifest, payload["commenting"])
        cleanup_json_artifacts.append(docx_audit_report_json)
        payload["artifacts"]["reviewed_docx"] = commented_docx
        payload["artifacts"]["validation_report"] = str(validation_report_json)
        payload["artifacts"]["docx_audit_report"] = str(docx_audit_report_json)
        positioning_errors = validate_comment_positioning_quality(payload["commenting"])
        if positioning_errors:
            validation_report.setdefault("errors", []).extend(positioning_errors)
            validation_report["passed"] = False
            write_json(validation_report_json, validation_report)
            write_json(review_json, payload)
            report_text = build_detailed_report(
                payload,
                validation_report,
                current_project=current_project,
                commented_docx=commented_docx,
            )
            write_text(report_path, report_text)
            print_summary(payload, validation_report)
            print("批注定位质量未达到上线门禁。")
            for item in positioning_errors:
                print(f"- {item}")
            print(f"详细报告: {report_path}")
            if commented_docx:
                print(f"批注文档: {commented_docx}")
            if not args.keep_json_artifacts:
                for artifact in cleanup_json_artifacts:
                    remove_if_exists(artifact)
            return 1
        write_json(review_json, payload)
        docx_audit = audit_docx_with_minimax(Path(commented_docx), docx_audit_report_json)
        if not docx_audit.get("success"):
            payload["artifacts"]["docx_audit_report"] = str(docx_audit_report_json)
            payload["docx_audit"] = {
                "success": False,
                "error": (
                    docx_audit.get("stderr", "").strip()
                    or docx_audit.get("stdout", "").strip()
                    or docx_audit.get("error", "")
                    or "unknown error"
                ),
            }
            write_json(review_json, payload)
            report_text = build_detailed_report(
                payload,
                validation_report,
                current_project=current_project,
                commented_docx=commented_docx,
            )
            write_text(report_path, report_text)
            print_summary(payload, validation_report)
            print("minimax-docx 文档有效性校验失败。")
            print(docx_audit.get("stderr", "").strip() or docx_audit.get("stdout", "").strip() or "unknown error")
            print(f"详细报告: {report_path}")
            if not args.keep_json_artifacts:
                for artifact in cleanup_json_artifacts:
                    remove_if_exists(artifact)
            return 1
        payload["docx_audit"] = {"success": True}
        write_json(review_json, payload)

    report_text = build_detailed_report(
        payload,
        validation_report,
        current_project=current_project,
        commented_docx=commented_docx,
    )
    write_text(report_path, report_text)
    print_summary(payload, validation_report)
    print(f"当前项目号: {current_project}")
    print(f"详细报告: {report_path}")
    if commented_docx:
        print(f"批注文档: {commented_docx}")
    if not args.keep_json_artifacts:
        for artifact in cleanup_json_artifacts:
            remove_if_exists(artifact)
    return 0


if __name__ == "__main__":
    sys.exit(main())

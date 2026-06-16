import sys
import json
import tempfile
import unittest
from pathlib import Path

from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt


SKILL_DIR = Path(__file__).resolve().parents[1]
SCRIPTS_DIR = SKILL_DIR / "scripts"
sys.path.insert(0, str(SCRIPTS_DIR))

import run_qa_review as qa  # noqa: E402
import validate_pipeline_gate as gate  # noqa: E402
import add_word_comments_v3 as comments  # noqa: E402


def paragraph(text: str, logical_index: int = 1) -> dict:
    return {
        "paragraph_id": f"p-{logical_index:05d}",
        "text": text,
        "logical_index": logical_index,
        "xml_index": logical_index,
        "type": "paragraph",
        "location_kind": "paragraph",
        "document_zone": "body",
        "anchor_locator": f"paragraph:{logical_index}",
    }


def table_cell(text: str, logical_index: int = 1, table_index: int = 7, row_index: int = 5, cell_index: int = 3) -> dict:
    item = paragraph(text, logical_index=logical_index)
    item.update(
        {
            "document_zone": "table",
            "location_kind": "table",
            "table_index": table_index,
            "row_index": row_index,
            "cell_index": cell_index,
        }
    )
    return item


def agent_issue(**overrides) -> dict:
    payload = {
        "category": "semantic_consistency",
        "severity": "主要",
        "paragraph_index": "1",
        "original": "原液稳定性样品在25 °C ± 2 °C条件下放置。",
        "issue": "温度条件与同段英文表述不一致。",
        "suggestion": "请将中英文温度条件统一为同一条件。",
        "evidence": "同一段内中文和英文温度条件不一致。",
        "confidence": "high",
        "coverage_domain": "data_consistency",
        "review_basis": "single_doc_internal",
        "requires_external_evidence": False,
        "external_evidence_type": "none",
        "comment_intent": "suggest_change",
    }
    payload.update(overrides)
    return payload


def word_paragraph(text: str):
    p = comments.etree.Element(f"{comments.W}p")
    r = comments.etree.SubElement(p, f"{comments.W}r")
    t = comments.etree.SubElement(r, f"{comments.W}t")
    t.text = text
    return p


class QualityGateTests(unittest.TestCase):
    def test_sample_shape_specific_fallback_constants_are_not_present(self):
        self.assertFalse(hasattr(qa, "AGITATION_CONDITION_PATTERN"))
        self.assertFalse(hasattr(qa, "DS_EXAMINED_PATTERN"))
        self.assertFalse(any(name.startswith("REDUCED_CE_SDS_") for name in vars(qa)))

    def test_short_source_replacement_typo_is_accepted(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("The Cilent sample met the acceptance criteria.")],
            agent_issue(
                category="en_language",
                original="Cilent",
                issue="Typographical error: `Cilent` should be `Client`.",
                suggestion="Replace `Cilent` with `Client`.",
                evidence="The source sentence contains `Cilent`.",
                coverage_domain="en_language",
                review_basis="agent_semantic",
            ),
        )

        self.assertIsNotNone(issue)
        self.assertEqual(issue["anchor_text"], "Cilent")

    def test_short_duplicate_token_replacement_is_accepted(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("The HMW HMW result was reported in Table 1.")],
            agent_issue(
                category="en_language",
                original="HMW HMW",
                issue="Duplicate adjacent token: `HMW HMW`.",
                suggestion="Replace `HMW HMW` with `HMW`.",
                evidence="The source sentence contains `HMW HMW`.",
                coverage_domain="en_language",
                review_basis="agent_semantic",
            ),
        )

        self.assertIsNotNone(issue)
        self.assertEqual(issue["anchor_text"], "HMW HMW")

    def test_agent_runtime_error_text_is_filtered_from_word_comments(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("The visible document sentence is valid.")],
            agent_issue(
                category="semantic_consistency",
                paragraph_index="",
                original="agent review JSON parse error",
                issue="SDK JSON parse error should be reviewed.",
                suggestion="Check the agent runtime log.",
                evidence="agent review JSON parse error",
                coverage_domain="semantic",
                review_basis="agent_runtime",
            ),
        )

        self.assertIsNone(issue)

    def test_non_global_agent_issue_without_stable_anchor_is_filtered(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("本段没有模型输出中的原文锚点。")],
            agent_issue(paragraph_index="1"),
        )

        self.assertIsNone(issue)

    def test_low_confidence_agent_issue_is_filtered_even_with_anchor(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("原液稳定性样品在25 °C ± 2 °C条件下放置。")],
            agent_issue(confidence="low"),
        )

        self.assertIsNone(issue)

    def test_agent_issue_without_explicit_evidence_is_filtered(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("原液稳定性样品在25 °C ± 2 °C条件下放置。")],
            agent_issue(evidence=""),
        )

        self.assertIsNone(issue)

    def test_agent_issue_with_vague_suggestion_is_filtered(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("原液稳定性样品在25 °C ± 2 °C条件下放置。")],
            agent_issue(suggestion="请核对是否一致。"),
        )

        self.assertIsNone(issue)

    def test_english_actionable_suggestion_is_localized_in_visible_comments(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("Almost colorless, not more intensely colored than EP Y7Standard Solution.")],
            agent_issue(
                category="en_language",
                original="EP Y7Standard Solution",
                issue="Missing space between Y7 and Standard Solution.",
                suggestion="Insert a space between Y7 and Standard Solution.",
                evidence="Original text contains `EP Y7Standard Solution` without a space.",
                coverage_domain="en_language",
                review_basis="agent_semantic",
            ),
        )

        self.assertIsNotNone(issue)
        self.assertEqual(issue["anchor_text"], "EP Y7Standard Solution")
        self.assertEqual(issue["comments_added"], 1)
        self.assertIn("发现：Y7 与 Standard Solution 之间缺少空格。", issue["comment_text"])
        self.assertIn("建议：在 Y7 和 Standard Solution 之间补充空格。", issue["comment_text"])
        self.assertNotIn("Missing space", issue["comment_text"])
        self.assertNotIn("Insert a space", issue["comment_text"])

    def test_english_source_with_chinese_issue_and_suggestion_is_accepted(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("Almost colorless, not more intensely colored than EP Y7Standard Solution.")],
            agent_issue(
                category="en_language",
                original="EP Y7Standard Solution",
                issue="英文中 Y7 与 Standard Solution 之间缺少空格。",
                suggestion="将原文改为 `EP Y7 Standard Solution`。",
                evidence="原文 `EP Y7Standard Solution` 中 Y7 与 Standard Solution 连写。",
                coverage_domain="en_language",
                review_basis="agent_semantic",
            ),
        )

        self.assertIsNotNone(issue)
        self.assertEqual(issue["anchor_text"], "EP Y7Standard Solution")
        self.assertIn("发现：英文中", issue["comment_text"])
        self.assertIn("建议：将原文改为", issue["comment_text"])

    def test_direct_bare_replacement_suggestions_are_accepted_when_source_local(self):
        cases = [
            {
                "text": "审批信息：客户Cilent。",
                "original": "Cilent",
                "issue": "Spelling error: 'Cilent' should be 'Client'.",
                "suggestion": "Client",
                "evidence": "The same source cell shows '客户Cilent'.",
            },
            {
                "text": "洗脱方式：等度洗脱Isometric elution。",
                "original": "Isometric elution",
                "issue": "Scientific terminology error: 'Isometric' is incorrect in chromatography context.",
                "suggestion": "Isocratic elution",
                "evidence": "The same table cell reads '等度洗脱Isometric elution'.",
            },
            {
                "text": "The results included monomers, HMW HMW, and LMW.",
                "original": "HMW HMW",
                "issue": "Duplicate abbreviation: 'HMW' appears twice consecutively.",
                "suggestion": "HMW",
                "evidence": "The source sentence reads 'monomers, HMW HMW, and LMW'.",
            },
            {
                "text": "十二水合磷酸氢二钠（Na2HPO4·12H2O）Disodium Hydrogen Phosphate Decahydrate",
                "original": "Disodium Hydrogen Phosphate Decahydrate",
                "issue": "Chemical hydrate number is incorrect because Na2HPO4·12H2O contradicts Decahydrate.",
                "suggestion": "Disodium Hydrogen Phosphate Dodecahydrate",
                "evidence": "The same table cell states Na2HPO4·12H2O and the English says Decahydrate.",
            },
            {
                "text": "组成compose",
                "original": "compose",
                "issue": "Incorrect word form: 'compose' is a verb used as a table heading.",
                "suggestion": "Composition",
                "evidence": "The same table heading reads '组成compose'.",
            },
            {
                "text": "filter with 0.22 μm filter membrane after mix well.",
                "original": "filter with 0.22 μm filter membrane after mix well",
                "issue": "Grammar error: 'after mix well' should use the gerund form.",
                "suggestion": "filter with 0.22 μm filter membrane after mixing well",
                "evidence": "The source phrase is 'filter with 0.22 μm filter membrane after mix well'.",
            },
        ]

        for index, case in enumerate(cases, start=1):
            with self.subTest(original=case["original"]):
                issue = qa.make_llm_issue(
                    index,
                    [table_cell(case["text"], logical_index=index)],
                    agent_issue(
                        category="en_language",
                        paragraph_index=f"P{index}",
                        original=case["original"],
                        issue=case["issue"],
                        suggestion=case["suggestion"],
                        evidence=case["evidence"],
                        coverage_domain="en_language",
                        review_basis="single_doc_internal",
                    ),
                )

                self.assertIsNotNone(issue)
                self.assertEqual(issue["anchor_text"], case["original"])
                self.assertEqual(issue["comments_added"], 1)

    def test_bare_replacement_title_spacing_preference_is_filtered(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("样品信息Table 2.Sample information")],
            agent_issue(
                category="en_language",
                original="Table 2.Sample information",
                issue="Missing space after the period following the table number.",
                suggestion="Table 2. Sample information",
                evidence="The paragraph text reads '样品信息Table 2.Sample information'.",
                coverage_domain="en_language",
                review_basis="single_doc_internal",
            ),
        )

        self.assertIsNone(issue)

    def test_semantic_table_title_punctuation_format_claim_is_filtered_without_template_rule(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("积分示例参数Table 9:Integration example parameters")],
            agent_issue(
                category="semantic_consistency",
                original="Table 9:Integration example parameters",
                issue="表格标题标点不一致：此处使用冒号且无空格，文档其余表格标题均使用句点加空格格式。",
                suggestion='将 "Table 9:Integration" 改为 "Table 9. Integration"，与文档其余表格标题格式保持一致。',
                evidence='P20: "Table 7. Sequence parameters"; P30: "Table 9:Integration example parameters"。',
                coverage_domain="structure",
                review_basis="single_doc_internal",
            ),
        )

        self.assertIsNone(issue)

    def test_semantic_table_title_meaning_error_is_not_filtered_as_title_style(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("样品A检测结果Table 9. Summary of Sample B results")],
            agent_issue(
                category="semantic_consistency",
                original="Table 9. Summary of Sample B results",
                issue="表格标题语义不一致：中文标题为样品A检测结果，英文标题写为Sample B results。",
                suggestion="将英文标题改为 Table 9. Summary of Sample A results。",
                evidence="同一标题中中文为样品A检测结果，英文为Sample B results。",
                coverage_domain="structure",
                review_basis="single_doc_internal",
            ),
        )

        self.assertIsNotNone(issue)
        self.assertEqual(issue["anchor_text"], "Table 9. Summary of Sample B results")

    def test_generic_missing_bilingual_counterpart_claim_is_filtered(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("The stability samples were stored at 25 °C ± 2 °C.")],
            agent_issue(
                category="bilingual_consistency",
                original="The stability samples were stored at 25 °C ± 2 °C.",
                issue="英文存在明确的试验条件句，但中文缺少对应内容。",
                suggestion="补齐中文对应句，保持与英文在信息层级和顺序上基本一致。",
                evidence="该英文段落有试验条件，但相邻位置未见中文对应句。",
                coverage_domain="bilingual",
                review_basis="agent_semantic",
            ),
        )

        self.assertIsNone(issue)

    def test_bare_replacement_common_noun_capitalization_preference_is_filtered(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("The area percentage of Monomer peak was calculated.")],
            agent_issue(
                category="en_language",
                original="Monomer peak",
                issue="Improper capitalization: 'Monomer' is a common noun mid-sentence.",
                suggestion="monomer peak",
                evidence="The sentence reads 'The area percentage of Monomer peak'.",
                coverage_domain="en_language",
                review_basis="single_doc_internal",
            ),
        )

        self.assertIsNone(issue)

    def test_labeled_bilingual_fragments_can_locate_record_without_paragraph_index(self):
        issue = qa.make_llm_issue(
            1,
            [
                paragraph("中文：有关物质限度为2.0%。 English: The related substances limit is 2.1%.", logical_index=4)
            ],
            agent_issue(
                category="bilingual_consistency",
                paragraph_index="",
                original="Chinese: 2.0%; English: 2.1%",
                issue="同一限度的中英文百分比不一致。",
                suggestion="将英文百分比改为 2.0%，使中英文限度保持一致。",
                evidence="同一段中文为 2.0%，英文为 2.1%。",
                coverage_domain="bilingual",
                review_basis="single_doc_internal",
            ),
        )

        self.assertIsNotNone(issue)
        self.assertTrue(issue["location"].startswith("第4段"))
        self.assertIn(issue["anchor_text"], {"2.0%", "2.1%"})
        self.assertEqual(issue["comments_added"], 1)

    def test_table_cell_temperature_tolerance_mismatch_is_insertable(self):
        issue = qa.make_llm_issue(
            1,
            [
                table_cell(
                    "分别用0.1 mol/L盐酸和0.1 mol/L NaOH调节样品的pH为4.0和8.0，"
                    "将样品置于25°C±2°C，60%RH±5%RH、遮光条件下。"
                    "The pH of the samples was adjusted to 4.0 and 8.0 with 0.1 mol/L "
                    "hydrochloric acid and 0.1 mol/L NaOH, respectively, and the samples "
                    "were stored at 25°C±5°C, 60%RH±5%RH and protected from light.",
                    logical_index=571,
                )
            ],
            agent_issue(
                paragraph_index="P571 (T7R5C3)",
                original='Chinese: "将样品置于25°C±2°C，60%RH±5%RH、遮光条件下"\n'
                'English: "the samples were stored at 25°C±5°C, 60%RH±5%RH and protected from light."',
                issue="English temperature tolerance differs from Chinese. Chinese states 25°C±2°C while English states 25°C±5°C for the same extreme pH test condition.",
                suggestion="Unify to 25°C±2°C in both Chinese and English, consistent with other test conditions in Table 5.",
                evidence="Chinese cell T7R5C3: 25°C±2°C; English cell T7R5C3: 25°C±5°C.",
                coverage_domain="data_consistency",
                review_basis="single_doc_internal",
            ),
        )

        self.assertIsNotNone(issue)
        self.assertEqual(issue["location_kind"], "table")
        self.assertEqual(issue["document_zone"], "table")
        self.assertEqual(issue["comments_added"], 1)
        self.assertIn("25°C±", issue["anchor_text"])

    def test_table_cell_micro_sign_variant_still_anchors(self):
        self.assertIsNotNone(qa.find_exact_span("0.22 μm sterile filter", "0.22 µm sterile filter"))

        issue = qa.make_llm_issue(
            1,
            [
                table_cell(
                    "Samples containing 0.2% tBHP were filter through a 0.22 μm sterile filter, "
                    "stored at 25°C±5°C, 60%RH±5%RH and protected from light.",
                    logical_index=574,
                    row_index=6,
                )
            ],
            agent_issue(
                category="en_language",
                paragraph_index="P574",
                original="Samples containing 0.2% tBHP were filter through a 0.22 µm sterile filter",
                issue='Incorrect verb form: "were filter" should be "were filtered".',
                suggestion='Change "were filter through" to "were filtered through".',
                evidence='The table cell source text uses "were filter through" in a passive construction.',
                coverage_domain="en_language",
                review_basis="single_doc_internal",
            ),
        )

        self.assertIsNotNone(issue)
        self.assertEqual(issue["comments_added"], 1)
        self.assertIn("were filter", issue["anchor_text"])

    def test_adjacent_bilingual_note_can_anchor_when_target_points_to_paired_note(self):
        issue = qa.make_llm_issue(
            1,
            [
                paragraph("注：*1表示引用放行数据", logical_index=843),
                paragraph("Note: *1 indicates this data reference to the releasing data.", logical_index=844),
            ],
            agent_issue(
                category="en_language",
                paragraph_index="P843",
                original="Note: *1 indicates this data reference to the releasing data.",
                issue='"data reference to the releasing data" is grammatically awkward.',
                suggestion='Change to "Note: *1 indicates data referenced from the lot release testing."',
                evidence='The paired English note says "data reference to the releasing data".',
                coverage_domain="en_language",
                review_basis="single_doc_internal",
            ),
        )

        self.assertIsNotNone(issue)
        self.assertEqual(issue["location"], "第844段（XML:844）")
        self.assertEqual(issue["anchor_text"], "Note: *1 indicates this data reference to the releasing data")

    def test_bilingual_title_mismatch_anchors_to_wrong_english_title(self):
        issue = qa.make_llm_issue(
            1,
            [
                paragraph(
                    "根据《Project X Protocol》（DOC-001），"
                    "According to < Project X Summary > (DOC-001), the test plan is shown in Table 5.",
                    logical_index=554,
                )
            ],
            agent_issue(
                category="bilingual_consistency",
                paragraph_index="P554",
                original="Project X Protocol vs Project X Summary",
                issue="The Chinese document title says Protocol, but English says Summary.",
                suggestion="Change the English title to `Project X Protocol`.",
                evidence="P554 Chinese title uses Protocol; English title uses Project X Summary.",
                confidence="high",
                coverage_domain="bilingual",
                review_basis="single_doc_internal",
            ),
        )

        self.assertIsNotNone(issue)
        self.assertEqual(issue["anchor_text"], "Project X Summary")

    def test_bilingual_value_mismatch_anchors_to_wrong_english_value(self):
        issue = qa.make_llm_issue(
            1,
            [
                paragraph(
                    "中文样品编号为 Sample A。English sample ID: Sample B.",
                    logical_index=846,
                )
            ],
            agent_issue(
                category="semantic_consistency",
                paragraph_index="P846",
                original="English sample ID: Sample B",
                issue="Chinese and English sample IDs differ.",
                suggestion="Change the English sample ID to `Sample A`.",
                evidence="P846 Chinese says Sample A; English says Sample B.",
                confidence="high",
                coverage_domain="data_consistency",
                review_basis="single_doc_internal",
            ),
        )

        self.assertIsNotNone(issue)
        self.assertEqual(issue["anchor_text"], "English sample ID: Sample B")

    def test_deterministic_content_branch_catches_table_temperature_tolerance_mismatch(self):
        issues = qa.run_content_consistency_branch(
            [
                table_cell(
                    "分别用0.1 mol/L盐酸和0.1 mol/L NaOH调节样品的pH为4.0和8.0，"
                    "将样品置于25°C±2°C，60%RH±5%RH、遮光条件下。"
                    "The pH of the samples was adjusted to 4.0 and 8.0 with 0.1 mol/L "
                    "hydrochloric acid and 0.1 mol/L NaOH, respectively, and the samples "
                    "were stored at 25°C±5°C, 60%RH±5%RH and protected from light.",
                    logical_index=571,
                ),
                table_cell(
                    "25°C±2°C，60%RH±5%RH。25°C ± 2°C, 60%RH±5%RH.",
                    logical_index=572,
                    row_index=7,
                ),
            ],
            1,
        )

        temp_issues = [item for item in issues if item["rule_id"] == "CONTENT-BI-TEMP-001"]
        self.assertEqual(len(temp_issues), 1)
        self.assertEqual(temp_issues[0]["anchor_text"], "25°C±5°C")
        self.assertEqual(temp_issues[0]["comments_added"], 1)
        self.assertFalse(temp_issues[0]["requires_external_evidence"])

    def test_deterministic_content_branch_catches_passive_filter_grammar(self):
        issues = qa.run_content_consistency_branch(
            [
                table_cell(
                    "Samples containing 0.2% tBHP were filter through a 0.22 μm sterile filter.",
                    logical_index=574,
                    row_index=6,
                )
            ],
            1,
        )

        grammar_issues = [item for item in issues if item["rule_id"] == "CONTENT-EN-GRAMMAR-001"]
        self.assertEqual(len(grammar_issues), 1)
        self.assertEqual(grammar_issues[0]["anchor_text"], "were filter")
        self.assertIn("were filtered", grammar_issues[0]["suggestion"])
        self.assertEqual(grammar_issues[0]["comments_added"], 1)

    def test_deterministic_content_branch_does_not_emit_removed_stress_fallback(self):
        text = (
            "将样品按对应试验条件进行6次循环冻融后，结果无明显变化。"
            "After the samples were processed by Method B under Condition B for five days, "
            "no significant changes were observed. Therefore, avoid freeze-thawing during storage."
        )
        issues = qa.run_content_consistency_branch([paragraph(text, logical_index=1398)], 1)

        self.assertFalse(any(item["rule_id"].startswith("CONTENT-BI-STRESS") for item in issues))

    def test_deterministic_content_branch_does_not_emit_removed_method_fallback(self):
        text = (
            "高温试验结果显示，两批样品Method A的指标分别下降了1.8%、2.0%。"
            "High-temperature study results: the same indicator by Method B "
            "decreased by 1.8% and 2.1%, respectively."
        )
        issues = qa.run_content_consistency_branch([paragraph(text, logical_index=846)], 1)

        self.assertFalse(any(item["rule_id"].startswith("CONTENT-BI-METHOD") for item in issues))
        self.assertFalse(any(item["rule_id"] == "CONTENT-BI-003" for item in issues))

        lc_ms_text = "去糖基化完整分子量（LC-MS法） Deglycosylated intact mass by LC-MS"
        lc_ms_issues = qa.run_content_consistency_branch([paragraph(lc_ms_text, logical_index=847)], 1)
        self.assertFalse(any(item["rule_id"] == "CONTENT-BI-003" for item in lc_ms_issues))

        lc_ms_adjacent_issues = qa.run_content_consistency_branch(
            [
                table_cell("去糖基化完整分子量（LC-MS法）", logical_index=848, table_index=1, row_index=1, cell_index=1),
                table_cell("Deglycosylated intact mass by LC-MS", logical_index=849, table_index=1, row_index=1, cell_index=1),
            ],
            1,
        )
        self.assertFalse(any(item["rule_id"] == "CONTENT-BI-003" for item in lc_ms_adjacent_issues))

    def test_deterministic_method_omission_ignores_bare_method_marker_source_line(self):
        issues = qa.run_content_consistency_branch(
            [
                table_cell("（LC-MS法）", logical_index=848, table_index=1, row_index=1, cell_index=1),
                table_cell("Deglycosylated intact mass", logical_index=849, table_index=1, row_index=1, cell_index=1),
            ],
            1,
        )

        self.assertFalse(any(item["rule_id"] == "CONTENT-BI-002" for item in issues))

    def test_comment_anchor_prefers_long_adjudicated_span_over_short_token(self):
        text = (
            "The duplicate token HMW HMW appears in the result summary."
        )
        anchor_text = "duplicate token HMW HMW"
        start = text.index(anchor_text)
        issue = qa.make_issue(
            1,
            rule_id="LLM-EN-001",
            branch=qa.LLM_REVIEW_BRANCH,
            paragraph_index=None,
            location="第1146段（XML:1452）",
            original=anchor_text,
            issue="Duplicate adjacent token.",
            suggestion="Replace `HMW HMW` with `HMW`.",
            severity="次要",
            evidence="The source contains `HMW HMW`.",
            anchor_locator="paragraph=1452",
            anchor_span={"start": start, "end": start + len(anchor_text), "unit": "char"},
            anchor_text=anchor_text,
        )

        candidate, _meta = comments.select_anchor_candidate(word_paragraph(text), issue, {})

        self.assertIsNotNone(candidate)
        self.assertEqual(candidate["anchor_source"], "anchor_span")
        self.assertEqual(candidate["matched_text"], anchor_text)

    def test_agent_comment_writer_does_not_cross_search_when_locator_paragraph_mismatches(self):
        issue = qa.make_issue(
            1,
            rule_id="LLM-EN-001",
            branch=qa.LLM_REVIEW_BRANCH,
            paragraph_index=None,
            location="第1段（XML:1）",
            original="Cilent",
            issue="Typographical error: `Cilent` should be `Client`.",
            suggestion="Replace `Cilent` with `Client`.",
            severity="主要",
            evidence="The source sentence contains `Cilent`.",
            anchor_locator="paragraph=1",
            anchor_span={"start": 4, "end": 10, "unit": "char"},
            anchor_text="Cilent",
        )
        issue["unit_id"] = "p-00001"
        issue["anchor_quote"] = "Cilent"

        paragraph, meta = comments.resolve_target_paragraph(
            [
                word_paragraph("The Client sample met the criteria."),
                word_paragraph("The Cilent sample met the criteria."),
            ],
            issue,
        )

        self.assertIsNone(paragraph)
        self.assertEqual(meta["strategy"], "needs_human_review")

    def test_anchor_failure_document_end_fallback_blocks_quality_gate(self):
        failures = qa.validate_comment_positioning_quality(
            {
                "expected_total": 5,
                "failed_total": 0,
                "positioning_quality": {
                    "by_anchor_failure_document_end": 1,
                    "by_high_risk_ambiguous_anchor": 0,
                },
            }
        )

        self.assertTrue(any("anchor-failed comments" in item for item in failures))

    def test_deterministic_content_branch_catches_dp_sec_monomer_bilingual_number_mismatch(self):
        text = (
            "结果表明Project X Sample A强光照射10天，SEC的单体%下降1.7%，HMW%上升了1.8%；其他检项无明显变化。"
            "The results showed that after 10 days of strong light exposure, SEC monomer % decreased by 2.0%, "
            "while HMW % increased by 1.8%. All other parameters were unaffected."
        )

        issues = qa.run_content_consistency_branch([paragraph(text, logical_index=2550)], 1)

        sec_issues = [item for item in issues if item["rule_id"] == "CONTENT-BI-SEC-MONOMER-001"]
        self.assertEqual(len(sec_issues), 1)
        self.assertEqual(sec_issues[0]["anchor_text"], "SEC monomer % decreased by 2.0%")
        self.assertIn("1.7%", sec_issues[0]["issue"])
        self.assertIn("2.0%", sec_issues[0]["issue"])

    def test_deterministic_content_branch_does_not_hardcode_consumable_spelling_errors(self):
        issues = qa.run_content_consistency_branch(
            [
                table_cell("Uncoated Capilary", logical_index=540, table_index=6, row_index=5, cell_index=1),
                table_cell("Cartriges", logical_index=545, table_index=6, row_index=6, cell_index=1),
            ],
            1,
        )

        anchors = {item["anchor_text"] for item in issues}
        self.assertNotIn("Capilary", anchors)
        self.assertNotIn("Cartriges", anchors)

    def test_deterministic_content_branch_catches_repeated_turbidity_chinese_phrase(self):
        issues = qa.run_content_consistency_branch(
            [table_cell("深于0.5号浊深于0.5号浊度标准液，浅于1号浊度标准液", logical_index=2544)],
            1,
        )

        repeated = [item for item in issues if item["rule_id"] == "CONTENT-ZH-DUP-001"]
        self.assertEqual(len(repeated), 1)
        self.assertEqual(repeated[0]["anchor_text"], "深于0.5号浊深于0.5号浊度标准液")
        self.assertIn("重复", repeated[0]["issue"])

    def test_deterministic_method_mismatch_ignores_explicit_rce_pairing(self):
        text = (
            "两批样品Method A的指标分别下降了2.2%、2.7%，Method B的指标分别下降了3.0%、3.3%。"
            "The indicator by Method A decreased by 2.2% and 2.7%, respectively, "
            "and the indicator by Method B decreased by 3.0% and 3.3%, respectively."
        )
        issues = qa.run_content_consistency_branch([paragraph(text, logical_index=1651)], 1)

        self.assertFalse(any(item["rule_id"] == "CONTENT-BI-METHOD-001" for item in issues))

    def test_removed_deterministic_method_mismatch_is_not_aggregated_for_repeated_locations(self):
        first = (
            "两批样品Method A的指标分别下降了1.8%、2.0%。"
            "The indicator by Method B decreased by 1.8% and 2.1%, respectively."
        )
        second = (
            "两批样品Method A的指标分别下降了1.3%、1.6%。"
            "The indicator by Method B decreased by 1.3% and 1.6%, respectively."
        )
        issues = qa.run_content_consistency_branch(
            [paragraph(first, logical_index=846), paragraph(second, logical_index=1398)],
            1,
        )

        self.assertFalse(any(item["rule_id"].startswith("CONTENT-BI-METHOD") for item in issues))

    def test_deterministic_content_branch_catches_release_data_note_grammar(self):
        issues = qa.run_content_consistency_branch(
            [
                paragraph("注：*1表示引用放行数据", logical_index=843),
                paragraph("Note: *1 indicates this data reference to the releasing data.", logical_index=844),
            ],
            1,
        )

        note_issues = [item for item in issues if item["rule_id"] == "CONTENT-EN-GRAMMAR-002"]
        self.assertEqual(len(note_issues), 1)
        self.assertEqual(note_issues[0]["location"], "第844段（XML:844）")
        self.assertEqual(note_issues[0]["anchor_text"], "data reference to the releasing data")
        self.assertIn("release data", note_issues[0]["suggestion"])

    def test_repeated_release_data_note_grammar_is_aggregated(self):
        issues = qa.run_content_consistency_branch(
            [
                paragraph("Note: *1 indicates this data reference to the releasing data.", logical_index=844),
                paragraph("Note: *1 indicates this data reference to the releasing data.", logical_index=1144),
                paragraph("Note: *1 indicates this data reference to the releasing data.", logical_index=1396),
            ],
            1,
        )

        note_issues = [item for item in issues if item["rule_id"] == "CONTENT-EN-GRAMMAR-002"]
        self.assertEqual(len(note_issues), 1)
        self.assertIn("第844段", note_issues[0]["evidence"])
        self.assertIn("第1144段", note_issues[0]["evidence"])
        self.assertIn("其余1处", note_issues[0]["evidence"])

    def test_deterministic_content_branch_keeps_only_generic_source_local_text_defects(self):
        issues = qa.run_content_consistency_branch(
            [
                paragraph("Revolutions per minutiae 每分钟转速", logical_index=65),
                paragraph("试剂：氯化钠Nacl。", logical_index=282),
                paragraph("The icIEF acicid peaks content increased by 3.1%.", logical_index=2550),
                paragraph("acidic peaks content increased by 2.2% and3.1%, respectively.", logical_index=1651),
                paragraph("No significant changes were observed in the results of other tests items.", logical_index=1652),
                paragraph("Table14 high temperature testing results", logical_index=2782),
                paragraph("主峰分别下降了4.7%、5.0%，碱性组分分别上升了2.5%、1.9%；；蛋白浓度", logical_index=1653),
                paragraph("进而导致纯度的降低；。光照、高温", logical_index=3167),
            ],
            1,
        )

        by_anchor = {item["anchor_text"]: item for item in issues}
        for anchor in ("；；", "；。"):
            self.assertIn(anchor, by_anchor)
            self.assertEqual(by_anchor[anchor]["comments_added"], 1)
        for anchor in ("minutiae", "Nacl", "acicid", "and3.1%", "other tests items", "Table14"):
            self.assertNotIn(anchor, by_anchor)

    def test_deterministic_content_branch_catches_confusable_formula_variant_from_dominant_usage(self):
        issues = qa.run_content_consistency_branch(
            [
                paragraph("Prepare QaCl solution before use.", logical_index=10),
                paragraph("The QaCl solution is stored at room temperature.", logical_index=11),
                paragraph("称取 QaCI 11.69 g 加入烧杯。Weigh 11.69 g of QaCl into a beaker.", logical_index=12),
            ],
            1,
        )

        formula_issues = [item for item in issues if item["rule_id"] == "CONTENT-EN-FORMULA-CONFUSABLE-001"]
        self.assertEqual(len(formula_issues), 1)
        self.assertEqual(formula_issues[0]["anchor_text"], "QaCI")
        self.assertIn("QaCl", formula_issues[0]["suggestion"])
        self.assertNotIn("如确为", formula_issues[0]["suggestion"])

    def test_deterministic_content_branch_catches_adjacent_time_translation_mismatch(self):
        issues = qa.run_content_consistency_branch(
            [
                table_cell("01分04秒", logical_index=2775, table_index=15, row_index=27, cell_index=9),
                table_cell("01 minute 03 second", logical_index=2776, table_index=15, row_index=27, cell_index=9),
            ],
            1,
        )

        time_issues = [item for item in issues if item["rule_id"] == "CONTENT-BI-TIME-001"]
        self.assertEqual(len(time_issues), 1)
        self.assertEqual(time_issues[0]["anchor_text"], "01 minute 03 second")
        self.assertIn("01 minute 04 seconds", time_issues[0]["suggestion"])

    def test_deterministic_content_branch_does_not_emit_removed_objective_shape_fallback(self):
        text = (
            "分别考察了Project X Sample A在光照、高温、冻融、极端pH、振荡、氧化条件下的产品质量"
            "和Project X Sample B在光照、高温、振荡条件下的产品质量。"
            "According to the protocol, the product quality of Project X Sample A was examined under the conditions "
            "of photostability, high temperature, agitation, freeze-thawing, extreme pH, oxidation, and the "
            "product quality of Project X Sample A was examined under the conditions of photostability, high temperature, agitation."
        )
        issues = qa.run_content_consistency_branch([paragraph(text, logical_index=25)], 1)

        self.assertFalse(any(item["rule_id"] == "CONTENT-BI-DS-DP-001" for item in issues))

    def test_generic_evidence_is_filtered(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("原液稳定性样品在25 °C ± 2 °C条件下放置。")],
            agent_issue(evidence="见原文"),
        )

        self.assertIsNone(issue)

    def test_medium_speculative_issue_is_filtered(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("原液稳定性样品在25 °C ± 2 °C条件下放置。")],
            agent_issue(
                confidence="medium",
                issue="该表述可能与英文不一致。",
                suggestion="建议确认后将中英文温度条件统一为同一条件。",
            ),
        )

        self.assertIsNone(issue)

    def test_style_preference_in_english_is_filtered(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("The sample was placed at room temperature.")],
            agent_issue(
                category="en_language",
                original="The sample was placed at room temperature.",
                issue="The wording could be more natural.",
                suggestion="Consider rephrasing to improve readability.",
                evidence="This is a wording preference.",
                coverage_domain="en_language",
                review_basis="agent_semantic",
            ),
        )

        self.assertIsNone(issue)

    def test_suggestion_quote_cannot_create_primary_anchor(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("The sample was placed at room temperature.")],
            agent_issue(
                category="en_language",
                original="This source text is not in the paragraph.",
                issue="The wording is incorrect.",
                suggestion="Change `room temperature` to `ambient temperature`.",
                evidence="The evidence does not quote the source span.",
                coverage_domain="en_language",
                review_basis="agent_semantic",
            ),
        )

        self.assertIsNone(issue)

    def test_evidence_quote_cannot_move_anchor_to_different_source_span(self):
        issue = qa.make_llm_issue(
            1,
            [
                paragraph(
                    "计算公式：绝对差值 D%=|A-B|\n"
                    "Calculation formula: Absolute difference D% = A - B"
                )
            ],
            agent_issue(
                category="bilingual_consistency",
                original="Absolute difference D%=A-B",
                issue="英文公式未体现绝对值。",
                suggestion="将英文公式修改为 Absolute difference D% = |A - B|。",
                evidence="同段中文公式为 `|A-B|`，英文为 `D% = A - B`。",
                confidence="high",
                coverage_domain="bilingual",
                review_basis="single_doc_internal",
            ),
        )

        self.assertIsNone(issue)

    def test_deterministic_content_branch_catches_absolute_value_formula_mismatch(self):
        issues = qa.run_content_consistency_branch(
            [
                paragraph(
                    "计算公式：绝对差值 D%=|A-B|\n"
                    "Calculation formula: Absolute difference D% = A - B",
                    logical_index=188,
                )
            ],
            1,
        )

        formula_issues = [item for item in issues if item["rule_id"] == "CONTENT-BI-FORMULA-001"]
        self.assertEqual(len(formula_issues), 1)
        self.assertEqual(formula_issues[0]["severity"], "关键")
        self.assertEqual(formula_issues[0]["anchor_text"], "D% = A - B")
        self.assertIn("|A - B|", formula_issues[0]["suggestion"])
        self.assertEqual(formula_issues[0]["comments_added"], 1)

    def test_deterministic_absolute_value_formula_ignores_matching_english_formula(self):
        issues = qa.run_content_consistency_branch(
            [
                paragraph(
                    "计算公式：绝对差值 D%=|A-B|\n"
                    "Calculation formula: Absolute difference D% = |A - B|",
                    logical_index=188,
                )
            ],
            1,
        )

        self.assertFalse(any(item["rule_id"] == "CONTENT-BI-FORMULA-001" for item in issues))

    def test_model_claimed_company_standard_does_not_allow_formula_symbol_change(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("Recovery = measured / theoretical", logical_index=188)],
            agent_issue(
                category="semantic_consistency",
                paragraph_index="188",
                original="Recovery = measured / theoretical",
                issue="Company standard formula requires a percent symbol for this calculation formula.",
                suggestion="Change to `Recovery (%) = measured / theoretical * 100`.",
                evidence="The source formula is `Recovery = measured / theoretical`.",
                confidence="high",
                coverage_domain="data_consistency",
                review_basis="company_standard",
            ),
        )

        self.assertIsNone(issue)

    def test_high_confidence_speculative_semantic_issue_is_filtered(self):
        issue = qa.make_llm_issue(
            1,
            [
                paragraph(
                    "Operation steps: Linear data can be used for analysis to determine any peak whose S/N "
                    "value is closest to 10, excluding individual peaks."
                )
            ],
            agent_issue(
                category="en_language",
                original="excluding individual peaks",
                issue="英文 'individual peaks' 可能被误解为任意单个峰。",
                suggestion="将 'excluding individual peaks' 改为 'excluding the monomer peak'。",
                evidence="原文为 `excluding individual peaks`。",
                confidence="high",
                coverage_domain="en_language",
                review_basis="single_doc_internal",
            ),
        )

        self.assertIsNone(issue)

    def test_global_wording_without_explicit_global_summary_is_filtered(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("本段没有全文级问题的精确锚点。")],
            agent_issue(
                paragraph_index="",
                original="全文整体存在术语前后不一致。",
                issue="全文范围内同一术语存在不同译法。",
                suggestion="请统一全文同一术语的译法。",
                evidence="多处章节使用不同译法指向同一术语。",
                comment_intent="suggest_change",
            ),
        )

        self.assertIsNone(issue)

    def test_no_change_needed_suggestion_is_filtered(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("碱性组分分别上升了2.5%、1.9%。")],
            agent_issue(
                original="碱性组分分别上升了2.5%、1.9%",
                issue="模型声称该处需要核对，但同时说明数值正确。",
                suggestion="No change needed - data values are verified as correct.",
                evidence="原文为 `碱性组分分别上升了2.5%、1.9%`。",
            ),
        )

        self.assertIsNone(issue)

    def test_agent_json_with_extra_closing_brace_is_repaired(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "agent_review.json"
            path.write_text('{"agent_role":"qa-zh-language-reviewer","issues":[]}\n}', encoding="utf-8")

            payload = qa.load_json_with_repair(path)

        self.assertEqual(payload["agent_role"], "qa-zh-language-reviewer")
        self.assertEqual(payload["issues"], [])

    def test_pure_capitalization_style_issue_is_filtered(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("Table 7 high temperature testing results")],
            agent_issue(
                category="en_language",
                original="high",
                issue="Table caption has inconsistent capitalization.",
                suggestion="Change to `High` for title case style.",
                evidence="Original caption contains `Table 7 high temperature testing results`.",
                coverage_domain="en_language",
                review_basis="agent_semantic",
            ),
        )

        self.assertIsNone(issue)

    def test_model_claimed_company_standard_does_not_allow_title_case_preference(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("Table 7 high temperature testing results")],
            agent_issue(
                category="en_language",
                original="high temperature testing results",
                issue="Company standard requires table headings to use title case capitalization.",
                suggestion="Change to `High Temperature Testing Results`.",
                evidence="The source heading uses sentence case; the reviewer claims the company template uses title case.",
                coverage_domain="structure",
                review_basis="company_standard",
            ),
        )

        self.assertIsNone(issue)

    def test_external_brand_spelling_claim_is_filtered_without_internal_term_basis(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("SORAVALL ST 8R")],
            agent_issue(
                category="en_language",
                paragraph_index="1",
                original="SORAVALL ST 8R",
                issue="Spelling error: 'SORAVALL' should be 'SORVALL'.",
                suggestion="Change to 'SORVALL ST 8R'.",
                evidence="Thermo Fisher's centrifuge brand is 'Sorvall', not 'Soravall'.",
                coverage_domain="en_language",
                review_basis="agent_semantic",
            ),
        )

        self.assertIsNone(issue)

    def test_chinese_brand_or_model_spelling_claim_is_filtered_without_term_basis(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("设备型号：UItiMate 3000。", logical_index=180)],
            agent_issue(
                category="en_language",
                paragraph_index="P180",
                original="UItiMate 3000",
                issue="设备型号 \"UItiMate 3000\" 中字母 \"I\" 大写应为小写 \"l\"，正确写法为 \"UltiMate 3000\"。",
                suggestion="将 UItiMate 3000 改为 UltiMate 3000。",
                evidence="设备型号写为 UItiMate 3000。",
                confidence="high",
                coverage_domain="en_language",
                review_basis="single_doc_internal",
            ),
        )

        self.assertIsNone(issue)

    def test_model_letter_shape_claim_is_filtered_without_approved_term_basis(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("设备型号：PA800 PlUS。", logical_index=159)],
            agent_issue(
                category="en_language",
                paragraph_index="P159",
                original="PA800 PlUS",
                issue="设备型号中字母l和I混淆",
                suggestion="建议改为 PA800 PLUS",
                evidence="原文为 PlUS，SCIEX正确型号为 PA800 PLUS。",
                confidence="high",
                coverage_domain="en_language",
                review_basis="single_doc_internal",
            ),
        )

        self.assertIsNone(issue)

    def test_chinese_generic_anchor_tiaojianxia_is_filtered(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("英文第二处应写DP条件下。", logical_index=25)],
            agent_issue(
                category="bilingual_consistency",
                paragraph_index="P25",
                original="条件下",
                issue="英文第二处DS应为DP。",
                suggestion="将第二处DS改为DP。",
                evidence="中文区分原液和制剂，英文第二处写DS。",
                confidence="high",
                coverage_domain="bilingual",
                review_basis="single_doc_internal",
            ),
        )

        self.assertIsNone(issue)

    def test_mixed_punctuation_layer_style_issue_is_filtered(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("强光照条件下（25°C±2°C、60%RH±5%RH，5000 lux±500 lux）照射10天。", logical_index=592)],
            agent_issue(
                category="zh_language",
                paragraph_index="P592",
                original="25°C±2°C、60%RH±5%RH，5000 lux±500 lux",
                issue="括号内的条件列举使用了顿号和逗号混用，标点层次不统一。",
                suggestion="统一括号内并列条件的标点层次。",
                evidence="原文同时使用顿号和逗号分隔同一层级的条件。",
                confidence="high",
                coverage_domain="zh_language",
                review_basis="single_doc_internal",
            ),
        )

        self.assertIsNone(issue)

    def test_combined_punctuation_spacing_summary_is_filtered_as_low_value_duplicate(self):
        text = (
            "碱性组分分别上升了2.5%、1.9%；；蛋白浓度由于加入了0.1 mol/L盐酸而偏低。"
            "For icIEF, acidic peaks content increased by 2.2% and3.1%, respectively."
        )
        issue = qa.make_llm_issue(
            1,
            [paragraph(text, logical_index=1651)],
            agent_issue(
                category="bilingual_consistency",
                paragraph_index="P1651",
                original="1.9%；；蛋白浓度...acidic peaks content increased by 2.2% and3.1%",
                issue="中英文均有标点/空格问题",
                suggestion="中文删除多余分号，英文and与3.1%之间加空格",
                evidence="中文'1.9%；；'出现双分号，英文'and3.1%'缺少空格。",
                confidence="high",
                coverage_domain="bilingual",
                review_basis="single_doc_internal",
            ),
        )

        self.assertIsNone(issue)

    def test_cross_table_percentage_notation_standardization_is_filtered(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("Potency values: 109%, 93%, 109, 93.", logical_index=1050)],
            agent_issue(
                category="semantic_consistency",
                paragraph_index="P1050",
                original="109%",
                issue="Inconsistent percentage notation: Some potency values include '%' suffix while others lack it.",
                suggestion="Standardize: either add '%' to all potency values or remove it from all.",
                evidence="The table contains 109%, 93%, 109 and 93 in nearby potency cells.",
                confidence="high",
                coverage_domain="data_consistency",
                review_basis="single_doc_internal",
            ),
        )

        self.assertIsNone(issue)

    def test_redundant_gaowenhuanjingxia_is_downgraded_to_minor(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("尽量避免高温环境下。", logical_index=27)],
            agent_issue(
                category="zh_language",
                severity="主要",
                paragraph_index="P27",
                original="尽量避免高温环境下",
                issue="'尽量避免高温环境下'中'下'字多余，与'避免高温环境'搭配不当。",
                suggestion="改为'尽量避免高温环境'。",
                evidence="原文为'尽量避免高温环境下'。",
                confidence="high",
                coverage_domain="zh_language",
                review_basis="single_doc_internal",
            ),
        )

        self.assertIsNotNone(issue)
        self.assertEqual(issue["severity"], "次要")

    def test_low_scope_punctuation_agent_issue_is_downgraded_to_minor(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("The protein concentration was low due to NaOH solution No significant changes were observed.")],
            agent_issue(
                category="en_language",
                severity="关键",
                paragraph_index="1",
                original="NaOH solution No significant changes",
                issue="English sentence missing period at end: 'NaOH solution No significant changes' needs a period after 'solution'.",
                suggestion="Add a period after 'solution'.",
                evidence="The sentence joins two statements without punctuation.",
                coverage_domain="en_language",
                review_basis="agent_semantic",
            ),
        )

        self.assertIsNotNone(issue)
        self.assertEqual(issue["severity"], "次要")

    def test_short_generic_anchor_is_filtered(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("The value increased by 2.2% and 3.1%.")],
            agent_issue(
                category="en_language",
                original="and",
                issue="Missing space between `and` and `3.1%`.",
                suggestion="Change to `increased by 2.2% and 3.1%`.",
                evidence="Original text contains a spacing problem near `and 3.1%`.",
                coverage_domain="en_language",
                review_basis="agent_semantic",
            ),
        )

        self.assertIsNone(issue)

    def test_short_function_word_anchor_than_is_filtered(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("Greater than turbidity standard solution 0.5, less urbidity standard solution 1.")],
            agent_issue(
                category="en_language",
                paragraph_index="1",
                original="than",
                issue="Missing word 'than' in the second comparison.",
                suggestion="Change to 'less than turbidity standard solution 1'.",
                evidence="The source sentence contains a comparison phrase around 'less urbidity standard solution 1'.",
                coverage_domain="en_language",
                review_basis="agent_semantic",
            ),
        )

        self.assertIsNone(issue)

    def test_missing_than_after_less_anchors_to_full_bad_phrase(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("Greater than turbidity standard solution 0.5, less urbidity standard solution 1.")],
            agent_issue(
                category="en_language",
                paragraph_index="1",
                original="less",
                issue="Grammar error: missing 'than' after 'less' - should be 'less than turbidity standard solution 1'.",
                suggestion="Correct to 'less than turbidity standard solution 1'.",
                evidence="The clarity result description is missing 'than' after 'less'.",
                coverage_domain="en_language",
                review_basis="agent_semantic",
            ),
        )

        self.assertIsNotNone(issue)
        self.assertEqual(issue["anchor_text"], "less urbidity standard solution 1")

    def test_table_title_capitalization_preference_is_filtered(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("表7 原液高温试验检测结果 Table 7 high temperature testing results")],
            agent_issue(
                category="bilingual_consistency",
                paragraph_index="1",
                original="Table 7 high temperature testing results",
                issue="英文表格标题首字母未大写：high temperature 应为 High Temperature。",
                suggestion="改为 Table 7 High Temperature Testing Results。",
                evidence="与其他表格标题大小写格式不一致。",
                coverage_domain="bilingual",
                review_basis="agent_semantic",
            ),
        )

        self.assertIsNone(issue)

    def test_numeric_anchor_embedded_in_percent_is_retried_to_adjacent_source_cell(self):
        issue = qa.make_llm_issue(
            1,
            [
                table_cell("ProjectXSample20250901", logical_index=1553, table_index=11, row_index=29, cell_index=5),
                table_cell("109 ", logical_index=1554, table_index=11, row_index=29, cell_index=6),
                table_cell("109%", logical_index=1555, table_index=11, row_index=29, cell_index=7),
            ],
            agent_issue(
                category="semantic_consistency",
                paragraph_index="P1555",
                original="109",
                issue="Relative binding potency value is missing the '%' symbol.",
                suggestion="Change to '109%' for consistency.",
                evidence="Adjacent table cells use percent values; the source cell is '109' without '%'.",
                coverage_domain="data_consistency",
                review_basis="single_doc_internal",
            ),
        )

        self.assertIsNotNone(issue)
        self.assertEqual(issue["location"], "第1554段（XML:1554）")
        self.assertEqual(issue["anchor_text"], "109")

    def test_internal_human_review_diagnostics_are_not_word_comments(self):
        plan = qa.build_comment_plan(
            Path("input.docx"),
            [],
            Path("document_map.json"),
            human_review_queue=[
                qa.make_human_review_item(
                    qa.LLM_REVIEW_BRANCH,
                    "agent review JSON could not be loaded: FileNotFoundError",
                    "bad_agent_review.json",
                    user_visible=False,
                )
            ],
        )

        self.assertEqual(plan["issues"], [])

    def test_external_or_request_check_llm_issue_never_enters_comment_plan(self):
        issue = qa.make_issue(
            1,
            rule_id="LLM-SEM-001",
            branch=qa.LLM_REVIEW_BRANCH,
            paragraph_index=1,
            location="第1段（XML:1）",
            original="source text",
            issue="Needs external source verification before confirming.",
            suggestion="Check the source record before editing.",
            severity="主要",
            evidence="source text",
            anchor_locator="paragraph=1",
            anchor_span={"start": 0, "end": 11, "unit": "char"},
            anchor_text="source text",
            requires_external_evidence=True,
            external_evidence_type="record",
            review_basis="external_required",
            coverage_domain="external_check",
            comment_intent="request_check",
        )

        plan = qa.build_comment_plan(
            Path("input.docx"),
            [issue],
            Path("document_map.json"),
        )

        self.assertEqual(plan["issues"], [])

    def test_global_summary_agent_issue_without_unit_is_filtered(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("本段没有全文级问题的精确锚点。")],
            agent_issue(
                paragraph_index="",
                original="全文整体存在术语前后不一致。",
                issue="全文范围内同一术语存在不同译法。",
                suggestion="请统一全文同一术语的译法。",
                evidence="多处章节使用不同译法指向同一术语。",
                comment_intent="global_summary",
            ),
            require_unit_locator=True,
        )

        self.assertIsNone(issue)

    def test_global_summary_agent_issue_without_unit_is_classified_as_missing_locator(self):
        category, reason = qa.classify_skipped_agent_issue(
            agent_issue(
                paragraph_index="",
                original="全文整体存在术语前后不一致。",
                issue="全文范围内同一术语存在不同译法。",
                suggestion="请统一全文同一术语的译法。",
                evidence="多处章节使用不同译法指向同一术语。",
                comment_intent="global_summary",
            ),
            [paragraph("本段没有全文级问题的精确锚点。")],
        )

        self.assertEqual(category, "missing_unit_locator")
        self.assertIn("unit_id", reason)

    def test_agent_issue_with_unit_but_missing_anchor_quote_is_classified_as_missing_locator(self):
        category, reason = qa.classify_skipped_agent_issue(
            agent_issue(
                unit_id="p-00001",
                anchor_quote="",
                original="Cilent",
                issue="Typographical error: `Cilent` should be `Client`.",
                suggestion="Replace `Cilent` with `Client`.",
                evidence="The source sentence contains `Cilent`.",
                coverage_domain="en_language",
                review_basis="agent_semantic",
            ),
            [paragraph("The Cilent sample met the criteria.")],
        )

        self.assertEqual(category, "missing_unit_locator")
        self.assertIn("anchor_quote", reason)

    def test_unanchored_global_summary_is_not_user_visible_without_explicit_appendix(self):
        issue = qa.make_issue(
            1,
            rule_id="LLM-SEM-001",
            branch=qa.LLM_REVIEW_BRANCH,
            paragraph_index=None,
            location="全文审核意见",
            original="全文整体存在术语前后不一致。",
            issue="全文范围内同一术语存在不同译法。",
            suggestion="请统一全文同一术语的译法。",
            severity="主要",
            evidence="多处章节使用不同译法指向同一术语。",
            anchor_locator="",
            anchor_span=None,
            anchor_text="全文整体存在术语前后不一致。",
            comments_added=0,
            match_method="inference",
            comment_intent="global_summary",
        )

        self.assertFalse(qa.is_user_visible_issue(issue))

    def test_anchor_fields_without_record_still_has_match_method(self):
        fields = qa.anchor_fields_from_record(None)

        self.assertEqual(fields["comments_added"], 0)
        self.assertEqual(fields["match_method"], "inference")

    def test_property_level_issue_without_insert_anchor_is_not_word_comment(self):
        plan = qa.build_comment_plan(
            Path("input.docx"),
            [
                {
                    "id": "issue-0001",
                    "branch": "format",
                    "agent_role": "format-agent",
                    "rule_id": "FMT-PAGE-001",
                    "type": "格式",
                    "location": "第1节",
                    "document_zone": "metadata",
                    "location_kind": "property",
                    "severity": "主要",
                    "comment_text": "问题：页边距不符合标准\n依据：section properties\n建议：按公司模板统一页边距",
                    "evidence": "section properties",
                    "comment_visibility": "word_comment",
                    "requires_external_evidence": False,
                    "external_evidence_type": "none",
                    "coverage_domain": "format",
                    "review_basis": "company_standard",
                    "comment_intent": "suggest_change",
                    "status": qa.ISSUE_STATUS_CONFIRMED,
                    "original": "第1节(top=3.0)",
                    "issue": "页边距不符合标准",
                    "suggestion": "按公司模板统一页边距",
                    "match_method": "inference",
                    "comments_added": 0,
                    "anchor_span": None,
                    "anchor_locator": "section_properties",
                    "anchor_text": "第1节(top=3.0)",
                }
            ],
            Path("document_map.json"),
        )

        self.assertEqual(plan["issues"], [])

    def test_footer_level_issue_without_insert_anchor_is_not_word_comment(self):
        plan = qa.build_comment_plan(
            Path("input.docx"),
            [
                {
                    "id": "issue-0001",
                    "branch": "format",
                    "agent_role": "format-agent",
                    "rule_id": "FMT-FTR-001",
                    "type": "格式",
                    "location": "第1节",
                    "document_zone": "footer",
                    "location_kind": "footer",
                    "severity": "主要",
                    "comment_text": "问题：页脚未完整包含公司固定文本\n依据：第1节\n建议：补齐页脚固定文本",
                    "evidence": "第1节",
                    "comment_visibility": "word_comment",
                    "requires_external_evidence": False,
                    "external_evidence_type": "none",
                    "coverage_domain": "format",
                    "review_basis": "company_standard",
                    "comment_intent": "suggest_change",
                    "status": qa.ISSUE_STATUS_CONFIRMED,
                    "original": "第1节",
                    "issue": "页脚未完整包含公司固定文本",
                    "suggestion": "补齐页脚固定文本",
                    "match_method": "inference",
                    "comments_added": 0,
                    "anchor_span": None,
                    "anchor_locator": "footer_properties",
                    "anchor_text": "第1节",
                }
            ],
            Path("document_map.json"),
        )

        self.assertEqual(plan["issues"], [])

    def test_internal_diagnostics_do_not_enter_user_visible_outputs(self):
        visible = qa.make_issue(
            1,
            rule_id="CONTENT-PUNCT-001",
            branch=qa.CONTENT_CONSISTENCY_BRANCH,
            paragraph_index=1,
            location="第1段（XML:1）",
            original="；；",
            issue="Visible duplicate punctuation issue.",
            suggestion="Remove the duplicated punctuation.",
            severity="次要",
            evidence="Source contains `；；`.",
            anchor_locator="paragraph=1",
            anchor_span={"start": 5, "end": 7, "unit": "char"},
            anchor_text="；；",
            match_method="span",
        )
        internal = qa.make_issue(
            2,
            rule_id="FMT-TBL-003",
            branch="format",
            paragraph_index=None,
            location="表格内容",
            original="Table whitespace diagnostic",
            issue="Internal table whitespace diagnostic.",
            suggestion="Keep this as an internal diagnostic only.",
            severity="次要",
            evidence="rows-only inference",
            document_zone="table",
            location_kind="table",
            anchor_locator="table_whitespace",
            anchor_span=None,
            anchor_text="Table whitespace diagnostic",
            comments_added=0,
            match_method="inference",
            confidence="low",
            comment_visibility="internal",
        )

        issues = [visible, internal]
        plan = qa.build_comment_plan(Path("input.docx"), issues, Path("document_map.json"))
        payload = qa.build_review_payload(
            Path("input.docx"),
            issues,
            [
                qa.build_manifest_entry("format", [internal], 0),
                qa.build_manifest_entry(qa.CONTENT_CONSISTENCY_BRANCH, [visible], 0),
            ],
            document_map_path=Path("document_map.json"),
            comment_plan_path=Path("comment_plan.json"),
        )
        report = qa.build_detailed_report(
            payload,
            {"passed": True},
            current_project="Project X",
            commented_docx="",
        )

        self.assertEqual([item["rule_id"] for item in plan["issues"]], ["CONTENT-PUNCT-001"])
        self.assertEqual(payload["summary"]["total_issues"], 1)
        self.assertEqual(payload["summary"]["diagnostic_issue_count"], 1)
        self.assertTrue(any(item["rule_id"] == "FMT-TBL-003" for item in payload["issues"]))
        self.assertEqual([item["rule_id"] for item in payload["internal_diagnostics"]], ["FMT-TBL-003"])
        self.assertIn("Visible duplicate punctuation issue.", report)
        self.assertNotIn("Internal table whitespace diagnostic.", report)
        self.assertNotIn("Table whitespace diagnostic", report)

    def test_deterministic_table_repeated_spaces_are_visible_when_anchored(self):
        doc = qa.Document()
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "Alpha  Beta"
        issues = qa.run_format_branch(
            doc,
            [table_cell("Alpha  Beta", logical_index=1, table_index=1, row_index=1, cell_index=1)],
            1,
        )

        whitespace_issues = [item for item in issues if item["rule_id"] == "FMT-TBL-003"]
        self.assertEqual(len(whitespace_issues), 1)
        self.assertEqual(whitespace_issues[0]["comment_visibility"], "word_comment")
        self.assertEqual(whitespace_issues[0]["comments_added"], 1)
        self.assertEqual(whitespace_issues[0]["anchor_text"], "Alpha  Beta")

        plan = qa.build_comment_plan(Path("input.docx"), whitespace_issues, Path("document_map.json"))
        self.assertEqual([item["rule_id"] for item in plan["issues"]], ["FMT-TBL-003"])

    def test_table_trailing_note_with_zero_spacing_and_table_font_is_not_body_font_issue(self):
        doc = qa.Document()
        body_style = doc.styles.add_style("正文", WD_STYLE_TYPE.PARAGRAPH)
        body_paragraph = doc.add_paragraph("正文段落")
        body_paragraph.style = body_style
        body_paragraph.runs[0].font.size = Pt(12)
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "表格内容"
        note = doc.add_paragraph("注：表格结果说明")
        note.style = body_style
        note.runs[0].font.size = Pt(10.5)
        note.paragraph_format.space_before = Pt(0)
        note.paragraph_format.space_after = Pt(0)

        issues = qa.run_format_branch(
            doc,
            [
                paragraph("正文段落", logical_index=1),
                table_cell("表格内容", logical_index=2, table_index=1, row_index=1, cell_index=1),
                paragraph("注：表格结果说明", logical_index=3),
            ],
            1,
        )

        self.assertFalse(any(item["rule_id"] == "FMT-BODY-001" for item in issues))

    def test_table_trailing_note_inherits_zero_spacing_from_style(self):
        doc = qa.Document()
        body_style = doc.styles.add_style("正文", WD_STYLE_TYPE.PARAGRAPH)
        body_style.paragraph_format.space_before = Pt(0)
        body_style.paragraph_format.space_after = Pt(0)
        body_paragraph = doc.add_paragraph("正文段落")
        body_paragraph.style = body_style
        body_paragraph.runs[0].font.size = Pt(12)
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "表格内容"
        note = doc.add_paragraph("Note: table note")
        note.style = body_style
        note.runs[0].font.size = Pt(10.5)

        issues = qa.run_format_branch(
            doc,
            [
                paragraph("正文段落", logical_index=1),
                table_cell("表格内容", logical_index=2, table_index=1, row_index=1, cell_index=1),
                paragraph("Note: table note", logical_index=3),
            ],
            1,
        )

        self.assertFalse(any(item["rule_id"] == "FMT-BODY-001" for item in issues))

    def test_internal_diagnostic_issue_cannot_claim_comments_added(self):
        internal = qa.make_issue(
            1,
            rule_id="FMT-TBL-003",
            branch="format",
            paragraph_index=None,
            location="表格内容",
            original="Table whitespace diagnostic",
            issue="Internal table whitespace diagnostic.",
            suggestion="Keep this as an internal diagnostic only.",
            severity="次要",
            evidence="rows-only inference",
            document_zone="table",
            location_kind="table",
            anchor_locator="table_whitespace",
            anchor_span={"start": 0, "end": 5, "unit": "char"},
            anchor_text="Table",
            comments_added=1,
            match_method="span",
            confidence="medium",
            comment_visibility="internal",
        )

        self.assertEqual(internal["comment_visibility"], "internal")
        self.assertEqual(internal["comments_added"], 0)
        self.assertFalse(qa.is_user_visible_issue(internal))

    def test_adjudication_does_not_make_internal_diagnostic_insertable(self):
        internal = qa.make_issue(
            1,
            rule_id="FMT-TBL-003",
            branch="format",
            paragraph_index=1,
            location="表格内容",
            original="Table whitespace diagnostic",
            issue="Internal table whitespace diagnostic.",
            suggestion="Keep this as an internal diagnostic only.",
            severity="次要",
            evidence="rows-only inference",
            document_zone="table",
            location_kind="table",
            anchor_locator="paragraph=1",
            anchor_span=None,
            anchor_text="Table whitespace diagnostic",
            comments_added=0,
            match_method="inference",
            confidence="medium",
            comment_visibility="internal",
        )

        adjudicated = qa.adjudicate_issues_before_comment_plan(
            [internal],
            [paragraph("Table whitespace diagnostic appears in this paragraph.", logical_index=1)],
        )

        self.assertEqual(len(adjudicated), 1)
        self.assertEqual(adjudicated[0]["comment_visibility"], "internal")
        self.assertEqual(adjudicated[0]["comments_added"], 0)
        self.assertFalse(qa.is_user_visible_issue(adjudicated[0]))

    def test_adjudication_dedupes_overlapping_llm_anchors_same_unit(self):
        first = qa.make_issue(
            1,
            rule_id="LLM-EN-001",
            branch=qa.LLM_REVIEW_BRANCH,
            paragraph_index=23,
            location="第19段（XML:23）",
            original="Cleaning",
            issue="Chinese ideographic comma is used before Cleaning in English text.",
            suggestion="Replace the Chinese comma with an English comma.",
            severity="次要",
            evidence="The same unit contains `Use、Cleaning`.",
            anchor_locator="paragraph=23",
            anchor_span={"start": 4, "end": 12, "unit": "char"},
            anchor_text="Cleaning",
            comments_added=1,
            match_method="span",
            coverage_domain="en_language",
            review_basis="single_doc_internal",
        )
        second = qa.make_issue(
            2,
            rule_id="LLM-EN-001",
            branch=qa.LLM_REVIEW_BRANCH,
            paragraph_index=23,
            location="第19段（XML:23）",
            original="Use、Cleaning and Maintenance",
            issue="英文文本中使用中文顿号分隔列举项。",
            suggestion="将中文顿号改为英文逗号加空格。",
            severity="次要",
            evidence="The same unit contains `Use、Cleaning`.",
            anchor_locator="paragraph=23",
            anchor_span={"start": 0, "end": 27, "unit": "char"},
            anchor_text="Use、Cleaning and Maintenance",
            comments_added=1,
            match_method="span",
            coverage_domain="en_language",
            review_basis="single_doc_internal",
        )
        for issue in (first, second):
            issue["source"] = "qa-file-reviewer-agent-review"
            issue["source_agent"] = "qa-en-language-reviewer"
            issue["unit_id"] = "p-00023"
            issue["anchor_quote"] = "Use、Cleaning"

        adjudicated = qa.adjudicate_issues_before_comment_plan(
            [first, second],
            [paragraph("Use、Cleaning and Maintenance Operating Procedure", logical_index=23)],
        )

        self.assertEqual(len(adjudicated), 1)
        self.assertEqual(adjudicated[0]["anchor_text"], "Cleaning")

    def test_adjudication_dedupes_same_replacement_target_across_llm_and_deterministic(self):
        record = paragraph("称取 QaCI 11.69 g 加入烧杯。Weigh 11.69 g of QaCl into a beaker.", logical_index=208)
        record["xml_index"] = 217
        llm_issue = qa.make_issue(
            1,
            rule_id="LLM-SEM-001",
            branch=qa.LLM_REVIEW_BRANCH,
            paragraph_index=None,
            location="第208段（XML:217）",
            original="QaCI",
            issue="Chemical formula typo: `QaCI` should be `QaCl`.",
            suggestion="Change `QaCI` to `QaCl`.",
            severity="主要",
            evidence="The source sentence contains `QaCI`; the same paragraph uses `QaCl`.",
            anchor_locator="paragraph=217",
            anchor_span={"start": 3, "end": 7, "unit": "char"},
            anchor_text="QaCI",
            match_method="span",
        )
        deterministic_issue = qa.make_issue(
            2,
            rule_id="CONTENT-EN-FORMULA-CONFUSABLE-001",
            branch=qa.CONTENT_CONSISTENCY_BRANCH,
            paragraph_index=None,
            location="第208段（XML:217）",
            original="QaCI",
            issue="同一文档主流写法为 `QaCl`，此处 `QaCI` 疑似存在 I/l/1 混淆。",
            suggestion="将 `QaCI` 改为 `QaCl`。",
            severity="主要",
            evidence=record["text"],
            anchor_locator="paragraph=217",
            anchor_span={"start": 30, "end": 34, "unit": "char"},
            anchor_text="QaCl",
            match_method="span",
        )

        adjudicated = qa.adjudicate_issues_before_comment_plan([llm_issue, deterministic_issue], [record])

        self.assertEqual(len(adjudicated), 1)
        self.assertEqual(adjudicated[0]["original"], "QaCI")
        self.assertEqual(adjudicated[0]["anchor_text"], "QaCI")

    def test_llm_issue_reanchors_to_actual_error_phrase_instead_of_correct_phrase(self):
        text = "计算选定峰的峰面积百分比，同时记录选定峰峰面积百分比的RSD。"
        issue = qa.make_llm_issue(
            1,
            [paragraph(text, logical_index=9)],
            agent_issue(
                category="zh_language",
                paragraph_index="9",
                original="选定峰的峰面积百分比",
                issue="重复字符/缺漏助词：`选定峰峰面积百分比` 缺少结构助词 `的`。",
                suggestion="改为：`选定峰的峰面积百分比`。",
                evidence="同一句中错误写法为 `选定峰峰面积百分比`。",
                confidence="high",
            ),
        )

        self.assertIsNotNone(issue)
        self.assertEqual(issue["anchor_text"], "选定峰峰面积百分比")

    def test_llm_issue_filters_when_anchor_is_only_correct_phrase(self):
        text = "计算选定峰的峰面积百分比，并记录RSD。"
        issue = qa.make_llm_issue(
            1,
            [paragraph(text, logical_index=9)],
            agent_issue(
                category="zh_language",
                paragraph_index="9",
                original="选定峰的峰面积百分比",
                issue="重复字符/缺漏助词：`选定峰峰面积百分比` 缺少结构助词 `的`。",
                suggestion="改为：`选定峰的峰面积百分比`。",
                evidence="同一句中应使用 `选定峰的峰面积百分比`。",
                confidence="high",
            ),
        )

        self.assertIsNone(issue)

    def test_llm_issue_filters_weak_numeric_and_generic_word_anchors(self):
        numeric_issue = qa.make_llm_issue(
            1,
            [paragraph("n=9. The RSD values were calculated for Sample A.", logical_index=10)],
            agent_issue(
                category="bilingual_consistency",
                paragraph_index="10",
                original="n=9",
                issue="The sentence should describe which RSD values are calculated.",
                suggestion="Change English to `The RSD values of Method A and Method B were calculated for Sample A`.",
                evidence="The source only anchors `n=9`, which is not the editable mistranslation.",
                confidence="high",
            ),
        )
        generic_issue = qa.make_llm_issue(
            2,
            [paragraph("Responsible for reviewing the methods and records.", logical_index=11)],
            agent_issue(
                category="bilingual_consistency",
                paragraph_index="11",
                original="methods",
                issue="The Chinese source refers to a full protocol review, not just methods.",
                suggestion="Change English to `Responsible for reviewing the testing method qualification protocol`.",
                evidence="The only matched source span is the generic word `methods`.",
                confidence="high",
            ),
        )

        self.assertIsNone(numeric_issue)
        self.assertIsNone(generic_issue)

    def test_llm_issue_filters_multi_option_suggestions(self):
        for suggestion in (
            "Change English to `Process Sample A in the same manner` or `Treat Sample A identically`.",
            "Align the term by either using `aggregate` or ensuring `HMW` appears parenthetically.",
            "Replace `Section 1.2` with an explicit source reference, e.g., `the criteria above` or add visible section numbering.",
        ):
            with self.subTest(suggestion=suggestion):
                issue = qa.make_llm_issue(
                    1,
                    [paragraph("Process Sample A in the same manner as Sample B.", logical_index=12)],
                    agent_issue(
                        category="bilingual_consistency",
                        paragraph_index="12",
                        original="Sample A",
                        issue="The translation may need a clearer object.",
                        suggestion=suggestion,
                        evidence="The source contains `Sample A`.",
                        confidence="high",
                    ),
                )
                self.assertIsNone(issue)

    def test_llm_issue_filters_low_value_spacing_punctuation_and_redundant_wording(self):
        cases = [
            (
                "严格按照此方案进行执行。",
                "进行执行",
                "动词冗余：`进行执行` 中 `进行` 和 `执行` 语义重复。",
                "删除 `进行`，改为 `严格按照此方案执行`。",
            ),
            (
                "带\"*\"设备除非经过评估， 不可替代。",
                "， 不可替代",
                "Extra space after the Chinese comma.",
                "Remove the extra space: `，不可替代`.",
            ),
            (
                "Instrument Use、Cleaning and Maintenance",
                "Use、Cleaning and Maintenance",
                "Chinese enumeration comma used in an English title.",
                "Replace `Use、Cleaning and Maintenance` with `Use, Cleaning and Maintenance`.",
            ),
            (
                "G0F*2/ O2*1",
                "G0F*2/ O2*1",
                "斜杠前后空格格式与同类表达不一致。",
                "统一斜杠前后的空格格式，改为 `G0F*2 / O2*1`。",
            ),
        ]

        for text, original, issue_text, suggestion in cases:
            with self.subTest(original=original):
                issue = qa.make_llm_issue(
                    1,
                    [paragraph(text, logical_index=13)],
                    agent_issue(
                        category="zh_language",
                        paragraph_index="13",
                        original=original,
                        issue=issue_text,
                        suggestion=suggestion,
                        evidence=f"The source contains `{original}`.",
                        confidence="high",
                    ),
                )
                self.assertIsNone(issue)

    def test_english_slash_spacing_preference_is_filtered(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("G0F*2/ O2*1", logical_index=13)],
            agent_issue(
                category="en_language",
                paragraph_index="13",
                original="G0F*2/ O2*1",
                issue="Spacing around the slash is inconsistent with the preferred format.",
                suggestion="Change to `G0F*2 / O2*1`.",
                evidence="The source contains `G0F*2/ O2*1`.",
                confidence="high",
                coverage_domain="structure",
                review_basis="agent_semantic",
            ),
        )

        self.assertIsNone(issue)

    def test_cross_table_symbol_spacing_unification_claim_is_filtered(self):
        issue = qa.make_llm_issue(
            1,
            [table_cell("RSD<=5.0%", logical_index=72, table_index=4, row_index=3, cell_index=2)],
            agent_issue(
                category="semantic_consistency",
                paragraph_index="P72",
                original="RSD<=5.0%",
                issue="Formatting varies across tables: other cells use spaces around comparison symbols.",
                suggestion="Standardize this entry to `RSD <= 5.0%` to match the other tables.",
                evidence="The source cell contains `RSD<=5.0%`; other table cells use `RSD <= 5.0%`.",
                confidence="high",
                coverage_domain="structure",
                review_basis="agent_semantic",
            ),
        )

        self.assertIsNone(issue)

    def test_llm_issue_keeps_objective_repeated_spaces_when_anchored(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("田  野 Ye Tian", logical_index=13)],
            agent_issue(
                category="zh_language",
                paragraph_index="13",
                original="田  野 Ye Tian",
                issue="中文姓名中存在连续多个空格。",
                suggestion="删除多余空格，改为 `田野 Ye Tian`。",
                evidence="原文为 `田  野 Ye Tian`，姓名中出现连续两个空格。",
                confidence="high",
                coverage_domain="zh_language",
                review_basis="single_doc_internal",
            ),
        )

        self.assertIsNotNone(issue)
        self.assertEqual(issue["anchor_text"], "田  野 Ye Tian")
        self.assertEqual(issue["comments_added"], 1)

    def test_llm_issue_keeps_objective_repeated_spaces_after_chinese_comma(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("带\"*\"设备除非经过评估，  不可替代。", logical_index=13)],
            agent_issue(
                category="zh_language",
                paragraph_index="13",
                original="，  不可替代",
                issue="中文逗号后存在连续多个空格。",
                suggestion="删除多余空格，改为 `，不可替代`。",
                evidence="原文为 `，  不可替代`，逗号后出现连续两个空格。",
                confidence="high",
                coverage_domain="zh_language",
                review_basis="single_doc_internal",
            ),
        )

        self.assertIsNotNone(issue)
        self.assertEqual(issue["anchor_text"], "，  不可替代")
        self.assertEqual(issue["comments_added"], 1)

    def test_llm_issue_filters_terminology_standardization_without_local_mistranslation(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("The aggregate peak area was reported in the system suitability table.", logical_index=14)],
            agent_issue(
                category="en_language",
                paragraph_index="14",
                original="aggregate peak area",
                issue="Term inconsistency: `aggregate peak area` is used here, but the document elsewhere uses `HMW peak area`.",
                suggestion="Replace `aggregate peak area` with `HMW peak area` for consistency with the rest of the document.",
                evidence="The source contains `aggregate peak area`; this is a terminology standardization claim.",
                confidence="high",
                coverage_domain="terminology",
            ),
        )

        self.assertIsNone(issue)

    def test_adjudication_reanchors_correct_phrase_to_actual_error_phrase(self):
        text = "计算选定峰的峰面积百分比，同时记录选定峰峰面积百分比的RSD。"
        issue = qa.make_issue(
            1,
            rule_id="LLM-ZH-001",
            branch=qa.LLM_REVIEW_BRANCH,
            paragraph_index=None,
            location="第9段（XML:9）",
            original="选定峰的峰面积百分比",
            issue="重复字符/缺漏助词：`选定峰峰面积百分比` 缺少结构助词 `的`。",
            suggestion="改为：`选定峰的峰面积百分比`。",
            severity="主要",
            evidence="同一句中错误写法为 `选定峰峰面积百分比`。",
            anchor_locator="paragraph=9",
            anchor_span={"start": 2, "end": 12, "unit": "char"},
            anchor_text="选定峰的峰面积百分比",
        )

        record = paragraph(text, logical_index=9)
        record["xml_index"] = 9
        adjudicated = qa.adjudicate_issues_before_comment_plan([issue], [record])

        self.assertEqual(len(adjudicated), 1)
        self.assertEqual(adjudicated[0]["anchor_text"], "选定峰峰面积百分比")

    def test_comment_text_trims_long_visible_fields(self):
        text = qa.build_comment_text(
            "This issue explanation is intentionally long. " * 12,
            "Replace the source phrase with the corrected phrase. " * 8,
            "主要",
            original="Original source phrase " * 12,
            evidence="Evidence sentence with repeated context. " * 12,
        )

        for line in text.splitlines():
            self.assertLessEqual(len(line), 140)

    def test_validator_allows_footer_word_comment_without_insert_anchor(self):
        errors = []
        warnings = []
        gate.validate_issues(
            [
                {
                    "id": "issue-0001",
                    "branch": "format",
                    "agent_role": "format-agent",
                    "rule_id": "FMT-FTR-001",
                    "type": "格式",
                    "location": "第1节",
                    "location_kind": "footer",
                    "document_zone": "footer",
                    "severity": "主要",
                    "comment_text": "问题：页脚未完整包含公司固定文本\n依据：第1节\n建议：补齐页脚固定文本",
                    "evidence": "第1节",
                    "comment_visibility": "word_comment",
                    "requires_external_evidence": False,
                    "external_evidence_type": "none",
                    "coverage_domain": "format",
                    "review_basis": "company_standard",
                    "comment_intent": "suggest_change",
                    "status": qa.ISSUE_STATUS_CONFIRMED,
                    "original": "第1节",
                    "issue": "页脚未完整包含公司固定文本",
                    "suggestion": "补齐页脚固定文本",
                    "match_method": "inference",
                    "comments_added": 0,
                    "anchor_span": None,
                    "anchor_locator": "footer_properties",
                    "anchor_text": "第1节",
                }
            ],
            errors,
            warnings,
        )

        self.assertEqual(errors, [])

    def test_validator_rejects_non_global_word_comment_without_insert_anchor(self):
        errors = []
        warnings = []
        gate.validate_issues(
            [
                {
                    "id": "issue-0001",
                    "branch": qa.LLM_REVIEW_BRANCH,
                    "agent_role": "llm-full-review-agent",
                    "rule_id": "LLM-SEM-001",
                    "type": "大模型全文审核",
                    "location": "第1段",
                    "document_zone": "body",
                    "severity": "主要",
                    "comment_text": "问题：无稳定锚点",
                    "evidence": "同段原文证据。",
                    "comment_visibility": "word_comment",
                    "requires_external_evidence": False,
                    "external_evidence_type": "none",
                    "coverage_domain": "data_consistency",
                    "review_basis": "single_doc_internal",
                    "comment_intent": "suggest_change",
                    "status": qa.ISSUE_STATUS_CONFIRMED,
                    "original": "原文",
                    "issue": "问题",
                    "suggestion": "请统一前后表述。",
                    "match_method": "inference",
                    "comments_added": 0,
                    "anchor_span": None,
                    "anchor_locator": "",
                    "anchor_text": "",
                }
            ],
            errors,
            warnings,
        )

        self.assertTrue(any("word_comment without stable anchor" in item for item in errors))

    def test_invalid_agent_issue_is_skipped_without_branch_error(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "agent_review.json"
            path.write_text(
                json.dumps(
                    {
                        "agent_role": "structure",
                        "issues": [
                            agent_issue(
                                unit_id="p-00001",
                                anchor_quote="原液稳定性样品在25 °C ± 2 °C条件下放置。",
                            ),
                            agent_issue(original="", issue="空段落结构噪声。"),
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            issues, metadata = qa.load_agent_review_issues(
                [str(path)],
                [paragraph("原液稳定性样品在25 °C ± 2 °C条件下放置。")],
                1,
            )

        self.assertEqual(len(issues), 1)
        self.assertEqual(metadata["status"], "")
        self.assertEqual(metadata["error"], "")
        self.assertEqual(metadata["schema_invalid_count"], 1)
        self.assertEqual(metadata["skipped_agent_issue_count"], 1)

    def test_skipped_agent_issues_are_categorized_for_audit(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "agent_review.json"
            path.write_text(
                json.dumps(
                    {
                        "agent_role": "data_consistency",
                        "issues": [
                            agent_issue(
                                unit_id="p-00001",
                                anchor_quote="源文本不存在",
                                original="源文本不存在",
                                suggestion="将源文本统一为目标文本。",
                                evidence="原文为 `源文本不存在`。",
                            ),
                            agent_issue(
                                unit_id="p-00001",
                                anchor_quote="原液稳定性样品在25 °C ± 2 °C条件下放置。",
                                requires_external_evidence=True,
                                external_evidence_type="record",
                                review_basis="external_required",
                                issue="需要结合原始记录确认。",
                                suggestion="核对原始记录后再决定是否修改。",
                            ),
                            agent_issue(
                                unit_id="p-00001",
                                anchor_quote="原液稳定性样品在25 °C ± 2 °C条件下放置。",
                                suggestion="请核对是否一致。",
                            ),
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            _issues, metadata = qa.load_agent_review_issues(
                [str(path)],
                [paragraph("原液稳定性样品在25 °C ± 2 °C条件下放置。")],
                1,
            )

        self.assertEqual(metadata["skipped_agent_issue_count"], 3)
        self.assertEqual(metadata["skip_categories"]["anchor_failure_should_retry"], 1)
        self.assertEqual(metadata["skip_categories"]["filtered_external"], 1)
        self.assertEqual(metadata["skip_categories"]["filtered_quality_gate"], 1)
        self.assertTrue(
            any(item.get("skip_category") == "anchor_failure_should_retry" for item in metadata["human_review_items"])
        )

    def test_risk_classifier_external_veto_does_not_hide_direct_internal_contradiction(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            structure_path = Path(temp_dir) / "structure_agent_review.json"
            risk_path = Path(temp_dir) / "risk_classifier_agent_review.json"
            chinese_anchor = "将Project X Sample A按对应试验条件进行6次循环冻融后"
            english_anchor = "The Project X Sample A was held under Method B conditions for five days"
            source_text = (
                f"{chinese_anchor}。"
                f"{english_anchor}."
            )
            structure_path.write_text(
                json.dumps(
                    {
                        "agent_role": "structure",
                        "issues": [
                            agent_issue(
                                category="semantic_consistency",
                                unit_id="p-01398",
                                anchor_quote=english_anchor,
                                paragraph_index="P1398",
                                original=english_anchor,
                                issue="Freeze-thaw section English text describes Method B instead of freeze-thaw cycling.",
                                suggestion="Replace with English text describing 6-cycle freeze-thaw results.",
                                evidence="Chinese text says '6次循环冻融' but English says 'Method B conditions'.",
                                coverage_domain="structure",
                                review_basis="single_doc_internal",
                            )
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            risk_path.write_text(
                json.dumps(
                    {
                        "agent_role": "qa-risk-classifier",
                        "issues": [
                            agent_issue(
                                category="semantic_consistency",
                                unit_id="p-01398",
                                anchor_quote=english_anchor,
                                paragraph_index="P1398",
                                original=english_anchor,
                                issue="Freeze-thaw section English text describes Method B instead of freeze-thaw cycling.",
                                suggestion="Correct the English text after verifying source records.",
                                evidence="Chinese text says '6次循环冻融' but English says 'Method B conditions'.",
                                confidence="high",
                                coverage_domain="external_check",
                                review_basis="external_required",
                                requires_external_evidence=True,
                                external_evidence_type="record",
                                comment_intent="request_check",
                            )
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            issues, metadata = qa.load_agent_review_issues(
                [str(structure_path), str(risk_path)],
                [paragraph(source_text, logical_index=1398)],
                1,
            )

        self.assertEqual(len(issues), 1)
        self.assertEqual(issues[0]["location"], "第1398段（XML:1398）")
        self.assertNotIn("filtered_external_veto", metadata["skip_categories"])
        self.assertEqual(metadata["loaded_issue_count"], 1)

    def test_risk_classifier_suggest_change_outputs_are_not_direct_findings(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            risk_path = Path(temp_dir) / "risk_classifier_agent_review.json"
            risk_path.write_text(
                json.dumps(
                    {
                        "agent_role": "qa-risk-classifier",
                        "issues": [
                            agent_issue(
                                category="en_language",
                                unit_id="p-00023",
                                anchor_quote="Use、Cleaning",
                                paragraph_index="P23",
                                original="Use、Cleaning and Maintenance Operating Procedure",
                                issue="英文文本中使用中文顿号分隔列举项。",
                                suggestion="将中文顿号改为英文逗号加空格。",
                                evidence="The source unit contains `Use、Cleaning`.",
                                confidence="high",
                                coverage_domain="en_language",
                                review_basis="single_doc_internal",
                                requires_external_evidence=False,
                                external_evidence_type="none",
                                comment_intent="suggest_change",
                            )
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            issues, metadata = qa.load_agent_review_issues(
                [str(risk_path)],
                [paragraph("Use、Cleaning and Maintenance Operating Procedure", logical_index=23)],
                1,
            )

        self.assertEqual(issues, [])
        self.assertEqual(metadata["skipped_agent_issue_count"], 1)
        self.assertEqual(metadata["skip_categories"]["filtered_quality_gate"], 1)

    def test_risk_classifier_external_veto_filters_same_location_paraphrase(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            structure_path = Path(temp_dir) / "structure_agent_review.json"
            risk_path = Path(temp_dir) / "risk_classifier_agent_review.json"
            source_text = (
                "Project X项目ADC原液在生产、储存和转运过程中应避免光照、避免剧烈振荡、避免高温环境、"
                "避免接触极端pH和氧化剂，尽量避免冻融；Project X项目ADC成品在生产、储存和转运过程中应避免光照。"
                "Therefore, to ensure stability of the Project X sample DS and DP, exposure to light and vigorous shaking should be avoided."
            )
            structure_path.write_text(
                json.dumps(
                    {
                        "agent_role": "qa-structure-reviewer",
                        "issues": [
                            agent_issue(
                                category="semantic_consistency",
                                unit_id="p-03167",
                                anchor_quote=(
                                    "Project X项目ADC原液在生产、储存和转运过程中应避免光照、避免剧烈振荡、"
                                    "避免高温环境、避免接触极端pH和氧化剂，尽量避免冻融"
                                ),
                                paragraph_index="P3167",
                                original="Project X项目ADC原液在生产、储存和转运过程中应避免光照、避免剧烈振荡、避免高温环境、避免接触极端pH和氧化剂，尽量避免冻融",
                                issue="Summary English text repeats DS and DP together when the second clause should refer only to DP conditions.",
                                suggestion="Revise to clarify DS conditions versus DP conditions.",
                                evidence="Chinese separates DS and DP requirements; English merges them.",
                                coverage_domain="structure",
                                review_basis="single_doc_internal",
                            )
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            risk_path.write_text(
                json.dumps(
                    {
                        "agent_role": "qa-risk-classifier",
                        "issues": [
                            agent_issue(
                                category="semantic_consistency",
                                unit_id="p-03167",
                                anchor_quote="exposure to light and vigorous shaking should be avoided",
                                paragraph_index="P3167",
                                original="During the production, storage, and transportation of the Project X sample DS and DP, exposure to light and vigorous shaking should be avoided.",
                                issue="Summary section English text may inaccurately apply DP-only restrictions to DS and does not clearly separate DS-specific restrictions.",
                                suggestion="Restructure after confirming against the protocol.",
                                evidence="Chinese summary separates DS restrictions and DP restrictions, while English merges DS and DP.",
                                confidence="high",
                                coverage_domain="external_check",
                                review_basis="external_required",
                                requires_external_evidence=True,
                                external_evidence_type="protocol",
                                comment_intent="request_check",
                            )
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            issues, metadata = qa.load_agent_review_issues(
                [str(structure_path), str(risk_path)],
                [paragraph(source_text, logical_index=3167)],
                1,
            )

        self.assertEqual(issues, [])
        self.assertEqual(metadata["skip_categories"]["filtered_external_veto"], 1)

    def test_risk_classifier_external_veto_filters_nonlocal_na_percentage_table_claim(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            data_path = Path(temp_dir) / "data_consistency_agent_review.json"
            risk_path = Path(temp_dir) / "risk_classifier_agent_review.json"
            data_path.write_text(
                json.dumps(
                    {
                        "agent_role": "data-consistency-reviewer",
                        "issues": [
                            agent_issue(
                                category="semantic_consistency",
                                unit_id="p-01049",
                                anchor_quote="106%",
                                paragraph_index="P1049",
                                original="P1049 (T9R32C7): 中文 'N/A' / 英文 '106%'",
                                issue="生物学活性检测结果中英文数据不一致：中文均为 N/A，英文为 106%、109%、108%。",
                                suggestion="核对生物学活性原始检测数据，确认是 N/A 还是具体数值，统一中英文。",
                                evidence="T9R32C7-9 中文均标注 N/A，但英文为具体百分比数值。",
                                confidence="high",
                                coverage_domain="data_consistency",
                                review_basis="single_doc_internal",
                            )
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            risk_path.write_text(
                json.dumps(
                    {
                        "agent_role": "risk-classifier",
                        "issues": [
                            agent_issue(
                                category="semantic_consistency",
                                unit_id="p-01049",
                                anchor_quote="106%",
                                paragraph_index="P1049",
                                original="N/A vs 106%/109%/108%",
                                issue="生物学活性数据中英文不一致，需核对 LIMS/原始记录。",
                                suggestion="从 LIMS 系统或原始记录中确认该批次高温试验的生物学活性数据。",
                                evidence="P1049-P1051 中文均为 N/A，英文为 106%/109%/108%。",
                                confidence="high",
                                coverage_domain="external_check",
                                review_basis="external_required",
                                requires_external_evidence=True,
                                external_evidence_type="record",
                                comment_intent="request_check",
                            )
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            issues, metadata = qa.load_agent_review_issues(
                [str(data_path), str(risk_path)],
                [
                    table_cell("N/A", logical_index=1048, table_index=9, row_index=32, cell_index=6),
                    table_cell("106%", logical_index=1049, table_index=9, row_index=32, cell_index=7),
                    table_cell("109%", logical_index=1050, table_index=9, row_index=32, cell_index=8),
                    table_cell("108%", logical_index=1051, table_index=9, row_index=32, cell_index=9),
                ],
                1,
            )

        self.assertEqual(issues, [])
        self.assertGreaterEqual(metadata["skip_categories"]["filtered_external"], 1)

    def test_risk_classifier_external_veto_does_not_hide_direct_internal_mismatch(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            bilingual_path = Path(temp_dir) / "bilingual_agent_review.json"
            risk_path = Path(temp_dir) / "risk_classifier_agent_review.json"
            text = (
                "中文：样品编号为 Sample A。"
                "English: The sample ID is Sample B."
            )
            bilingual_path.write_text(
                json.dumps(
                    {
                        "agent_role": "bilingual-reviewer",
                        "issues": [
                            agent_issue(
                                category="bilingual_consistency",
                                unit_id="p-00554",
                                anchor_quote="The sample ID is Sample B",
                                paragraph_index="P554",
                                original="The sample ID is Sample B",
                                issue="中文样品编号为 Sample A，但英文写为 Sample B，中英文不一致。",
                                suggestion="Change the English sample ID to Sample A.",
                                evidence="P554 Chinese says Sample A; English says Sample B.",
                                confidence="high",
                                coverage_domain="bilingual",
                                review_basis="single_doc_internal",
                            )
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )
            risk_path.write_text(
                json.dumps(
                    {
                        "agent_role": "risk-classifier",
                        "issues": [
                            agent_issue(
                                category="bilingual_consistency",
                                unit_id="p-00554",
                                anchor_quote="The sample ID is Sample B",
                                paragraph_index="P554",
                                original="Sample A vs Sample B",
                                issue="样品编号中英文不一致，需核对源记录。",
                                suggestion="核对源记录后统一样品编号。",
                                evidence="P554 Chinese says Sample A; English says Sample B.",
                                confidence="high",
                                coverage_domain="external_check",
                                review_basis="external_required",
                                requires_external_evidence=True,
                                external_evidence_type="protocol",
                                comment_intent="request_check",
                            )
                        ],
                    },
                    ensure_ascii=False,
                ),
                encoding="utf-8",
            )

            issues, metadata = qa.load_agent_review_issues(
                [str(bilingual_path), str(risk_path)],
                [paragraph(text, logical_index=554)],
                1,
            )

        self.assertEqual(len(issues), 1)
        self.assertEqual(issues[0]["anchor_text"], "The sample ID is Sample B")
        self.assertNotIn("filtered_external_veto", metadata["skip_categories"])

    def test_lux_definition_style_claim_is_filtered_as_low_value(self):
        issue = qa.make_llm_issue(
            1,
            [table_cell("lumen per square meter光照强度的单位", logical_index=47, table_index=2, row_index=8, cell_index=2)],
            agent_issue(
                category="semantic_consistency",
                paragraph_index="P47",
                original="Lux: lumen per square meter",
                issue="Lux 定义不准确：Lux 是照度单位，lumen per square meter 是 lux 的定义方式，但表述不够规范。",
                suggestion="考虑改为 'Lux (lx): unit of illuminance, equal to one lumen per square meter'。",
                evidence="缩略语表中 Lux 的英文描述为 'lumen per square meter'，不够完整规范。",
                confidence="medium",
                coverage_domain="data_consistency",
                review_basis="single_doc_internal",
            ),
        )

        self.assertIsNone(issue)

    def test_table_of_contents_spacing_claim_is_filtered_as_low_value_template_preference(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("目 录 Contents", logical_index=23)],
            agent_issue(
                category="semantic_consistency",
                paragraph_index="P23",
                original="目 录 Contents",
                issue="目录标题中包含多余空格，不符合排版规范。",
                suggestion="删除多余空格，改为'目录 Contents'。",
                evidence="P23 标题为 '目 录 Contents'。",
                confidence="high",
                coverage_domain="structure",
                review_basis="agent_semantic",
            ),
        )

        self.assertIsNone(issue)

    def test_packaging_material_official_term_rewrite_is_filtered_without_approved_terminology(self):
        source = "注射制剂用氯化丁基橡胶塞，注射剂用铝塑组合盖。Chlorinated butyl rubber stoppers for injectable preparations"
        issue = qa.make_llm_issue(
            1,
            [table_cell(source, logical_index=88, table_index=3, row_index=4, cell_index=4)],
            agent_issue(
                category="semantic_consistency",
                paragraph_index="P88",
                original="注射制剂用",
                issue="包材描述中标点使用不一致：前句为'注射制剂用'，后句为'注射剂用'；且中文后直接衔接英文缺少统一格式。",
                suggestion="统一表述为'注射用氯化丁基橡胶塞，注射用铝塑组合盖'。",
                evidence="P88 出现 '注射制剂用氯化丁基橡胶塞，注射剂用铝塑组合盖'。",
                confidence="medium",
                coverage_domain="structure",
                review_basis="agent_semantic",
            ),
        )

        self.assertIsNone(issue)

    def test_external_acronym_full_name_claim_is_filtered_without_controlled_source(self):
        record = paragraph("CHMP Committees for human medicinal products人用药委员会", logical_index=104)
        record["paragraph_id"] = "p-00104"

        issue = qa.make_llm_issue(
            1,
            [record],
            agent_issue(
                category="semantic_consistency",
                unit_id="p-00104",
                anchor_quote="Committees for human medicinal products",
                original="Committees for human medicinal products",
                issue=(
                    "缩写CHMP的英文全称不正确。CHMP的正确全称为"
                    "'Committee for Medicinal Products for Human Use'。"
                ),
                suggestion="更正为'Committee for Medicinal Products for Human Use'。",
                evidence="CHMP是EMA的官方委员会，标准全称为'Committee for Medicinal Products for Human Use'。",
                confidence="high",
                coverage_domain="terminology",
                review_basis="company_standard",
            ),
        )

        self.assertIsNone(issue)

    def test_agent_claiming_approved_terminology_does_not_create_controlled_source(self):
        record = paragraph("CHMP Committees for human medicinal products人用药委员会", logical_index=104)
        record["paragraph_id"] = "p-00104"

        issue = qa.make_llm_issue(
            1,
            [record],
            agent_issue(
                category="semantic_consistency",
                unit_id="p-00104",
                anchor_quote="Committees for human medicinal products",
                original="Committees for human medicinal products",
                issue="Approved terminology says CHMP should be written as Committee for Medicinal Products for Human Use.",
                suggestion="Replace with Committee for Medicinal Products for Human Use.",
                evidence="Controlled terminology source requires this official full name.",
                confidence="high",
                coverage_domain="terminology",
                review_basis="company_standard",
            ),
        )

        self.assertIsNone(issue)

    def test_internal_dp_ds_table_header_mismatch_is_not_filtered_as_external_acronym_claim(self):
        record = table_cell(
            "制剂批号DS Batch Number 对应原液批号Corresponding DS Batch Number IP166DP001",
            logical_index=286,
            table_index=2,
            row_index=3,
            cell_index=2,
        )
        record["paragraph_id"] = "p-00286"

        issue = qa.make_llm_issue(
            1,
            [record],
            agent_issue(
                category="bilingual_consistency",
                unit_id="p-00286",
                anchor_quote="制剂批号DS Batch Number",
                original="制剂批号DS Batch Number",
                issue="列标题中英文不一致：中文为制剂批号，但英文写为DS Batch Number；同表另有对应原液批号列。",
                suggestion="将该列英文改为DP Batch Number。",
                evidence="同一单元格为'制剂批号DS Batch Number'，同表另有'对应原液批号Corresponding DS Batch Number'。",
                confidence="high",
                coverage_domain="bilingual",
                review_basis="single_doc_internal",
            ),
        )

        self.assertIsNotNone(issue)
        self.assertEqual(issue["anchor_text"], "制剂批号DS Batch Number")

    def test_same_paragraph_product_name_typo_is_not_filtered_as_external_acronym_claim(self):
        record = paragraph("依苏帕格鲁肽α为目标产品。依苏帕鲁格肽α注射液获得批准上市。", logical_index=66)
        record["paragraph_id"] = "p-00066"

        issue = qa.make_llm_issue(
            1,
            [record],
            agent_issue(
                category="zh_language",
                unit_id="p-00066",
                anchor_quote="依苏帕鲁格肽α注射液",
                original="依苏帕鲁格肽α注射液",
                issue="前文使用'依苏帕格鲁肽α'，本句使用'依苏帕鲁格肽α'，产品名存在前后不一致。",
                suggestion="统一为'依苏帕格鲁肽α注射液'。",
                evidence="同一段前文为'依苏帕格鲁肽α'，后文为'依苏帕鲁格肽α注射液'。",
                confidence="high",
                coverage_domain="zh_language",
                review_basis="single_doc_internal",
            ),
        )

        self.assertIsNotNone(issue)
        self.assertEqual(issue["anchor_text"], "依苏帕鲁格肽α注射液")

    def test_cross_cell_percent_symbol_style_claim_is_filtered_as_low_value(self):
        issue = qa.make_llm_issue(
            1,
            [table_cell("106%", logical_index=910, table_index=9, row_index=12, cell_index=6)],
            agent_issue(
                category="semantic_consistency",
                paragraph_index="P910",
                original="106%",
                issue="表格数据中部分数值带百分号而其他格无百分号，格式不一致。",
                suggestion="统一同一表格中百分号的使用格式。",
                evidence="同一表格中有的单元格写为 106%，有的仅写 106。",
                confidence="high",
                coverage_domain="structure",
                review_basis="agent_semantic",
            ),
        )

        self.assertIsNone(issue)

    def test_cross_cell_percent_sign_style_claim_is_filtered_as_low_value(self):
        issue = qa.make_llm_issue(
            1,
            [table_cell("108%", logical_index=1102, table_index=10, row_index=29, cell_index=7)],
            agent_issue(
                category="semantic_consistency",
                paragraph_index="P1102",
                original="108%",
                issue="数值格式不一致：同表格中部分数值带%号部分不带，T10R29C7中109不带%而108带%。",
                suggestion="统一同一表格中是否使用 % 号。",
                evidence="同一表格相邻数值中部分带 % 号，部分不带 % 号。",
                confidence="high",
                coverage_domain="structure",
                review_basis="agent_semantic",
            ),
        )

        self.assertIsNone(issue)

    def test_signature_date_placeholder_column_claim_is_filtered_as_low_value(self):
        issue = qa.make_llm_issue(
            1,
            [table_cell("确认内容", logical_index=1, table_index=1, row_index=1, cell_index=1)],
            agent_issue(
                category="semantic_consistency",
                unit_id="p-00001",
                anchor_quote="确认内容",
                original=(
                    "The header row (b-00001) has 5 columns "
                    "(确认内容, 确认结果, 是否合格, 签字/日期, 复核/日期), "
                    "but data rows only have 3 columns "
                    "(确认内容, 确认结果, 是否合格)."
                ),
                issue=(
                    "表头行定义了5列（确认内容、确认结果、是否合格、签字/日期、复核/日期），"
                    "但数据行只有3列，缺少签字/日期和复核/日期两列。"
                ),
                suggestion="统一表头与数据行列数，补充签字/日期和复核/日期列或缩减表头。",
                evidence=(
                    "表头包含签字/日期、复核/日期，数据行仅包含确认内容、确认结果、是否合格。"
                ),
                confidence="high",
                coverage_domain="structure",
                review_basis="single_doc_internal",
            ),
        )

        self.assertIsNone(issue)

    def test_non_signature_table_column_mismatch_is_not_filtered_as_placeholder(self):
        issue = qa.make_llm_issue(
            1,
            [table_cell("样品编号", logical_index=2, table_index=2, row_index=1, cell_index=1)],
            agent_issue(
                category="semantic_consistency",
                unit_id="p-00002",
                anchor_quote="样品编号",
                original="表头定义4列（样品编号、检测项目、结果、限度），但数据行仅包含样品编号和检测项目。",
                issue="数据行缺少结果和限度两列，表格无法承载应记录的检测结论。",
                suggestion="补充数据行中的结果和限度列。",
                evidence="同一表格表头包含结果和限度，数据行仅保留样品编号和检测项目。",
                confidence="high",
                coverage_domain="structure",
                review_basis="single_doc_internal",
            ),
        )

        self.assertIsNotNone(issue)

    def test_cross_cell_spacing_style_claim_on_correct_anchor_is_filtered(self):
        issue = qa.make_llm_issue(
            1,
            [table_cell("EPY7 Standard Solution", logical_index=830, table_index=8, row_index=40, cell_index=8)],
            agent_issue(
                category="en_language",
                paragraph_index="P830",
                original="EPY7 Standard Solution",
                issue="前后空格格式不一致，同一表格部分单元格有空格部分无空格。",
                suggestion="统一同一表格中 EPY7 Standard Solution 的空格格式。",
                evidence="同一表格中部分单元格写为 EPY7 Standard Solution，部分写法空格不同。",
                confidence="high",
                coverage_domain="structure",
                review_basis="agent_semantic",
            ),
        )

        self.assertIsNone(issue)

    def test_cross_document_standard_format_unification_claim_is_filtered_when_not_exact_typo(self):
        issue = qa.make_llm_issue(
            1,
            [table_cell("EP Y5 standard solution.", logical_index=1615, table_index=11, row_index=12, cell_index=8)],
            agent_issue(
                category="en_language",
                paragraph_index="P1615",
                original="EP Y5 standard solution.",
                issue="EP 标准格式不一致: 同一文档中引用欧洲药典标准时混用 EP Y5, EPY7, EP Y7 等多种格式, 应统一。",
                suggestion="建议统一格式, 例如统一为 EPY5 standard solution。",
                evidence="文档不同表格中 EP 标准表述不一致: P1366 使用 EP Y7, P1618 使用 EPY7Standard, P1615 使用 EP Y5。",
                confidence="medium",
                coverage_domain="structure",
                review_basis="agent_semantic",
            ),
        )

        self.assertIsNone(issue)

    def test_heading_numbering_prefix_claim_is_filtered_without_template_basis(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("目的PURPOSE", logical_index=24)],
            agent_issue(
                category="semantic_consistency",
                paragraph_index="P24",
                original="目的PURPOSE",
                issue="Section 1 heading is missing the '1.' numbering prefix, breaking consistency with other headings.",
                suggestion="Add the prefix '1.' before the PURPOSE heading.",
                evidence="Other top-level headings use numeric prefixes, while this heading appears as PURPOSE.",
                confidence="high",
                coverage_domain="structure",
                review_basis="agent_semantic",
            ),
        )

        self.assertIsNone(issue)

    def test_sentence_initial_capitalization_after_period_is_filtered_as_low_value(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("The Project X Sample A were stored at pH 8.0. for the two batches of drug substance.")],
            agent_issue(
                category="en_language",
                paragraph_index="1",
                original="pH 8.0",
                issue="Minor English punctuation/capitalization issue: after 'pH 8.0.' the following word 'for' should be capitalized.",
                suggestion="Change 'for' to 'For' after the period.",
                evidence="The sentence continues with lowercase 'for' after a period.",
                confidence="high",
                coverage_domain="en_language",
                review_basis="agent_semantic",
            ),
        )

        self.assertIsNone(issue)

    def test_spelling_issue_prefers_actual_misspelled_source_anchor(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("Greater than turbidity standard solution 0.5, less urbidity standard solution 1.")],
            agent_issue(
                category="en_language",
                paragraph_index="1",
                original="turbidity",
                issue="英文'turbidity'拼写为'urbidity'（缺首字母't'）。",
                suggestion="改为'less than turbidity standard solution 1'。",
                evidence="原文为'less urbidity standard solution 1'。",
                coverage_domain="en_language",
                review_basis="agent_semantic",
            ),
        )

        self.assertIsNotNone(issue)
        self.assertEqual(issue["anchor_text"], "urbidity")

    def test_source_local_typo_with_speculative_explanation_is_kept_when_anchored(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("Greater than turbidity standard solution 0.5, less urbidity standard solution 1.")],
            agent_issue(
                category="en_language",
                paragraph_index="1",
                original="less urbidity standard solution 1.",
                issue="Typo: 'less urbidity' is missing 'than'.",
                suggestion="Correct to 'less than turbidity standard solution 1.'",
                evidence=(
                    "P831 contains 'less urbidity standard solution 1.'; "
                    "the word 'urbidity' appears to be a truncation of 'turbidity'."
                ),
                coverage_domain="en_language",
                review_basis="agent_semantic",
            ),
        )

        self.assertIsNotNone(issue)
        self.assertEqual(issue["anchor_text"], "less urbidity standard solution 1")

    def test_unique_source_local_typo_can_reanchor_when_paragraph_index_is_wrong(self):
        issue = qa.make_llm_issue(
            1,
            [
                paragraph("Uncoated Capilary", logical_index=540),
                paragraph("Other nearby consumable text.", logical_index=541),
                paragraph("More consumable metadata.", logical_index=542),
                paragraph("Capillary method description without the typo.", logical_index=543),
            ],
            agent_issue(
                category="en_language",
                paragraph_index="P543",
                original="Uncoated Capilary",
                issue="Spelling error: 'Capilary' should be 'Capillary'.",
                suggestion="Change 'Capilary' to 'Capillary'.",
                evidence="P540 shows 'Uncoated Capilary' - missing one 'l'.",
                coverage_domain="en_language",
                review_basis="agent_semantic",
            ),
        )

        self.assertIsNotNone(issue)
        self.assertEqual(issue["location"], "第540段（XML:540）")
        self.assertEqual(issue["anchor_text"], "Uncoated Capilary")

    def test_chemical_formula_case_issue_is_not_filtered_as_low_value_style(self):
        issue = qa.make_llm_issue(
            1,
            [paragraph("试剂：氯化钠Nacl。", logical_index=282)],
            agent_issue(
                category="en_language",
                paragraph_index="P282",
                original="Nacl",
                issue="英文化学式Nacl应写为NaCl，大小写不规范。",
                suggestion="将Nacl改为NaCl。",
                evidence="P282中英文化学式Nacl应为NaCl。",
                coverage_domain="en_language",
                review_basis="single_doc_internal",
            ),
        )

        self.assertIsNotNone(issue)
        self.assertEqual(issue["anchor_text"], "Nacl")

    def test_adjudication_dedupes_same_location_same_cause_with_different_anchors(self):
        record = paragraph(
            "将Project X项目ADC成品在避光、300rpm、室温条件下振荡5天后，"
            "其余各检项均无明显变化。因此，为保证样品稳定性，"
            "Project X单抗在生产、储存、转运等过程中建议避免剧烈振荡。",
            logical_index=2974,
        )
        record["xml_index"] = 3948
        issues = [
            qa.make_issue(
                1,
                rule_id="LLM-BI-001",
                branch=qa.LLM_REVIEW_BRANCH,
                paragraph_index=None,
                location="第2974段（XML:3948）",
                original='Chinese: "Project X单抗在生产、储存、转运等过程中建议避免剧烈振荡"\nEnglish: "Project X monoclonal antibody"',
                issue="Both Chinese and English refer to 'Project X单抗/Project X monoclonal antibody' but the document is about Project X Sample.",
                suggestion="Change both Chinese and English to refer to 'Project X项目ADC成品/Project X Sample B'.",
                severity="主要",
                evidence="同段使用 Project X单抗/Project X monoclonal antibody。",
                anchor_locator="paragraph=3948",
                anchor_span={"start": 13, "end": 23, "unit": "char"},
                anchor_text="Project X项目ADC",
            ),
            qa.make_issue(
                2,
                rule_id="LLM-SEM-001",
                branch=qa.LLM_REVIEW_BRANCH,
                paragraph_index=None,
                location="第2974段（XML:3948）",
                original="Project X单抗在生产、储存、转运等过程中建议避免剧烈振荡",
                issue="Section text refers to 'Project X单抗' but the document is about Project X Sample.",
                suggestion="Change 'Project X单抗' to 'Project X项目ADC成品'.",
                severity="次要",
                evidence="原文包含 Project X单抗。",
                anchor_locator="paragraph=3948",
                anchor_span={"start": 13, "end": 20, "unit": "char"},
                anchor_text="Project X单抗",
            ),
            qa.make_issue(
                3,
                rule_id="LLM-ZH-001",
                branch=qa.LLM_REVIEW_BRANCH,
                paragraph_index=None,
                location="第2974段（XML:3948）",
                original="Project X单抗在生产、储存、转运等过程中建议避免剧烈振荡",
                issue="文中误将'Project X项目ADC'写为'Project X单抗'。",
                suggestion="将'Project X单抗'修改为'Project X项目ADC成品'，与全文其他地方保持一致。",
                severity="次要",
                evidence="原文包含 Project X单抗。",
                anchor_locator="paragraph=3948",
                anchor_span={"start": 13, "end": 20, "unit": "char"},
                anchor_text="Project X项目ADC成品",
            ),
        ]

        adjudicated = qa.adjudicate_issues_before_comment_plan(issues, [record])

        self.assertEqual(len(adjudicated), 1)
        self.assertIn("Project X单抗", adjudicated[0]["issue"])

    def test_adjudication_dedupes_same_location_semantic_paraphrases(self):
        chinese_anchor = "将Project X Sample A按对应试验条件进行6次循环冻融后"
        english_anchor = "The Project X Sample A was held under Method B conditions for five days"
        record_text = f"{chinese_anchor}。{english_anchor}."
        record = paragraph(record_text, logical_index=1398)
        record["xml_index"] = 1820
        first = qa.make_issue(
            1,
            rule_id="LLM-SEM-001",
            branch=qa.LLM_REVIEW_BRANCH,
            paragraph_index=None,
            location="第1398段（XML:1820）",
            original=english_anchor,
            issue="Freeze-thaw test section has English text describing Method B instead of freeze-thaw cycling.",
            suggestion="Replace with English text describing 6-cycle freeze-thaw results matching the Chinese description.",
            severity="主要",
            evidence="Chinese text says '6次循环冻融' but English says 'Method B conditions'.",
            anchor_locator="paragraph=1820",
            anchor_span={
                "start": record_text.index(english_anchor),
                "end": record_text.index(english_anchor) + len(english_anchor),
                "unit": "char",
            },
            anchor_text=english_anchor,
        )
        second = qa.make_issue(
            2,
            rule_id="LLM-BI-001",
            branch=qa.LLM_REVIEW_BRANCH,
            paragraph_index=None,
            location="第1398段（XML:1820）",
            original=chinese_anchor,
            issue="Freeze-thaw section English text incorrectly describes Method B instead of the Chinese 6-cycle freeze-thaw test.",
            suggestion="Correct the English text to match the Chinese freeze-thaw description.",
            severity="主要",
            evidence="Chinese text describes 6 freeze-thaw cycles; English mentions Method B conditions.",
            anchor_locator="paragraph=1820",
            anchor_span={"start": 0, "end": len(chinese_anchor), "unit": "char"},
            anchor_text=chinese_anchor,
        )

        adjudicated = qa.adjudicate_issues_before_comment_plan([first, second], [record])

        self.assertEqual(len(adjudicated), 1)

    def test_adjudication_dedupes_deterministic_and_agent_temperature_findings(self):
        record = table_cell(
            "将样品置于25°C±2°C，60%RH±5%RH、遮光条件下。"
            "the samples were stored at 25°C±5°C, 60%RH±5%RH and protected from light.",
            logical_index=571,
        )
        record["xml_index"] = 639
        deterministic = qa.content_issue(
            1,
            record,
            rule_id="CONTENT-BI-TEMP-001",
            original="25°C±5°C",
            issue="同一表格单元格内温度允差中英文不一致：同一温度条件同时出现 `25°C±2°C` 和 `25°C±5°C`。",
            suggestion="将该处温度允差统一为 `25°C±2°C`。",
            evidence=record["text"],
        )
        agent = qa.make_issue(
            2,
            rule_id="LLM-SEM-001",
            branch=qa.LLM_REVIEW_BRANCH,
            paragraph_index=None,
            location="第571段（XML:639）",
            original='Chinese: "25°C±2°C"; English: "25°C±5°C"',
            issue="English temperature tolerance differs from Chinese. Chinese states 25°C±2°C while English states 25°C±5°C.",
            suggestion="Unify to 25°C±2°C in both Chinese and English.",
            severity="主要",
            evidence="Chinese cell T7R5C3: 25°C±2°C; English cell T7R5C3: 25°C±5°C.",
            anchor_locator="paragraph=639",
            anchor_span={"start": 55, "end": 62, "unit": "char"},
            anchor_text="25°C±2°C",
        )

        adjudicated = qa.adjudicate_issues_before_comment_plan([deterministic, agent], [record])

        self.assertEqual(len(adjudicated), 1)
        self.assertEqual(adjudicated[0]["rule_id"], "CONTENT-BI-TEMP-001")

    def test_adjudication_keeps_title_mismatch_on_wrong_english_title(self):
        record = paragraph(
            "根据《Project X Protocol》（DOC-X-003），"
            "According to the < Project X Summary > (DOC-X-003), the test plan is shown in Table 5.",
            logical_index=554,
        )
        issue = qa.make_llm_issue(
            1,
            [record],
            agent_issue(
                category="en_language",
                paragraph_index="P554",
                original="Project X Summary",
                issue=(
                    "Document reference name mismatch: Chinese document title refers to Protocol, "
                    "but English translates it as Summary."
                ),
                suggestion=(
                    "Correct English reference to `Project X Protocol`."
                ),
                evidence=(
                    "Chinese references `Project X Protocol`, but English calls it `Project X Summary`."
                ),
                coverage_domain="en_language",
                review_basis="agent_semantic",
            ),
        )

        self.assertIsNotNone(issue)
        adjudicated = qa.adjudicate_issues_before_comment_plan([issue], [record])
        self.assertEqual(len(adjudicated), 1)
        self.assertEqual(adjudicated[0]["anchor_text"], "Project X Summary")

    def test_adjudication_dedupes_deterministic_and_agent_english_grammar(self):
        record = table_cell(
            "Samples containing 0.2% tBHP were filter through a 0.22 μm sterile filter.",
            logical_index=574,
            row_index=6,
        )
        record["xml_index"] = 643
        deterministic = qa.content_issue(
            1,
            record,
            rule_id="CONTENT-EN-GRAMMAR-001",
            original="were filter",
            issue='英文被动语态动词形式错误：`were filter` 应为 `were filtered`。',
            suggestion='将 `were filter` 改为 `were filtered`。',
            severity="次要",
            evidence=record["text"],
        )
        agent = qa.make_llm_issue(
            2,
            [record],
            agent_issue(
                category="en_language",
                paragraph_index="P574",
                original="were filter",
                issue='Incorrect verb form: "were filter" should be "were filtered".',
                suggestion='Change "were filter" to "were filtered".',
                evidence='The source text contains "were filter" after passive auxiliary "were".',
                coverage_domain="en_language",
                review_basis="single_doc_internal",
            ),
        )

        adjudicated = qa.adjudicate_issues_before_comment_plan([deterministic, agent], [record])

        self.assertEqual(len(adjudicated), 1)
        self.assertEqual(adjudicated[0]["rule_id"], "CONTENT-EN-GRAMMAR-001")

    def test_adjudication_keeps_distinct_temperature_and_grammar_issues_same_cell(self):
        record = table_cell(
            "将含0.2%叔丁基过氧化氢的样品于25°C±2°C放置。"
            "Samples containing 0.2% tBHP were filter through a 0.22 μm sterile filter, "
            "stored at 25°C±5°C.",
            logical_index=574,
            row_index=6,
        )
        record["xml_index"] = 643
        temp_issue = qa.content_issue(
            1,
            record,
            rule_id="CONTENT-BI-TEMP-001",
            original="25°C±5°C",
            issue="同一表格单元格内温度允差中英文不一致：同一温度条件同时出现 `25°C±2°C` 和 `25°C±5°C`。",
            suggestion="将该处温度允差统一为 `25°C±2°C`。",
            evidence=record["text"],
        )
        grammar_issue = qa.content_issue(
            2,
            record,
            rule_id="CONTENT-EN-GRAMMAR-001",
            original="were filter",
            issue='英文被动语态动词形式错误：`were filter` 应为 `were filtered`。',
            suggestion='将 `were filter` 改为 `were filtered`。',
            severity="次要",
            evidence=record["text"],
        )

        adjudicated = qa.adjudicate_issues_before_comment_plan([temp_issue, grammar_issue], [record])

        self.assertEqual({item["rule_id"] for item in adjudicated}, {"CONTENT-BI-TEMP-001", "CONTENT-EN-GRAMMAR-001"})

    def test_adjudication_dedupes_overlapping_release_note_agent_anchor(self):
        record = paragraph("Note: *1 indicates this data reference to the releasing data.", logical_index=844)
        record["xml_index"] = 1034
        deterministic = qa.content_issue(
            1,
            record,
            rule_id="CONTENT-EN-GRAMMAR-002",
            original="data reference to the releasing data",
            issue="英文脚注表达不清：`data reference to the releasing data` 不是规范英语表达。",
            suggestion="改为 `data references the release data` 或 `data refers to the release data`。",
            severity="次要",
            evidence="同类英文脚注出现于第844段、第1144段，其余8处。",
        )
        agent = qa.make_issue(
            2,
            rule_id="LLM-EN-001",
            branch=qa.LLM_REVIEW_BRANCH,
            paragraph_index=None,
            location="第844段（XML:1034）",
            original="this data reference to the releasing data",
            issue="英文语法不当，'this data reference to'应改为'this data references'。",
            suggestion="改为'this data references the release data'",
            severity="次要",
            evidence="原文为 this data reference to the releasing data。",
            anchor_locator="paragraph=1034",
            anchor_span={"start": 19, "end": 41, "unit": "char"},
            anchor_text="this data reference to",
        )

        adjudicated = qa.adjudicate_issues_before_comment_plan([deterministic, agent], [record])

        self.assertEqual(len(adjudicated), 1)
        self.assertEqual(adjudicated[0]["rule_id"], "CONTENT-EN-GRAMMAR-002")

    def test_adjudication_dedupes_llm_restatement_of_deterministic_duplicate_semicolon(self):
        record = paragraph("碱性组分分别上升了2.5%、1.9%；；蛋白浓度由于加入盐酸而偏低。", logical_index=1651)
        record["xml_index"] = 2189
        deterministic = qa.content_issue(
            1,
            record,
            rule_id="CONTENT-PUNCT-001",
            original="；；",
            issue="中文正文出现连续两个分号，属于多余标点。",
            suggestion="删除多余分号，保留一个分号。",
            severity="次要",
            evidence=record["text"],
        )
        agent = qa.make_issue(
            2,
            rule_id="LLM-ZH-001",
            branch=qa.LLM_REVIEW_BRANCH,
            paragraph_index=None,
            location="第1651段（XML:2189）",
            original="蛋白浓度由于加入盐酸而偏低",
            issue="中文标点重复：出现连续两个分号 '；；'。",
            suggestion="删除多余的分号，改为：`；蛋白浓度由于加入盐酸而偏低`。",
            severity="次要",
            evidence="原文包含连续两个分号。",
            anchor_locator="paragraph=2189",
            anchor_span={"start": 18, "end": 30, "unit": "char"},
            anchor_text="蛋白浓度由于加入盐酸而偏低",
        )

        adjudicated = qa.adjudicate_issues_before_comment_plan([deterministic, agent], [record])

        self.assertEqual(len(adjudicated), 1)
        self.assertEqual(adjudicated[0]["rule_id"], "CONTENT-PUNCT-001")

    def test_adjudication_dedupes_llm_restatement_of_deterministic_semicolon_period(self):
        record = paragraph("分子大小变异体发生变化，进而导致纯度的降低；。光照、高温也会影响成品。", logical_index=3167)
        record["xml_index"] = 4184
        deterministic = qa.content_issue(
            1,
            record,
            rule_id="CONTENT-PUNCT-001",
            original="；。",
            issue="中文正文出现分号后紧跟句号的异常标点组合。",
            suggestion="删除分号或改为单个句号，保持句读清晰。",
            severity="次要",
            evidence=record["text"],
        )
        agent = qa.make_issue(
            2,
            rule_id="LLM-ZH-001",
            branch=qa.LLM_REVIEW_BRANCH,
            paragraph_index=None,
            location="第3167段（XML:4184）",
            original="进而导致纯度的降低；。光照、高温",
            issue="中文标点错误：分号后紧跟句号 '；。'。",
            suggestion="删除多余的分号，改为：`进而导致纯度的降低。光照、高温`。",
            severity="次要",
            evidence="原文包含 '；。'。",
            anchor_locator="paragraph=4184",
            anchor_span={"start": 13, "end": 29, "unit": "char"},
            anchor_text="进而导致纯度的降低；。光照、高温",
        )

        adjudicated = qa.adjudicate_issues_before_comment_plan([deterministic, agent], [record])

        self.assertEqual(len(adjudicated), 1)
        self.assertEqual(adjudicated[0]["rule_id"], "CONTENT-PUNCT-001")

    def test_adjudication_dedupes_cross_branch_temperature_semantic_duplicate(self):
        record = paragraph(
            "中文条件为25°C±2°C。English condition was stored at 25°C±5°C.",
            logical_index=574,
        )
        record["xml_index"] = 643
        deterministic = qa.content_issue(
            1,
            record,
            rule_id="CONTENT-BI-TEMP-001",
            original="25°C±5°C",
            issue="同一表格单元格内温度允差中英文不一致：同一温度条件同时出现 `25°C±2°C` 和 `25°C±5°C`。",
            suggestion="将该处温度允差统一为 `25°C±2°C`。",
            evidence=record["text"],
        )
        agent = qa.make_issue(
            2,
            rule_id="LLM-SEM-001",
            branch=qa.LLM_REVIEW_BRANCH,
            paragraph_index=None,
            location="第574段（XML:643）",
            original="中文条件为25°C±2°C",
            issue="英文条件温度公差为±5°C，但中文为±2°C。",
            suggestion="将英文25°C±5°C改为25°C±2°C。",
            severity="主要",
            evidence="中文条件为25°C±2°C，英文写25°C±5°C。",
            anchor_locator="paragraph=643",
            anchor_span={"start": 0, "end": 13, "unit": "char"},
            anchor_text="中文条件为25°C±2°C",
        )

        adjudicated = qa.adjudicate_issues_before_comment_plan([deterministic, agent], [record])

        self.assertEqual(len(adjudicated), 1)
        self.assertEqual(adjudicated[0]["rule_id"], "CONTENT-BI-TEMP-001")

    def test_adjudication_dedupes_same_location_replacement_findings(self):
        record = paragraph("The Cilent value was copied into the table.", logical_index=846)
        record["xml_index"] = 1036
        first = qa.make_issue(
            1,
            rule_id="LLM-EN-001",
            branch=qa.LLM_REVIEW_BRANCH,
            paragraph_index=None,
            location="?846??XML:1036?",
            original="Cilent",
            issue="Typographical error: `Cilent` should be `Client`.",
            suggestion="Replace `Cilent` with `Client`.",
            severity="??",
            evidence="The source sentence contains `Cilent`.",
            anchor_locator="paragraph=1036",
            anchor_span={"start": 4, "end": 10, "unit": "char"},
            anchor_text="Cilent",
        )
        second = qa.make_issue(
            2,
            rule_id="LLM-EN-001",
            branch=qa.LLM_REVIEW_BRANCH,
            paragraph_index=None,
            location="?846??XML:1036?",
            original="Cilent value",
            issue="The same typo appears in the source phrase.",
            suggestion="Change `Cilent` to `Client`.",
            severity="??",
            evidence="The source sentence contains `Cilent value`.",
            anchor_locator="paragraph=1036",
            anchor_span={"start": 4, "end": 16, "unit": "char"},
            anchor_text="Cilent value",
        )

        adjudicated = qa.adjudicate_issues_before_comment_plan([first, second], [record])

        self.assertEqual(len(adjudicated), 1)
        self.assertIn("Cilent", adjudicated[0]["anchor_text"])

    def test_adjudication_refines_broad_anchor_to_source_side_replacement_phrase(self):
        record = paragraph("The HMW HMW result was reported in Table 1.", logical_index=1146)
        record["xml_index"] = 1452
        issue = qa.make_issue(
            1,
            rule_id="LLM-EN-001",
            branch=qa.LLM_REVIEW_BRANCH,
            paragraph_index=None,
            location="?1146??XML:1452?",
            original=record["text"],
            issue="Duplicate adjacent token in the source sentence.",
            suggestion="Replace `HMW HMW` with `HMW`.",
            severity="??",
            evidence="The source sentence contains `HMW HMW`.",
            anchor_locator="paragraph=1452",
            anchor_span={"start": 0, "end": len(record["text"]), "unit": "char"},
            anchor_text=record["text"],
        )

        adjudicated = qa.adjudicate_issues_before_comment_plan([issue], [record])

        self.assertEqual(len(adjudicated), 1)
        self.assertEqual(adjudicated[0]["anchor_text"], "HMW HMW")

    def test_adjudication_keeps_distinct_issues_same_location(self):
        record = paragraph("The Cilent HMW HMW result was reviewed.", logical_index=1200)
        record["xml_index"] = 1500
        typo = qa.make_issue(
            1,
            rule_id="LLM-EN-001",
            branch=qa.LLM_REVIEW_BRANCH,
            paragraph_index=None,
            location="?1200??XML:1500?",
            original="Cilent",
            issue="Typographical error: `Cilent` should be `Client`.",
            suggestion="Replace `Cilent` with `Client`.",
            severity="??",
            evidence="The source sentence contains `Cilent`.",
            anchor_locator="paragraph=1500",
            anchor_span={"start": 4, "end": 10, "unit": "char"},
            anchor_text="Cilent",
        )
        duplicate = qa.make_issue(
            2,
            rule_id="LLM-EN-001",
            branch=qa.LLM_REVIEW_BRANCH,
            paragraph_index=None,
            location="?1200??XML:1500?",
            original="HMW HMW",
            issue="Duplicate adjacent token: `HMW HMW`.",
            suggestion="Replace `HMW HMW` with `HMW`.",
            severity="??",
            evidence="The source sentence contains `HMW HMW`.",
            anchor_locator="paragraph=1500",
            anchor_span={"start": 11, "end": 18, "unit": "char"},
            anchor_text="HMW HMW",
        )

        adjudicated = qa.adjudicate_issues_before_comment_plan([typo, duplicate], [record])

        self.assertEqual(len(adjudicated), 2)
        self.assertEqual({item["anchor_text"] for item in adjudicated}, {"Cilent", "HMW HMW"})

    def test_adjudication_clusters_recurring_turbidity_findings_across_locations(self):
        first = qa.make_issue(
            1,
            rule_id="LLM-BI-001",
            branch=qa.LLM_REVIEW_BRANCH,
            paragraph_index=None,
            location="第831段（XML:1017）",
            original="turbidity",
            issue='浊度标准液数字不一致且英文不完整，出现"less urbidity"。',
            suggestion="统一浊度标准液编号并改为less than turbidity。",
            severity="主要",
            evidence="P831、P833等多处出现less urbidity。",
            anchor_locator="paragraph=1017",
            anchor_span={"start": 0, "end": 9, "unit": "char"},
            anchor_text="turbidity",
        )
        second = qa.make_issue(
            2,
            rule_id="LLM-BI-001",
            branch=qa.LLM_REVIEW_BRANCH,
            paragraph_index=None,
            location="第833段（XML:1019）",
            original="less urbidity",
            issue="浊度标准液数字不一致，英文错误写为0.5号且语法不完整。",
            suggestion="统一浊度标准液编号并补全less than turbidity。",
            severity="主要",
            evidence="原文含less urbidity standard solution。",
            anchor_locator="paragraph=1019",
            anchor_span={"start": 0, "end": 13, "unit": "char"},
            anchor_text="less urbidity",
        )

        adjudicated = qa.adjudicate_issues_before_comment_plan([first, second], [])

        self.assertEqual(len(adjudicated), 1)
        self.assertEqual(adjudicated[0]["location"], "第831段（XML:1017）")

    def test_turbidity_less_urbidity_agent_issue_is_not_critical(self):
        severity = qa.normalized_agent_visible_severity(
            "关键",
            agent_issue(
                category="bilingual_consistency",
                original="中文：深于2号浊度标准液，浅于3号浊度标准液；英文：Greater than turbidity standard solution 0.5, less urbidity",
                issue="中英文浊度标准液编号不一致，英文描述为 0.5 且出现 less urbidity。",
                suggestion="统一中英文浊度标准液编号，并补全 less than turbidity。",
                evidence="P1129 中英文浊度结果不一致，英文含 less urbidity。",
                confidence="high",
            ),
        )

        self.assertEqual(severity, "主要")

    def test_zh_sentence_flow_polish_is_filtered_as_low_value(self):
        issue = qa.make_llm_issue(
            1,
            [
                paragraph(
                    "按照《Project X项目ADC原液和制剂影响因素考察试验方案》（TP-001）根据产品特性开展了Project X项目ADC的影响因素试验。",
                    logical_index=25,
                )
            ],
            agent_issue(
                category="zh_language",
                paragraph_index="P25",
                original="根据产品特性",
                issue="两个介词短语之间缺少连接标点和合理衔接，导致句子结构堆叠、语意不清。",
                suggestion="建议改为：根据试验方案，结合产品特性，开展影响因素试验。",
                evidence="原文 P25 中 '按照...方案' 和 '根据产品特性' 衔接不自然。",
                confidence="high",
                coverage_domain="zh_language",
                review_basis="agent_semantic",
            ),
        )

        self.assertIsNone(issue)

    def test_zh_redundant_direction_word_is_minor_not_major(self):
        severity = qa.normalized_agent_visible_severity(
            "主要",
            agent_issue(
                category="zh_language",
                original="尽量避免高温环境下",
                issue="'尽量避免高温环境下'中'下'字赘余，语义不完整。",
                suggestion="改为'避免高温环境'。",
                evidence="原文为 '尽量避免高温环境下'。",
                confidence="high",
            ),
        )

        self.assertEqual(severity, "次要")

    def test_adjudication_dedupes_release_note_agent_issue_across_repeated_locations(self):
        deterministic_record = paragraph("Note: *1 indicates this data reference to the releasing data.", logical_index=844)
        deterministic_record["xml_index"] = 1034
        agent_record = paragraph("Note: *1 indicates this data reference to the releasing data.", logical_index=1648)
        agent_record["xml_index"] = 2186
        deterministic = qa.content_issue(
            1,
            deterministic_record,
            rule_id="CONTENT-EN-GRAMMAR-002",
            original="data reference to the releasing data",
            issue="英文脚注表达不清：`data reference to the releasing data` 不是规范英语表达。",
            suggestion="改为 `data references the release data` 或 `data refers to the release data`。",
            severity="次要",
            evidence="同类英文脚注出现于第844段、第1144段，其余8处。",
        )
        agent = qa.make_issue(
            2,
            rule_id="LLM-EN-001",
            branch=qa.LLM_REVIEW_BRANCH,
            paragraph_index=None,
            location="第1648段（XML:2186）",
            original="Note: *1 indicates this data reference to the releasing data.",
            issue="英文语法需要调整，indicates this data reference to应为indicates that this data refers to。",
            suggestion="改为Note: *1 indicates that this data refers to the release data.",
            severity="主要",
            evidence="P1648等多处注释中indicates this data reference to the releasing data语法不规范。",
            anchor_locator="paragraph=2186",
            anchor_span={"start": 0, "end": 58, "unit": "char"},
            anchor_text="Note: *1 indicates this data reference to the releasing data",
        )

        adjudicated = qa.adjudicate_issues_before_comment_plan([agent, deterministic], [deterministic_record, agent_record])

        self.assertEqual(len(adjudicated), 1)
        self.assertEqual(adjudicated[0]["rule_id"], "CONTENT-EN-GRAMMAR-002")

    def test_comment_text_is_concise_and_hides_internal_basis(self):
        text = qa.build_comment_text(
            "English temperature tolerance differs from Chinese.",
            "Unify to 25°C±2°C in both Chinese and English.",
            "主要",
            issue_type="大模型全文审核",
            original='Chinese: "25°C±2°C"; English: "25°C±5°C"',
            evidence="Chinese cell T7R5C3: 25°C±2°C; English cell T7R5C3: 25°C±5°C.",
            review_basis="single_doc_internal",
        )

        self.assertNotIn("主要 / 大模型全文审核", text)
        self.assertNotIn("审核依据", text)
        self.assertIn("发现：", text)
        self.assertIn("原文：", text)
        self.assertIn("建议：", text)
        self.assertNotIn("问题：", text)
        self.assertNotIn("依据：", text)
        self.assertLessEqual(len(text.splitlines()), 3)

    def test_comment_text_hides_internal_paragraph_locators(self):
        text = qa.build_comment_text(
            "偏差列表编号前后不一致。",
            "核对正确的偏差列表编号，并统一全文引用。",
            "主要",
            issue_type="大模型全文审核",
            original="P342原文：「验证执行过程中发生的任何偏差应记录在MV(C)-IP312-P-020-R05《偏差列表》中」。",
            evidence="P342原文：「验证执行过程中发生的任何偏差应记录在MV(C)-IP312-P-020-R05《偏差列表》中」。P354（记录章节）列出：「MV(C)-IP312-P-027-R04《偏差列表》」。",
            review_basis="single_doc_internal",
        )

        self.assertIn("发现：偏差列表编号前后不一致。", text)
        self.assertIn("原文：「验证执行过程中发生的任何偏差应记录在MV(C)-IP312-P-020-R05《偏差列表》中」。", text)
        self.assertIn("建议：核对正确的偏差列表编号，并统一全文引用。", text)
        self.assertNotRegex(text, r"\bP\d+\b")
        self.assertNotIn("依据：", text)

    def test_comment_text_lines_are_compact_for_long_format_examples(self):
        text = qa.build_comment_text(
            "页面设置中的纸张尺寸、装订线或页眉页脚距离不符合公司标准。",
            "将纸张统一为 A4，装订线设为左侧 0.5cm，页眉 1.5cm，页脚 1.75cm。",
            "主要",
            issue_type="格式",
            original=(
                "第1节(page=21.0x29.7, header=1.5, footer=1.75, gutter=0.0)；"
                "第2节(page=21.0x29.7, header=1.5, footer=1.75, gutter=0.0)；"
                "第3节(page=21.0x29.7, header=1.5, footer=1.75, gutter=0.0)；"
                "第4节(page=29.7x21.0, header=1.5, footer=1.75, gutter=0.0)"
            ),
            evidence=(
                "第1节(page=21.0x29.7, header=1.5, footer=1.75, gutter=0.0)；"
                "第2节(page=21.0x29.7, header=1.5, footer=1.75, gutter=0.0)；"
                "第3节(page=21.0x29.7, header=1.5, footer=1.75, gutter=0.0)"
            ),
            review_basis="company_standard",
        )

        self.assertLessEqual(max(len(line) for line in text.splitlines()), 150)
        self.assertLessEqual(len(text), 520)

    def test_hidden_human_review_items_do_not_mark_manifest_human_review(self):
        entry = qa.build_manifest_entry(
            qa.LLM_REVIEW_BRANCH,
            [],
            1,
            human_review_items=[
                qa.make_human_review_item(
                    qa.LLM_REVIEW_BRANCH,
                    "filtered low-quality agent issue",
                    "P1",
                    user_visible=False,
                )
            ],
            branch_details={"loaded_issue_count": 1, "schema_invalid_count": 0},
        )

        self.assertEqual(entry["status"], "completed")
        self.assertEqual(entry["loaded_issue_count"], 1)

    def test_manifest_issue_counts_sync_after_cross_branch_dedupe(self):
        manifest = [
            qa.build_manifest_entry("format", [], 1),
            qa.build_manifest_entry("project_number", [], 1),
            qa.build_manifest_entry(
                qa.CONTENT_CONSISTENCY_BRANCH,
                [
                    {"branch": qa.CONTENT_CONSISTENCY_BRANCH},
                    {"branch": qa.CONTENT_CONSISTENCY_BRANCH},
                ],
                1,
            ),
            qa.build_manifest_entry(
                qa.LLM_REVIEW_BRANCH,
                [{"branch": qa.LLM_REVIEW_BRANCH}],
                1,
                branch_details={"loaded_issue_count": 1},
            ),
        ]

        qa.sync_manifest_issue_counts(
            manifest,
            [
                {"branch": qa.CONTENT_CONSISTENCY_BRANCH},
                {"branch": qa.LLM_REVIEW_BRANCH},
            ],
        )

        by_branch = {item["branch"]: item for item in manifest}
        self.assertEqual(by_branch[qa.CONTENT_CONSISTENCY_BRANCH]["issue_count"], 1)
        self.assertEqual(by_branch[qa.LLM_REVIEW_BRANCH]["issue_count"], 1)
        self.assertEqual(by_branch[qa.LLM_REVIEW_BRANCH]["loaded_issue_count"], 1)


if __name__ == "__main__":
    unittest.main()

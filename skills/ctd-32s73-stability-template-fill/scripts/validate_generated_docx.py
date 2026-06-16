from __future__ import annotations

import argparse
import json
import re
from pathlib import Path
from typing import Any

from runtime_guard import require_internal_context

from docx import Document
from docx.oxml.ns import qn


DEFAULT_PLACEHOLDERS = ["XXX", "IPXXX", "单抗/原液", "{{", "}}", "已完成的申报资料"]
DEFAULT_REQUIRED_SECTIONS = {
    "long_term": ["长期稳定性", "长期条件考察"],
    "accelerated": ["加速稳定性", "加速条件考察"],
    "stress_study": ["影响因素", "强制降解", "应力"],
}
DEFAULT_TABLE_INDEXES = [1]
CAPTION_RE = re.compile(r"^(表|图)\s*3\.2\.S\.7\.3-\s*(\d+)")
REFERENCE_RE = re.compile(r"(表|图)\s*3\.2\.S\.7\.3-\s*(\d+)")
ROLE_ALIASES = {
    "light_stress": "stress",
    "agitation_stress": "stress",
    "freeze_thaw_stress": "stress",
    "high_temperature_stress": "stress",
    "low_ph_stress": "stress",
    "high_ph_stress": "stress",
    "oxidation_stress": "stress",
}
STRESS_SECTION_RULES = [
    {
        "key": "light",
        "roles": {"light_stress"},
        "heading_markers": ["光照试验结果"],
        "aliases": ["光照", "light"],
    },
    {
        "key": "agitation",
        "roles": {"agitation_stress"},
        "heading_markers": ["振荡试验结果", "震荡试验结果"],
        "aliases": ["振荡", "震荡", "agitation", "shaking", "shake"],
    },
    {
        "key": "freeze_thaw",
        "roles": {"freeze_thaw_stress"},
        "heading_markers": ["冻融试验结果", "反复冻融试验结果"],
        "aliases": ["冻融", "反复冻融", "freeze", "thaw", "freeze_thaw"],
    },
    {
        "key": "high_temperature",
        "roles": {"high_temperature_stress"},
        "heading_markers": ["高温试验结果"],
        "aliases": ["高温", "high temperature", "high_temperature"],
    },
    {
        "key": "ph",
        "roles": {"low_ph_stress", "high_ph_stress"},
        "heading_markers": ["pH试验结果", "PH试验结果", "酸碱试验结果"],
        "aliases": ["ph", "pH", "PH", "低pH", "高pH", "低ph", "高ph", "酸碱", "low ph", "high ph"],
    },
    {
        "key": "oxidation",
        "roles": {"oxidation_stress"},
        "heading_markers": ["氧化试验结果"],
        "aliases": ["氧化", "oxidation", "oxidative", "H2O2", "TBHP"],
    },
]

STRESS_ROLES = frozenset(role for rule in STRESS_SECTION_RULES for role in rule["roles"])


def cell_text(tc) -> str:
    texts: list[str] = []
    for paragraph in tc.p_lst:
        texts.append("".join(node.text or "" for node in paragraph.iter(qn("w:t"))))
    return "/".join(text for text in texts if text)


def grid_span(tc) -> int:
    tc_pr = tc.tcPr
    if tc_pr is not None and tc_pr.gridSpan is not None:
        return int(tc_pr.gridSpan.val)
    return 1


def v_merge(tc) -> str:
    tc_pr = tc.tcPr
    if tc_pr is not None and tc_pr.vMerge is not None:
        return str(tc_pr.vMerge.val or "continue")
    return ""


def document_text(doc: Document) -> str:
    paragraph_text = [p.text for p in doc.paragraphs]
    table_text = [cell_text(tc) for table in doc.tables for tr in table._tbl.tr_lst for tc in tr.tc_lst]
    return "\n".join(paragraph_text + table_text)


def body_paragraph_text(doc: Document) -> str:
    skipped = {"toc 1", "toc 2", "toc 3", "table of figures"}
    return "\n".join(p.text for p in doc.paragraphs if str(p.style.name or "").lower() not in skipped)


def parse_csv_ints(values: list[str]) -> list[int]:
    parsed: list[int] = []
    for value in values:
        for part in value.split(","):
            part = part.strip()
            if part:
                parsed.append(int(part))
    return parsed


def normalize_role(role: str) -> str:
    return ROLE_ALIASES.get(role, role)


def table_render_inputs(study: Any) -> list[dict[str, Any]]:
    if not isinstance(study, dict):
        return []
    values = study.get("table_render_inputs")
    return [item for item in values if isinstance(item, dict)] if isinstance(values, list) else []


def test_name_value(item: Any) -> str:
    if isinstance(item, dict):
        for key in ("role", "template_table_role", "test", "test_name", "name"):
            value = item.get(key)
            if value:
                return str(value)
        return ""
    return str(item or "")


def normalize_test_name(value: str) -> str:
    return re.sub(r"[\s_\-（）()／/]+", "", value).lower()


def active_stress_roles_from_fact_packet(packet: dict[str, Any]) -> set[str]:
    stress = packet.get("stress_study")
    roles: set[str] = set()
    for table in table_render_inputs(stress):
        role = str(table.get("role") or "").strip()
        if role in STRESS_ROLES:
            roles.add(role)
    if isinstance(stress, dict) and isinstance(stress.get("included_tests"), list):
        for item in stress["included_tests"]:
            raw = test_name_value(item)
            normalized = normalize_test_name(raw)
            if raw in STRESS_ROLES:
                roles.add(raw)
                continue
            for rule in STRESS_SECTION_RULES:
                for alias in rule["aliases"]:
                    alias_normalized = normalize_test_name(str(alias))
                    if alias_normalized and alias_normalized in normalized:
                        roles.update(rule["roles"])
    return roles


def active_stress_roles_from_manifest(table_render_manifest: dict[str, Any]) -> set[str]:
    roles: set[str] = set()
    for key in ("generated_table_slots", "tables_rendered"):
        values = table_render_manifest.get(key)
        if not isinstance(values, list):
            continue
        for item in values:
            if not isinstance(item, dict):
                continue
            role = str(item.get("role") or "").strip()
            if role in STRESS_ROLES:
                roles.add(role)
    return roles


def stress_section_checks(text: str, packet: dict[str, Any], table_render_manifest: dict[str, Any]) -> dict[str, Any]:
    fact_packet_roles = active_stress_roles_from_fact_packet(packet)
    manifest_roles = active_stress_roles_from_manifest(table_render_manifest)
    active_roles = fact_packet_roles or manifest_roles
    section_results: list[dict[str, Any]] = []
    issues: list[str] = []
    for rule in STRESS_SECTION_RULES:
        present = any(marker in text for marker in rule["heading_markers"])
        expected = bool(active_roles & rule["roles"])
        section_results.append(
            {
                "key": rule["key"],
                "roles": sorted(rule["roles"]),
                "expected": expected,
                "present": present,
            }
        )
        if expected and not present:
            issues.append(f"Expected stress-study section is missing: {rule['key']} ({sorted(rule['roles'])})")
        if not expected and present:
            issues.append(f"Inactive stress-study section remains in DOCX: {rule['key']} ({sorted(rule['roles'])})")
    return {
        "provided": bool(packet or table_render_manifest),
        "active_roles": sorted(active_roles),
        "fact_packet_roles": sorted(fact_packet_roles),
        "manifest_roles": sorted(manifest_roles),
        "sections": section_results,
        "issues": issues,
        "passed": not issues,
    }


def roles_compatible(table_role: str, slot_role: str) -> bool:
    if table_role == slot_role:
        return True
    return normalize_role(table_role) == normalize_role(slot_role) and slot_role == "stress"


def batch_no_from_entry(entry: Any) -> str:
    if isinstance(entry, dict):
        return str(entry.get("batch_no") or "").strip()
    return str(entry or "").strip()


def render_scope_action(scope: Any) -> str:
    if isinstance(scope, dict):
        return str(scope.get("batch_structure_action") or "")
    return ""


def count_placeholders(text: str, placeholders: list[str], allowlist: list[str]) -> dict[str, dict[str, int | bool]]:
    results = {}
    for placeholder in placeholders:
        raw_count = text.count(placeholder)
        allowed_count = sum(text.count(allowed) for allowed in allowlist if placeholder in allowed)
        unresolved_count = max(raw_count - allowed_count, 0)
        results[placeholder] = {
            "raw_count": raw_count,
            "allowed_count": allowed_count,
            "unresolved_count": unresolved_count,
            "passed": unresolved_count == 0,
        }
    return results


def section_presence(text: str, required_sections: list[str]) -> dict[str, bool]:
    selected = required_sections or list(DEFAULT_REQUIRED_SECTIONS)
    return {
        name: any(marker in text for marker in DEFAULT_REQUIRED_SECTIONS.get(name, [name]))
        for name in selected
    }


def table_samples(doc: Document, indexes: list[int], max_rows: int) -> list[dict[str, Any]]:
    samples: list[dict[str, Any]] = []
    for table_idx in indexes:
        if table_idx >= len(doc.tables):
            samples.append({"table_index": table_idx, "exists": False})
            continue
        table = doc.tables[table_idx]
        rows = []
        for row_idx, tr in enumerate(table._tbl.tr_lst[:max_rows]):
            cells = []
            for tc in tr.tc_lst:
                cells.append(
                    {
                        "text": cell_text(tc)[:120],
                        "gridSpan": grid_span(tc),
                        "vMerge": v_merge(tc),
                    }
                )
            rows.append({"row_index": row_idx, "physical_cells": len(tr.tc_lst), "cells": cells})
        samples.append(
            {
                "table_index": table_idx,
                "exists": True,
                "logical_rows": len(table.rows),
                "logical_cols": len(table.columns),
                "sample_rows": rows,
            }
        )
    return samples


def table_structure_summary(samples: list[dict[str, Any]]) -> dict[str, Any]:
    per_table = []
    for sample in samples:
        if not sample.get("exists"):
            per_table.append({"table_index": sample["table_index"], "exists": False, "has_gridSpan_gt_1": False, "has_vMerge": False})
            continue
        cells = [cell for row in sample["sample_rows"] for cell in row["cells"]]
        per_table.append(
            {
                "table_index": sample["table_index"],
                "exists": True,
                "has_gridSpan_gt_1": any(cell["gridSpan"] > 1 for cell in cells),
                "has_vMerge": any(bool(cell["vMerge"]) for cell in cells),
            }
        )
    return {
        "checked_table_indexes": [item["table_index"] for item in per_table],
        "missing_table_indexes": [item["table_index"] for item in per_table if not item["exists"]],
        "tables_without_gridSpan": [item["table_index"] for item in per_table if item["exists"] and not item["has_gridSpan_gt_1"]],
        "tables_without_vMerge": [item["table_index"] for item in per_table if item["exists"] and not item["has_vMerge"]],
        "per_table": per_table,
    }


def is_month_time_table(table) -> bool:
    if len(table._tbl.tr_lst) < 3:
        return False
    header_text = "\n".join(cell_text(tc) for tr in table._tbl.tr_lst[:2] for tc in tr.tc_lst)
    return "时间（月" in header_text


def month_data_start_col(table) -> int:
    for idx, tc in enumerate(table._tbl.tr_lst[0].tc_lst):
        if "时间（月" in cell_text(tc):
            return idx
    for idx, tc in enumerate(table._tbl.tr_lst[1].tc_lst):
        if re.match(r"^\s*\d+(?:\.\d+)?(?:\s|$)", cell_text(tc)):
            return idx
    return 4


def stability_marker_checks(doc: Document, text: str) -> dict[str, Any]:
    checked_tables: list[int] = []
    blank_cells: list[dict[str, int]] = []
    marker_note_issues: list[str] = []
    uses_dash = False
    uses_na = False

    for table_idx, table in enumerate(doc.tables):
        if not is_month_time_table(table):
            continue
        checked_tables.append(table_idx)
        data_start_col = month_data_start_col(table)
        for row_idx, tr in enumerate(table._tbl.tr_lst[2:], start=2):
            if len(tr.tc_lst) <= data_start_col:
                continue
            leading_cells = tr.tc_lst[:data_start_col]
            data_cells = tr.tc_lst[data_start_col:]
            if not any(cell_text(tc).strip() for tc in leading_cells):
                continue
            for col_offset, tc in enumerate(data_cells, start=data_start_col):
                value = cell_text(tc).strip()
                if value == "":
                    blank_cells.append({"table_index": table_idx, "row_index": row_idx, "col_index": col_offset})
                elif value == "---":
                    uses_dash = True
                elif value == "N/A":
                    uses_na = True

    if uses_dash and "---表示" not in text:
        marker_note_issues.append("Document uses --- in month-based stability tables but does not explain ---表示.")
    if uses_na and "N/A表示" not in text:
        marker_note_issues.append("Document uses N/A in month-based stability tables but does not explain N/A表示.")

    return {
        "checked_table_indexes": checked_tables,
        "blank_result_cells": blank_cells,
        "marker_note_issues": marker_note_issues,
        "uses_dash": uses_dash,
        "uses_na": uses_na,
        "passed": not blank_cells and not marker_note_issues,
    }


def paragraph_records(doc: Document) -> list[tuple[str, str]]:
    return [(p.style.name if p.style is not None else "", p.text.strip()) for p in doc.paragraphs if p.text.strip()]


def caption_reference_checks(doc: Document) -> dict[str, Any]:
    captions: dict[str, set[int]] = {"表": set(), "图": set()}
    references: dict[str, set[int]] = {"表": set(), "图": set()}
    body_reference_paragraphs: list[str] = []

    for style_name, text in paragraph_records(doc):
        is_directory_or_toc = style_name.lower() in {"toc 1", "toc 2", "toc 3", "table of figures"}
        caption_match = None if is_directory_or_toc else CAPTION_RE.match(text)
        if caption_match:
            captions[caption_match.group(1)].add(int(caption_match.group(2)))
        if is_directory_or_toc:
            continue
        for kind, number in REFERENCE_RE.findall(text):
            references[kind].add(int(number))
            if not caption_match:
                body_reference_paragraphs.append(text[:240])

    missing_referenced_captions = {
        kind: sorted(numbers - captions[kind])
        for kind, numbers in references.items()
    }
    duplicate_caption_numbers = find_duplicate_caption_numbers(doc)
    non_continuous_caption_numbers = {
        kind: [number for number in range(1, max(numbers) + 1) if number not in numbers] if numbers else []
        for kind, numbers in captions.items()
    }
    return {
        "caption_numbers": {kind: sorted(numbers) for kind, numbers in captions.items()},
        "reference_numbers": {kind: sorted(numbers) for kind, numbers in references.items()},
        "missing_referenced_captions": missing_referenced_captions,
        "duplicate_caption_numbers": duplicate_caption_numbers,
        "non_continuous_caption_numbers": non_continuous_caption_numbers,
        "body_reference_examples": body_reference_paragraphs[:20],
        "passed": not any(missing_referenced_captions.values()) and not any(duplicate_caption_numbers.values()) and not any(non_continuous_caption_numbers.values()),
    }


def find_duplicate_caption_numbers(doc: Document) -> dict[str, list[int]]:
    seen: dict[str, list[int]] = {"表": [], "图": []}
    duplicates: dict[str, set[int]] = {"表": set(), "图": set()}
    for style_name, text in paragraph_records(doc):
        if style_name.lower() in {"toc 1", "toc 2", "toc 3", "table of figures"}:
            continue
        match = CAPTION_RE.match(text)
        if not match:
            continue
        kind = match.group(1)
        number = int(match.group(2))
        if number in seen[kind]:
            duplicates[kind].add(number)
        seen[kind].append(number)
    return {kind: sorted(numbers) for kind, numbers in duplicates.items()}


def read_text_file(path: Path | None) -> str:
    if path is None or not path.exists():
        return ""
    return path.read_text(encoding="utf-8", errors="ignore")


def read_json_file(path: Path | None) -> dict[str, Any]:
    if path is None or not path.exists():
        return {}
    try:
        value = json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return {}
    return value if isinstance(value, dict) else {}


def table_signature(doc: Document, table_idx: int, max_rows: int) -> dict[str, Any]:
    if table_idx >= len(doc.tables):
        return {"exists": False}
    table = doc.tables[table_idx]
    return {
        "exists": True,
        "logical_rows": len(table.rows),
        "logical_cols": len(table.columns),
        "rows": [
            {
                "physical_cells": len(tr.tc_lst),
                "gridSpan": [grid_span(tc) for tc in tr.tc_lst],
                "vMerge": [v_merge(tc) for tc in tr.tc_lst],
            }
            for tr in table._tbl.tr_lst[:max_rows]
        ],
    }


def template_comparison(doc: Document, template_path: Path | None, indexes: list[int], max_rows: int) -> dict[str, Any]:
    if template_path is None:
        return {"provided": False, "passed": True, "table_differences": []}
    template_doc = Document(template_path)
    differences = []
    for table_idx in indexes:
        generated_signature = table_signature(doc, table_idx, max_rows)
        template_signature = table_signature(template_doc, table_idx, max_rows)
        if not generated_signature.get("exists") or not template_signature.get("exists"):
            differences.append(
                {
                    "table_index": table_idx,
                    "issue": "missing_table",
                    "generated_exists": generated_signature.get("exists", False),
                    "template_exists": template_signature.get("exists", False),
                }
            )
            continue
        generated_rows = generated_signature["rows"]
        template_rows = template_signature["rows"]
        rows_to_compare = min(len(generated_rows), len(template_rows), max_rows)
        for row_idx in range(rows_to_compare):
            generated_row = generated_rows[row_idx]
            template_row = template_rows[row_idx]
            if generated_row != template_row:
                differences.append(
                    {
                        "table_index": table_idx,
                        "row_index": row_idx,
                        "issue": "sampled_ooxml_signature_changed",
                        "generated": generated_row,
                        "template": template_row,
                    }
                )
    return {
        "provided": True,
        "template_docx": str(template_path),
        "checked_table_indexes": indexes,
        "table_differences": differences,
        "passed": not differences,
    }


def text_mentions_all(text: str, values: list[str]) -> dict[str, bool]:
    return {value: value in text for value in values}


def validation_report_checks(validation_report_text: str, warnings: list[str]) -> dict[str, Any]:
    if not validation_report_text:
        return {"provided": False, "warnings_explained": {}, "passed": not warnings}
    checks = {}
    for warning in warnings:
        tokens = [token for token in re.split(r"[\s:=,;，。；：\[\]<>]+", warning) if len(token) >= 3]
        checks[warning] = any(token in validation_report_text for token in tokens[:6])
    return {"provided": True, "warnings_explained": checks, "passed": all(checks.values())}


def split_expected_warnings(warnings: list[str], expected_patterns: list[str]) -> tuple[list[str], list[dict[str, str]]]:
    active_warnings: list[str] = []
    expected_warnings: list[dict[str, str]] = []
    patterns = [pattern for pattern in expected_patterns if pattern]
    for warning in warnings:
        matched_pattern = next((pattern for pattern in patterns if pattern in warning), "")
        if matched_pattern:
            expected_warnings.append({"warning": warning, "matched_pattern": matched_pattern})
        else:
            active_warnings.append(warning)
    return active_warnings, expected_warnings


def scaffold_checks(evidence_text: str, validation_report_text: str) -> dict[str, Any]:
    combined = f"{evidence_text}\n{validation_report_text}"
    markers = [
        "pending_fact_extraction",
        "pending extraction",
        "workspace initialized",
        "Initial scaffold validation",
        "not a final project validation report",
    ]
    found = [marker for marker in markers if marker in combined]
    return {"scaffold_markers": found, "passed": not found}


def evidence_checks(evidence_text: str, forbidden_sources: list[str], required_sources: list[str]) -> dict[str, Any]:
    if not evidence_text:
        return {
            "provided": False,
            "mentions_active_template": False,
            "mentions_excluded_sources": False,
            "required_sources_listed": text_mentions_all("", required_sources),
            "forbidden_sources_marked_excluded": {},
        }

    lowered = evidence_text.lower()
    forbidden_results = {}
    exclusion_markers = ["排除", "未使用", "禁止", "excluded", "not used"]
    for source in forbidden_sources:
        source_present = source in evidence_text
        exclusion_marker_present = any(marker in evidence_text for marker in exclusion_markers[:3]) or any(marker in lowered for marker in exclusion_markers[3:])
        forbidden_results[source] = bool(source_present and exclusion_marker_present)

    return {
        "provided": True,
        "mentions_active_template": "active template" in lowered or "模板" in evidence_text,
        "mentions_excluded_sources": any(marker in evidence_text for marker in exclusion_markers[:3]) or any(marker in lowered for marker in exclusion_markers[3:]),
        "required_sources_listed": text_mentions_all(evidence_text, required_sources),
        "forbidden_sources_marked_excluded": forbidden_results,
    }


def figure_caption_and_reference_sets(doc: Document) -> tuple[set[int], set[int], set[int]]:
    captions: set[int] = set()
    body_references: set[int] = set()
    directory_references: set[int] = set()
    for style_name, text in paragraph_records(doc):
        is_directory = style_name.lower() in {"toc 1", "toc 2", "toc 3", "table of figures"}
        caption_match = None if is_directory else CAPTION_RE.match(text)
        if caption_match and caption_match.group(1) == "图":
            captions.add(int(caption_match.group(2)))
            continue
        for kind, number in REFERENCE_RE.findall(text):
            if kind != "图":
                continue
            if is_directory:
                directory_references.add(int(number))
            else:
                body_references.add(int(number))
    return captions, body_references, directory_references


def trend_chart_manifest_checks(manifest_path: Path | None, doc: Document | None = None) -> dict[str, Any]:
    if manifest_path is None:
        return {"provided": False, "passed": True, "warnings": [], "missing_files": [], "chart_count": 0, "consistency_issues": []}
    if not manifest_path.exists():
        return {"provided": False, "passed": False, "warnings": [f"Trend chart manifest not found: {manifest_path}"], "missing_files": [], "chart_count": 0, "consistency_issues": []}

    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    chart_records = manifest.get("charts", [])
    warnings = list(manifest.get("warnings", []))
    missing_files = []
    figure_numbers: set[int] = set()
    for chart in chart_records:
        chart_path = Path(str(chart.get("path", "")))
        if not chart_path.exists() or chart_path.stat().st_size == 0:
            missing_files.append(str(chart_path))
        try:
            figure_numbers.add(int(chart.get("figure_no")))
        except (TypeError, ValueError):
            warnings.append(f"Trend chart record lacks an integer figure_no: {chart.get('title') or chart.get('path')}")
    consistency_issues: list[str] = []
    figure_captions: set[int] = set()
    figure_body_refs: set[int] = set()
    figure_directory_refs: set[int] = set()
    if doc is not None:
        figure_captions, figure_body_refs, figure_directory_refs = figure_caption_and_reference_sets(doc)
        for number in sorted(figure_numbers):
            if number not in figure_captions:
                consistency_issues.append(f"Trend chart figure_no has no matching DOCX figure caption: {number}")
            if number not in figure_directory_refs:
                consistency_issues.append(f"Trend chart figure_no has no matching DOCX figure directory entry: {number}")
    return {
        "provided": True,
        "manifest_path": str(manifest_path),
        "chart_count": len(chart_records),
        "figure_numbers": sorted(figure_numbers),
        "docx_figure_captions": sorted(figure_captions),
        "docx_figure_body_references": sorted(figure_body_refs),
        "docx_figure_directory_references": sorted(figure_directory_refs),
        "warnings": warnings,
        "missing_files": missing_files,
        "consistency_issues": consistency_issues,
        "passed": not warnings and not missing_files and not consistency_issues,
    }


def read_json_if_exists(path: Path | None) -> dict[str, Any]:
    if path is None or not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def table_caption_and_reference_sets(doc: Document) -> tuple[set[int], set[int], set[int]]:
    captions: set[int] = set()
    body_references: set[int] = set()
    directory_references: set[int] = set()
    for style_name, text in paragraph_records(doc):
        is_directory = style_name.lower() in {"toc 1", "toc 2", "toc 3", "table of figures"}
        caption_match = None if is_directory else CAPTION_RE.match(text)
        if caption_match and caption_match.group(1) == "表":
            captions.add(int(caption_match.group(2)))
            continue
        for kind, number in REFERENCE_RE.findall(text):
            if kind != "表":
                continue
            if is_directory:
                directory_references.add(int(number))
            else:
                body_references.add(int(number))
    return captions, body_references, directory_references


def slot_table_index(slot: dict[str, Any]) -> int | None:
    try:
        return int(slot.get("table_index"))
    except (TypeError, ValueError):
        return None


def generated_table_slots_from_manifest(manifest: dict[str, Any]) -> list[dict[str, Any]]:
    slots = manifest.get("generated_table_slots")
    return [slot for slot in slots if isinstance(slot, dict)] if isinstance(slots, list) else []


def rendered_batch_numbers_from_scope(scope: Any) -> list[str]:
    if not isinstance(scope, dict):
        return []
    rendered = scope.get("rendered_batches")
    if not isinstance(rendered, list):
        return []
    return [batch_no for item in rendered if (batch_no := batch_no_from_entry(item))]


def omitted_batch_numbers_from_scope(scope: Any) -> list[str]:
    if not isinstance(scope, dict):
        return []
    omitted = scope.get("omitted_batches")
    if not isinstance(omitted, list):
        return []
    return [batch_no for item in omitted if (batch_no := batch_no_from_entry(item))]


REPEATED_TEXT_PATTERNS = [
    re.compile(r"(.{1,12})\1{2,}"),
    re.compile(r"(检测项目|检测指标|时间（月）|时间（天）|时间（循环）)\1+"),
]


def merged_cell_artifact_checks(doc: Document, indexes: list[int]) -> dict[str, Any]:
    artifacts: list[dict[str, Any]] = []
    for table_idx in indexes:
        if table_idx >= len(doc.tables):
            continue
        table = doc.tables[table_idx]
        for row_idx, tr in enumerate(table._tbl.tr_lst):
            for cell_idx, tc in enumerate(tr.tc_lst):
                text = cell_text(tc).strip()
                if len(text) < 3:
                    continue
                if text in {"---", "N/A"}:
                    continue
                if any(pattern.search(text) for pattern in REPEATED_TEXT_PATTERNS):
                    artifacts.append(
                        {
                            "table_index": table_idx,
                            "row_index": row_idx,
                            "cell_index": cell_idx,
                            "text": text[:120],
                        }
                    )
    return {"artifacts": artifacts, "passed": not artifacts}


def table_render_manifest_checks(manifest_path: Path | None, doc: Document) -> dict[str, Any]:
    if manifest_path is None:
        return {"provided": False, "passed": True, "warnings": [], "tables_rendered": [], "stress_header_issues": [], "structure_consistency_issues": []}
    if not manifest_path.exists():
        return {"provided": False, "passed": False, "warnings": [f"Table render manifest not found: {manifest_path}"], "tables_rendered": [], "stress_header_issues": [], "structure_consistency_issues": []}

    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    warnings = list(manifest.get("warnings", []))
    table_issues: list[str] = []
    stress_header_issues: list[str] = []
    structure_consistency_issues: list[str] = []
    render_scope = manifest.get("render_scope")
    action = render_scope_action(render_scope)
    slots = generated_table_slots_from_manifest(manifest)
    slots_by_index = {slot_table_index(slot): slot for slot in slots if slot_table_index(slot) is not None}
    caption_numbers, body_reference_numbers, directory_reference_numbers = table_caption_and_reference_sets(doc)
    text = document_text(doc)

    if action in {"shrink", "expand"}:
        if not slots:
            structure_consistency_issues.append("Dynamic batch structure requires table-render-manifest.json generated_table_slots.")

    for batch in omitted_batch_numbers_from_scope(render_scope):
        if batch and batch in text:
            structure_consistency_issues.append(f"Omitted batch still appears in generated DOCX text: {batch}")

    for record in manifest.get("tables_rendered", []):
        table_idx = record.get("table_index")
        if not isinstance(table_idx, int) or table_idx >= len(doc.tables):
            table_issues.append(f"Rendered table index is missing from DOCX: {table_idx}")
            continue
        if slots:
            slot = slots_by_index.get(table_idx)
            if not slot:
                structure_consistency_issues.append(f"Rendered table index is not declared in table-render-manifest generated_table_slots: {table_idx}")
            else:
                record_role = str(record.get("role") or "")
                record_batch = str(record.get("batch_no") or "")
                slot_role = str(slot.get("role") or "")
                slot_batch = str(slot.get("batch_no") or "")
                if not roles_compatible(record_role, slot_role) or record_batch != slot_batch:
                    structure_consistency_issues.append(
                        f"Rendered table {table_idx} role/batch does not match generated table slot: "
                        f"{record_role}/{record_batch} vs {slot_role}/{slot_batch}"
                    )
                if not record.get("generated_slot"):
                    structure_consistency_issues.append(f"Rendered table {table_idx} does not record its generated_slot.")
                try:
                    caption_no = int(slot.get("caption_no"))
                except (TypeError, ValueError):
                    caption_no = None
                if caption_no is None:
                    structure_consistency_issues.append(f"Rendered table {table_idx} generated table slot does not declare caption_no.")
                else:
                    if caption_no not in caption_numbers:
                        structure_consistency_issues.append(f"Rendered table {table_idx} has no matching caption number: {caption_no}")
                    if caption_no not in body_reference_numbers and caption_no not in directory_reference_numbers:
                        structure_consistency_issues.append(f"Caption number {caption_no} has neither a body reference nor a table directory entry.")
        table = doc.tables[table_idx]
        time_points = [str(point) for point in record.get("time_points", [])]
        header_cells = table._tbl.tr_lst[1].tc_lst if len(table._tbl.tr_lst) > 1 else []
        if header_cells and time_points:
            header_texts = [cell_text(tc).strip() for tc in header_cells]
            for start in range(0, len(header_texts) - len(time_points) + 1):
                if header_texts[start : start + len(time_points)] == time_points:
                    break
            else:
                table_issues.append(f"Rendered table {table_idx} time headers differ from manifest: {header_texts} does not contain {time_points}")
        if record.get("normalized_role") == "stress":
            if record.get("data_row_count", 0) <= 0:
                table_issues.append(f"Rendered stress table {table_idx} has no data rows in manifest.")
            if len(table._tbl.tr_lst) < 2:
                stress_header_issues.append(f"Stress table {table_idx} has fewer than two header rows.")
                continue
            row0 = table._tbl.tr_lst[0].tc_lst
            row1 = table._tbl.tr_lst[1].tc_lst
            if not row0 or not row1:
                stress_header_issues.append(f"Stress table {table_idx} missing physical header cells.")
                continue
            left0 = row0[0]
            left1 = row1[0]
            if "检测指标" not in cell_text(left0).strip():
                stress_header_issues.append(f"Stress table {table_idx} row 0 col 0 header is not 检测指标: {cell_text(left0).strip()!r}")
            left_span = grid_span(left0)
            if left_span < 2 or v_merge(left0) != "restart":
                stress_header_issues.append(f"Stress table {table_idx} row 0 col 0 must preserve prototype merged header; got gridSpan={left_span} vMerge={v_merge(left0)!r}")
            if grid_span(left1) != left_span or v_merge(left1) != "continue":
                stress_header_issues.append(f"Stress table {table_idx} row 1 col 0 must continue prototype merged header; got gridSpan={grid_span(left1)} vMerge={v_merge(left1)!r}")
            if len(table._tbl.tr_lst) > 2:
                data_row = table._tbl.tr_lst[2].tc_lst
                if len(data_row) < 2:
                    stress_header_issues.append(f"Stress table {table_idx} first data row lacks the two-column left data area.")
                else:
                    if not cell_text(data_row[0]).strip():
                        stress_header_issues.append(f"Stress table {table_idx} first data row first column is empty; expected item plus method.")
                    if not cell_text(data_row[1]).strip():
                        stress_header_issues.append(f"Stress table {table_idx} first data row second column is empty; expected quality standard.")
    return {
        "provided": True,
        "manifest_path": str(manifest_path),
        "tables_rendered": manifest.get("tables_rendered", []),
        "generated_table_slots": slots,
        "render_scope": render_scope,
        "warnings": warnings,
        "table_issues": table_issues,
        "stress_header_issues": stress_header_issues,
        "structure_consistency_issues": structure_consistency_issues,
        "passed": not warnings and not table_issues and not stress_header_issues and not structure_consistency_issues,
    }


def rendered_table_indexes(table_render_manifest: dict[str, Any]) -> set[int]:
    indexes: set[int] = set()
    for record in table_render_manifest.get("tables_rendered", []):
        table_idx = record.get("table_index")
        if isinstance(table_idx, int):
            indexes.add(table_idx)
    return indexes


def build_report(args: argparse.Namespace) -> dict[str, Any]:
    doc = Document(args.docx)
    text = document_text(doc)
    body_text = body_paragraph_text(doc)
    packet = read_json_file(args.fact_packet)
    requested_table_indexes = parse_csv_ints(args.table_index)
    placeholders = count_placeholders(text, args.placeholder, args.allow_placeholder)
    marker_checks = stability_marker_checks(doc, text)
    evidence_text = read_text_file(args.evidence_summary)
    evidence = evidence_checks(evidence_text, args.forbidden_source, args.required_source)
    trend_charts = trend_chart_manifest_checks(args.trend_chart_manifest, doc)
    table_render_manifest = table_render_manifest_checks(args.table_render_manifest, doc)
    stress_sections = stress_section_checks(body_text, packet, table_render_manifest)
    rendered_indexes = rendered_table_indexes(table_render_manifest)
    table_indexes = requested_table_indexes or sorted(rendered_indexes) or DEFAULT_TABLE_INDEXES
    samples = table_samples(doc, table_indexes, args.max_sample_rows)
    table_structure = table_structure_summary(samples)
    merged_cell_artifacts = merged_cell_artifact_checks(doc, table_indexes)
    sections = section_presence(text, args.required_section)
    captions = caption_reference_checks(doc) if args.check_caption_references else {"passed": True}
    template = template_comparison(doc, args.template_docx, table_indexes, args.max_sample_rows)
    validation_report_text = read_text_file(args.validation_report)
    scaffold = scaffold_checks(evidence_text, validation_report_text)

    warnings = []
    for placeholder, result in placeholders.items():
        if result["unresolved_count"]:
            warnings.append(f"Unresolved placeholder remains: {placeholder}={result['unresolved_count']}")
    if args.disallow_hour_time and "时间（小时）" in text:
        warnings.append("Document still contains 时间（小时）; confirm oxidation or stress-study time unit against source evidence.")
    for section, present in sections.items():
        if not present:
            warnings.append(f"Required section marker not found: {section}")
    for batch in args.expected_batch:
        if batch not in text:
            warnings.append(f"Expected batch marker not found in generated DOCX: {batch}")
    if args.min_tables is not None and len(doc.tables) < args.min_tables:
        warnings.append(f"Document has fewer tables than expected: {len(doc.tables)} < {args.min_tables}")
    if args.max_tables is not None and len(doc.tables) > args.max_tables:
        warnings.append(f"Document has more tables than expected: {len(doc.tables)} > {args.max_tables}")
    for table_idx in table_structure["missing_table_indexes"]:
        warnings.append(f"Requested table index is missing: {table_idx}")
    for table_idx in table_structure["tables_without_gridSpan"]:
        warnings.append(f"Representative table lacks gridSpan in sampled rows: {table_idx}")
    for table_idx in table_structure["tables_without_vMerge"]:
        warnings.append(f"Representative table lacks vMerge in sampled rows: {table_idx}")
    if marker_checks["blank_result_cells"]:
        warnings.append(
            "Month-based stability result tables contain blank data cells; render source blanks as --- for future timepoints or N/A for completed timepoints without the test item."
        )
    for issue in marker_checks["marker_note_issues"]:
        warnings.append(issue)
    if args.evidence_summary and not evidence["mentions_active_template"]:
        warnings.append("Evidence summary does not clearly mention the active template.")
    for source, marked_excluded in evidence.get("forbidden_sources_marked_excluded", {}).items():
        if not marked_excluded:
            warnings.append(f"Evidence summary does not clearly mark forbidden source as excluded: {source}")
    for source, listed in evidence.get("required_sources_listed", {}).items():
        if not listed:
            warnings.append(f"Evidence summary does not list required source: {source}")
    for warning in trend_charts.get("warnings", []):
        warnings.append(f"Trend chart manifest warning: {warning}")
    for missing_file in trend_charts.get("missing_files", []):
        warnings.append(f"Trend chart file listed in manifest is missing or empty: {missing_file}")
    for issue in trend_charts.get("consistency_issues", []):
        warnings.append(f"Trend chart DOCX consistency issue: {issue}")
    if args.trend_chart_manifest and not trend_charts.get("provided"):
        warnings.append(f"Trend chart manifest not found: {args.trend_chart_manifest}")
    for artifact in merged_cell_artifacts["artifacts"]:
        warnings.append(
            f"Possible merged-cell logical write artifact at table {artifact['table_index']}, row {artifact['row_index']}, cell {artifact['cell_index']}: {artifact['text']}"
        )
    for warning in table_render_manifest.get("warnings", []):
        warnings.append(f"Table render manifest warning: {warning}")
    for issue in table_render_manifest.get("table_issues", []):
        warnings.append(f"Table render manifest issue: {issue}")
    for issue in table_render_manifest.get("stress_header_issues", []):
        warnings.append(f"Stress table header contract issue: {issue}")
    for issue in table_render_manifest.get("structure_consistency_issues", []):
        warnings.append(f"Dynamic batch structure issue: {issue}")
    if args.table_render_manifest and not table_render_manifest.get("provided"):
        warnings.append(f"Table render manifest not found: {args.table_render_manifest}")
    for issue in stress_sections.get("issues", []):
        warnings.append(issue)
    for difference in template.get("table_differences", []):
        table_idx = difference.get("table_index")
        if table_idx in rendered_indexes:
            continue
        row_idx = difference.get("row_index")
        if row_idx is None:
            warnings.append(f"Template table signature differs at table {table_idx}: {difference.get('issue')}")
        else:
            warnings.append(f"Template table signature differs at table {table_idx}, row {row_idx}: {difference.get('issue')}")
    if args.check_caption_references:
        for kind, numbers in captions.get("missing_referenced_captions", {}).items():
            if numbers:
                warnings.append(f"Referenced {kind} caption numbers are missing: {numbers}")
        for kind, numbers in captions.get("duplicate_caption_numbers", {}).items():
            if numbers:
                warnings.append(f"Duplicate {kind} caption numbers found: {numbers}")
        for kind, numbers in captions.get("non_continuous_caption_numbers", {}).items():
            if numbers:
                warnings.append(f"{kind} caption numbers are not continuous; missing: {numbers}")

    for marker in scaffold["scaffold_markers"]:
        warnings.append(f"Output still contains scaffold-only marker and is not final: {marker}")

    validation_report = validation_report_checks(validation_report_text, warnings)
    if args.validation_report and not validation_report["passed"]:
        warnings.append("Validation report does not explain every validation warning.")

    active_warnings, expected_warnings = split_expected_warnings(warnings, args.expected_warning)

    return {
        "docx": str(args.docx),
        "summary": {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "inline_shapes": len(doc.inline_shapes),
        },
        "placeholder_counts": placeholders,
        "section_presence": sections,
        "table_structure": table_structure,
        "stability_table_marker_semantics": marker_checks,
        "table_samples": samples,
        "caption_reference_checks": captions,
        "template_comparison": template,
        "evidence_summary_checks": evidence,
        "trend_chart_manifest_checks": trend_charts,
        "table_render_manifest_checks": table_render_manifest,
        "stress_section_checks": stress_sections,
        "merged_cell_artifact_checks": merged_cell_artifacts,
        "scaffold_checks": scaffold,
        "validation_report_checks": validation_report,
        "warnings": active_warnings,
        "expected_warnings": expected_warnings,
        "all_warnings": warnings,
        "passed": not active_warnings,
    }


def print_human_report(report: dict[str, Any]) -> None:
    print(f"DOCX={report['docx']}")
    print("SUMMARY")
    for key, value in report["summary"].items():
        print(f"  {key}={value}")
    print("PLACEHOLDERS")
    for key, value in report["placeholder_counts"].items():
        print(f"  {key}: unresolved={value['unresolved_count']} raw={value['raw_count']} allowed={value['allowed_count']}")
    print("SECTIONS")
    for key, value in report["section_presence"].items():
        print(f"  {key}={'PASS' if value else 'WARN'}")
    print("TABLE_STRUCTURE")
    for key in ["checked_table_indexes", "missing_table_indexes", "tables_without_gridSpan", "tables_without_vMerge"]:
        print(f"  {key}={report['table_structure'][key]}")
    print("STABILITY_TABLE_MARKERS")
    marker_checks = report["stability_table_marker_semantics"]
    print(f"  checked_table_indexes={marker_checks.get('checked_table_indexes')}")
    print(f"  blank_result_cells={len(marker_checks.get('blank_result_cells', []))}")
    print(f"  marker_note_issues={marker_checks.get('marker_note_issues')}")
    print("CAPTION_REFERENCES")
    print(f"  passed={report['caption_reference_checks'].get('passed')}")
    print(f"  missing_referenced_captions={report['caption_reference_checks'].get('missing_referenced_captions')}")
    print(f"  duplicate_caption_numbers={report['caption_reference_checks'].get('duplicate_caption_numbers')}")
    print("TEMPLATE_COMPARISON")
    print(f"  provided={report['template_comparison'].get('provided')}")
    print(f"  passed={report['template_comparison'].get('passed')}")
    print(f"  table_differences={len(report['template_comparison'].get('table_differences', []))}")
    print("EVIDENCE_SUMMARY")
    for key, value in report["evidence_summary_checks"].items():
        print(f"  {key}={value}")
    print("TREND_CHART_MANIFEST")
    print(f"  provided={report['trend_chart_manifest_checks'].get('provided')}")
    print(f"  passed={report['trend_chart_manifest_checks'].get('passed')}")
    print(f"  chart_count={report['trend_chart_manifest_checks'].get('chart_count')}")
    print(f"  warnings={report['trend_chart_manifest_checks'].get('warnings')}")
    print(f"  missing_files={report['trend_chart_manifest_checks'].get('missing_files')}")
    print(f"  consistency_issues={report['trend_chart_manifest_checks'].get('consistency_issues')}")
    print("TABLE_RENDER_MANIFEST")
    print(f"  provided={report['table_render_manifest_checks'].get('provided')}")
    print(f"  passed={report['table_render_manifest_checks'].get('passed')}")
    print(f"  tables_rendered={len(report['table_render_manifest_checks'].get('tables_rendered', []))}")
    print(f"  warnings={report['table_render_manifest_checks'].get('warnings')}")
    print(f"  table_issues={report['table_render_manifest_checks'].get('table_issues')}")
    print(f"  stress_header_issues={report['table_render_manifest_checks'].get('stress_header_issues')}")
    print(f"  structure_consistency_issues={report['table_render_manifest_checks'].get('structure_consistency_issues')}")
    print("STRESS_SECTIONS")
    print(f"  active_roles={report['stress_section_checks'].get('active_roles')}")
    print(f"  issues={report['stress_section_checks'].get('issues')}")
    print("MERGED_CELL_ARTIFACTS")
    print(f"  passed={report['merged_cell_artifact_checks'].get('passed')}")
    print(f"  artifacts={len(report['merged_cell_artifact_checks'].get('artifacts', []))}")
    print("SCAFFOLD_CHECK")
    print(f"  passed={report['scaffold_checks'].get('passed')}")
    print(f"  scaffold_markers={report['scaffold_checks'].get('scaffold_markers')}")
    print("VALIDATION_REPORT")
    print(f"  provided={report['validation_report_checks'].get('provided')}")
    print(f"  passed={report['validation_report_checks'].get('passed')}")
    print("WARNINGS")
    if report["warnings"]:
        for warning in report["warnings"]:
            print(f"  - {warning}")
    else:
        print("  none")
    print("EXPECTED_WARNINGS")
    if report.get("expected_warnings"):
        for item in report["expected_warnings"]:
            print(f"  - {item['warning']} (matched {item['matched_pattern']!r})")
    else:
        print("  none")
    print(f"PASSED={'true' if report['passed'] else 'false'}")


def main() -> None:
    require_internal_context("validate_generated_docx.py")
    parser = argparse.ArgumentParser(description="Validate generated CTD 3.2.S.7.3 DOCX output.")
    parser.add_argument("docx", type=Path)
    parser.add_argument("--evidence-summary", type=Path)
    parser.add_argument("--validation-report", type=Path)
    parser.add_argument("--template-docx", type=Path)
    parser.add_argument("--fact-packet", type=Path)
    parser.add_argument("--trend-chart-manifest", type=Path)
    parser.add_argument("--table-render-manifest", type=Path)
    parser.add_argument("--expected-batch", action="append", default=[])
    parser.add_argument("--required-source", action="append", default=[])
    parser.add_argument("--forbidden-source", action="append", default=[])
    parser.add_argument("--placeholder", action="append", default=DEFAULT_PLACEHOLDERS)
    parser.add_argument("--allow-placeholder", action="append", default=[])
    parser.add_argument("--required-section", action="append", default=[])
    parser.add_argument("--table-index", action="append", default=[])
    parser.add_argument("--max-sample-rows", type=int, default=2)
    parser.add_argument("--min-tables", type=int)
    parser.add_argument("--max-tables", type=int)
    parser.add_argument("--check-caption-references", action="store_true")
    parser.add_argument("--disallow-hour-time", action="store_true")
    parser.add_argument(
        "--expected-warning",
        action="append",
        default=[],
        help="Substring of a validation warning that is expected and explained in the validation report.",
    )
    parser.add_argument("--json-out", type=Path)
    args = parser.parse_args()

    report = build_report(args)
    print_human_report(report)

    if args.json_out:
        args.json_out.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")


# Internal module — do not invoke directly. Use skill_step.py as the public entry point.
if __name__ == "__main__":
    main()

from __future__ import annotations

import argparse
import json
import os
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document

from runtime_guard import require_internal_context


SOURCE_INDEX_SCHEMA = "ctd-32s73-source-index-v1"
DOCX_READER = {
    "kind": "builtin_source_indexer",
    "name": "python-docx",
    "purpose": "docx paragraph/table preview for fact localization",
}
XLSX_READER = {
    "kind": "builtin_source_indexer",
    "name": "openpyxl",
    "purpose": "xlsx workbook/sheet/cell preview for fact localization",
}
TEXT_SUFFIXES = {".txt", ".md", ".csv", ".tsv", ".json"}
DOCX_SUFFIXES = {".docx"}
XLSX_SUFFIXES = {".xlsx", ".xlsm"}
PDF_SUFFIXES = {".pdf"}
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}
INDEXABLE_SUFFIXES = TEXT_SUFFIXES | DOCX_SUFFIXES | XLSX_SUFFIXES | PDF_SUFFIXES | IMAGE_SUFFIXES


def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def write_json(path: Path, data: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def source_path(source: Any) -> str:
    if isinstance(source, dict):
        return str(source.get("path") or source.get("file") or "")
    return str(source)


def user_path(value: Any) -> Path:
    text = str(value)
    path = Path(text)
    if text == "~" or text.startswith("~/") or text.startswith("~\\"):
        return path.expanduser()
    return path


def state_sources(state: dict[str, Any], key: str) -> list[str]:
    sources = state.get("sources")
    if not isinstance(sources, dict):
        return []
    values = sources.get(key)
    if not isinstance(values, list):
        return []
    return [path for item in values if (path := source_path(item))]


def resolved_key(path: Path) -> str:
    try:
        return str(path.resolve()).casefold()
    except OSError:
        return str(path.absolute()).casefold()


def is_under(path: Path, parent: Path) -> bool:
    try:
        path.resolve().relative_to(parent.resolve())
        return True
    except (ValueError, OSError):
        return False


def default_forbidden(path: Path) -> bool:
    name = path.name
    return name.startswith("已完成的申报资料") or name.startswith("~$")


def classify(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in DOCX_SUFFIXES:
        return "docx"
    if suffix in XLSX_SUFFIXES:
        return "xlsx"
    if suffix in PDF_SUFFIXES:
        return "pdf"
    if suffix in IMAGE_SUFFIXES:
        return "image"
    if suffix in TEXT_SUFFIXES:
        return "text"
    return "other"


def file_modified_at(path: Path) -> str | None:
    try:
        return datetime.fromtimestamp(path.stat().st_mtime, timezone.utc).isoformat()
    except OSError:
        return None


def common_record(path: Path) -> dict[str, Any]:
    try:
        size = path.stat().st_size
    except OSError:
        size = None
    return {
        "path": str(path),
        "name": path.name,
        "suffix": path.suffix.lower(),
        "type": classify(path),
        "size_bytes": size,
        "modified_at": file_modified_at(path),
    }


def preview_text(value: Any, limit: int = 240) -> str:
    text = "" if value is None else str(value)
    normalized = " ".join(text.replace("\u3000", " ").split())
    if len(normalized) <= limit:
        return normalized
    return normalized[: limit - 3] + "..."


def json_cell(value: Any) -> Any:
    if value is None or isinstance(value, (bool, int, float)):
        return value
    if hasattr(value, "isoformat"):
        try:
            return value.isoformat()
        except TypeError:
            pass
    return preview_text(value, 200)


def index_docx(path: Path, max_paragraphs: int, max_table_rows: int, max_table_cols: int) -> dict[str, Any]:
    doc = Document(path)
    paragraphs = []
    for idx, paragraph in enumerate(doc.paragraphs):
        text = preview_text(paragraph.text)
        if text:
            paragraphs.append(
                {
                    "paragraph_index": idx,
                    "style": paragraph.style.name if paragraph.style else None,
                    "text": text,
                }
            )
        if len(paragraphs) >= max_paragraphs:
            break

    tables = []
    for table_idx, table in enumerate(doc.tables):
        rows_preview = []
        for row_idx, row in enumerate(list(table.rows)[:max_table_rows]):
            cells = [preview_text(cell.text, 120) for cell in list(row.cells)[:max_table_cols]]
            rows_preview.append({"row_index": row_idx, "cells": cells})
        tables.append(
            {
                "table_index": table_idx,
                "row_count": len(table.rows),
                "column_count": len(table.columns),
                "preview_rows": rows_preview,
            }
        )

    return {
        **common_record(path),
        "status": "indexed",
        "reader": DOCX_READER,
        "paragraph_count": len(doc.paragraphs),
        "nonempty_paragraph_preview_count": len(paragraphs),
        "paragraphs_truncated": len(paragraphs) >= max_paragraphs,
        "paragraphs_preview": paragraphs,
        "table_count": len(doc.tables),
        "tables": tables,
        "inline_shape_count": len(doc.inline_shapes),
    }


def index_xlsx(path: Path, max_excel_rows: int, max_excel_cols: int) -> dict[str, Any]:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return {
            **common_record(path),
            "status": "dependency_missing",
            "reader": XLSX_READER,
            "warning": "openpyxl is required to index Excel workbooks.",
        }

    workbook = load_workbook(path, read_only=True, data_only=True)
    try:
        sheets = []
        for sheet_name in workbook.sheetnames:
            worksheet = workbook[sheet_name]
            preview_rows = []
            for row_idx, row in enumerate(
                worksheet.iter_rows(
                    min_row=1,
                    max_row=min(worksheet.max_row or 0, max_excel_rows),
                    max_col=min(worksheet.max_column or 0, max_excel_cols),
                    values_only=True,
                ),
                start=1,
            ):
                values = [json_cell(value) for value in row]
                if any(value not in (None, "") for value in values):
                    preview_rows.append({"row_index": row_idx, "values": values})
            sheets.append(
                {
                    "name": sheet_name,
                    "max_row": worksheet.max_row,
                    "max_column": worksheet.max_column,
                    "preview_rows": preview_rows,
                    "preview_rows_truncated": bool((worksheet.max_row or 0) > max_excel_rows),
                    "preview_columns_truncated": bool((worksheet.max_column or 0) > max_excel_cols),
                }
            )
    finally:
        workbook.close()

    return {
        **common_record(path),
        "status": "indexed",
        "reader": XLSX_READER,
        "sheet_count": len(sheets),
        "sheets": sheets,
    }


def find_ocr_artifacts(output_dir: Path, source: Path) -> list[dict[str, Any]]:
    artifacts = []
    if not output_dir.exists():
        return artifacts
    source_name = source.name.casefold()
    source_resolved = resolved_key(source)
    for manifest_path in output_dir.rglob("*ocr-manifest*.json"):
        try:
            manifest = read_json(manifest_path)
        except (json.JSONDecodeError, OSError):
            continue
        manifest_source = str(manifest.get("source") or "")
        source_matches = (
            manifest_source.casefold() == str(source).casefold()
            or Path(manifest_source).name.casefold() == source_name
            or resolved_key(Path(manifest_source)) == source_resolved
        )
        if not source_matches:
            continue
        combined = manifest.get("combined_markdown")
        combined_path = Path(combined) if combined else None
        if combined_path and not combined_path.is_absolute():
            combined_path = (manifest_path.parent / combined_path).resolve()
        preview = None
        if combined_path and combined_path.exists():
            preview = preview_text(combined_path.read_text(encoding="utf-8", errors="replace"), 2000)
        artifacts.append(
            {
                "manifest_path": str(manifest_path),
                "combined_markdown": str(combined_path) if combined_path else combined,
                "page_count": manifest.get("page_count") or manifest.get("extract_progress", {}).get("totalPages"),
                "markdown_preview": preview,
            }
        )
    return artifacts


def index_pdf_or_image(path: Path, output_dir: Path) -> dict[str, Any]:
    ocr_artifacts = find_ocr_artifacts(output_dir, path)
    if ocr_artifacts:
        status = "indexed_existing_ocr"
        ocr_required = False
    else:
        status = "ocr_required"
        ocr_required = True
    return {
        **common_record(path),
        "status": status,
        "ocr_required": ocr_required,
        "existing_ocr_artifacts": ocr_artifacts,
        "paddleocr_token_available": bool(os.environ.get("PADDLEOCR_TOKEN")),
        "note": "Source index does not submit OCR jobs. Use project-approved OCR evidence before writing fact-packet.json when needed.",
    }


def index_text(path: Path) -> dict[str, Any]:
    for encoding in ("utf-8", "utf-8-sig", "gb18030"):
        try:
            text = path.read_text(encoding=encoding)
            return {
                **common_record(path),
                "status": "indexed",
                "encoding": encoding,
                "line_count_preview": len(text.splitlines()),
                "text_preview": preview_text(text, 3000),
            }
        except UnicodeDecodeError:
            continue
    return {**common_record(path), "status": "unreadable_text_encoding"}


def index_source(path: Path, output_dir: Path, args: argparse.Namespace) -> dict[str, Any]:
    if not path.exists():
        return {
            "path": str(path),
            "name": path.name,
            "suffix": path.suffix.lower(),
            "type": classify(path),
            "status": "missing",
        }
    try:
        source_type = classify(path)
        if source_type == "docx":
            return index_docx(path, args.max_paragraphs, args.max_docx_table_rows, args.max_docx_table_cols)
        if source_type == "xlsx":
            return index_xlsx(path, args.max_excel_rows, args.max_excel_cols)
        if source_type in {"pdf", "image"}:
            return index_pdf_or_image(path, output_dir)
        if source_type == "text":
            return index_text(path)
        return {**common_record(path), "status": "not_indexed_unsupported_type"}
    except Exception as exc:  # noqa: BLE001 - index should keep going and record per-source failure.
        return {**common_record(path), "status": "index_error", "error": f"{type(exc).__name__}: {exc}"}


def expand_allowed_sources(
    allowed_values: list[str],
    forbidden_values: list[str],
    output_dir: Path,
    max_files: int,
) -> tuple[list[Path], list[dict[str, Any]], list[str]]:
    warnings: list[str] = []
    discovered: list[Path] = []
    seen: set[str] = set()
    forbidden_records: list[dict[str, Any]] = []
    forbidden_keys = set()

    for value in forbidden_values:
        path = user_path(value)
        key = resolved_key(path)
        forbidden_keys.add(key)
        forbidden_records.append({"path": str(path), "status": "excluded", "reason": "registered forbidden source"})

    def add_forbidden(path: Path, reason: str) -> None:
        key = resolved_key(path)
        if key in forbidden_keys:
            return
        forbidden_keys.add(key)
        forbidden_records.append({"path": str(path), "status": "excluded", "reason": reason})

    def add_source(path: Path) -> None:
        key = resolved_key(path)
        if key in seen:
            return
        if key in forbidden_keys:
            return
        if default_forbidden(path):
            add_forbidden(path, "default forbidden source pattern")
            return
        if len(discovered) >= max_files:
            warnings.append(f"Source index reached max_files={max_files}; remaining files were skipped.")
            return
        seen.add(key)
        discovered.append(path)

    for value in allowed_values:
        path = user_path(value)
        if default_forbidden(path):
            add_forbidden(path, "default forbidden source pattern")
            continue
        if not path.exists():
            add_source(path)
            continue
        if path.is_file():
            add_source(path)
            continue
        if path.is_dir():
            for child in sorted(path.rglob("*")):
                if not child.is_file():
                    continue
                if is_under(child, output_dir):
                    continue
                if child.suffix.lower() not in INDEXABLE_SUFFIXES:
                    continue
                add_source(child)
            continue
        warnings.append(f"Source is neither file nor directory: {path}")

    return discovered, forbidden_records, warnings


def build_source_index(args: argparse.Namespace) -> dict[str, Any]:
    output_dir = args.output_dir.resolve()
    state_path = args.workflow_state.resolve()
    state = read_json(state_path) if state_path.exists() else {}
    allowed_values = state_sources(state, "allowed") + list(args.source)
    forbidden_values = state_sources(state, "forbidden") + list(args.forbidden_source)
    discovered_sources, forbidden_records, warnings = expand_allowed_sources(
        allowed_values,
        forbidden_values,
        output_dir,
        args.max_files,
    )
    indexed_sources = [index_source(path, output_dir, args) for path in discovered_sources]
    for record in indexed_sources:
        if record.get("status") in {"dependency_missing", "index_error", "missing", "ocr_required"}:
            warnings.append(f"{record.get('name')}: {record.get('status')}")

    return {
        "schema_version": SOURCE_INDEX_SCHEMA,
        "generated_at": now_iso(),
        "workflow_state": str(state_path),
        "output_dir": str(output_dir),
        "source_count": len(indexed_sources),
        "forbidden_source_count": len(forbidden_records),
        "sources": indexed_sources,
        "forbidden_sources": forbidden_records,
        "warnings": warnings,
    }


def markdown_table_row(values: list[Any]) -> str:
    escaped = [preview_text(value, 80).replace("|", "\\|") for value in values]
    return "| " + " | ".join(escaped) + " |"


def source_index_markdown(index: dict[str, Any]) -> str:
    lines = [
        "# Source Index",
        "",
        f"- Schema: {index['schema_version']}",
        f"- Generated at: {index['generated_at']}",
        f"- Workflow state: {index['workflow_state']}",
        f"- Indexed sources: {index['source_count']}",
        f"- Excluded sources: {index['forbidden_source_count']}",
        "",
    ]
    warnings = index.get("warnings") or []
    if warnings:
        lines.extend(["## Warnings", ""])
        lines.extend(f"- {warning}" for warning in warnings)
        lines.append("")

    forbidden = index.get("forbidden_sources") or []
    if forbidden:
        lines.extend(["## Excluded Sources", ""])
        for item in forbidden:
            lines.append(f"- `{item.get('path')}`: {item.get('reason')}")
        lines.append("")

    lines.extend(["## Indexed Sources", ""])
    for idx, source in enumerate(index.get("sources") or [], start=1):
        lines.extend(
            [
                f"### {idx}. {source.get('name')}",
                "",
                f"- Path: `{source.get('path')}`",
                f"- Type: `{source.get('type')}`",
                f"- Status: `{source.get('status')}`",
                "",
            ]
        )
        reader = source.get("reader")
        if isinstance(reader, dict):
            lines.extend(
                [
                    f"- Reader: `{reader.get('kind')}/{reader.get('name')}`",
                    f"- Reader purpose: {reader.get('purpose')}",
                    "",
                ]
            )
        if source.get("type") == "docx":
            lines.extend(
                [
                    f"- Paragraphs: {source.get('paragraph_count')}",
                    f"- Tables: {source.get('table_count')}",
                    "",
                    "#### Paragraph Preview",
                    "",
                ]
            )
            for paragraph in source.get("paragraphs_preview", [])[:30]:
                lines.append(
                    f"- [{paragraph.get('paragraph_index')}] {paragraph.get('style')}: "
                    f"{paragraph.get('text')}"
                )
            lines.extend(["", "#### Table Preview", ""])
            for table in source.get("tables", [])[:20]:
                lines.append(
                    f"- Table {table.get('table_index')}: "
                    f"{table.get('row_count')} rows x {table.get('column_count')} columns"
                )
                preview_rows = table.get("preview_rows") or []
                if preview_rows:
                    lines.append("")
                    lines.append(markdown_table_row(["row", "cells"]))
                    lines.append("| --- | --- |")
                    for row in preview_rows:
                        lines.append(markdown_table_row([row.get("row_index"), " / ".join(row.get("cells") or [])]))
                    lines.append("")
        elif source.get("type") == "xlsx":
            for sheet in source.get("sheets", []):
                lines.extend(
                    [
                        f"#### Sheet: {sheet.get('name')}",
                        "",
                        f"- Size: {sheet.get('max_row')} rows x {sheet.get('max_column')} columns",
                        "",
                    ]
                )
                preview_rows = sheet.get("preview_rows") or []
                if preview_rows:
                    lines.append(markdown_table_row(["row", "values"]))
                    lines.append("| --- | --- |")
                    for row in preview_rows[:20]:
                        lines.append(markdown_table_row([row.get("row_index"), " / ".join(preview_text(v, 60) for v in row.get("values") or [])]))
                    lines.append("")
        elif source.get("type") in {"pdf", "image"}:
            lines.append(f"- OCR required: {source.get('ocr_required')}")
            artifacts = source.get("existing_ocr_artifacts") or []
            if artifacts:
                lines.append("- Existing OCR artifacts:")
                for artifact in artifacts:
                    lines.append(f"  - `{artifact.get('manifest_path')}` pages={artifact.get('page_count')}")
                    if artifact.get("markdown_preview"):
                        lines.append(f"    - Preview: {artifact.get('markdown_preview')}")
            lines.append("")
        elif source.get("type") == "text":
            lines.extend(["#### Text Preview", "", source.get("text_preview") or "", ""])
        elif source.get("warning"):
            lines.extend([f"- Warning: {source.get('warning')}", ""])
        elif source.get("error"):
            lines.extend([f"- Error: {source.get('error')}", ""])

    return "\n".join(lines).rstrip() + "\n"


def main() -> None:
    require_internal_context("extract_source_index.py")
    parser = argparse.ArgumentParser(description="Build a source index for CTD 3.2.S.7.3 fact extraction.")
    parser.add_argument("--workflow-state", type=Path, required=True)
    parser.add_argument("--output-dir", type=Path, required=True)
    parser.add_argument("--source", action="append", default=[])
    parser.add_argument("--forbidden-source", action="append", default=[])
    parser.add_argument("--json-out", type=Path)
    parser.add_argument("--markdown-out", type=Path)
    parser.add_argument("--max-files", type=int, default=500)
    parser.add_argument("--max-paragraphs", type=int, default=120)
    parser.add_argument("--max-docx-table-rows", type=int, default=6)
    parser.add_argument("--max-docx-table-cols", type=int, default=8)
    parser.add_argument("--max-excel-rows", type=int, default=30)
    parser.add_argument("--max-excel-cols", type=int, default=16)
    args = parser.parse_args()

    output_dir = args.output_dir.resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    json_out = args.json_out or (output_dir / "source-index.json")
    markdown_out = args.markdown_out or (output_dir / "source-index.md")
    index = build_source_index(args)
    write_json(json_out, index)
    markdown_out.parent.mkdir(parents=True, exist_ok=True)
    markdown_out.write_text(source_index_markdown(index), encoding="utf-8")
    print(json.dumps({"source_index": str(json_out), "source_index_markdown": str(markdown_out)}, ensure_ascii=False, indent=2))


# Internal module — do not invoke directly. Use skill_step.py as the public entry point.
if __name__ == "__main__":
    main()

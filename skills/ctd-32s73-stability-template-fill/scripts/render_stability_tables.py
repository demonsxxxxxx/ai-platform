from __future__ import annotations

import argparse
import json
import re
from collections import defaultdict
from copy import deepcopy
from pathlib import Path
from typing import Any

from runtime_guard import require_internal_context

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell

SKILL_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_TABLE_PROTOTYPE = SKILL_ROOT / "assets" / "table-prototypes" / "ctd-32s73-table-prototypes.docx"
TABLE_DIRECTORY_PLACEHOLDER = "{{TABLE_DIRECTORY}}"
CAPTION_PREFIX_RE = re.compile(r"^表\s*3\.2\.S\.7\.3-\s*\d+\s*")


ROLE_PROFILES: dict[str, dict[str, Any]] = {
    "long_term": {
        "kind": "stability",
        "seed_table_indexes": [1, 2],
        "header_rows": 2,
        "time_label_row": 0,
        "time_points_row": 1,
        "time_start_col": 4,
        "time_label_col": 3,
        "time_points_col": 3,
        "data_start_row": 2,
        "data_template_row": 2,
        "left_columns": ["group", "item", "method", "acceptance"],
        "vmerge_columns": [0, 1, 2],
    },
    "accelerated": {
        "kind": "stability",
        "seed_table_indexes": [3, 4],
        "header_rows": 2,
        "time_label_row": 0,
        "time_points_row": 1,
        "time_start_col": 4,
        "time_label_col": 3,
        "time_points_col": 3,
        "data_start_row": 2,
        "data_template_row": 2,
        "left_columns": ["group", "item", "method", "acceptance"],
        "vmerge_columns": [0, 1, 2],
    },
    "stress": {
        "kind": "stress",
        "seed_table_indexes": [5, 6, 7, 8, 9, 10, 11, 12, 13, 14, 15, 16, 17, 18],
        "header_rows": 2,
        "time_label_row": 0,
        "time_points_row": 1,
        "time_start_col": 2,
        "time_label_col": 1,
        "time_points_col": 1,
        "data_start_row": 2,
        "data_template_row": 2,
        "left_columns": ["indicator_method", "acceptance"],
        "vmerge_columns": [],
    },
}

ROLE_ALIASES = {
    "light_stress": "stress",
    "agitation_stress": "stress",
    "freeze_thaw_stress": "stress",
    "high_temperature_stress": "stress",
    "low_ph_stress": "stress",
    "high_ph_stress": "stress",
    "oxidation_stress": "stress",
}

ROLE_SEED_TABLE_INDEXES = {
    "long_term": [1, 2],
    "accelerated": [3, 4],
    "light_stress": [5, 6],
    "agitation_stress": [7, 8],
    "freeze_thaw_stress": [9, 10],
    "high_temperature_stress": [11, 12],
    "low_ph_stress": [13, 14],
    "high_ph_stress": [15, 16],
    "oxidation_stress": [17, 18],
    "stress": ROLE_PROFILES["stress"]["seed_table_indexes"],
}

TABLE_BLOCK_PLACEHOLDERS = {
    "long_term": "{{LONG_TERM_TABLE_BLOCK}}",
    "accelerated": "{{ACCELERATED_TABLE_BLOCK}}",
    "light_stress": "{{STRESS_LIGHT_TABLE_BLOCK}}",
    "agitation_stress": "{{STRESS_AGITATION_TABLE_BLOCK}}",
    "freeze_thaw_stress": "{{STRESS_FREEZE_THAW_TABLE_BLOCK}}",
    "high_temperature_stress": "{{STRESS_HIGH_TEMPERATURE_TABLE_BLOCK}}",
    "low_ph_stress": "{{STRESS_LOW_PH_TABLE_BLOCK}}",
    "high_ph_stress": "{{STRESS_HIGH_PH_TABLE_BLOCK}}",
    "oxidation_stress": "{{STRESS_OXIDATION_TABLE_BLOCK}}",
}

STRESS_ROLE_TITLES = {
    "light_stress": "光照试验结果",
    "agitation_stress": "振荡试验结果",
    "freeze_thaw_stress": "冻融试验结果",
    "high_temperature_stress": "高温试验结果",
    "low_ph_stress": "低pH试验结果",
    "high_ph_stress": "高pH试验结果",
    "oxidation_stress": "氧化试验结果",
}


def normalize_role(role: str) -> str:
    return ROLE_ALIASES.get(role, role)


def role_seed_table_indexes(role: str) -> list[int]:
    return list(ROLE_SEED_TABLE_INDEXES.get(role, ROLE_PROFILES[normalize_role(role)]["seed_table_indexes"]))


def roles_compatible(table_role: str, slot_role: str) -> bool:
    if table_role == slot_role:
        return True
    return normalize_role(table_role) == normalize_role(slot_role) and slot_role == "stress"


def require_profile(role: str) -> dict[str, Any]:
    normalized = normalize_role(role)
    if normalized not in ROLE_PROFILES:
        raise ValueError(f"Unsupported table role: {role}")
    return ROLE_PROFILES[normalized]


def table_block_placeholder(role: str) -> str:
    if role not in TABLE_BLOCK_PLACEHOLDERS:
        raise ValueError(f"Unsupported table block role: {role}")
    return TABLE_BLOCK_PLACEHOLDERS[role]


def paragraph_text(paragraph_element) -> str:
    texts: list[str] = []
    for node in paragraph_element.iter(qn("w:t")):
        texts.append(node.text or "")
    return "".join(texts)


def set_paragraph_text(paragraph_element, text: str) -> None:
    paragraph_properties = paragraph_element.pPr
    for child in list(paragraph_element):
        if child is not paragraph_properties:
            paragraph_element.remove(child)
    run = OxmlElement("w:r")
    text_node = OxmlElement("w:t")
    text_node.set(qn("xml:space"), "preserve")
    text_node.text = text
    run.append(text_node)
    paragraph_element.append(run)


def clone_paragraph_with_text(source_paragraph, text: str):
    paragraph = deepcopy(source_paragraph)
    set_paragraph_text(paragraph, text)
    return paragraph


def new_paragraph_with_text(text: str):
    paragraph = OxmlElement("w:p")
    set_paragraph_text(paragraph, text)
    return paragraph


def find_placeholder_paragraphs(doc: Document, placeholders: set[str]) -> dict[str, Any]:
    found: dict[str, Any] = {}
    for child in doc.element.body.iterchildren():
        if child.tag != qn("w:p"):
            continue
        text = paragraph_text(child).strip()
        if text not in placeholders:
            continue
        if text in found:
            raise ValueError(f"Duplicate structural placeholder in DOCX template: {text}")
        found[text] = child
    return found


def insert_before(anchor, element) -> None:
    parent = anchor.getparent()
    parent.insert(parent.index(anchor), element)


def remove_element(element) -> None:
    parent = element.getparent()
    parent.remove(element)


def table_index_for_tbl(doc: Document, tbl_element) -> int:
    for idx, table in enumerate(doc.tables):
        if table._tbl is tbl_element:
            return idx
    raise ValueError("Inserted table is missing from python-docx table collection.")


def prototype_table_index(role: str, batch_position: int) -> int:
    indexes = role_seed_table_indexes(role)
    if not indexes:
        raise ValueError(f"No prototype table indexes for role {role}")
    position = max(batch_position, 1)
    return indexes[min(position - 1, len(indexes) - 1)]


def profile_value(config: dict[str, Any], key: str) -> str:
    profile = config.get("project_profile")
    if isinstance(profile, dict) and profile.get(key):
        return str(profile[key])
    return ""


def caption_subject(config: dict[str, Any]) -> str:
    return (
        profile_value(config, "product_name")
        or profile_value(config, "product_expression")
        or profile_value(config, "sample_type")
        or "样品"
    )


def caption_condition_fragment(table_spec: dict[str, Any]) -> str:
    condition = str(table_spec.get("condition") or table_spec.get("study_condition") or "")
    return f"，{condition}" if condition else ""


def caption_text(caption_no: int, table_spec: dict[str, Any], config: dict[str, Any]) -> str:
    custom = table_spec.get("caption")
    if custom:
        text = CAPTION_PREFIX_RE.sub("", str(custom)).strip()
        return f"表3.2.S.7.3- {caption_no} {text}"
    role = str(table_spec.get("role") or "")
    batch_no = str(table_spec.get("batch_no") or "")
    subject = caption_subject(config)
    condition = caption_condition_fragment(table_spec)
    if role == "long_term":
        title = f"{subject}长期稳定性试验结果"
        detail = f"（批号：{batch_no}{condition}）" if batch_no or condition else ""
    elif role == "accelerated":
        title = f"{subject}加速稳定性试验结果"
        detail = f"（批号：{batch_no}{condition}）" if batch_no or condition else ""
    else:
        title = STRESS_ROLE_TITLES.get(role, "影响因素试验结果")
        prefix = batch_no
        detail = f"（{condition.lstrip('，')}）" if condition else ""
        title = f"{prefix}{title}" if prefix else title
    return f"表3.2.S.7.3- {caption_no} {title}{detail}"


def values_contain_marker(table_spec: dict[str, Any], marker: str) -> bool:
    for row in table_spec.get("rows", []):
        if not isinstance(row, dict):
            continue
        values = row.get("values")
        if not isinstance(values, dict):
            continue
        for value in values.values():
            rendered = value.get("value") if isinstance(value, dict) else value
            if str(rendered) == marker:
                return True
    return False


def table_note_text(table_spec: dict[str, Any]) -> str:
    explicit = table_spec.get("note") or table_spec.get("table_note")
    if explicit:
        return str(explicit)
    role = str(table_spec.get("role") or "")
    notes: list[str] = []
    if role in {"long_term", "accelerated"}:
        if values_contain_marker(table_spec, "---"):
            notes.append("---表示尚未进行到此时间点，样品在持续考察中")
        if values_contain_marker(table_spec, "N/A"):
            notes.append("N/A表示该时间点无此检项")
        if not notes:
            notes.append("N/A表示该时间点无此检项")
    return f"表注：{'；'.join(notes)}。" if notes else ""


def physical_tc(table, row_idx: int, col_idx: int):
    return table._tbl.tr_lst[row_idx].tc_lst[col_idx]


def physical_cell(table, row_idx: int, col_idx: int) -> _Cell:
    return _Cell(physical_tc(table, row_idx, col_idx), table)


def cell_text(tc) -> str:
    texts: list[str] = []
    for paragraph in tc.p_lst:
        texts.append("".join(node.text or "" for node in paragraph.iter(qn("w:t"))))
    return "/".join(text for text in texts if text)


def get_or_add_tc_pr(tc):
    tc_pr = tc.tcPr
    if tc_pr is None:
        tc_pr = OxmlElement("w:tcPr")
        tc.insert(0, tc_pr)
    return tc_pr


def remove_child(parent, tag: str) -> None:
    child = parent.find(qn(tag))
    if child is not None:
        parent.remove(child)


def set_grid_span(tc, span: int) -> None:
    tc_pr = get_or_add_tc_pr(tc)
    remove_child(tc_pr, "w:gridSpan")
    if span > 1:
        grid_span = OxmlElement("w:gridSpan")
        grid_span.set(qn("w:val"), str(span))
        tc_pr.append(grid_span)


def set_v_merge(tc, value: str | None) -> None:
    tc_pr = get_or_add_tc_pr(tc)
    remove_child(tc_pr, "w:vMerge")
    if value:
        v_merge = OxmlElement("w:vMerge")
        if value != "continue":
            v_merge.set(qn("w:val"), value)
        tc_pr.append(v_merge)


def tc_grid_span(tc) -> int:
    tc_pr = tc.tcPr
    if tc_pr is not None and tc_pr.gridSpan is not None:
        try:
            return int(tc_pr.gridSpan.val)
        except (TypeError, ValueError):
            return 1
    return 1


def clear_tc_content_keep_properties(tc) -> None:
    for child in list(tc):
        if child.tag != qn("w:tcPr"):
            tc.remove(child)


def set_cell_text_preserve_format(cell: _Cell, text: Any) -> None:
    value = "" if text is None else str(text)
    tc = cell._tc
    first_paragraph = cell.paragraphs[0] if cell.paragraphs else None
    p_pr = deepcopy(first_paragraph._p.pPr) if first_paragraph is not None and first_paragraph._p.pPr is not None else None
    run_props = None
    if first_paragraph is not None:
        for run in first_paragraph.runs:
            if run._r.rPr is not None:
                run_props = deepcopy(run._r.rPr)
                break
    tc.clear_content()
    paragraph = OxmlElement("w:p")
    if p_pr is not None:
        paragraph.append(p_pr)
    run = OxmlElement("w:r")
    if run_props is not None:
        run.append(run_props)
    text_node = OxmlElement("w:t")
    text_node.set(qn("xml:space"), "preserve")
    text_node.text = value
    run.append(text_node)
    paragraph.append(run)
    tc.append(paragraph)


def ensure_row_count(table, target_rows: int, template_row_idx: int) -> tuple[int, int]:
    cloned = 0
    deleted = 0
    while len(table._tbl.tr_lst) < target_rows:
        source_idx = min(template_row_idx, len(table._tbl.tr_lst) - 1)
        table._tbl.append(deepcopy(table._tbl.tr_lst[source_idx]))
        cloned += 1
    while len(table._tbl.tr_lst) > target_rows:
        table._tbl.remove(table._tbl.tr_lst[-1])
        deleted += 1
    return cloned, deleted


def sync_tbl_grid(table, target_cols: int) -> None:
    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    while len(tbl_grid.gridCol_lst) < target_cols:
        source = tbl_grid.gridCol_lst[-1] if tbl_grid.gridCol_lst else None
        tbl_grid.append(deepcopy(source) if source is not None else OxmlElement("w:gridCol"))
    while len(tbl_grid.gridCol_lst) > target_cols:
        tbl_grid.remove(tbl_grid.gridCol_lst[-1])


def time_grid_start_col(table, profile: dict[str, Any]) -> int:
    point_row = table._tbl.tr_lst[profile["time_points_row"]]
    return sum(tc_grid_span(tc) for tc in point_row.tc_lst[: profile["time_points_col"]])


def normalize_table_columns(table, profile: dict[str, Any], time_point_count: int) -> dict[str, int]:
    cloned = 0
    deleted = 0
    label_row_idx = profile["time_label_row"]
    points_row_idx = profile["time_points_row"]
    data_start_row = profile["data_start_row"]

    for row_idx in range(len(table._tbl.tr_lst)):
        tr = table._tbl.tr_lst[row_idx]
        if row_idx == label_row_idx:
            target_cells = profile["time_label_col"] + 1
            resize_col = profile["time_label_col"]
        elif row_idx == points_row_idx:
            target_cells = profile.get("time_points_col", profile["time_start_col"]) + time_point_count
            resize_col = profile.get("time_points_col", profile["time_start_col"])
        elif row_idx >= data_start_row:
            target_cells = profile["time_start_col"] + time_point_count
            resize_col = profile["time_start_col"]
        else:
            continue

        clone_idx = min(max(resize_col, 0), len(tr.tc_lst) - 1)
        while len(tr.tc_lst) < target_cells:
            tr.append(deepcopy(tr.tc_lst[clone_idx]))
            cloned += 1
        while len(tr.tc_lst) > target_cells:
            remove_idx = min(max(resize_col, 0), len(tr.tc_lst) - 1)
            tr.remove(tr.tc_lst[remove_idx])
            deleted += 1

    sync_tbl_grid(table, time_grid_start_col(table, profile) + time_point_count)
    return {"cloned_cells": cloned, "deleted_cells": deleted}


def row_values(row_spec: dict[str, Any], profile: dict[str, Any], time_points: list[str]) -> list[str]:
    values = [str(row_spec.get(key, "")) for key in profile["left_columns"]]
    value_map = row_spec.get("values", {})
    values.extend(str(value_map.get(point, "")) for point in time_points)
    return values


def render_stability_time_headers(table, profile: dict[str, Any], time_points: list[str]) -> None:
    label_row = table._tbl.tr_lst[profile["time_label_row"]]
    point_row = table._tbl.tr_lst[profile["time_points_row"]]

    time_label_tc = label_row.tc_lst[profile["time_label_col"]]
    set_grid_span(time_label_tc, len(time_points))
    set_v_merge(time_label_tc, None)

    time_points_col = profile["time_points_col"]
    for idx, tc in enumerate(point_row.tc_lst):
        if idx < time_points_col:
            continue
        set_grid_span(tc, 1)
        set_v_merge(tc, None)
        point_idx = idx - time_points_col
        set_cell_text_preserve_format(_Cell(tc, table), time_points[point_idx] if point_idx < len(time_points) else "")


def render_stress_time_headers(table, profile: dict[str, Any], time_points: list[str], requested_time_label: str = "") -> None:
    label_row = table._tbl.tr_lst[profile["time_label_row"]]
    point_row = table._tbl.tr_lst[profile["time_points_row"]]
    left_span = int(profile["time_start_col"])

    left_header = label_row.tc_lst[0]
    set_grid_span(left_header, left_span)
    set_v_merge(left_header, "restart")

    left_point_header = point_row.tc_lst[0]
    set_grid_span(left_point_header, left_span)
    set_v_merge(left_point_header, "continue")
    set_cell_text_preserve_format(_Cell(left_point_header, table), "")

    time_header = label_row.tc_lst[profile["time_label_col"]]
    set_grid_span(time_header, len(time_points))
    set_v_merge(time_header, None)
    if requested_time_label:
        set_cell_text_preserve_format(_Cell(time_header, table), requested_time_label)

    time_points_col = profile["time_points_col"]
    for idx, tc in enumerate(point_row.tc_lst):
        if idx < time_points_col:
            continue
        set_grid_span(tc, 1)
        set_v_merge(tc, None)
        point_idx = idx - time_points_col
        set_cell_text_preserve_format(_Cell(tc, table), time_points[point_idx] if point_idx < len(time_points) else "")


def render_time_headers(table, profile: dict[str, Any], time_points: list[str], requested_time_label: str = "") -> None:
    if profile.get("kind") == "stress":
        render_stress_time_headers(table, profile, time_points, requested_time_label)
    else:
        render_stability_time_headers(table, profile, time_points)


def prototype_time_header_text(table, profile: dict[str, Any]) -> str:
    return cell_text(physical_tc(table, profile["time_label_row"], profile["time_label_col"]))


def stress_row_parts(row_spec: dict[str, Any]) -> tuple[str, str]:
    item = str(row_spec.get("item", row_spec.get("sub_item", "")) or "")
    method = str(row_spec.get("method", "") or "")
    indicator_method = str(row_spec.get("indicator_method", "") or "")
    if not indicator_method:
        indicator_method = "\n".join(part for part in [item, method] if part.strip())
    acceptance = str(row_spec.get("acceptance", row_spec.get("quality_standard", "")) or "")
    return indicator_method, acceptance


def ensure_stress_row_shape(table, row_idx: int, time_point_count: int) -> dict[str, int]:
    tr = table._tbl.tr_lst[row_idx]
    cloned = 0
    deleted = 0
    target_cells = 2 + time_point_count

    while len(tr.tc_lst) < 2:
        source = tr.tc_lst[0] if tr.tc_lst else OxmlElement("w:tc")
        tr.insert(len(tr.tc_lst), deepcopy(source))
        cloned += 1
    while len(tr.tc_lst) < target_cells:
        source_idx = min(len(tr.tc_lst) - 1, max(1, len(tr.tc_lst) - 1))
        tr.append(deepcopy(tr.tc_lst[source_idx]))
        cloned += 1
    while len(tr.tc_lst) > target_cells:
        tr.remove(tr.tc_lst[-1])
        deleted += 1

    return {"cloned_cells": cloned, "deleted_cells": deleted}


def render_stability_data_rows(table, profile: dict[str, Any], rows: list[dict[str, Any]], time_points: list[str]) -> dict[str, int]:
    start = profile["data_start_row"]
    for offset, row_spec in enumerate(rows):
        row_idx = start + offset
        values = row_values(row_spec, profile, time_points)
        for col_idx, value in enumerate(values):
            tc = physical_tc(table, row_idx, col_idx)
            set_grid_span(tc, 1)
            set_v_merge(tc, None)
            set_cell_text_preserve_format(_Cell(tc, table), value)
    return {"cloned_cells": 0, "deleted_cells": 0}


def render_stress_data_rows(table, profile: dict[str, Any], rows: list[dict[str, Any]], time_points: list[str]) -> dict[str, int]:
    start = profile["data_start_row"]
    cloned = 0
    deleted = 0
    for offset, row_spec in enumerate(rows):
        row_idx = start + offset
        indicator_method, acceptance = stress_row_parts(row_spec)
        changes = ensure_stress_row_shape(table, row_idx, len(time_points))
        cloned += changes["cloned_cells"]
        deleted += changes["deleted_cells"]
        tr = table._tbl.tr_lst[row_idx]
        value_map = row_spec.get("values", {})

        indicator_tc = tr.tc_lst[0]
        acceptance_tc = tr.tc_lst[1]
        set_grid_span(indicator_tc, 1)
        set_v_merge(indicator_tc, None)
        set_cell_text_preserve_format(_Cell(indicator_tc, table), indicator_method)
        set_grid_span(acceptance_tc, 1)
        set_v_merge(acceptance_tc, None)
        set_cell_text_preserve_format(_Cell(acceptance_tc, table), acceptance)
        time_offset = 2

        for point_idx, point in enumerate(time_points):
            tc = tr.tc_lst[time_offset + point_idx]
            set_grid_span(tc, 1)
            set_v_merge(tc, None)
            set_cell_text_preserve_format(_Cell(tc, table), str(value_map.get(point, "")))

    return {"cloned_cells": cloned, "deleted_cells": deleted}


def render_data_rows(table, profile: dict[str, Any], rows: list[dict[str, Any]], time_points: list[str]) -> dict[str, int]:
    if profile.get("kind") == "stress":
        return render_stress_data_rows(table, profile, rows, time_points)
    return render_stability_data_rows(table, profile, rows, time_points)


def rebuild_stability_vmerge(table, profile: dict[str, Any], rows: list[dict[str, Any]]) -> None:
    start = profile["data_start_row"]
    for col_idx in profile.get("vmerge_columns", []):
        offset = 0
        while offset < len(rows):
            key_parts = [rows[offset].get(profile["left_columns"][idx], "") for idx in range(min(col_idx + 1, len(profile["left_columns"])))]
            if not all(str(part) for part in key_parts):
                set_v_merge(physical_tc(table, start + offset, col_idx), None)
                offset += 1
                continue
            end = offset + 1
            while end < len(rows):
                next_parts = [rows[end].get(profile["left_columns"][idx], "") for idx in range(min(col_idx + 1, len(profile["left_columns"])))]
                if tuple(next_parts) != tuple(key_parts):
                    break
                end += 1
            if end - offset > 1:
                set_v_merge(physical_tc(table, start + offset, col_idx), "restart")
                for inner in range(offset + 1, end):
                    tc = physical_tc(table, start + inner, col_idx)
                    set_v_merge(tc, "continue")
                    set_cell_text_preserve_format(_Cell(tc, table), "")
            else:
                set_v_merge(physical_tc(table, start + offset, col_idx), None)
            offset = end


def rebuild_stress_vmerge(table, profile: dict[str, Any], rows: list[dict[str, Any]]) -> None:
    start = profile["data_start_row"]
    for offset in range(len(rows)):
        for col_idx in range(2):
            set_v_merge(physical_tc(table, start + offset, col_idx), None)


def rebuild_vmerge(table, profile: dict[str, Any], rows: list[dict[str, Any]]) -> None:
    if profile.get("kind") == "stress":
        rebuild_stress_vmerge(table, profile, rows)
    else:
        rebuild_stability_vmerge(table, profile, rows)


def read_json_if_exists(path: Path | None) -> dict[str, Any]:
    if path is None or not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def resolve_config_path(value: str, base_dir: Path) -> Path:
    path = Path(value)
    if path.is_absolute():
        return path
    if path.exists():
        return path.resolve()
    return (base_dir / path).resolve()


def batch_structure_action(config: dict[str, Any]) -> str:
    scope = config.get("render_scope")
    if isinstance(scope, dict) and scope.get("batch_structure_action"):
        return str(scope.get("batch_structure_action"))
    return ""


def render_table_at_index(
    doc: Document,
    table_spec: dict[str, Any],
    table_idx: int,
    *,
    caption_no: int,
    caption: str,
    placeholder: str,
    prototype_table_index_value: int,
) -> dict[str, Any]:
    profile = require_profile(str(table_spec["role"]))
    table = doc.tables[table_idx]
    time_points = [str(point) for point in table_spec.get("time_points") or table_spec.get("time_header", {}).get("points", [])]
    if not time_points:
        raise ValueError(f"Table {table_idx} has no time points")
    requested_time_label = str(table_spec.get("time_label") or table_spec.get("time_header", {}).get("label") or "")
    rows = list(table_spec.get("rows", []))
    target_rows = profile["data_start_row"] + len(rows)

    row_cloned, row_deleted = ensure_row_count(table, target_rows, profile["data_template_row"])
    column_changes = normalize_table_columns(table, profile, len(time_points))
    render_time_headers(table, profile, time_points, requested_time_label)
    sync_tbl_grid(table, time_grid_start_col(table, profile) + len(time_points))
    preserved_time_label = prototype_time_header_text(table, profile)
    target_cols = time_grid_start_col(table, profile) + len(time_points)
    data_row_changes = render_data_rows(table, profile, rows, time_points)
    rebuild_vmerge(table, profile, rows)

    return {
        "role": table_spec["role"],
        "normalized_role": normalize_role(str(table_spec["role"])),
        "batch_no": table_spec.get("batch_no"),
        "table_index": table_idx,
        "time_header_label": preserved_time_label,
        "requested_time_header_label": requested_time_label,
        "time_points": time_points,
        "data_row_count": len(rows),
        "target_columns": target_cols,
        "target_rows": target_rows,
        "cloned_rows": row_cloned,
        "deleted_rows": row_deleted,
        "cloned_cells": column_changes["cloned_cells"] + data_row_changes["cloned_cells"],
        "deleted_cells": column_changes["deleted_cells"] + data_row_changes["deleted_cells"],
        "gridSpan_updated": True,
        "tblGrid_updated": True,
        "vMerge_rebuilt": profile.get("kind") == "stress" or bool(profile.get("vmerge_columns")),
        "header_contract": "stress-prototype-fixed-header" if profile.get("kind") == "stress" else "stability-prototype-fixed-header",
        "header_text_policy": "stress time label is rendered from table_spec.time_header.label when provided; stability fixed prototype label is preserved",
        "placeholder": placeholder,
        "caption_no": caption_no,
        "caption": caption,
        "prototype_table_index": prototype_table_index_value,
        "warnings": [],
    }


def grouped_tables_by_placeholder(tables: list[dict[str, Any]]) -> dict[str, list[dict[str, Any]]]:
    grouped: dict[str, list[dict[str, Any]]] = defaultdict(list)
    for table in tables:
        role = str(table.get("role") or "")
        grouped[table_block_placeholder(role)].append(table)
    return grouped


def directory_entry_text(record: dict[str, Any]) -> str:
    return str(record.get("caption") or "")


def render_table_blocks(doc: Document, prototype_doc: Document, config: dict[str, Any]) -> tuple[list[dict[str, Any]], list[str]]:
    all_placeholders = set(TABLE_BLOCK_PLACEHOLDERS.values()) | {TABLE_DIRECTORY_PLACEHOLDER}
    placeholder_paragraphs = find_placeholder_paragraphs(doc, all_placeholders)
    grouped = grouped_tables_by_placeholder([table for table in config.get("tables", []) if isinstance(table, dict)])
    missing = sorted(placeholder for placeholder in grouped if placeholder not in placeholder_paragraphs)
    if missing:
        raise ValueError(f"Template is missing table block placeholder(s): {missing}")

    records: list[dict[str, Any]] = []
    generated_slots: list[dict[str, Any]] = []
    warnings: list[str] = []
    caption_no = 1

    for placeholder, anchor in list(placeholder_paragraphs.items()):
        if placeholder == TABLE_DIRECTORY_PLACEHOLDER:
            continue
        table_specs = grouped.get(placeholder, [])
        for position, table_spec in enumerate(table_specs, start=1):
            role = str(table_spec.get("role") or "")
            proto_idx = prototype_table_index(role, position)
            if proto_idx >= len(prototype_doc.tables):
                raise ValueError(f"Prototype table index {proto_idx} for role {role} is missing from prototype DOCX.")
            caption = caption_text(caption_no, table_spec, config)
            insert_before(anchor, clone_paragraph_with_text(anchor, caption))
            table_element = deepcopy(prototype_doc.tables[proto_idx]._tbl)
            insert_before(anchor, table_element)
            table_idx = table_index_for_tbl(doc, table_element)
            record = render_table_at_index(
                doc,
                table_spec,
                table_idx,
                caption_no=caption_no,
                caption=caption,
                placeholder=placeholder,
                prototype_table_index_value=proto_idx,
            )
            note = table_note_text(table_spec)
            if note:
                insert_before(anchor, new_paragraph_with_text(note))
                record["table_note"] = note
            slot = {
                "placeholder": placeholder,
                "role": table_spec.get("role"),
                "normalized_role": normalize_role(str(table_spec.get("role") or "")),
                "batch_no": table_spec.get("batch_no"),
                "batch_position": position,
                "table_index": table_idx,
                "caption_no": caption_no,
                "caption": caption,
                "prototype_table_index": proto_idx,
                "structure_action": "generated_from_table_block_placeholder",
            }
            record["generated_slot"] = slot
            records.append(record)
            generated_slots.append(slot)
            caption_no += 1
        remove_element(anchor)

    directory_anchor = placeholder_paragraphs.get(TABLE_DIRECTORY_PLACEHOLDER)
    if directory_anchor is None:
        warnings.append(f"Template has no {TABLE_DIRECTORY_PLACEHOLDER}; table directory was not generated.")
    else:
        for record in records:
            insert_before(directory_anchor, clone_paragraph_with_text(directory_anchor, directory_entry_text(record)))
        remove_element(directory_anchor)
    for record, slot in zip(records, generated_slots, strict=True):
        record["generated_slot"] = slot
    return records, warnings


def render_document(config: dict[str, Any]) -> dict[str, Any]:
    template_docx = Path(config["template_docx"])
    output_docx = Path(config["output_docx"])
    doc = Document(template_docx)
    prototype_docx = Path(str(config.get("table_prototype_docx") or DEFAULT_TABLE_PROTOTYPE))
    prototype_doc = Document(prototype_docx)
    records, warnings = render_table_blocks(doc, prototype_doc, config)
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)
    manifest = {
        "renderer": "render_stability_tables.py",
        "template_docx": str(template_docx),
        "table_prototype_docx": str(prototype_docx),
        "output_docx": str(output_docx),
        "template_contract": config.get("template_contract", "ctd-32s73-stability-template-v2"),
        "render_scope": config.get("render_scope"),
        "tables_rendered": records,
        "generated_table_slots": [record["generated_slot"] for record in records],
        "warnings": [*warnings, *[warning for record in records for warning in record.get("warnings", [])]],
    }
    manifest_path = Path(config.get("manifest_path") or output_docx.with_name("table-render-manifest.json"))
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    manifest["manifest_path"] = str(manifest_path)
    return manifest


def main() -> None:
    require_internal_context("render_stability_tables.py")
    parser = argparse.ArgumentParser(description="Render CTD 3.2.S.7.3 stability tables into a DOCX template.")
    parser.add_argument("config", type=Path, help="JSON config containing template_docx, output_docx, and tables.")
    args = parser.parse_args()

    config = json.loads(args.config.read_text(encoding="utf-8"))
    for key in ["template_docx", "table_prototype_docx"]:
        if config.get(key):
            config[key] = str(resolve_config_path(str(config[key]), args.config.parent))
    for key in ["output_docx", "manifest_path"]:
        if config.get(key):
            path = Path(str(config[key]))
            config[key] = str(path if path.is_absolute() else path.resolve())
    manifest = render_document(config)
    print(json.dumps(manifest, ensure_ascii=False, indent=2))


# Internal module — do not invoke directly. Use skill_step.py as the public entry point.
if __name__ == "__main__":
    main()

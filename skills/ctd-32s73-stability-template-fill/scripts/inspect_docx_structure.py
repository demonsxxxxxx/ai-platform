from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document


def cell_text(tc) -> str:
    return "".join(tc.itertext()).replace("\n", "/")


def grid_span(tc) -> str:
    tc_pr = tc.tcPr
    if tc_pr is not None and tc_pr.gridSpan is not None:
        return str(tc_pr.gridSpan.val)
    return "1"


def v_merge(tc) -> str:
    tc_pr = tc.tcPr
    if tc_pr is not None and tc_pr.vMerge is not None:
        return str(tc_pr.vMerge.val or "continue")
    return ""


def parse_indexes(value: str | None) -> list[int]:
    if not value:
        return []
    result: list[int] = []
    for part in value.split(","):
        part = part.strip()
        if not part:
            continue
        result.append(int(part))
    return result


def print_doc_summary(doc: Document, path: Path):
    print(f"FILE={path}")
    print(f"PARAGRAPHS={len(doc.paragraphs)}")
    print(f"TABLES={len(doc.tables)}")
    print(f"INLINE_SHAPES={len(doc.inline_shapes)}")


def print_placeholder_counts(doc: Document, placeholders: list[str]):
    body = "\n".join(
        [p.text for p in doc.paragraphs]
        + [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    )
    print("PLACEHOLDERS")
    for placeholder in placeholders:
        print(f"  {placeholder}={body.count(placeholder)}")


def print_paragraphs(doc: Document, indexes: list[int]):
    if not indexes:
        return
    print("PARAGRAPHS")
    for idx in indexes:
        if idx >= len(doc.paragraphs):
            print(f"  [{idx}] OUT_OF_RANGE")
            continue
        paragraph = doc.paragraphs[idx]
        print(
            f"  [{idx}] style={paragraph.style.name!r} "
            f"alignment={paragraph.alignment} text={paragraph.text[:120]!r}"
        )


def print_caption_paragraphs(doc: Document):
    print("CAPTION_AND_SECTION_PARAGRAPHS")
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip().replace("\t", " ")
        if not text:
            continue
        style_name = paragraph.style.name
        is_candidate = (
            style_name in {"Caption", "table of figures", "toc 1", "toc 2"}
            or text.startswith(("表", "图"))
            or "稳定性" in text
            or "影响因素" in text
        )
        if is_candidate:
            print(f"  [{idx}] style={style_name!r} text={text[:160]!r}")


def print_tables(doc: Document, indexes: list[int], max_rows: int):
    if not indexes:
        return
    print("TABLES")
    for idx in indexes:
        if idx >= len(doc.tables):
            print(f"  TABLE[{idx}] OUT_OF_RANGE")
            continue
        table = doc.tables[idx]
        print(f"  TABLE[{idx}] logical_rows={len(table.rows)} logical_cols={len(table.columns)}")
        for row_idx, tr in enumerate(table._tbl.tr_lst[:max_rows]):
            cells = []
            for tc in tr.tc_lst:
                text = cell_text(tc)[:40]
                cells.append(f"span={grid_span(tc)},v={v_merge(tc)},text={text!r}")
            print(f"    row[{row_idx}] physical_cells={len(tr.tc_lst)} {cells}")


def main():
    parser = argparse.ArgumentParser(description="Inspect DOCX paragraph styles and table OOXML structure.")
    parser.add_argument("docx", type=Path)
    parser.add_argument("--paragraph-indexes", default="", help="Comma-separated paragraph indexes to inspect.")
    parser.add_argument("--table-indexes", default="", help="Comma-separated table indexes to inspect.")
    parser.add_argument("--all-tables", action="store_true", help="Inspect all tables.")
    parser.add_argument(
        "--caption-paragraphs",
        action="store_true",
        help="Print likely section, table-caption, figure-caption, and table-of-figures paragraphs.",
    )
    parser.add_argument("--max-rows", type=int, default=3)
    parser.add_argument(
        "--placeholders",
        default="XXX,IPXXX,单抗/原液,已完成的申报资料,时间（小时）",
        help="Comma-separated placeholder strings to count.",
    )
    args = parser.parse_args()

    doc = Document(args.docx)
    print_doc_summary(doc, args.docx)
    print_placeholder_counts(doc, [x for x in args.placeholders.split(",") if x])
    print_paragraphs(doc, parse_indexes(args.paragraph_indexes))
    if args.caption_paragraphs:
        print_caption_paragraphs(doc)
    table_indexes = list(range(len(doc.tables))) if args.all_tables else parse_indexes(args.table_indexes)
    print_tables(doc, table_indexes, args.max_rows)


# Read-only diagnostic helper. It may be invoked directly, but it must not be used to advance the workflow.
if __name__ == "__main__":
    main()

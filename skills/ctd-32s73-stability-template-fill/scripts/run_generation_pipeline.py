from __future__ import annotations

import argparse
import json
import re
import shutil
import subprocess
import sys
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from runtime_guard import make_internal_env, require_internal_context
from docx import Document
from docx.shared import Inches


SKILL_ROOT = Path(__file__).resolve().parents[1]
DEFAULT_TEMPLATE = SKILL_ROOT / "assets" / "templates" / "ctd-32s73-stability-template-v2.docx"
DEFAULT_TEMPLATE_ID = "ctd-32s73-stability-template-v2"
DEFAULT_TABLE_PROTOTYPE = SKILL_ROOT / "assets" / "table-prototypes" / "ctd-32s73-table-prototypes.docx"
INSPECT_SCRIPT = SKILL_ROOT / "scripts" / "inspect_docx_structure.py"
RENDER_SCRIPT = SKILL_ROOT / "scripts" / "render_stability_tables.py"
TREND_CHART_SCRIPT = SKILL_ROOT / "scripts" / "generate_stability_trend_charts.py"
VALIDATE_DOCX_SCRIPT = SKILL_ROOT / "scripts" / "validate_generated_docx.py"
VALIDATE_FACT_SCRIPT = SKILL_ROOT / "scripts" / "validate_fact_packet.py"
FIGURE_DIRECTORY_PLACEHOLDER = "{{FIGURE_DIRECTORY}}"
LONG_TERM_FIGURE_BLOCK_PLACEHOLDER = "{{LONG_TERM_FIGURE_BLOCK}}"
ACCELERATED_FIGURE_BLOCK_PLACEHOLDER = "{{ACCELERATED_FIGURE_BLOCK}}"
STRESS_SECTION_BLOCKS_PLACEHOLDER = "{{STRESS_SECTION_BLOCKS}}"

STRESS_SECTION_SPECS = [
    {
        "key": "light",
        "roles": ["light_stress"],
        "title": "光照试验结果",
        "placeholders": ["{{STRESS_LIGHT_REFERENCE}}", "{{STRESS_LIGHT_SUMMARY}}", "{{STRESS_LIGHT_TABLE_BLOCK}}"],
        "aliases": ["光照", "light"],
    },
    {
        "key": "agitation",
        "roles": ["agitation_stress"],
        "title": "振荡试验结果",
        "placeholders": ["{{STRESS_AGITATION_SUMMARY}}", "{{STRESS_AGITATION_TABLE_BLOCK}}"],
        "aliases": ["振荡", "震荡", "agitation", "shaking", "shake"],
    },
    {
        "key": "freeze_thaw",
        "roles": ["freeze_thaw_stress"],
        "title": "冻融试验结果",
        "placeholders": ["{{STRESS_FREEZE_THAW_SUMMARY}}", "{{STRESS_FREEZE_THAW_TABLE_BLOCK}}"],
        "aliases": ["冻融", "反复冻融", "freeze", "thaw", "freeze_thaw"],
    },
    {
        "key": "high_temperature",
        "roles": ["high_temperature_stress"],
        "title": "高温试验结果",
        "placeholders": ["{{STRESS_HIGH_TEMPERATURE_SUMMARY}}", "{{STRESS_HIGH_TEMPERATURE_TABLE_BLOCK}}"],
        "aliases": ["高温", "high temperature", "high_temperature"],
    },
    {
        "key": "ph",
        "roles": ["low_ph_stress", "high_ph_stress"],
        "title": "pH试验结果",
        "placeholders": ["{{STRESS_PH_SUMMARY}}", "{{STRESS_LOW_PH_TABLE_BLOCK}}", "{{STRESS_HIGH_PH_TABLE_BLOCK}}"],
        "aliases": ["ph", "pH", "PH", "低pH", "高pH", "低ph", "高ph", "酸碱", "low ph", "high ph"],
    },
    {
        "key": "oxidation",
        "roles": ["oxidation_stress"],
        "title": "氧化试验结果",
        "placeholders": ["{{STRESS_OXIDATION_SUMMARY}}", "{{STRESS_OXIDATION_TABLE_BLOCK}}"],
        "aliases": ["氧化", "oxidation", "oxidative", "H2O2", "TBHP"],
    },
]


def read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def write_json(path: Path, data: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def write_text(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")


def set_paragraph_text(paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def find_placeholder_paragraph(doc: Document, placeholder: str):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            return paragraph
    return None


def insert_paragraph_before(anchor, text: str, style: str | None = None):
    paragraph = anchor.insert_paragraph_before(text)
    if style:
        try:
            paragraph.style = style
        except KeyError:
            pass
    return paragraph


def run_command(command: list[str], stdout_path: Path) -> dict[str, Any]:
    completed = subprocess.run(
        command,
        text=True,
        capture_output=True,
        env=make_internal_env("run_generation_pipeline.py", "run child command"),
    )
    write_text(stdout_path, completed.stdout)
    stderr_path = None
    if completed.stderr:
        stderr_path = stdout_path.with_suffix(stdout_path.suffix + ".stderr.txt")
        write_text(stderr_path, completed.stderr)
    return {
        "command": command,
        "returncode": completed.returncode,
        "stdout": str(stdout_path),
        "stderr": str(stderr_path) if stderr_path else None,
    }


def source_path(source: Any) -> str:
    if isinstance(source, dict):
        return str(source.get("path") or source.get("file") or source)
    return str(source)


def source_status(source: Any) -> str:
    if isinstance(source, dict):
        return str(source.get("status") or "")
    return ""


def source_reason(source: Any) -> str:
    if isinstance(source, dict):
        return str(source.get("reason") or source.get("note") or "")
    return ""


def packet_sources(packet: dict[str, Any], key: str) -> list[Any]:
    sources = packet.get("sources")
    if not isinstance(sources, dict):
        return []
    values = sources.get(key)
    return values if isinstance(values, list) else []


def batch_numbers(packet: dict[str, Any]) -> list[str]:
    profile = packet.get("project_profile")
    if not isinstance(profile, dict):
        return []
    batches = profile.get("batches")
    if not isinstance(batches, list):
        return []
    numbers = []
    for batch in batches:
        if isinstance(batch, dict) and batch.get("batch_no"):
            numbers.append(str(batch["batch_no"]))
    return numbers


def render_scope(packet: dict[str, Any]) -> dict[str, Any]:
    plan = packet.get("docx_render_plan")
    if not isinstance(plan, dict):
        return {}
    scope = plan.get("render_scope")
    return scope if isinstance(scope, dict) else {}


def batch_structure_action(packet: dict[str, Any]) -> str:
    scope = render_scope(packet)
    if scope.get("batch_structure_action"):
        return str(scope.get("batch_structure_action"))
    plan = packet.get("docx_render_plan")
    if isinstance(plan, dict):
        return str(plan.get("batch_count_action") or "")
    return ""


def table_render_inputs(study: Any) -> list[dict[str, Any]]:
    if not isinstance(study, dict):
        return []
    values = study.get("table_render_inputs")
    return [item for item in values if isinstance(item, dict)] if isinstance(values, list) else []


def normalize_test_name(value: str) -> str:
    return re.sub(r"[\s_\-（）()／/]+", "", value).lower()


def test_name_value(item: Any) -> str:
    if isinstance(item, dict):
        for key in ("role", "template_table_role", "test", "test_name", "name"):
            value = item.get(key)
            if value:
                return str(value)
        return ""
    return str(item or "")


def active_stress_roles(packet: dict[str, Any]) -> set[str]:
    stress = packet.get("stress_study")
    roles: set[str] = set()
    for table in table_render_inputs(stress):
        role = str(table.get("role") or "").strip()
        if any(role in spec["roles"] for spec in STRESS_SECTION_SPECS):
            roles.add(role)
    if isinstance(stress, dict) and isinstance(stress.get("included_tests"), list):
        for item in stress["included_tests"]:
            raw = test_name_value(item)
            normalized = normalize_test_name(raw)
            for spec in STRESS_SECTION_SPECS:
                if raw in spec["roles"]:
                    roles.add(raw)
                    continue
                for alias in spec["aliases"]:
                    alias_normalized = normalize_test_name(str(alias))
                    if alias_normalized and alias_normalized in normalized:
                        roles.update(str(role) for role in spec["roles"])
    return roles


def render_stress_sections_in_docx(packet: dict[str, Any], source_docx: Path, output_docx: Path, report_path: Path) -> dict[str, Any]:
    doc = Document(source_docx)
    anchor = find_placeholder_paragraph(doc, STRESS_SECTION_BLOCKS_PLACEHOLDER)
    active_roles = active_stress_roles(packet)
    generated: list[dict[str, Any]] = []
    if anchor is not None:
        section_no = 1
        for spec in STRESS_SECTION_SPECS:
            roles = [str(role) for role in spec["roles"]]
            if not (active_roles & set(roles)):
                continue
            heading = f"3.1.{section_no} {spec['title']}"
            insert_paragraph_before(anchor, heading, style=anchor.style.name)
            for placeholder in spec["placeholders"]:
                paragraph_style = "Caption" if placeholder.endswith("_TABLE_BLOCK}}") else anchor.style.name
                insert_paragraph_before(anchor, placeholder, style=paragraph_style)
            generated.append(
                {
                    "key": spec["key"],
                    "roles": roles,
                    "heading": heading,
                    "placeholders": list(spec["placeholders"]),
                }
            )
            section_no += 1
        remove_paragraph(anchor)
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)
    report = {
        "schema_version": "ctd-32s73-stress-section-render-v1",
        "input_docx": str(source_docx),
        "output_docx": str(output_docx),
        "placeholder_found": anchor is not None,
        "active_roles": sorted(active_roles),
        "generated_sections": generated,
        "policy": "Generate influence-factor subsections from stress_study.included_tests and stress_study.table_render_inputs; omitted roles are not materialized in the DOCX.",
    }
    write_json(report_path, report)
    return report


def placeholder_for_section(section_key: str) -> str:
    return "{{" + section_key.upper() + "}}"


def body_sections(packet: dict[str, Any]) -> dict[str, dict[str, Any]]:
    values = packet.get("body_sections")
    if not isinstance(values, dict):
        plan = packet.get("docx_render_plan")
        values = plan.get("body_sections") if isinstance(plan, dict) else {}
    if not isinstance(values, dict):
        return {}
    sections: dict[str, dict[str, Any]] = {}
    for key, value in values.items():
        section_key = str(key)
        if isinstance(value, str):
            sections[section_key] = {
                "placeholder": placeholder_for_section(section_key),
                "text": value,
            }
            continue
        if not isinstance(value, dict):
            continue
        text_value = value.get("text", value.get("content", ""))
        if isinstance(text_value, list):
            text = "\n".join(str(item) for item in text_value)
        else:
            text = str(text_value or "")
        sections[section_key] = {
            **value,
            "placeholder": str(value.get("placeholder") or placeholder_for_section(section_key)),
            "text": text,
        }
    return sections


def replace_body_section_paragraph(paragraph, sections: dict[str, dict[str, Any]]) -> list[str]:
    original = paragraph.text
    changed: list[str] = []
    new_text = original
    for key, section in sections.items():
        placeholder = str(section.get("placeholder") or placeholder_for_section(key))
        if placeholder in new_text:
            new_text = new_text.replace(placeholder, str(section.get("text") or ""))
            changed.append(key)
    if not changed or new_text == original:
        return []
    paragraph.clear()
    paragraph.add_run(new_text)
    return changed


def apply_body_sections_to_docx(
    packet: dict[str, Any],
    trend_manifest: dict[str, Any],
    source_docx: Path,
    output_docx: Path,
    report_path: Path,
) -> dict[str, Any]:
    doc = Document(source_docx)
    sections, trend_token_counts, trend_tokens = render_body_sections(packet, trend_manifest)
    counts: dict[str, int] = {}
    for paragraph in doc.paragraphs:
        for key in replace_body_section_paragraph(paragraph, sections):
            counts[key] = counts.get(key, 0) + 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key in replace_body_section_paragraph(paragraph, sections):
                        counts[key] = counts.get(key, 0) + 1
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)
    report = {
        "schema_version": "ctd-32s73-body-section-application-v1",
        "input_docx": str(source_docx),
        "output_docx": str(output_docx),
        "body_section_counts": counts,
        "body_section_placeholders": {
            key: section.get("placeholder") for key, section in sections.items()
        },
        "trend_token_counts": trend_token_counts,
        "trend_tokens": trend_tokens,
        "trend_manifest_provided": bool(trend_manifest),
        "policy": "Semantic body_sections are applied in run_generation_pipeline.py after trend chart manifest generation and dynamic influence-factor section rendering, before table rendering.",
    }
    write_json(report_path, report)
    return report


def collect_table_inputs(packet: dict[str, Any]) -> list[dict[str, Any]]:
    tables: list[dict[str, Any]] = []
    for study_key in ["long_term", "accelerated", "stress_study"]:
        study = packet.get(study_key)
        if not isinstance(study, dict):
            continue
        inputs = study.get("table_render_inputs")
        if not isinstance(inputs, list):
            continue
        for item in inputs:
            if isinstance(item, dict):
                tables.append(normalize_table_for_renderer(item, study, study_key))
    return tables


def renderer_value(value: Any) -> Any:
    if isinstance(value, dict):
        return value.get("value", "")
    return value


def normalize_table_for_renderer(table: dict[str, Any], study: dict[str, Any], study_key: str) -> dict[str, Any]:
    normalized = deepcopy(table)
    normalized.setdefault("study_key", study_key)
    if study.get("condition") and not normalized.get("condition"):
        normalized["condition"] = study["condition"]
    for row in normalized.get("rows", []):
        if not isinstance(row, dict):
            continue
        if "acceptance" not in row and "quality_standard" in row:
            row["acceptance"] = row["quality_standard"]
        values = row.get("values")
        if isinstance(values, dict):
            row["values"] = {str(point): renderer_value(value) for point, value in values.items()}
    return normalized


def build_evidence_summary(
    packet: dict[str, Any],
    template_path: Path,
    template_id: str,
    starting_docx: Path,
    render_starting_docx: Path,
    fact_packet_path: Path,
) -> str:
    timestamp = datetime.now(timezone.utc).isoformat()
    lines = [
        "# Evidence Summary",
        "",
        f"Generated at: {timestamp}",
        "",
        "## Active template",
        "",
        f"- Template path: {template_path}",
        f"- Template ID: {template_id}",
        f"- Body/skeleton DOCX: {starting_docx}",
        f"- Starting DOCX for table render: {render_starting_docx}",
        f"- Fact packet: {fact_packet_path}",
        "",
        "## Used sources",
        "",
    ]
    allowed = packet_sources(packet, "allowed")
    if allowed:
        for source in allowed:
            status = source_status(source)
            suffix = f" (status: {status})" if status else ""
            lines.append(f"- {source_path(source)}{suffix}")
    else:
        lines.append("- No allowed sources listed in fact packet.")

    lines.extend(["", "## Excluded sources", ""])
    forbidden = packet_sources(packet, "forbidden")
    if forbidden:
        for source in forbidden:
            reason = source_reason(source)
            suffix = f"; reason: {reason}" if reason else ""
            lines.append(f"- {source_path(source)} (status: excluded{suffix})")
    else:
        lines.append("- No forbidden sources listed in fact packet.")

    lines.extend(
        [
            "",
            "## Pipeline boundary",
            "",
            "- This summary is generated from fact-packet metadata.",
            "- The pipeline validates structure and renders declared tables; it does not independently prove that source extraction was correct.",
            "- Project-specific body text comes from body_sections; trend chart references may be resolved from trend-chart-manifest.json; table captions, table blocks, table directory entries, and figure directory entries are generated from fact-packet inputs.",
            "- Word fields are not refreshed automatically; figure directory entries are rendered from trend-chart-manifest.json for generated long-term and accelerated figures.",
            "",
        ]
    )
    return "\n".join(lines)


def build_missing_evidence_report(packet: dict[str, Any]) -> str:
    missing = packet.get("missing_evidence")
    manual = packet.get("manual_review_items")
    lines = ["# Missing Evidence Report", ""]
    lines.append("## Missing evidence")
    lines.append("")
    if isinstance(missing, list) and missing:
        for item in missing:
            lines.append(f"- {json.dumps(item, ensure_ascii=False) if not isinstance(item, str) else item}")
    else:
        lines.append("- None recorded in fact-packet.json.")
    lines.extend(["", "## Manual review items", ""])
    if isinstance(manual, list) and manual:
        for item in manual:
            lines.append(f"- {json.dumps(item, ensure_ascii=False) if not isinstance(item, str) else item}")
    else:
        lines.append("- None recorded in fact-packet.json.")
    lines.append("")
    return "\n".join(lines)


def build_table_render_config(
    packet: dict[str, Any],
    starting_docx: Path,
    output_docx: Path,
    manifest_path: Path,
    template_id: str,
) -> dict[str, Any]:
    config = {
        "template_docx": str(starting_docx),
        "table_prototype_docx": str(DEFAULT_TABLE_PROTOTYPE),
        "output_docx": str(output_docx),
        "manifest_path": str(manifest_path),
        "template_contract": template_id,
        "render_scope": render_scope(packet),
        "project_profile": packet.get("project_profile") if isinstance(packet.get("project_profile"), dict) else {},
        "tables": collect_table_inputs(packet),
    }
    return config


def normalize_trend_chart_config(packet: dict[str, Any]) -> dict[str, Any] | None:
    trend_charts = packet.get("trend_charts")
    if trend_charts is None:
        return None
    if isinstance(trend_charts, dict):
        charts = trend_charts.get("charts")
        if not isinstance(charts, list) or not charts:
            return None
        return {
            key: value
            for key, value in trend_charts.items()
            if key in {"charts", "style", "dpi"}
        }
    if isinstance(trend_charts, list) and trend_charts:
        # Backward-compatible runtime support; validator now requires the object form.
        return {"charts": trend_charts}
    return None


def read_json_optional(path: Path | None) -> dict[str, Any]:
    if path is None or not path.exists():
        return {}
    return read_json(path)


def chart_group(chart: dict[str, Any]) -> str:
    explicit = str(chart.get("study_key") or chart.get("group") or chart.get("section") or "").strip()
    if explicit:
        if explicit in {"long", "long_term", "长期", "长期稳定性"}:
            return "long_term"
        if explicit in {"accelerated", "accel", "加速", "加速稳定性"}:
            return "accelerated"
        if explicit in {"stress", "stress_study", "影响因素", "强制降解"}:
            return "stress_study"
    title = str(chart.get("title") or "")
    if "长期" in title or "long" in title.lower():
        return "long_term"
    if "加速" in title or "accelerated" in title.lower() or "accel" in title.lower():
        return "accelerated"
    if any(marker in title for marker in ["影响因素", "光照", "振荡", "冻融", "高温", "低pH", "高pH", "氧化"]):
        return "stress_study"
    return "all"


def figure_no_value(chart: dict[str, Any]) -> int | None:
    try:
        return int(chart.get("figure_no"))
    except (TypeError, ValueError):
        return None


def figure_ref(number: int) -> str:
    return f"图3.2.S.7.3- {number}"


def figure_range(numbers: list[int]) -> str:
    if not numbers:
        return ""
    if len(numbers) == 1:
        return figure_ref(numbers[0])
    return f"{figure_ref(numbers[0])} ~ {figure_ref(numbers[-1])}"


def trend_body_context(trend_manifest: dict[str, Any]) -> dict[str, Any]:
    charts = [chart for chart in trend_manifest.get("charts", []) if isinstance(chart, dict)]
    groups: dict[str, list[dict[str, Any]]] = {"all": []}
    for chart in charts:
        number = figure_no_value(chart)
        if number is None:
            continue
        enriched = {**chart, "figure_no": number}
        groups.setdefault("all", []).append(enriched)
        groups.setdefault(chart_group(chart), []).append(enriched)

    tokens: dict[str, str] = {}
    group_tokens = {
        "all": "TREND",
        "long_term": "LONG_TERM_TREND",
        "accelerated": "ACCELERATED_TREND",
    }
    for group, prefix in group_tokens.items():
        group_charts = sorted(groups.get(group, []), key=lambda item: int(item["figure_no"]))
        numbers = [int(chart["figure_no"]) for chart in group_charts]
        titles = [str(chart.get("title") or "") for chart in group_charts if chart.get("title")]
        tokens[f"{{{{{prefix}_FIGURE_RANGE}}}}"] = figure_range(numbers)
        tokens[f"{{{{{prefix}_FIGURE_REFS}}}}"] = "、".join(figure_ref(number) for number in numbers)
        tokens[f"{{{{{prefix}_FIGURE_TITLES}}}}"] = "；".join(titles)
    return {"charts": charts, "groups": groups, "tokens": tokens}


def inferred_trend_group_for_section(section_key: str) -> str:
    key = section_key.lower()
    if key.startswith("long_term"):
        return "long_term"
    if key.startswith("accelerated"):
        return "accelerated"
    if key.startswith("stress"):
        return "stress_study"
    return "all"


def apply_trend_tokens(text: str, section_key: str, context: dict[str, Any]) -> tuple[str, list[str]]:
    tokens = dict(context.get("tokens") or {})
    group = inferred_trend_group_for_section(section_key)
    if group == "long_term":
        tokens["{{TREND_FIGURE_RANGE}}"] = tokens.get("{{LONG_TERM_TREND_FIGURE_RANGE}}", "")
        tokens["{{TREND_FIGURE_REFS}}"] = tokens.get("{{LONG_TERM_TREND_FIGURE_REFS}}", "")
        tokens["{{TREND_FIGURE_TITLES}}"] = tokens.get("{{LONG_TERM_TREND_FIGURE_TITLES}}", "")
    elif group == "accelerated":
        tokens["{{TREND_FIGURE_RANGE}}"] = tokens.get("{{ACCELERATED_TREND_FIGURE_RANGE}}", "")
        tokens["{{TREND_FIGURE_REFS}}"] = tokens.get("{{ACCELERATED_TREND_FIGURE_REFS}}", "")
        tokens["{{TREND_FIGURE_TITLES}}"] = tokens.get("{{ACCELERATED_TREND_FIGURE_TITLES}}", "")
    elif group == "stress_study":
        tokens["{{TREND_FIGURE_RANGE}}"] = ""
        tokens["{{TREND_FIGURE_REFS}}"] = ""
        tokens["{{TREND_FIGURE_TITLES}}"] = ""
    changed: list[str] = []
    result = text
    for token, value in tokens.items():
        if token in result:
            result = result.replace(token, value)
            changed.append(token)
    return result, changed


def render_body_sections(packet: dict[str, Any], trend_manifest: dict[str, Any]) -> tuple[dict[str, dict[str, Any]], dict[str, list[str]], dict[str, str]]:
    context = trend_body_context(trend_manifest)
    rendered: dict[str, dict[str, Any]] = {}
    token_counts: dict[str, list[str]] = {}
    for key, section in body_sections(packet).items():
        text, changed = apply_trend_tokens(str(section.get("text") or ""), key, context)
        rendered[key] = {**section, "text": text}
        if changed:
            token_counts[key] = changed
    return rendered, token_counts, dict(context.get("tokens") or {})


def figure_caption(chart: dict[str, Any]) -> str:
    number = figure_no_value(chart)
    title = str(chart.get("caption") or chart.get("title") or "")
    if number is None:
        return title
    return f"{figure_ref(number)}{title}"


def grouped_trend_charts(trend_manifest: dict[str, Any]) -> dict[str, list[dict[str, Any]]]:
    groups: dict[str, list[dict[str, Any]]] = {"long_term": [], "accelerated": []}
    for chart in trend_manifest.get("charts", []):
        if not isinstance(chart, dict):
            continue
        group = chart_group(chart)
        if group in groups:
            groups[group].append(chart)
    for group in groups:
        groups[group].sort(key=lambda item: figure_no_value(item) or 0)
    return groups


def render_figure_blocks_in_docx(
    trend_manifest: dict[str, Any],
    source_docx: Path,
    output_docx: Path,
    report_path: Path,
) -> dict[str, Any]:
    doc = Document(source_docx)
    groups = grouped_trend_charts(trend_manifest)
    generated: list[dict[str, Any]] = []
    warnings: list[str] = []
    placeholders = {
        "long_term": LONG_TERM_FIGURE_BLOCK_PLACEHOLDER,
        "accelerated": ACCELERATED_FIGURE_BLOCK_PLACEHOLDER,
    }
    for group, placeholder in placeholders.items():
        anchor = find_placeholder_paragraph(doc, placeholder)
        if anchor is None:
            if groups.get(group):
                warnings.append(f"Figure block placeholder not found for {group}: {placeholder}")
            continue
        for chart in groups.get(group, []):
            caption = figure_caption(chart)
            insert_paragraph_before(anchor, caption, style="Caption")
            image_paragraph = insert_paragraph_before(anchor, "", style="Normal")
            image_path = Path(str(chart.get("path") or ""))
            if image_path.exists():
                image_paragraph.add_run().add_picture(str(image_path), width=Inches(5.8))
            else:
                warnings.append(f"Trend chart image missing for figure block: {image_path}")
            generated.append(
                {
                    "group": group,
                    "figure_no": chart.get("figure_no"),
                    "caption": caption,
                    "path": str(image_path),
                    "placeholder": placeholder,
                }
            )
        remove_paragraph(anchor)

    directory_anchor = find_placeholder_paragraph(doc, FIGURE_DIRECTORY_PLACEHOLDER)
    if directory_anchor is None:
        if generated:
            warnings.append(f"Figure directory placeholder not found: {FIGURE_DIRECTORY_PLACEHOLDER}")
    else:
        for record in generated:
            insert_paragraph_before(directory_anchor, str(record["caption"]), style=directory_anchor.style.name)
        remove_paragraph(directory_anchor)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)
    report = {
        "schema_version": "ctd-32s73-figure-render-v1",
        "input_docx": str(source_docx),
        "output_docx": str(output_docx),
        "generated_figures": generated,
        "warnings": warnings,
        "policy": "Render only long-term and accelerated trend charts; influence-factor studies do not generate trend figures.",
    }
    write_json(report_path, report)
    return report


def warning_lines(validation_json: Path) -> list[str]:
    if not validation_json.exists():
        return [f"Validation JSON was not generated: {validation_json}"]
    report = read_json(validation_json)
    warnings = report.get("all_warnings") or report.get("warnings") or []
    return [str(warning) for warning in warnings]


def build_validation_report(
    output_docx: Path,
    evidence_summary: Path,
    fact_packet_validation: Path,
    table_render_manifest: Path,
    trend_chart_manifest: Path | None,
    structure_snapshot: Path,
    validation_json: Path,
    first_pass_warnings: list[str],
) -> str:
    timestamp = datetime.now(timezone.utc).isoformat()
    lines = [
        "# Validation Report",
        "",
        f"Generated at: {timestamp}",
        "",
        "## Output bundle",
        "",
        f"- Filled DOCX: {output_docx}",
        f"- Evidence summary: {evidence_summary}",
        f"- Fact-packet validation: {fact_packet_validation}",
        f"- Table render manifest: {table_render_manifest}",
        f"- Trend chart manifest: {trend_chart_manifest}" if trend_chart_manifest else "- Trend chart manifest: not generated",
        f"- Structure snapshot: {structure_snapshot}",
        f"- Validation JSON: {validation_json}",
        "",
        "## Pipeline boundary",
        "",
        "- This report is generated by `scripts/run_generation_pipeline.py` after fact-packet validation, table rendering, structure inspection, and DOCX validation.",
        "- Remaining warnings must be resolved by project-specific generation code or explicitly justified as expected differences with source evidence.",
        "",
        "## Validation warnings",
        "",
    ]
    if first_pass_warnings:
        for warning in first_pass_warnings:
            lines.extend(
                [
                    f"- Warning: {warning}",
                    "  Explanation: surfaced by the reusable validator; resolve it or mark it expected only when the fact packet and source evidence justify the difference.",
                ]
            )
    else:
        lines.append("- No warnings reported by the first validation pass.")
    lines.append("")
    return "\n".join(lines)


def build_validate_command(
    args: argparse.Namespace,
    packet: dict[str, Any],
    output_docx: Path,
    template_path: Path,
    evidence_summary: Path,
    table_render_manifest: Path,
    trend_chart_manifest: Path | None,
    validation_json: Path,
    validation_report: Path | None,
) -> list[str]:
    command = [
        sys.executable,
        str(VALIDATE_DOCX_SCRIPT),
        str(output_docx),
        "--fact-packet",
        str(args.fact_packet),
        "--template-docx",
        str(template_path),
        "--evidence-summary",
        str(evidence_summary),
        "--table-render-manifest",
        str(table_render_manifest),
        "--json-out",
        str(validation_json),
    ]
    if trend_chart_manifest is not None and trend_chart_manifest.exists():
        command.extend(["--trend-chart-manifest", str(trend_chart_manifest)])
    if not args.skip_caption_reference_check:
        command.append("--check-caption-references")
    if args.disallow_hour_time:
        command.append("--disallow-hour-time")
    if args.min_tables is not None:
        command.extend(["--min-tables", str(args.min_tables)])
    if args.max_tables is not None:
        command.extend(["--max-tables", str(args.max_tables)])
    if validation_report is not None:
        command.extend(["--validation-report", str(validation_report)])

    required_sources = [source_path(source) for source in packet_sources(packet, "allowed")]
    required_sources.extend(args.required_source)
    for source in required_sources:
        if source:
            command.extend(["--required-source", source])

    forbidden_sources = [source_path(source) for source in packet_sources(packet, "forbidden")]
    forbidden_sources.extend(args.forbidden_source)
    for source in forbidden_sources:
        if source:
            command.extend(["--forbidden-source", source])

    if not args.skip_expected_batch_checks:
        for batch in batch_numbers(packet):
            command.extend(["--expected-batch", batch])
    for batch in args.expected_batch:
        command.extend(["--expected-batch", batch])

    for placeholder in args.allow_placeholder:
        command.extend(["--allow-placeholder", placeholder])
    for warning in args.expected_warning:
        command.extend(["--expected-warning", warning])

    return command


def copy_fact_packet(source: Path, destination: Path) -> Path:
    source = source.resolve()
    destination = destination.resolve()
    if source != destination:
        shutil.copyfile(source, destination)
    return destination


def copy_starting_docx(source: Path, destination: Path) -> None:
    if source.resolve() != destination.resolve():
        shutil.copyfile(source, destination)


def main() -> None:
    require_internal_context("run_generation_pipeline.py")
    parser = argparse.ArgumentParser(
        description=(
            "Run the reusable CTD 3.2.S.7.3 post-extraction pipeline from a fact-packet.json. "
            "This is an internal implementation; use scripts/skill_step.py for generation, resume, validation, and delivery decisions."
        )
    )
    parser.add_argument("--fact-packet", type=Path, required=True)
    parser.add_argument("--output-dir", type=Path, required=True)
    parser.add_argument("--template", type=Path, default=DEFAULT_TEMPLATE)
    parser.add_argument("--template-id", default=DEFAULT_TEMPLATE_ID)
    parser.add_argument(
        "--working-docx",
        type=Path,
        help="Optional project-specific DOCX after body/skeleton edits; defaults to the bundled template.",
    )
    parser.add_argument("--filled-name", default="filled.docx")
    parser.add_argument("--skip-table-render", action="store_true")
    parser.add_argument("--skip-caption-reference-check", action="store_true")
    parser.add_argument("--skip-expected-batch-checks", action="store_true")
    parser.add_argument("--disallow-hour-time", action="store_true")
    parser.add_argument("--min-tables", type=int)
    parser.add_argument("--max-tables", type=int)
    parser.add_argument("--required-source", action="append", default=[])
    parser.add_argument("--forbidden-source", action="append", default=[])
    parser.add_argument("--expected-batch", action="append", default=[])
    parser.add_argument("--allow-placeholder", action="append", default=[])
    parser.add_argument("--expected-warning", action="append", default=[])
    parser.add_argument("--max-recovery-attempts", type=int, default=2)
    parser.add_argument("--fail-on-validation-warnings", action="store_true")
    args = parser.parse_args()

    output_dir = args.output_dir.resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    template_path = args.template.resolve()
    starting_docx = (args.working_docx or template_path).resolve()

    fact_packet_path = copy_fact_packet(args.fact_packet, output_dir / "fact-packet.json")
    fact_packet_validation = output_dir / "fact-packet-validation.json"
    fact_packet_validation_output = output_dir / "fact-packet-validation-output.txt"
    evidence_summary = output_dir / "evidence-summary.md"
    missing_evidence_report = output_dir / "missing-evidence-report.md"
    stress_sections_docx = output_dir / "stress-sections-rendered.docx"
    stress_section_render = output_dir / "stress-section-render.json"
    body_sections_docx = output_dir / "body-sections-applied.docx"
    body_section_application = output_dir / "body-section-application.json"
    table_render_config = output_dir / "table-render-config.json"
    table_render_manifest = output_dir / "table-render-manifest.json"
    table_render_output = output_dir / "table-render-output.txt"
    trend_chart_input = output_dir / "trend-chart-input.json"
    trend_chart_dir = output_dir / "trend-charts"
    trend_chart_manifest = output_dir / "trend-chart-manifest.json"
    trend_chart_output = output_dir / "trend-chart-output.txt"
    figure_render_manifest = output_dir / "figure-render-manifest.json"
    output_docx = output_dir / args.filled_name
    structure_snapshot = output_dir / "structure-snapshot.txt"
    validation_report = output_dir / "validation-report.md"
    validation_json = output_dir / "validation.json"
    validation_first_output = output_dir / "validation-first-pass-output.txt"
    validation_output = output_dir / "validation-output.txt"
    generation_manifest = output_dir / "generation-manifest.json"

    packet = read_json(fact_packet_path)
    write_text(missing_evidence_report, build_missing_evidence_report(packet))

    command_results: dict[str, Any] = {}
    require_trend_charts = bool(((packet.get("trend_charts") or {}).get("charts") or []))
    require_body_sections = bool(body_sections(packet))
    command_results["fact_packet_validation"] = run_command(
        [
            sys.executable,
            str(VALIDATE_FACT_SCRIPT),
            str(fact_packet_path),
            "--json-out",
            str(fact_packet_validation),
            "--max-recovery-attempts",
            str(args.max_recovery_attempts),
            *(["--require-trend-charts"] if require_trend_charts else []),
            *(["--require-body-sections"] if require_body_sections else []),
        ],
        fact_packet_validation_output,
    )
    if command_results["fact_packet_validation"]["returncode"] != 0:
        write_json(
            generation_manifest,
            {
                "status": "failed_fact_packet_validation",
                "output_dir": str(output_dir),
                "fact_packet": str(fact_packet_path),
                "fact_packet_validation": str(fact_packet_validation),
                "commands": command_results,
            },
        )
        print(json.dumps({"status": "failed_fact_packet_validation", "manifest": str(generation_manifest)}, ensure_ascii=False, indent=2))
        raise SystemExit(1)

    trend_chart_config = normalize_trend_chart_config(packet)
    if trend_chart_config is not None:
        write_json(trend_chart_input, trend_chart_config)
        command_results["trend_charts"] = run_command(
            [
                sys.executable,
                str(TREND_CHART_SCRIPT),
                str(trend_chart_input),
                "--output-dir",
                str(trend_chart_dir),
                "--manifest-out",
                str(trend_chart_manifest),
            ],
            trend_chart_output,
        )
        if command_results["trend_charts"]["returncode"] != 0:
            write_json(
                generation_manifest,
                {
                    "status": "failed_trend_chart_generation",
                    "output_dir": str(output_dir),
                    "fact_packet": str(fact_packet_path),
                    "trend_chart_input": str(trend_chart_input),
                    "commands": command_results,
                },
            )
            print(json.dumps({"status": "failed_trend_chart_generation", "manifest": str(generation_manifest)}, ensure_ascii=False, indent=2))
            raise SystemExit(1)
        trend_manifest_data = read_json_optional(trend_chart_manifest)
    else:
        trend_chart_manifest = None
        trend_manifest_data = {}

    stress_section_report = render_stress_sections_in_docx(packet, starting_docx, stress_sections_docx, stress_section_render)
    body_section_report = apply_body_sections_to_docx(packet, trend_manifest_data, stress_sections_docx, body_sections_docx, body_section_application)
    render_starting_docx = body_sections_docx
    write_text(
        evidence_summary,
        build_evidence_summary(packet, template_path, args.template_id, starting_docx, render_starting_docx, fact_packet_path),
    )

    if args.skip_table_render:
        copy_starting_docx(render_starting_docx, output_docx)
        write_json(
            table_render_manifest,
            {
                "renderer": "run_generation_pipeline.py",
                "template_docx": str(render_starting_docx),
                "output_docx": str(output_docx),
                "template_contract": args.template_id,
                "render_scope": render_scope(packet),
                "tables_rendered": [],
                "warnings": ["Table rendering skipped by --skip-table-render."],
            },
        )
    else:
        config = build_table_render_config(packet, render_starting_docx, output_docx, table_render_manifest, args.template_id)
        write_json(table_render_config, config)
        command_results["table_render"] = run_command(
            [sys.executable, str(RENDER_SCRIPT), str(table_render_config)],
            table_render_output,
        )
        if command_results["table_render"]["returncode"] != 0:
            write_json(
                generation_manifest,
                {
                    "status": "failed_table_render",
                    "output_dir": str(output_dir),
                    "fact_packet": str(fact_packet_path),
                    "table_render_config": str(table_render_config),
                    "commands": command_results,
                },
            )
            print(json.dumps({"status": "failed_table_render", "manifest": str(generation_manifest)}, ensure_ascii=False, indent=2))
            raise SystemExit(1)

    figure_render_report = render_figure_blocks_in_docx(trend_manifest_data, output_docx, output_docx, figure_render_manifest)

    command_results["structure_snapshot"] = run_command(
        [
            sys.executable,
            str(INSPECT_SCRIPT),
            str(output_docx),
            "--caption-paragraphs",
            "--all-tables",
            "--max-rows",
            "2",
        ],
        structure_snapshot,
    )

    first_validate_command = build_validate_command(
        args,
        packet,
        output_docx,
        template_path,
        evidence_summary,
        table_render_manifest,
        trend_chart_manifest,
        validation_json,
        None,
    )
    command_results["validation_first_pass"] = run_command(first_validate_command, validation_first_output)
    first_pass_warnings = warning_lines(validation_json)
    write_text(
        validation_report,
        build_validation_report(
            output_docx,
            evidence_summary,
            fact_packet_validation,
            table_render_manifest,
            trend_chart_manifest,
            structure_snapshot,
            validation_json,
            first_pass_warnings,
        ),
    )

    final_validate_command = build_validate_command(
        args,
        packet,
        output_docx,
        template_path,
        evidence_summary,
        table_render_manifest,
        trend_chart_manifest,
        validation_json,
        validation_report,
    )
    command_results["validation"] = run_command(final_validate_command, validation_output)
    final_validation = read_json(validation_json) if validation_json.exists() else {}
    validation_passed = bool(final_validation.get("passed"))

    status = "completed_validated" if validation_passed else "completed_with_validation_warnings"

    write_json(
        generation_manifest,
        {
            "status": status,
            "validation_passed": validation_passed,
            "output_dir": str(output_dir),
            "template": str(template_path),
            "template_id": args.template_id,
            "starting_docx": str(starting_docx),
            "render_starting_docx": str(render_starting_docx),
            "fact_packet": str(fact_packet_path),
            "fact_packet_validation": str(fact_packet_validation),
            "stress_section_render": str(stress_section_render),
            "stress_sections_docx": str(stress_sections_docx),
            "stress_sections": stress_section_report.get("generated_sections"),
            "body_section_application": str(body_section_application),
            "body_section_counts": body_section_report.get("body_section_counts"),
            "figure_render_manifest": str(figure_render_manifest),
            "figure_render_warnings": figure_render_report.get("warnings"),
            "filled_docx": str(output_docx),
            "evidence_summary": str(evidence_summary),
            "missing_evidence_report": str(missing_evidence_report),
            "table_render_config": str(table_render_config) if table_render_config.exists() else None,
            "table_render_manifest": str(table_render_manifest),
            "trend_chart_input": str(trend_chart_input) if trend_chart_input.exists() else None,
            "trend_chart_manifest": str(trend_chart_manifest) if trend_chart_manifest is not None and trend_chart_manifest.exists() else None,
            "trend_chart_dir": str(trend_chart_dir) if trend_chart_dir.exists() else None,
            "structure_snapshot": str(structure_snapshot),
            "validation_report": str(validation_report),
            "validation_json": str(validation_json),
            "commands": command_results,
            "pipeline_boundary": [
                "Consumes an existing fact-packet.json.",
                "Does not perform source extraction.",
                "Generates trend-chart-manifest.json before applying semantic body_sections.",
                "Generates dynamic influence-factor section blocks from stress_study facts.",
                "Applies semantic body_sections and trend chart reference tokens after dynamic influence-factor section rendering and before table rendering.",
                "Renders long-term and accelerated trend figure blocks after table rendering.",
                "Generates table blocks from structural placeholders and table_render_inputs.",
                "Generates static table and figure directory entries; does not refresh Word fields.",
            ],
            "content_provenance_policy": [
                "Project-specific facts and conclusions in filled.docx must be represented in fact-packet.json.",
                "Template, writing patterns, and scripts provide structure, style, fixed labels, table block placeholders, captions, and expression patterns only.",
                "A project-specific starting_docx must not introduce product names, batches, conditions, timepoints, results, trends, figure/table references, or conclusions outside fact-packet.json.",
            ],
        },
    )

    print(
        json.dumps(
            {
                "status": status,
                "validation_passed": validation_passed,
                "manifest": str(generation_manifest),
                "filled_docx": str(output_docx),
            },
            ensure_ascii=False,
            indent=2,
        )
    )
    if args.fail_on_validation_warnings and not validation_passed:
        raise SystemExit(2)


# Internal module — do not invoke directly. Use skill_step.py as the public entry point.
if __name__ == "__main__":
    main()

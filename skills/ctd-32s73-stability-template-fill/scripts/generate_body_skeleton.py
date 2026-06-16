from __future__ import annotations

import argparse
import json
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from runtime_guard import require_internal_context

from docx import Document


def read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def write_json(path: Path, data: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def project_profile(packet: dict[str, Any]) -> dict[str, Any]:
    profile = packet.get("project_profile")
    return profile if isinstance(profile, dict) else {}


def batches(packet: dict[str, Any]) -> list[dict[str, Any]]:
    values = project_profile(packet).get("batches")
    return [item for item in values if isinstance(item, dict)] if isinstance(values, list) else []


def batch_no_list(packet: dict[str, Any]) -> list[str]:
    return [str(batch.get("batch_no")) for batch in batches(packet) if batch.get("batch_no")]


def condition(packet: dict[str, Any], study_key: str) -> str:
    study = packet.get(study_key)
    if isinstance(study, dict) and study.get("condition"):
        return str(study["condition"])
    return ""


def max_completed_timepoint(packet: dict[str, Any], study_key: str) -> str:
    study = packet.get(study_key)
    if not isinstance(study, dict):
        return ""
    completed = study.get("completed_timepoints_by_batch")
    points: list[str] = []
    if isinstance(completed, dict):
        for values in completed.values():
            if isinstance(values, list):
                points.extend(str(value) for value in values)
    return points[-1] if points else ""


def render_scope(packet: dict[str, Any]) -> dict[str, Any]:
    plan = packet.get("docx_render_plan")
    if not isinstance(plan, dict):
        return {}
    scope = plan.get("render_scope")
    return scope if isinstance(scope, dict) else {}


def batch_structure_action(packet: dict[str, Any]) -> str:
    plan = packet.get("docx_render_plan")
    if not isinstance(plan, dict):
        return ""
    scope = render_scope(packet)
    if scope.get("batch_structure_action"):
        return str(scope.get("batch_structure_action"))
    return str(plan.get("batch_count_action") or "")


def rendered_batch_numbers(packet: dict[str, Any]) -> list[str]:
    scope = render_scope(packet)
    rendered = scope.get("rendered_batches")
    if isinstance(rendered, list) and rendered:
        numbers = []
        for item in rendered:
            if isinstance(item, dict) and item.get("batch_no"):
                numbers.append(str(item["batch_no"]))
            elif item:
                numbers.append(str(item))
        return numbers
    return batch_no_list(packet)


def placeholder_for_section(section_key: str) -> str:
    return "{{" + section_key.upper() + "}}"


def body_sections(packet: dict[str, Any]) -> dict[str, dict[str, Any]]:
    values = packet.get("body_sections")
    if not isinstance(values, dict):
        plan = packet.get("docx_render_plan")
        values = plan.get("body_sections") if isinstance(plan, dict) else {}
    if not isinstance(values, dict):
        return {}
    sections: dict[str, dict[str, Any]] = {}
    for key, value in values.items():
        section_key = str(key)
        if isinstance(value, str):
            sections[section_key] = {
                "placeholder": placeholder_for_section(section_key),
                "text": value,
            }
            continue
        if not isinstance(value, dict):
            continue
        text_value = value.get("text", value.get("content", ""))
        if isinstance(text_value, list):
            text = "\n".join(str(item) for item in text_value)
        else:
            text = str(text_value or "")
        sections[section_key] = {
            **value,
            "placeholder": str(value.get("placeholder") or placeholder_for_section(section_key)),
            "text": text,
        }
    return sections


def replacement_map(packet: dict[str, Any]) -> dict[str, str]:
    profile = project_profile(packet)
    project_code = str(profile.get("project_code") or "项目")
    display_name = str(profile.get("product_name") or profile.get("product_expression") or project_code)
    sample_type = str(profile.get("sample_type") or display_name)
    batch_numbers = batch_no_list(packet)
    first_batch = batch_numbers[0] if batch_numbers else project_code
    second_batch = batch_numbers[1] if len(batch_numbers) > 1 else first_batch
    long_condition = condition(packet, "long_term") or "来源条件"
    accelerated_condition = condition(packet, "accelerated") or "来源条件"
    long_completed = max_completed_timepoint(packet, "long_term") or "已完成"
    accelerated_completed = max_completed_timepoint(packet, "accelerated") or "已完成"

    return {
        "IPXXX-4": first_batch,
        "IPXXX-5": second_batch,
        "IPXXX": project_code,
        "XXX-2": display_name,
        "单抗/原液": sample_type,
        "XXX℃ ± XXX℃": long_condition,
        "XXX℃±XXX℃": long_condition,
        "XXX个月": f"{long_completed}个月",
        "XX个月": f"{accelerated_completed}个月",
        "XXX": project_code,
    }


KNOWN_REFERENCE_FIXES: dict[str, str] = {}


def replace_text(text: str, replacements: dict[str, str]) -> tuple[str, list[str]]:
    changed: list[str] = []
    result = text
    for old, new in {**KNOWN_REFERENCE_FIXES, **replacements}.items():
        if old in result:
            result = result.replace(old, new)
            changed.append(old)
    return result, changed


def replace_paragraph_text(paragraph, replacements: dict[str, str]) -> list[str]:
    original = paragraph.text
    new_text, changed = replace_text(original, replacements)
    if not changed or new_text == original:
        return []
    paragraph.clear()
    paragraph.add_run(new_text)
    return changed


def replace_body_section_paragraph(paragraph, sections: dict[str, dict[str, Any]]) -> list[str]:
    original = paragraph.text
    changed: list[str] = []
    new_text = original
    for key, section in sections.items():
        placeholder = str(section.get("placeholder") or placeholder_for_section(key))
        if placeholder in new_text:
            new_text = new_text.replace(placeholder, str(section.get("text") or ""))
            changed.append(key)
    if not changed or new_text == original:
        return []
    paragraph.clear()
    paragraph.add_run(new_text)
    return changed


def replace_cell_text(cell, replacements: dict[str, str], sections: dict[str, dict[str, Any]]) -> tuple[list[str], list[str]]:
    changed: list[str] = []
    section_changed: list[str] = []
    for paragraph in cell.paragraphs:
        section_changed.extend(replace_body_section_paragraph(paragraph, sections))
        changed.extend(replace_paragraph_text(paragraph, replacements))
    return changed, section_changed


def replace_document_text(doc: Document, replacements: dict[str, str], sections: dict[str, dict[str, Any]]) -> tuple[dict[str, int], dict[str, int]]:
    counts: dict[str, int] = {}
    section_counts: dict[str, int] = {}
    for paragraph in doc.paragraphs:
        for key in replace_body_section_paragraph(paragraph, sections):
            section_counts[key] = section_counts.get(key, 0) + 1
        for key in replace_paragraph_text(paragraph, replacements):
            counts[key] = counts.get(key, 0) + 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_counts, cell_section_counts = replace_cell_text(cell, replacements, sections)
                for key in cell_section_counts:
                    section_counts[key] = section_counts.get(key, 0) + 1
                for key in cell_counts:
                    counts[key] = counts.get(key, 0) + 1
    return counts, section_counts


def generate_body_skeleton(fact_packet: Path, template: Path, output_docx: Path, report_out: Path) -> dict[str, Any]:
    packet = read_json(fact_packet)
    doc = Document(template)
    replacements = replacement_map(packet)
    sections = body_sections(packet)
    replacement_counts, _ = replace_document_text(doc, replacements, {})
    body_section_counts: dict[str, int] = {}
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx)
    report = {
        "schema_version": "ctd-32s73-body-skeleton-report-v2",
        "generator": "scripts/generate_body_skeleton.py",
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "input_fact_packet": str(fact_packet),
        "template": str(template),
        "output_docx": str(output_docx),
        "render_scope": render_scope(packet),
        "replacement_counts": replacement_counts,
        "body_section_counts": body_section_counts,
        "body_section_placeholders": {
            key: section.get("placeholder") for key, section in sections.items()
        },
        "batch_numbers": batch_no_list(packet),
        "table_block_policy": "Table blocks remain structural placeholders in the working DOCX. render_stability_tables.py replaces them and records generated_table_slots in table-render-manifest.json.",
        "body_section_policy": "Semantic body_sections are intentionally deferred to run_generation_pipeline.py so skeleton generation only applies legacy text placeholders.",
        "stress_section_policy": "Influence-factor subsections are not pruned here. run_generation_pipeline.py materializes {{STRESS_SECTION_BLOCKS}} from fact-packet stress_study facts.",
        "warnings": [
            "Generic body/skeleton generation applies legacy placeholders only; semantic body_sections, dynamic influence-factor sections, tables, and figures are applied later by run_generation_pipeline.py."
        ],
    }
    write_json(report_out, report)
    return report


def main() -> None:
    require_internal_context("generate_body_skeleton.py")
    parser = argparse.ArgumentParser(description="Generate a project body/skeleton DOCX for CTD 3.2.S.7.3.")
    parser.add_argument("--workflow-state", type=Path)
    parser.add_argument("--output-dir", type=Path)
    parser.add_argument("--fact-packet", type=Path, required=True)
    parser.add_argument("--template", type=Path, required=True)
    parser.add_argument("--output-docx", type=Path, required=True)
    parser.add_argument("--manifest-out", type=Path, required=True, help="Compatibility name for the body skeleton report output path.")
    args = parser.parse_args()

    report = generate_body_skeleton(
        args.fact_packet.resolve(),
        args.template.resolve(),
        args.output_docx.resolve(),
        args.manifest_out.resolve(),
    )
    print(json.dumps(report, ensure_ascii=False, indent=2))


# Internal module. It is orchestrated by scripts/skill_step.py through the FSM.
if __name__ == "__main__":
    main()

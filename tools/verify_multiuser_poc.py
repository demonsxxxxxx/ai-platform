#!/usr/bin/env python3
"""Concurrent multi-user POC verification for ai-platform LambChat compatibility."""

from __future__ import annotations

import argparse
import concurrent.futures
import hashlib
import json
import re
import subprocess
import sys
import time
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from urllib import error, request

if __package__ in {None, ""}:
    sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from app.foundation_runtime_concurrency import (  # noqa: E402
    FOUNDATION_RUNTIME_CONCURRENCY_SCHEMA,
    build_foundation_runtime_concurrency_readiness,
)
from app.public_context_keys import safe_public_context_pack_version
from app.validation import assert_safe_id
from tools.verify_poc_gate import _context_public_projection_findings, _context_snapshot_payload_summary


DEFAULT_API_URL = "http://127.0.0.1:8020"
DEFAULT_SAMPLE_DOCX = "/tmp/ai-platform-multiuser-poc-sample.docx"
DEFAULT_POSTGRES_CONTAINER = "ai-platform-postgres"
DEFAULT_POSTGRES_USER = "ai_platform"
DEFAULT_POSTGRES_DB = "ai_platform"
DEFAULT_REDIS_CONTAINER = "ai-platform-redis"
DEFAULT_QUEUE_PREFIX = "ai-platform:runs"
DEFAULT_TEST_TENANT_PREFIX = "frc-test-"
DOWNLOAD_RE = re.compile(r"/api/ai/artifacts/(?P<artifact_id>art_[A-Za-z0-9_]+)/download")
DEFAULT_FIXTURE_TOOL_ID = "frc-test-tool-permission-probe"
PSQL_COMMAND_TAG_RE = re.compile(
    r"^(?:"
    r"DO|"
    r"INSERT \d+ \d+|"
    r"UPDATE \d+|"
    r"DELETE \d+|"
    r"SELECT \d+|"
    r"CREATE [A-Z ]+|"
    r"ALTER [A-Z ]+|"
    r"DROP [A-Z ]+|"
    r"TRUNCATE TABLE"
    r")$"
)
TERMINAL_LAMBCHAT_STATUSES = {"completed", "error", "cancelled", "canceled"}
CANCEL_EFFECT_STATUSES = {"cancel_requested", "cancelled", "canceled"}
QUEUE_PROBE_SOURCES = {"redis_metadata", "admin_runtime_queue"}
SANDBOX_LEASE_PROBE_SOURCE = "runtime_run_detail"
SYNTHETIC_SANDBOX_LEASE_PROBE_SOURCE = "post_run_sandbox_lease_probe"
DENIED_OR_CONFLICT_HTTP_STATUSES = {401, 403, 404, 409}
FORBIDDEN_PUBLIC_TERMS = (
    "authorization",
    "bearer ",
    "database_url",
    "executor_private_payload",
    "private_payload",
    "raw_storage_key",
    "redis_url",
    "runtime_private_payload",
    "sandbox_workdir",
    "storage_key",
)
FORBIDDEN_PUBLIC_KEY_ALIASES = {"".join(ch for ch in term if ch.isalnum()).lower() for term in FORBIDDEN_PUBLIC_TERMS}
FORBIDDEN_PUBLIC_VALUE_PATTERNS = (
    re.compile(r"\bauthorization\s*[:=]", re.IGNORECASE),
    re.compile(r"\bbearer\s+[A-Za-z0-9._~+/=-]{8,}\b", re.IGNORECASE),
    re.compile(r"\b(?:api[_-]?key|client[_-]?secret|database[_-]?url|password|redis[_-]?url|token)\s*[:=]", re.IGNORECASE),
    re.compile(r"\bsk-[A-Za-z0-9][A-Za-z0-9_-]{8,}\b", re.IGNORECASE),
)


@dataclass(frozen=True)
class Account:
    label: str
    username: str
    password: str
    tenant_id: str = "default"


@dataclass(frozen=True)
class CaseSpec:
    account: Account
    case_name: str
    scenario: str
    agent_id: str
    skill_id: str
    message: str
    uses_docx: bool
    workspace_id: str = "default"
    retry_source_run_id: str = ""


def ensure_default_sample_docx(docx_path: Path) -> Path:
    if docx_path.exists():
        return docx_path
    if str(docx_path) != DEFAULT_SAMPLE_DOCX:
        raise FileNotFoundError(f"sample docx not found: {docx_path}")
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        from docx import Document

        document = Document()
        document.add_heading("AI Platform POC Sample", level=1)
        document.add_paragraph("This document contains text for concurrent review and translation validation.")
        document.add_paragraph("请将这段中文内容翻译为英文，并保留原始含义。")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Field"
        table.cell(0, 1).text = "Value"
        table.cell(1, 0).text = "Purpose"
        table.cell(1, 1).text = "Multi-user POC validation"
        document.save(docx_path)
    except Exception:
        write_minimal_docx(docx_path)
    return docx_path


def write_minimal_docx(docx_path: Path) -> None:
    document_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    <w:p><w:r><w:t>AI Platform POC Sample</w:t></w:r></w:p>
    <w:p><w:r><w:t>This document contains text for concurrent review and translation validation.</w:t></w:r></w:p>
    <w:p><w:r><w:t>请将这段中文内容翻译为英文，并保留原始含义。</w:t></w:r></w:p>
    <w:sectPr/>
  </w:body>
</w:document>
"""
    with zipfile.ZipFile(docx_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr(
            "[Content_Types].xml",
            """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>
""",
        )
        archive.writestr(
            "_rels/.rels",
            """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>
""",
        )
        archive.writestr("word/document.xml", document_xml)


def json_request(method: str, url: str, payload: dict[str, Any] | None = None, headers: dict[str, str] | None = None, timeout: float = 30.0) -> tuple[int, Any]:
    data = json.dumps(payload or {}, ensure_ascii=False).encode("utf-8") if payload is not None else None
    req = request.Request(
        url,
        data=data,
        headers={"Accept": "application/json", "Content-Type": "application/json", **(headers or {})},
        method=method,
    )
    try:
        with request.urlopen(req, timeout=timeout) as response:
            status, body = response.status, response.read()
    except error.HTTPError as exc:
        status, body = exc.code, exc.read()
    try:
        return status, json.loads(body.decode("utf-8"))
    except Exception:
        return status, body.decode("utf-8", errors="replace")


def psql_json_rows(
    *,
    container: str,
    db_user: str,
    db_name: str,
    sql: str,
    timeout_seconds: float = 30.0,
) -> list[dict[str, Any]]:
    command = [
        "sudo",
        "-n",
        "docker",
        "exec",
        container,
        "psql",
        "-U",
        db_user,
        "-d",
        db_name,
        "-t",
        "-A",
        "-c",
        sql,
    ]
    completed = subprocess.run(command, check=False, capture_output=True, text=True, timeout=timeout_seconds)
    if completed.returncode != 0:
        raise RuntimeError(completed.stderr.strip() or completed.stdout.strip() or "psql query failed")
    rows: list[dict[str, Any]] = []
    for line in completed.stdout.splitlines():
        raw = line.strip()
        if not raw or PSQL_COMMAND_TAG_RE.match(raw):
            continue
        if not raw.startswith(("{", "[")):
            raise RuntimeError(f"unexpected psql output line: {raw}")
        rows.append(json.loads(raw))
    return rows


def redis_command(
    *,
    container: str,
    command: list[str],
    timeout_seconds: float = 30.0,
) -> list[str]:
    docker_command = [
        "sudo",
        "-n",
        "docker",
        "exec",
        container,
        "redis-cli",
        "--raw",
        *command,
    ]
    completed = subprocess.run(docker_command, check=False, capture_output=True, text=True, timeout=timeout_seconds)
    if completed.returncode != 0:
        raise RuntimeError(completed.stderr.strip() or completed.stdout.strip() or "redis command failed")
    return completed.stdout.splitlines()


def _sql_literal(value: str) -> str:
    return "'" + value.replace("'", "''") + "'"


def _safe_identifier(value: str, *, field_name: str) -> str:
    normalized = re.sub(r"[^A-Za-z0-9_.:-]+", "_", str(value).strip())
    normalized = normalized.strip("._:-")
    if not normalized or not re.fullmatch(r"[A-Za-z0-9][A-Za-z0-9_.:-]{0,127}", normalized):
        digest = hashlib.sha256(str(value).encode("utf-8")).hexdigest()[:12]
        normalized = f"{field_name}_{digest}"
    return normalized[:128]


def fixture_workspace_id(tenant_id: str) -> str:
    return _safe_identifier(f"{tenant_id}_default".replace("-", "_"), field_name="workspace")


def fixture_agent_id_for_skill(account: Account, skill_id: str) -> str:
    seed = f"{account.tenant_id}:{account.username}:{skill_id}"
    return _safe_identifier(f"frc_agent_{hashlib.sha256(seed.encode('utf-8')).hexdigest()[:16]}", field_name="agent")


def fixture_agent_id(account: Account) -> str:
    return fixture_agent_id_for_skill(account, "general-chat")


def _skill_id_for_agent(agent_id: str) -> str:
    return {
        "baoyu-translate": "baoyu-translate",
        "document-translation": "baoyu-translate",
        "qa-word-review": "qa-file-reviewer",
        "document-review": "qa-file-reviewer",
    }.get(agent_id, "general-chat")


def _agent_id_for_case(account: Account, agent_id: str, *, use_fixture_agents: bool) -> str:
    if not use_fixture_agents:
        return agent_id
    return fixture_agent_id_for_skill(account, _skill_id_for_agent(agent_id))


def fixture_retry_source_run_id(account: Account) -> str:
    return _safe_identifier(
        f"run_{account.tenant_id}_{account.username}_retry_source".replace("-", "_"),
        field_name="run",
    )


def fixture_retry_source_session_id(account: Account) -> str:
    return _safe_identifier(
        f"ses_{account.tenant_id}_{account.username}_retry_source".replace("-", "_"),
        field_name="ses",
    )


def fixture_context_snapshot_id(account: Account) -> str:
    return _safe_identifier(
        f"ctx_{account.tenant_id}_{account.username}_retry_source".replace("-", "_"),
        field_name="ctx",
    )


def fixture_sandbox_lease_id(account: Account) -> str:
    return _safe_identifier(
        f"lease_{account.tenant_id}_{account.username}_retry_source".replace("-", "_"),
        field_name="lease",
    )


def fixture_skill_snapshot_id(account: Account) -> str:
    return _safe_identifier(
        f"rss_{account.tenant_id}_{account.username}_retry_source".replace("-", "_"),
        field_name="rss",
    )


def _test_tenant_ids(tenant_ids: list[str], *, tenant_prefix: str) -> list[str]:
    clean = sorted({str(item).strip() for item in tenant_ids if str(item).strip()})
    if not clean:
        raise ValueError("cleanup requires at least one test tenant")
    if not tenant_prefix:
        raise ValueError("cleanup tenant prefix is required")
    unsafe = [tenant_id for tenant_id in clean if not tenant_id.startswith(tenant_prefix)]
    if unsafe:
        raise ValueError(f"cleanup only accepts test tenant ids with prefix {tenant_prefix!r}: {unsafe}")
    return clean


def build_foundation_runtime_cleanup_sql(
    tenant_ids: list[str],
    *,
    tenant_prefix: str = DEFAULT_TEST_TENANT_PREFIX,
) -> str:
    tenants = _test_tenant_ids(tenant_ids, tenant_prefix=tenant_prefix)
    tenant_array = ", ".join(_sql_literal(tenant_id) for tenant_id in tenants)
    return f"""
with target_tenants as (
  select unnest(array[{tenant_array}]::text[]) as tenant_id
),
deleted_run_events as (
  delete from run_events where tenant_id in (select tenant_id from target_tenants)
),
deleted_run_context_snapshots as (
  delete from run_context_snapshots where tenant_id in (select tenant_id from target_tenants)
),
deleted_run_tool_permission_requests as (
  delete from run_tool_permission_requests where tenant_id in (select tenant_id from target_tenants)
),
deleted_sandbox_leases as (
  delete from sandbox_leases where tenant_id in (select tenant_id from target_tenants)
),
deleted_run_skill_snapshots as (
  delete from run_skill_snapshots where tenant_id in (select tenant_id from target_tenants)
),
deleted_run_steps as (
  delete from run_steps where tenant_id in (select tenant_id from target_tenants)
),
deleted_artifacts as (
  delete from artifacts where tenant_id in (select tenant_id from target_tenants)
),
deleted_files as (
  delete from files where tenant_id in (select tenant_id from target_tenants)
),
deleted_messages as (
  delete from messages where tenant_id in (select tenant_id from target_tenants)
),
deleted_audit_logs as (
  delete from audit_logs where tenant_id in (select tenant_id from target_tenants)
),
deleted_memory_records as (
  delete from memory_records where tenant_id in (select tenant_id from target_tenants)
),
deleted_memory_policies as (
  delete from memory_policies where tenant_id in (select tenant_id from target_tenants)
),
deleted_runs as (
  delete from runs where tenant_id in (select tenant_id from target_tenants)
),
deleted_sessions as (
  delete from sessions where tenant_id in (select tenant_id from target_tenants)
),
deleted_tool_policies as (
  delete from tool_policies where tenant_id in (select tenant_id from target_tenants)
),
deleted_tenant_workbench_skills as (
  delete from tenant_workbench_skills where tenant_id in (select tenant_id from target_tenants)
),
deleted_agents as (
  delete from agents where tenant_id in (select tenant_id from target_tenants)
),
deleted_workspaces as (
  delete from workspaces where tenant_id in (select tenant_id from target_tenants)
),
deleted_users as (
  delete from users where tenant_id in (select tenant_id from target_tenants)
),
deleted_tenants as (
  delete from tenants where id in (select tenant_id from target_tenants)
)
select null::text where false;
""".strip()


def build_foundation_runtime_cleanup_count_sql(
    tenant_ids: list[str],
    *,
    tenant_prefix: str = DEFAULT_TEST_TENANT_PREFIX,
) -> str:
    tenants = _test_tenant_ids(tenant_ids, tenant_prefix=tenant_prefix)
    tenant_array = ", ".join(_sql_literal(tenant_id) for tenant_id in tenants)
    return f"""
with target_tenants as (
  select unnest(array[{tenant_array}]::text[]) as tenant_id
)
select json_build_object(
  'remaining_tenant_count', (select count(*) from tenants where id in (select tenant_id from target_tenants)),
  'remaining_run_count', (select count(*) from runs where tenant_id in (select tenant_id from target_tenants)),
  'remaining_artifact_count', (select count(*) from artifacts where tenant_id in (select tenant_id from target_tenants))
)::text;
""".strip()


def _redis_scalar(
    *,
    redis_container: str,
    command: list[str],
    default: str = "0",
) -> str:
    rows = redis_command(container=redis_container, command=command)
    return rows[0] if rows else default


def _json_payload_tenant_id(raw: str) -> str:
    try:
        payload = json.loads(raw)
    except (TypeError, json.JSONDecodeError):
        return ""
    if not isinstance(payload, dict):
        return ""
    return str(payload.get("tenant_id") or "")


def _json_payload_run_id(raw: str) -> str:
    try:
        payload = json.loads(raw)
    except (TypeError, json.JSONDecodeError):
        return ""
    if not isinstance(payload, dict):
        return ""
    return str(payload.get("run_id") or "")


def _hash_pairs(items: list[str]) -> list[tuple[str, str]]:
    return [(items[index], items[index + 1]) for index in range(0, len(items) - 1, 2)]


def _redis_hash_target_fields(
    *,
    redis_container: str,
    key: str,
    target_tenants: set[str],
) -> list[str]:
    if _redis_scalar(redis_container=redis_container, command=["TYPE", key], default="none") != "hash":
        return []
    fields: list[str] = []
    for field, value in _hash_pairs(redis_command(container=redis_container, command=["HGETALL", key])):
        if _json_payload_tenant_id(value) in target_tenants or field.split(":", 1)[0] in target_tenants:
            fields.append(field)
    return fields


def _message_ids_from_queued_index_value(raw: str) -> set[str]:
    try:
        decoded = json.loads(raw)
    except (TypeError, json.JSONDecodeError):
        decoded = raw
    if isinstance(decoded, list):
        return {str(item) for item in decoded if str(item)}
    if decoded:
        return {str(decoded)}
    return set()


def _count_target_list_items(
    *,
    redis_container: str,
    key: str,
    target_tenants: set[str],
) -> int:
    if _redis_scalar(redis_container=redis_container, command=["TYPE", key], default="none") != "list":
        return 0
    return sum(
        1
        for raw in redis_command(container=redis_container, command=["LRANGE", key, "0", "-1"])
        if _json_payload_tenant_id(raw) in target_tenants
    )


def cleanup_foundation_runtime_queue_residue(
    tenant_ids: list[str],
    *,
    redis_container: str = DEFAULT_REDIS_CONTAINER,
    tenant_prefix: str = DEFAULT_TEST_TENANT_PREFIX,
    queue_prefix: str = DEFAULT_QUEUE_PREFIX,
) -> dict[str, Any]:
    tenants = set(_test_tenant_ids(tenant_ids, tenant_prefix=tenant_prefix))
    keys = {
        "queued": f"{queue_prefix}:queued",
        "processing": f"{queue_prefix}:processing",
        "queued_meta": f"{queue_prefix}:queued-meta",
        "queued_run_index": f"{queue_prefix}:queued-run-index",
        "queued_order": f"{queue_prefix}:queued-order",
        "processing_meta": f"{queue_prefix}:processing-meta",
        "retry_meta": f"{queue_prefix}:retry-meta",
    }
    removed = {
        "queued_messages": 0,
        "processing_messages": 0,
        "queued_meta": 0,
        "queued_run_index": 0,
        "queued_order": 0,
        "processing_meta": 0,
        "retry_meta": 0,
    }
    message_ids: set[str] = set()
    queued_run_index_pairs = (
        _hash_pairs(redis_command(container=redis_container, command=["HGETALL", keys["queued_run_index"]]))
        if _redis_scalar(redis_container=redis_container, command=["TYPE", keys["queued_run_index"]], default="none")
        == "hash"
        else []
    )
    queued_run_index_by_field = dict(queued_run_index_pairs)

    for list_name in ("queued", "processing"):
        key = keys[list_name]
        if _redis_scalar(redis_container=redis_container, command=["TYPE", key], default="none") != "list":
            continue
        for raw in redis_command(container=redis_container, command=["LRANGE", key, "0", "-1"]):
            tenant_id = _json_payload_tenant_id(raw)
            if tenant_id not in tenants:
                continue
            message_ids.add(hashlib.sha256(raw.encode("utf-8")).hexdigest())
            run_id = _json_payload_run_id(raw)
            if run_id:
                message_ids.update(_message_ids_from_queued_index_value(queued_run_index_by_field.get(f"{tenant_id}:{run_id}", "")))
            result = redis_command(container=redis_container, command=["LREM", key, "0", raw])
            removed[f"{list_name}_messages"] += int(result[0]) if result and result[0].isdigit() else 0

    for hash_name in ("queued_meta", "processing_meta", "retry_meta"):
        for field in _redis_hash_target_fields(
            redis_container=redis_container,
            key=keys[hash_name],
            target_tenants=tenants,
        ):
            message_ids.add(field)
            result = redis_command(container=redis_container, command=["HDEL", keys[hash_name], field])
            removed[hash_name] += int(result[0]) if result and result[0].isdigit() else 0

    for field in _redis_hash_target_fields(
        redis_container=redis_container,
        key=keys["queued_run_index"],
        target_tenants=tenants,
    ):
        message_ids.update(_message_ids_from_queued_index_value(queued_run_index_by_field.get(field, "")))
        result = redis_command(container=redis_container, command=["HDEL", keys["queued_run_index"], field])
        removed["queued_run_index"] += int(result[0]) if result and result[0].isdigit() else 0

    if message_ids and _redis_scalar(redis_container=redis_container, command=["TYPE", keys["queued_order"]], default="none") == "zset":
        result = redis_command(container=redis_container, command=["ZREM", keys["queued_order"], *sorted(message_ids)])
        removed["queued_order"] += int(result[0]) if result and result[0].isdigit() else 0

    remaining_message_ids = {
        field
        for hash_name in ("queued_meta", "processing_meta", "retry_meta")
        for field in _redis_hash_target_fields(
            redis_container=redis_container,
            key=keys[hash_name],
            target_tenants=tenants,
        )
    }
    remaining_index_fields = _redis_hash_target_fields(
        redis_container=redis_container,
        key=keys["queued_run_index"],
        target_tenants=tenants,
    )
    remaining_message_ids.update(
        message_id
        for field in remaining_index_fields
        for message_id in _message_ids_from_queued_index_value(queued_run_index_by_field.get(field, ""))
    )
    remaining_queued_order_count = 0
    if remaining_message_ids and _redis_scalar(redis_container=redis_container, command=["TYPE", keys["queued_order"]], default="none") == "zset":
        for message_id in remaining_message_ids:
            if redis_command(container=redis_container, command=["ZSCORE", keys["queued_order"], message_id]):
                remaining_queued_order_count += 1

    remaining_counts = {
        "remaining_queued_count": _count_target_list_items(
            redis_container=redis_container,
            key=keys["queued"],
            target_tenants=tenants,
        ),
        "remaining_processing_count": _count_target_list_items(
            redis_container=redis_container,
            key=keys["processing"],
            target_tenants=tenants,
        ),
        "remaining_queued_meta_count": len(
            _redis_hash_target_fields(
                redis_container=redis_container,
                key=keys["queued_meta"],
                target_tenants=tenants,
            )
        ),
        "remaining_processing_meta_count": len(
            _redis_hash_target_fields(
                redis_container=redis_container,
                key=keys["processing_meta"],
                target_tenants=tenants,
            )
        ),
        "remaining_retry_meta_count": len(
            _redis_hash_target_fields(
                redis_container=redis_container,
                key=keys["retry_meta"],
                target_tenants=tenants,
            )
        ),
        "remaining_queued_run_index_count": len(remaining_index_fields),
        "remaining_queued_order_count": remaining_queued_order_count,
    }
    remaining_counts["remaining_queue_count"] = sum(remaining_counts.values())
    return {
        "status": "verified" if remaining_counts["remaining_queue_count"] == 0 else "remaining_records_detected",
        "redis_container": redis_container,
        "queue_prefix": queue_prefix,
        "removed_counts": removed,
        "remaining_counts": remaining_counts,
    }


def build_foundation_runtime_cleanup_proof(
    tenant_ids: list[str],
    *,
    postgres_container: str,
    postgres_user: str,
    postgres_db: str,
    redis_container: str = DEFAULT_REDIS_CONTAINER,
    tenant_prefix: str = DEFAULT_TEST_TENANT_PREFIX,
) -> dict[str, Any]:
    tenants = _test_tenant_ids(tenant_ids, tenant_prefix=tenant_prefix)
    psql_json_rows(
        container=postgres_container,
        db_user=postgres_user,
        db_name=postgres_db,
        sql=build_foundation_runtime_cleanup_sql(tenants, tenant_prefix=tenant_prefix),
    )
    rows = psql_json_rows(
        container=postgres_container,
        db_user=postgres_user,
        db_name=postgres_db,
        sql=build_foundation_runtime_cleanup_count_sql(tenants, tenant_prefix=tenant_prefix),
    )
    remaining = rows[0] if rows else {}
    remaining_counts = {
        "remaining_tenant_count": int(remaining.get("remaining_tenant_count") or 0),
        "remaining_run_count": int(remaining.get("remaining_run_count") or 0),
        "remaining_artifact_count": int(remaining.get("remaining_artifact_count") or 0),
    }
    redis_cleanup = cleanup_foundation_runtime_queue_residue(
        tenants,
        redis_container=redis_container,
        tenant_prefix=tenant_prefix,
    )
    remaining_counts["remaining_queue_count"] = int(
        redis_cleanup.get("remaining_counts", {}).get("remaining_queue_count") or 0
    )
    verified = all(value == 0 for value in remaining_counts.values())
    return {
        "schema_version": "ai-platform.foundation-runtime-cleanup-proof.v1",
        "status": "verified" if verified else "remaining_records_detected",
        "tenant_ids": tenants,
        "tenant_prefix": tenant_prefix,
        "remaining_counts": remaining_counts,
        "redis_cleanup": redis_cleanup,
    }


def build_foundation_runtime_fixture_sql(
    accounts: list[Account],
    *,
    tenant_prefix: str = DEFAULT_TEST_TENANT_PREFIX,
) -> str:
    tenants = _test_tenant_ids([account.tenant_id for account in accounts], tenant_prefix=tenant_prefix)
    unique_accounts = sorted(
        {(account.tenant_id, account.username, account.label) for account in accounts},
        key=lambda item: item,
    )
    if len({username for _tenant_id, username, _label in unique_accounts}) != len(unique_accounts):
        raise ValueError("foundation runtime fixture accounts require globally unique usernames")

    skill_names = {
        "general-chat": "General Chat",
        "qa-file-reviewer": "Document Review",
        "baoyu-translate": "Document Translation",
    }
    tenant_rows: list[str] = []
    workspace_rows: list[str] = []
    user_rows: list[str] = []
    workbench_rows: list[str] = []
    tool_policy_rows: list[str] = []
    agent_rows: list[str] = []
    session_rows: list[str] = []
    run_rows: list[str] = []
    event_rows: list[str] = []
    context_rows: list[str] = []
    sandbox_rows: list[str] = []
    skill_snapshot_rows: list[str] = []
    workspace_expectation_rows: list[str] = []
    user_expectation_rows: list[str] = []
    agent_expectation_rows: list[str] = []
    session_expectation_rows: list[str] = []
    run_expectation_rows: list[str] = []

    for tenant_id in tenants:
        workspace_id = fixture_workspace_id(tenant_id)
        tenant_rows.append(f"({_sql_literal(tenant_id)}, 'Foundation runtime test tenant')")
        workspace_rows.append(f"({_sql_literal(workspace_id)}, {_sql_literal(tenant_id)}, 'Foundation runtime test workspace')")
        workspace_expectation_rows.append(f"({_sql_literal(workspace_id)}, {_sql_literal(tenant_id)})")
        for skill_id in skill_names:
            workbench_rows.append(f"({_sql_literal(tenant_id)}, {_sql_literal(skill_id)}, 'active', true)")
        tool_policy_rows.append(
            f"({_sql_literal(tenant_id)}, {_sql_literal(DEFAULT_FIXTURE_TOOL_ID)}, 'active', false, 'low', true, 'Foundation runtime test tool policy')"
        )

    for tenant_id, username, label in unique_accounts:
        workspace_id = fixture_workspace_id(tenant_id)
        user_rows.append(f"({_sql_literal(username)}, {_sql_literal(tenant_id)}, {_sql_literal(label)})")
        user_expectation_rows.append(f"({_sql_literal(username)}, {_sql_literal(tenant_id)})")
        account = Account(label=label, username=username, password="", tenant_id=tenant_id)
        for skill_id, agent_name in skill_names.items():
            agent_id = fixture_agent_id_for_skill(account, skill_id)
            agent_rows.append(
                "("
                f"{_sql_literal(agent_id)}, {_sql_literal(tenant_id)}, {_sql_literal(agent_name + ' fixture')}, "
                f"'test', 'Foundation runtime isolated test agent', {_sql_literal(skill_id)}, 'active'"
                ")"
            )
            agent_expectation_rows.append(f"({_sql_literal(agent_id)}, {_sql_literal(tenant_id)})")

        session_id = fixture_retry_source_session_id(account)
        run_id = fixture_retry_source_run_id(account)
        agent_id = fixture_agent_id(account)
        trace_id = f"trace_{run_id}"
        context_snapshot_id = fixture_context_snapshot_id(account)
        sandbox_lease_id = fixture_sandbox_lease_id(account)
        skill_snapshot_id = fixture_skill_snapshot_id(account)
        input_json = {
            "input": {"message": "foundation runtime retry source"},
            "executor_type": "claude-agent-worker",
            "skill_version": "0.1.0",
            "release_decision": {"selected_version": "0.1.0", "policy_active": False},
        }
        result_json = {"message": "run_failed"}
        session_rows.append(
            "("
            f"{_sql_literal(session_id)}, {_sql_literal(tenant_id)}, {_sql_literal(workspace_id)}, "
            f"{_sql_literal(username)}, {_sql_literal(agent_id)}, 'Foundation retry source', 'active'"
            ")"
        )
        session_expectation_rows.append(
            "("
            f"{_sql_literal(session_id)}, {_sql_literal(tenant_id)}, {_sql_literal(workspace_id)}, "
            f"{_sql_literal(username)}, {_sql_literal(agent_id)}"
            ")"
        )
        run_rows.append(
            "("
            f"{_sql_literal(run_id)}, {_sql_literal(tenant_id)}, {_sql_literal(workspace_id)}, "
            f"{_sql_literal(session_id)}, {_sql_literal(username)}, {_sql_literal(agent_id)}, "
            f"'general-chat', {_sql_literal(trace_id)}, 'failed', "
            f"{_sql_literal(json.dumps(input_json, ensure_ascii=False))}::jsonb, "
            f"{_sql_literal(json.dumps(result_json, ensure_ascii=False))}::jsonb, "
            "'foundation_runtime_retry_source', 'run_failed', now(), now(), now()"
            ")"
        )
        run_expectation_rows.append(
            "("
            f"{_sql_literal(run_id)}, {_sql_literal(tenant_id)}, {_sql_literal(workspace_id)}, "
            f"{_sql_literal(session_id)}, {_sql_literal(username)}, {_sql_literal(agent_id)}"
            ")"
        )
        event_rows.extend(
            [
                "("
                f"{_sql_literal('evt_' + run_id + '_created')}, {_sql_literal(tenant_id)}, {_sql_literal(run_id)}, "
                f"{_sql_literal(trace_id)}, 1, 'run_created', 'control', 'Foundation runtime retry source created', true, "
                f"{_sql_literal(json.dumps({'visible_to_user': True}, ensure_ascii=False))}::jsonb"
                ")",
                "("
                f"{_sql_literal('evt_' + run_id + '_failed')}, {_sql_literal(tenant_id)}, {_sql_literal(run_id)}, "
                f"{_sql_literal(trace_id)}, 2, 'run_failed', 'executor', 'run_failed', true, "
                f"{_sql_literal(json.dumps({'visible_to_user': True, 'error_code': 'foundation_runtime_retry_source'}, ensure_ascii=False))}::jsonb"
                ")",
            ]
        )
        context_rows.append(
            "("
            f"{_sql_literal(context_snapshot_id)}, {_sql_literal(tenant_id)}, {_sql_literal(workspace_id)}, "
            f"{_sql_literal(username)}, {_sql_literal(session_id)}, {_sql_literal(run_id)}, {_sql_literal(trace_id)}, "
            "'executor', '[]'::jsonb, '[]'::jsonb, '[]'::jsonb, '[]'::jsonb, "
            f"{_sql_literal(json.dumps({'redacted': True}, ensure_ascii=False))}::jsonb, "
            f"{_sql_literal(json.dumps({'context_snapshot_id': context_snapshot_id, 'source': 'foundation_runtime_fixture'}, ensure_ascii=False))}::jsonb"
            ")"
        )
        sandbox_rows.append(
            "("
            f"{_sql_literal(sandbox_lease_id)}, {_sql_literal(tenant_id)}, {_sql_literal(workspace_id)}, "
            f"{_sql_literal(username)}, {_sql_literal(session_id)}, {_sql_literal(run_id)}, {_sql_literal(trace_id)}, "
            "'ephemeral', 'fake', 'released', false, '{}'::jsonb, "
            f"{_sql_literal(json.dumps({'workspace_id': workspace_id}, ensure_ascii=False))}::jsonb, "
            f"{_sql_literal(json.dumps({'source': 'foundation_runtime_fixture'}, ensure_ascii=False))}::jsonb, "
            "now(), now(), 'fixture_completed'"
            ")"
        )
        skill_snapshot_rows.append(
            "("
            f"{_sql_literal(skill_snapshot_id)}, {_sql_literal(tenant_id)}, {_sql_literal(run_id)}, "
            "'general-chat', '0.1.0', '0.1.0', "
            f"{_sql_literal(json.dumps({'kind': 'foundation_runtime_fixture'}, ensure_ascii=False))}::jsonb, "
            "'[]'::jsonb, true, true, true, 'fixture', false"
            ")"
        )

    run_id_filter = ", ".join(
        _sql_literal(fixture_retry_source_run_id(Account(label=label, username=username, password="", tenant_id=tenant_id)))
        for tenant_id, username, label in unique_accounts
    )
    tenant_filter = ", ".join(_sql_literal(tenant_id) for tenant_id in tenants)
    return f"""
do $$
begin
  if exists (
    select 1
    from workspaces
    join (values {", ".join(workspace_expectation_rows)}) as expected(id, tenant_id)
      on expected.id = workspaces.id
    where workspaces.tenant_id is distinct from expected.tenant_id
    limit 1
  ) then
    raise exception 'fixture_global_workspace_id_conflict';
  end if;
  if exists (
    select 1
    from users
    join (values {", ".join(user_expectation_rows)}) as expected(id, tenant_id)
      on expected.id = users.id
    where users.tenant_id is distinct from expected.tenant_id
    limit 1
  ) then
    raise exception 'fixture_global_user_id_conflict';
  end if;
  if exists (
    select 1
    from agents
    join (values {", ".join(agent_expectation_rows)}) as expected(id, tenant_id)
      on expected.id = agents.id
    where agents.tenant_id is distinct from expected.tenant_id
    limit 1
  ) then
    raise exception 'fixture_global_agent_id_conflict';
  end if;
  if exists (
    select 1
    from sessions
    join (values {", ".join(session_expectation_rows)})
      as expected(id, tenant_id, workspace_id, user_id, agent_id)
      on expected.id = sessions.id
    where sessions.tenant_id is distinct from expected.tenant_id
       or sessions.workspace_id is distinct from expected.workspace_id
       or sessions.user_id is distinct from expected.user_id
       or sessions.agent_id is distinct from expected.agent_id
    limit 1
  ) then
    raise exception 'fixture_global_session_id_conflict';
  end if;
  if exists (
    select 1
    from runs
    join (values {", ".join(run_expectation_rows)})
      as expected(id, tenant_id, workspace_id, session_id, user_id, agent_id)
      on expected.id = runs.id
    where runs.tenant_id is distinct from expected.tenant_id
       or runs.workspace_id is distinct from expected.workspace_id
       or runs.session_id is distinct from expected.session_id
       or runs.user_id is distinct from expected.user_id
       or runs.agent_id is distinct from expected.agent_id
    limit 1
  ) then
    raise exception 'fixture_global_run_id_conflict';
  end if;
end $$;

insert into tenants(id, name)
values {", ".join(tenant_rows)}
on conflict (id) do update set name = excluded.name, status = 'active';

insert into workspaces(id, tenant_id, name)
values {", ".join(workspace_rows)}
on conflict (id) do update set name = excluded.name, status = 'active';

insert into users(id, tenant_id, display_name)
values {", ".join(user_rows)}
on conflict (id) do nothing;

insert into tenant_workbench_skills(tenant_id, skill_id, status, visible_to_user)
values {", ".join(workbench_rows)}
on conflict (tenant_id, skill_id) do update set status = excluded.status, visible_to_user = excluded.visible_to_user;

insert into mcp_tools(
  id, server_id, name, description, transport_type, endpoint, auth_mode,
  allowed_tools, status, write_capable, risk_level, visible_to_user
)
values (
  {_sql_literal(DEFAULT_FIXTURE_TOOL_ID)},
  'foundation-runtime-fixture',
  'Foundation runtime tool permission probe',
  'Read-only fixture tool used only by the Foundation Runtime verifier.',
  'http',
  '',
  'platform-managed',
  '[]'::jsonb,
  'active',
  false,
  'low',
  true
)
on conflict (id) do update
set status = 'active',
    write_capable = false,
    risk_level = 'low',
    visible_to_user = true;

insert into tool_policies(tenant_id, tool_id, status, write_capable, risk_level, visible_to_user, reason)
values {", ".join(tool_policy_rows)}
on conflict (tenant_id, tool_id) do update
set status = excluded.status,
    write_capable = excluded.write_capable,
    risk_level = excluded.risk_level,
    visible_to_user = excluded.visible_to_user,
    reason = excluded.reason;

insert into agents(id, tenant_id, name, agent_type, description, default_skill_id, status)
values {", ".join(agent_rows)}
on conflict (id) do update
set name = excluded.name,
    agent_type = excluded.agent_type,
    description = excluded.description,
    default_skill_id = excluded.default_skill_id,
    status = excluded.status;

insert into sessions(id, tenant_id, workspace_id, user_id, agent_id, title, status)
values {", ".join(session_rows)}
on conflict (id) do update set title = excluded.title, status = excluded.status;

insert into runs(
  id, tenant_id, workspace_id, session_id, user_id, agent_id, skill_id, trace_id,
  status, input_json, result_json, error_code, error_message, queued_at, started_at, finished_at
)
values {", ".join(run_rows)}
on conflict (id) do update
set status = excluded.status,
    input_json = excluded.input_json,
    result_json = excluded.result_json,
    error_code = excluded.error_code,
    error_message = excluded.error_message,
    finished_at = excluded.finished_at;

insert into run_events(id, tenant_id, run_id, trace_id, sequence, event_type, stage, message, visible_to_user, payload_json)
values {", ".join(event_rows)}
on conflict (id) do nothing;

insert into run_context_snapshots(
  id, tenant_id, workspace_id, user_id, session_id, run_id, trace_id, context_kind,
  included_message_ids, included_file_ids, included_artifact_ids, included_memory_record_ids,
  redaction_summary_json, payload_json
)
values {", ".join(context_rows)}
on conflict (id) do nothing;

insert into sandbox_leases(
  id, tenant_id, workspace_id, user_id, session_id, run_id, trace_id, sandbox_mode,
  provider, status, browser_enabled, resource_limits_json, user_visible_payload_json,
  lease_payload_json, heartbeat_at, released_at, release_reason
)
values {", ".join(sandbox_rows)}
on conflict (id) do update set status = excluded.status, released_at = excluded.released_at, release_reason = excluded.release_reason;

insert into run_skill_snapshots(
  id, tenant_id, run_id, skill_id, skill_version, content_hash, source_json,
  dependency_ids, allowed, staged, used, used_skills_source, inferred_used
)
values {", ".join(skill_snapshot_rows)}
on conflict (tenant_id, run_id, skill_id) do update
set allowed = true,
    staged = true,
    used = true,
    used_skills_source = excluded.used_skills_source;

select json_build_object(
  'prepared_tenant_count', (select count(*) from tenants where id in ({tenant_filter})),
  'prepared_failed_run_count', (select count(*) from runs where id in ({run_id_filter}) and status = 'failed')
)::text;
""".strip()


def prepare_foundation_runtime_fixtures(
    accounts: list[Account],
    *,
    postgres_container: str,
    postgres_user: str,
    postgres_db: str,
    tenant_prefix: str = DEFAULT_TEST_TENANT_PREFIX,
) -> dict[str, Any]:
    tenants = _test_tenant_ids([account.tenant_id for account in accounts], tenant_prefix=tenant_prefix)
    rows = psql_json_rows(
        container=postgres_container,
        db_user=postgres_user,
        db_name=postgres_db,
        sql=build_foundation_runtime_fixture_sql(accounts, tenant_prefix=tenant_prefix),
    )
    prepared = rows[0] if rows else {}
    prepared_counts = {
        "prepared_tenant_count": int(prepared.get("prepared_tenant_count") or 0),
        "prepared_failed_run_count": int(prepared.get("prepared_failed_run_count") or 0),
    }
    return {
        "schema_version": "ai-platform.foundation-runtime-fixture-proof.v1",
        "status": "prepared" if all(value > 0 for value in prepared_counts.values()) else "fixture_incomplete",
        "tenant_ids": tenants,
        "tenant_prefix": tenant_prefix,
        "prepared_counts": prepared_counts,
    }


def multipart_file_post(
    url: str,
    *,
    filename: str,
    content: bytes,
    content_type: str,
    headers: dict[str, str],
    fields: dict[str, str] | None = None,
    timeout: float = 60.0,
) -> tuple[int, Any]:
    boundary = "----ai-platform-multiuser-poc"
    parts: list[bytes] = []
    for name, value in (fields or {}).items():
        parts.extend(
            [
                f"--{boundary}\r\n".encode("utf-8"),
                f'Content-Disposition: form-data; name="{name}"\r\n\r\n'.encode("utf-8"),
                str(value).encode("utf-8"),
                b"\r\n",
            ]
        )
    parts.extend(
        [
            f"--{boundary}\r\n".encode("utf-8"),
            f'Content-Disposition: form-data; name="file"; filename="{filename}"\r\n'.encode("utf-8"),
            f"Content-Type: {content_type}\r\n\r\n".encode("utf-8"),
            content,
            f"\r\n--{boundary}--\r\n".encode("utf-8"),
        ]
    )
    body = b"".join(parts)
    req = request.Request(
        url,
        data=body,
        headers={"Accept": "application/json", "Content-Type": f"multipart/form-data; boundary={boundary}", **headers},
        method="POST",
    )
    try:
        with request.urlopen(req, timeout=timeout) as response:
            status, raw = response.status, response.read()
    except error.HTTPError as exc:
        status, raw = exc.code, exc.read()
    try:
        return status, json.loads(raw.decode("utf-8"))
    except Exception:
        return status, raw.decode("utf-8", errors="replace")


def get_bytes(url: str, headers: dict[str, str], timeout: float = 60.0) -> tuple[int, bytes]:
    req = request.Request(url, headers={"Accept": "*/*", **headers}, method="GET")
    try:
        with request.urlopen(req, timeout=timeout) as response:
            return response.status, response.read()
    except error.HTTPError as exc:
        return exc.code, exc.read()


def fetch_context_snapshot_public_projection(api_url: str, headers: dict[str, str], run_id: str) -> dict[str, Any]:
    try:
        safe_run_id = assert_safe_id(str(run_id), "run_id")
    except ValueError as exc:
        return {"status": 0, "ok": False, "snapshot_count": 0, "error": str(exc)}

    status, payload = json_request(
        "GET",
        f"{api_url.rstrip('/')}/api/ai/runs/{safe_run_id}/context/snapshots",
        headers=headers,
        timeout=30,
    )
    snapshots = payload.get("context_snapshots") if isinstance(payload, dict) else []
    if not isinstance(snapshots, list):
        snapshots = []
    snapshot_payload_summaries = []
    for snapshot in snapshots:
        snapshot_payload = (
            snapshot.get("payload")
            if isinstance(snapshot, dict) and isinstance(snapshot.get("payload"), dict)
            else {}
        )
        snapshot_payload_summaries.append(_context_snapshot_payload_summary(snapshot_payload))
    primary_summary = (
        snapshot_payload_summaries[0]
        if snapshot_payload_summaries
        else _context_snapshot_payload_summary({})
    )
    missing_public_summary_fields = sorted(
        {
            field
            for summary in snapshot_payload_summaries
            for field in summary["missing_public_summary_fields"]
        }
        or set(primary_summary["missing_public_summary_fields"])
    )
    invalid_count_fields = sorted(
        {
            field
            for summary in snapshot_payload_summaries
            for field in summary["invalid_count_fields"]
        }
        or set(primary_summary["invalid_count_fields"])
    )
    unsafe_input_keys = sorted(
        {
            key
            for summary in snapshot_payload_summaries
            for key in summary["unsafe_input_keys"]
        }
        or set(primary_summary["unsafe_input_keys"])
    )
    raw_material_id_fields_present, forbidden_leaks = _context_public_projection_findings(payload)
    ok = (
        status == 200
        and bool(snapshots)
        and all(summary["counts_valid"] for summary in snapshot_payload_summaries)
        and raw_material_id_fields_present is False
        and not forbidden_leaks
        and not missing_public_summary_fields
    )
    evidence = {
        "status": status,
        "ok": ok,
        "snapshot_count": len(snapshots),
        "referenced_material_counts": primary_summary["counts"],
        "raw_material_id_fields_present": raw_material_id_fields_present,
        "forbidden_projection_leaks": forbidden_leaks,
        "summary_source": primary_summary["summary_source"],
        "input_keys": primary_summary["input_keys"],
        "memory_policy_source": primary_summary["memory_policy_source"],
        "long_term_memory_read": primary_summary["long_term_memory_read"],
        "execution_tier": primary_summary["execution_tier"],
        "context_pack_version": primary_summary["context_pack_version"],
        "context_pack_generated_at_present": primary_summary["context_pack_generated_at_present"],
    }
    if invalid_count_fields:
        evidence["invalid_referenced_material_count_fields"] = invalid_count_fields
    if unsafe_input_keys:
        evidence["unsafe_input_keys"] = unsafe_input_keys
    if missing_public_summary_fields:
        evidence["missing_public_summary_fields"] = missing_public_summary_fields
    return evidence


def foundation_runtime_memory_context_summary(results: list[dict[str, Any]]) -> dict[str, Any]:
    context_snapshot_count = 0
    context_snapshot_public_projection_count = 0
    context_pack_version_sample_count = 0
    missing_context_pack_version_count = 0
    unsafe_context_pack_version_count = 0
    context_scope_probe_count = 0
    cross_scope_context_leaks = 0
    long_term_cross_session_memory_read = False
    missing_public_summary_fields: set[str] = set()

    for item in results:
        projection = item.get("context_snapshot_public_projection")
        if not isinstance(projection, dict):
            missing_public_summary_fields.add("context_snapshot_public_projection")
            missing_context_pack_version_count += 1
            continue
        snapshot_count = projection.get("snapshot_count")
        if type(snapshot_count) is int and snapshot_count > 0:
            context_snapshot_count += snapshot_count
        if projection.get("ok") is True:
            context_snapshot_public_projection_count += 1
        raw_version = projection.get("context_pack_version")
        safe_version = safe_public_context_pack_version(raw_version)
        if safe_version is not None:
            context_pack_version_sample_count += 1
        elif raw_version is None:
            missing_context_pack_version_count += 1
        else:
            unsafe_context_pack_version_count += 1
        fields = projection.get("missing_public_summary_fields")
        if isinstance(fields, list):
            missing_public_summary_fields.update(str(field) for field in fields if str(field).strip())
        scope_probe = projection.get("scope_probe")
        if isinstance(scope_probe, dict):
            if scope_probe.get("same_run_snapshot") is True:
                context_scope_probe_count += 1
            if scope_probe.get("cross_scope_leak") is True:
                cross_scope_context_leaks += 1
            if scope_probe.get("long_term_cross_session_memory_read") is True:
                long_term_cross_session_memory_read = True

    failed = (
        context_snapshot_count < len(results)
        or context_snapshot_public_projection_count < len(results)
        or context_pack_version_sample_count < len(results)
        or missing_context_pack_version_count > 0
        or unsafe_context_pack_version_count > 0
        or context_scope_probe_count < len(results)
        or cross_scope_context_leaks > 0
        or long_term_cross_session_memory_read
        or bool(missing_public_summary_fields)
    )
    return {
        "status": "failed" if failed else "passed",
        "context_snapshot_count": context_snapshot_count,
        "context_snapshot_public_projection_count": context_snapshot_public_projection_count,
        "context_pack_version_sample_count": context_pack_version_sample_count,
        "missing_context_pack_version_count": missing_context_pack_version_count,
        "unsafe_context_pack_version_count": unsafe_context_pack_version_count,
        "missing_public_summary_fields": sorted(missing_public_summary_fields),
        "context_scope_probe_count": context_scope_probe_count,
        "cross_scope_context_leaks": cross_scope_context_leaks,
        "long_term_cross_session_memory_read": long_term_cross_session_memory_read,
    }


def login(api_url: str, account: Account) -> dict[str, str]:
    status, payload = json_request(
        "POST",
        f"{api_url.rstrip('/')}/api/auth/login",
        {"username": account.username, "password": account.password},
    )
    if status != 200 or not isinstance(payload, dict) or not payload.get("access_token"):
        raise RuntimeError(f"login failed for {account.label}: status={status} payload={payload}")
    token = str(payload["access_token"])
    return {"Authorization": f"Bearer {token}"}


def trusted_principal_headers(account: Account, *, role: str = "user") -> dict[str, str]:
    return {
        "X-AI-User-ID": account.username,
        "X-AI-User-Name": account.label,
        "X-AI-Tenant-ID": account.tenant_id,
        "X-AI-Roles": role,
    }


def auth_headers(
    api_url: str,
    account: Account,
    *,
    auth_mode: str = "login",
    trusted_header_role: str = "user",
) -> dict[str, str]:
    if auth_mode == "trusted-header":
        return trusted_principal_headers(account, role=trusted_header_role)
    return login(api_url, account)


def upload_docx(api_url: str, headers: dict[str, str], docx_path: Path, *, workspace_id: str = "default") -> dict[str, Any]:
    status, payload = multipart_file_post(
        f"{api_url.rstrip('/')}/api/upload/file?folder=uploads",
        filename=docx_path.name,
        content=docx_path.read_bytes(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
        fields={"workspace_id": workspace_id},
    )
    if status != 200 or not isinstance(payload, dict) or not str(payload.get("key") or "").startswith("file_"):
        raise RuntimeError(f"upload failed: status={status} payload={payload}")
    return payload


def submit_chat(
    api_url: str,
    headers: dict[str, str],
    *,
    agent_id: str,
    message: str,
    attachment: dict[str, Any] | None = None,
    workspace_id: str = "default",
    skill_id: str | None = None,
) -> dict[str, Any]:
    body: dict[str, Any] = {"message": message, "workspace_id": workspace_id}
    if skill_id:
        body["skill_id"] = skill_id
    if attachment:
        body["attachments"] = [attachment]
    status, payload = json_request("POST", f"{api_url.rstrip('/')}/api/chat/stream?agent_id={agent_id}", body, headers=headers)
    if status != 200 or not isinstance(payload, dict) or not payload.get("run_id"):
        raise RuntimeError(f"submit failed: status={status} payload={payload}")
    if not isinstance(payload.get("queue_position"), int) or int(payload["queue_position"]) < 1:
        raise RuntimeError(f"missing queue_position: payload={payload}")
    return payload


def wait_status(api_url: str, headers: dict[str, str], session_id: str, run_id: str, timeout_seconds: float = 240.0) -> dict[str, Any]:
    deadline = time.time() + timeout_seconds
    latest: dict[str, Any] = {}
    while time.time() < deadline:
        status, payload = json_request(
            "GET",
            f"{api_url.rstrip('/')}/api/chat/sessions/{session_id}/status?run_id={run_id}",
            headers=headers,
            timeout=20,
        )
        latest = payload if isinstance(payload, dict) else {"payload": payload}
        if status == 200 and str(latest.get("status") or "").lower() in TERMINAL_LAMBCHAT_STATUSES:
            return latest
        time.sleep(2)
    raise TimeoutError(f"run did not finish: session={session_id} run={run_id} latest={latest}")


def wait_runtime_sandbox_lease(
    api_url: str,
    headers: dict[str, str],
    run_id: str,
    *,
    timeout_seconds: float = 60.0,
) -> str:
    deadline = time.time() + timeout_seconds
    latest_status = 0
    while time.time() < deadline:
        status, payload = json_request(
            "GET",
            f"{api_url.rstrip()}/api/ai/admin/runs/{run_id}",
            headers=headers,
            timeout=20,
        )
        latest_status = status
        if status == 200 and isinstance(payload, dict):
            lease_id = _active_sandbox_lease_id(payload)
            if lease_id:
                return lease_id
        time.sleep(1)
    raise TimeoutError(f"runtime sandbox lease did not appear: run={run_id} latest_status={latest_status}")


def stream_answer(api_url: str, headers: dict[str, str], session_id: str, run_id: str) -> str:
    req = request.Request(
        f"{api_url.rstrip('/')}/api/chat/sessions/{session_id}/stream?run_id={run_id}",
        headers={"Accept": "text/event-stream", **headers},
        method="GET",
    )
    with request.urlopen(req, timeout=60) as response:
        return response.read().decode("utf-8", errors="replace")


def run_case(
    api_url: str,
    account: Account,
    case_name: str,
    agent_id: str,
    message: str,
    docx_path: Path | None,
    scenario: str = "execution",
    auth_mode: str = "login",
    trusted_header_role: str = "user",
    workspace_id: str = "default",
    skill_id: str | None = None,
    run_timeout_seconds: float = 240.0,
    retry_source_run_id: str = "",
) -> dict[str, Any]:
    started_at = time.perf_counter()
    headers = auth_headers(
        api_url,
        account,
        auth_mode=auth_mode,
        trusted_header_role=trusted_header_role,
    )
    attachment = None
    if docx_path is not None:
        upload = upload_docx(api_url, headers, docx_path, workspace_id=workspace_id)
        attachment = {
            "key": upload["key"],
            "name": upload["name"],
            "type": "uploads",
            "mimeType": upload["mimeType"],
            "size": upload["size"],
        }
    submitted = submit_chat(
        api_url,
        headers,
        agent_id=agent_id,
        message=message,
        attachment=attachment,
        workspace_id=workspace_id,
        skill_id=skill_id,
    )
    cancel_action_statuses: list[int] = []
    cancel_effect_statuses: list[str] = []
    if scenario == "cancel":
        if trusted_header_role == "developer":
            wait_runtime_sandbox_lease(api_url, headers, submitted["run_id"], timeout_seconds=run_timeout_seconds)
        status, payload = run_control_action(api_url, headers, submitted["run_id"], "cancel")
        cancel_action_statuses.append(status)
        effect_status = _run_control_payload_status(payload)
        if effect_status:
            cancel_effect_statuses.append(effect_status)
    final_status = wait_status(
        api_url,
        headers,
        submitted["session_id"],
        submitted["run_id"],
        timeout_seconds=run_timeout_seconds,
    )
    if scenario == "cancel":
        for candidate in (final_status.get("raw_status"), final_status.get("status")):
            normalized = str(candidate or "").strip().lower()
            if normalized in CANCEL_EFFECT_STATUSES and normalized not in cancel_effect_statuses:
                cancel_effect_statuses.append(normalized)
    answer = stream_answer(api_url, headers, submitted["session_id"], submitted["run_id"])
    artifact_ids = sorted(set(match.group("artifact_id") for match in DOWNLOAD_RE.finditer(answer)))
    downloads = []
    for artifact_id in artifact_ids:
        owner_status, owner_body = get_bytes(f"{api_url.rstrip('/')}/api/ai/artifacts/{artifact_id}/download", headers)
        downloads.append({"artifact_id": artifact_id, "owner_status": owner_status, "owner_bytes": len(owner_body)})
    retry_action_statuses: list[int] = []
    retry_created_run_ids: list[str] = []
    if scenario == "retry":
        retry_run_source = retry_source_run_id or str(submitted["run_id"])
        status, payload = run_control_action(api_url, headers, retry_run_source, "retry")
        retry_action_statuses.append(status)
        retry_run_id = _run_control_payload_run_id(payload)
        if retry_run_id and retry_run_id != retry_run_source:
            retry_created_run_ids.append(retry_run_id)
    context_projection = fetch_context_snapshot_public_projection(api_url, headers, str(submitted["run_id"]))
    finished_at = time.perf_counter()
    return {
        "tenant_id": account.tenant_id,
        "account": account.label,
        "case": case_name,
        "scenario": scenario,
        "agent_id": agent_id,
        "session_id": submitted["session_id"],
        "run_id": submitted["run_id"],
        "queue_position": submitted["queue_position"],
        "case_started_at_monotonic": started_at,
        "case_finished_at_monotonic": finished_at,
        "queue_probe": {
            "source": "submit_response",
            "queue_position": submitted["queue_position"],
            "submitted_queue_position": submitted["queue_position"],
            "stale_queue_entry": False,
            "cross_tenant_queue_leak": False,
            "admission_limit_violation": False,
        },
        "status": final_status.get("status"),
        "raw_status": final_status.get("raw_status"),
        "artifact_ids": artifact_ids,
        "downloads": downloads,
        "cancel_action_statuses": cancel_action_statuses,
        "cancel_effect_statuses": cancel_effect_statuses,
        "retry_action_statuses": retry_action_statuses,
        "retry_created_run_ids": retry_created_run_ids,
        "has_tmp_path": "/tmp/ai-platform-agent-workspaces/" in answer,
        "context_snapshot_public_projection": context_projection,
    }


def run_control_action(api_url: str, headers: dict[str, str], run_id: str, action: str) -> tuple[int, Any]:
    if action not in {"cancel", "retry"}:
        raise ValueError(f"unsupported run control action: {action}")
    return json_request("POST", f"{api_url.rstrip('/')}/api/ai/runs/{run_id}/{action}", {}, headers=headers, timeout=30)


def _run_control_payload_status(payload: Any) -> str:
    if not isinstance(payload, dict):
        return ""
    return str(payload.get("status") or "").strip().lower()


def _run_control_payload_run_id(payload: Any) -> str:
    if not isinstance(payload, dict):
        return ""
    return str(payload.get("run_id") or "").strip()


def _account_by_label(accounts: list[Account]) -> dict[str, Account]:
    return {account.label: account for account in accounts}


def _first_account(
    accounts: list[Account],
    *,
    tenant_id: str,
    excluded_label: str,
    same_tenant: bool,
) -> Account | None:
    for account in accounts:
        if account.label == excluded_label:
            continue
        if same_tenant and account.tenant_id == tenant_id:
            return account
        if not same_tenant and account.tenant_id != tenant_id:
            return account
    return None


def attach_artifact_acl_probe_results(
    api_url: str,
    results: list[dict[str, Any]],
    accounts: list[Account],
    *,
    auth_mode: str = "login",
    trusted_header_role: str = "user",
) -> None:
    by_label = _account_by_label(accounts)
    for item in results:
        account_label = str(item.get("account") or "")
        tenant_id = str(item.get("tenant_id") or "default")
        owner = by_label.get(account_label)
        cross_user = _first_account(accounts, tenant_id=tenant_id, excluded_label=account_label, same_tenant=True)
        cross_tenant = _first_account(accounts, tenant_id=tenant_id, excluded_label=account_label, same_tenant=False)
        cross_user_download_statuses: list[int] = []
        cross_tenant_download_statuses: list[int] = []
        cross_user_preview_statuses: list[int] = []
        cross_tenant_preview_statuses: list[int] = []
        artifact_ids = [str(value) for value in item.get("artifact_ids", []) if isinstance(value, str)]
        for artifact_id in artifact_ids:
            if cross_user is not None:
                cross_user_headers = auth_headers(
                    api_url,
                    cross_user,
                    auth_mode=auth_mode,
                    trusted_header_role=trusted_header_role,
                )
                status, _body = get_bytes(f"{api_url.rstrip('/')}/api/ai/artifacts/{artifact_id}/download", cross_user_headers)
                cross_user_download_statuses.append(status)
                status, _body = get_bytes(f"{api_url.rstrip('/')}/api/ai/artifacts/{artifact_id}/preview", cross_user_headers)
                cross_user_preview_statuses.append(status)
            if cross_tenant is not None:
                cross_tenant_headers = auth_headers(
                    api_url,
                    cross_tenant,
                    auth_mode=auth_mode,
                    trusted_header_role=trusted_header_role,
                )
                status, _body = get_bytes(f"{api_url.rstrip('/')}/api/ai/artifacts/{artifact_id}/download", cross_tenant_headers)
                cross_tenant_download_statuses.append(status)
                status, _body = get_bytes(f"{api_url.rstrip('/')}/api/ai/artifacts/{artifact_id}/preview", cross_tenant_headers)
                cross_tenant_preview_statuses.append(status)
        item["cross_user_download_statuses"] = cross_user_download_statuses
        item["cross_tenant_download_statuses"] = cross_tenant_download_statuses
        item["cross_user_preview_statuses"] = cross_user_preview_statuses
        item["cross_tenant_preview_statuses"] = cross_tenant_preview_statuses
        item["artifact_acl_probe_owner_present"] = owner is not None


def attach_context_scope_probe_results(
    api_url: str,
    results: list[dict[str, Any]],
    accounts: list[Account],
    *,
    auth_mode: str = "login",
    trusted_header_role: str = "user",
) -> None:
    """Attach same-run and cross-scope context snapshot probe outcomes."""
    by_label = _account_by_label(accounts)
    for item in results:
        account_label = str(item.get("account") or "")
        tenant_id = str(item.get("tenant_id") or "default")
        run_id = str(item.get("run_id") or "")
        owner = by_label.get(account_label)
        cross_user = _first_account(accounts, tenant_id=tenant_id, excluded_label=account_label, same_tenant=True)
        cross_tenant = _first_account(accounts, tenant_id=tenant_id, excluded_label=account_label, same_tenant=False)
        projection = item.get("context_snapshot_public_projection")
        if not isinstance(projection, dict):
            projection = {}
            item["context_snapshot_public_projection"] = projection
        if owner is None or not run_id:
            projection["scope_probe"] = {
                "same_run_snapshot": False,
                "cross_scope_leak": True,
                "long_term_cross_session_memory_read": projection.get("long_term_memory_read") is True,
                "owner_status": 0,
                "cross_user_statuses": [],
                "cross_tenant_statuses": [],
            }
            continue

        owner_headers = auth_headers(
            api_url,
            owner,
            auth_mode=auth_mode,
            trusted_header_role=trusted_header_role,
        )
        owner_status, owner_payload = json_request(
            "GET",
            f"{api_url.rstrip('/')}/api/ai/runs/{run_id}/context/snapshots",
            headers=owner_headers,
            timeout=30,
        )
        owner_snapshots = owner_payload.get("context_snapshots") if isinstance(owner_payload, dict) else []
        same_run_snapshot = owner_status == 200 and isinstance(owner_snapshots, list) and bool(owner_snapshots)

        cross_user_statuses: list[int] = []
        cross_tenant_statuses: list[int] = []
        for probe_account, target in ((cross_user, cross_user_statuses), (cross_tenant, cross_tenant_statuses)):
            if probe_account is None:
                continue
            headers = auth_headers(
                api_url,
                probe_account,
                auth_mode=auth_mode,
                trusted_header_role=trusted_header_role,
            )
            status, _payload = json_request(
                "GET",
                f"{api_url.rstrip('/')}/api/ai/runs/{run_id}/context/snapshots",
                headers=headers,
                timeout=30,
            )
            target.append(status)

        projection["scope_probe"] = {
            "same_run_snapshot": same_run_snapshot,
            "cross_scope_leak": any(status not in {403, 404} for status in cross_user_statuses + cross_tenant_statuses),
            "long_term_cross_session_memory_read": projection.get("long_term_memory_read") is True,
            "owner_status": owner_status,
            "cross_user_statuses": cross_user_statuses,
            "cross_tenant_statuses": cross_tenant_statuses,
        }


def attach_sandbox_lease_probe_results(
    api_url: str,
    results: list[dict[str, Any]],
    accounts: list[Account],
    *,
    auth_mode: str = "login",
    trusted_header_role: str = "user",
) -> None:
    """Attach sandbox lease create/release probe outcomes for each run."""
    by_label = _account_by_label(accounts)
    for item in results:
        account = by_label.get(str(item.get("account") or ""))
        run_id = str(item.get("run_id") or "")
        if account is None or not run_id:
            item["sandbox_lease_probe"] = {
                "create_status": 0,
                "release_status": 0,
                "lease_id": "",
                "source": SYNTHETIC_SANDBOX_LEASE_PROBE_SOURCE,
            }
            continue
        headers = auth_headers(
            api_url,
            account,
            auth_mode=auth_mode,
            trusted_header_role=trusted_header_role,
        )
        create_status, create_payload = json_request(
            "POST",
            f"{api_url.rstrip('/')}/api/ai/runs/{run_id}/sandbox/leases",
            {
                "sandbox_mode": "ephemeral",
                "provider": "fake",
                "ttl_seconds": 600,
                "resource_limits": {},
                "lease_payload": {"probe": "foundation_runtime"},
            },
            headers=headers,
            timeout=30,
        )
        lease_payload = create_payload.get("sandbox_lease") if isinstance(create_payload, dict) else {}
        lease_id = str(lease_payload.get("lease_id") or "") if isinstance(lease_payload, dict) else ""
        release_status = 0
        if create_status == 200 and lease_id:
            release_status, _release_payload = json_request(
                "POST",
                f"{api_url.rstrip('/')}/api/ai/runs/{run_id}/sandbox/leases/{lease_id}/release",
                {"reason": "foundation_runtime_probe_complete"},
                headers=headers,
                timeout=30,
            )
        item["sandbox_lease_probe"] = {
            "create_status": create_status,
            "release_status": release_status,
            "lease_id": lease_id,
            "source": SYNTHETIC_SANDBOX_LEASE_PROBE_SOURCE,
        }


def playback_private_payload_leak_count(payload: Any) -> int:
    leaks = 0

    def walk(value: Any, *, key: str = "") -> None:
        nonlocal leaks
        if key:
            key_alias = "".join(ch for ch in key if ch.isalnum()).lower()
            if key_alias in FORBIDDEN_PUBLIC_KEY_ALIASES:
                leaks += 1
                return
        if isinstance(value, dict):
            for child_key, child_value in value.items():
                walk(child_value, key=str(child_key))
            return
        if isinstance(value, list):
            for item in value:
                walk(item)
            return
        if isinstance(value, str) and any(pattern.search(value) for pattern in FORBIDDEN_PUBLIC_VALUE_PATTERNS):
            leaks += 1

    walk(payload)
    return leaks


def event_order_violations(payload: Any) -> int:
    if not isinstance(payload, dict):
        return 1
    events = payload.get("events")
    if not isinstance(events, list):
        return 1
    sequences = [
        event.get("sequence")
        for event in events
        if isinstance(event, dict) and type(event.get("sequence")) is int
    ]
    return 0 if sequences == sorted(sequences) else 1


def _workspace_fingerprint(run_payload: dict[str, Any], *, tenant_id: str, session_id: str, run_id: str) -> str:
    run = run_payload.get("run") if isinstance(run_payload.get("run"), dict) else {}
    input_payload = run_payload.get("input") if isinstance(run_payload.get("input"), dict) else {}
    if not input_payload and isinstance(run.get("input"), dict):
        input_payload = run["input"]
    workspace_id = str(run_payload.get("workspace_id") or run.get("workspace_id") or input_payload.get("workspace_id") or "default")
    return f"{tenant_id}:{workspace_id}:{session_id}:{run_id}"


def _context_snapshot_id(run_payload: dict[str, Any]) -> str:
    run = run_payload.get("run") if isinstance(run_payload.get("run"), dict) else {}
    candidates = [
        run_payload.get("context_snapshot_id"),
        (run_payload.get("input") or {}).get("context_snapshot_id") if isinstance(run_payload.get("input"), dict) else None,
        (run_payload.get("result") or {}).get("context_snapshot_id") if isinstance(run_payload.get("result"), dict) else None,
        (run.get("input") or {}).get("context_snapshot_id") if isinstance(run.get("input"), dict) else None,
        (run.get("result") or {}).get("context_snapshot_id") if isinstance(run.get("result"), dict) else None,
    ]
    events = run_payload.get("events")
    if isinstance(events, list):
        for event in events:
            payload = event.get("payload") if isinstance(event, dict) else None
            if isinstance(payload, dict):
                candidates.append(payload.get("context_snapshot_id"))
    for candidate in candidates:
        if candidate:
            return str(candidate)
    return ""


def _sandbox_lease_id(run_payload: dict[str, Any]) -> str:
    run = run_payload.get("run") if isinstance(run_payload.get("run"), dict) else {}
    for container in (run_payload.get("result"), run_payload.get("input"), run.get("result"), run.get("input")):
        if isinstance(container, dict) and container.get("sandbox_lease_id"):
            return str(container["sandbox_lease_id"])
    leases = run_payload.get("sandbox_leases")
    if isinstance(leases, list):
        for lease in leases:
            if not isinstance(lease, dict):
                continue
            lease_payload = lease.get("lease_payload") if isinstance(lease.get("lease_payload"), dict) else {}
            if lease_payload.get("probe") == "foundation_runtime":
                continue
            lease_id = str(lease.get("lease_id") or "").strip()
            if lease_id:
                return lease_id
    return ""


def _active_sandbox_lease_id(run_payload: dict[str, Any]) -> str:
    leases = run_payload.get("sandbox_leases")
    if not isinstance(leases, list):
        return ""
    for lease in leases:
        if not isinstance(lease, dict):
            continue
        lease_payload = lease.get("lease_payload") if isinstance(lease.get("lease_payload"), dict) else {}
        if lease_payload.get("probe") == "foundation_runtime":
            continue
        if str(lease.get("status") or "").strip().lower() != "active":
            continue
        lease_id = str(lease.get("lease_id") or "").strip()
        if lease_id:
            return lease_id
    return ""


def _queue_probe_from_run_detail(run_payload: dict[str, Any]) -> dict[str, Any] | None:
    events = run_payload.get("events")
    if not isinstance(events, list):
        return None
    for event in events:
        if not isinstance(event, dict):
            continue
        event_type = str(event.get("event_type") or event.get("type") or "")
        if event_type != "queued":
            continue
        payload = event.get("payload") if isinstance(event.get("payload"), dict) else {}
        source = str(payload.get("source") or "").strip()
        if source != "admin_runtime_queue":
            continue
        queue_position = payload.get("queue_position")
        queue_admission_ordinal = payload.get("queue_admission_ordinal", queue_position)
        if type(queue_admission_ordinal) is not int or queue_admission_ordinal < 1:
            continue
        return {
            "source": "admin_runtime_queue",
            "queue_position": queue_position if type(queue_position) is int and queue_position > 0 else queue_admission_ordinal,
            "queue_admission_ordinal": queue_admission_ordinal,
            "submitted_queue_position": None,
            "stale_queue_entry": False,
            "cross_tenant_queue_leak": False,
            "admission_limit_violation": False,
        }
    return None


def _tool_permission_summary(run_payload: dict[str, Any]) -> dict[str, int | str]:
    decided_ids: set[str] = set()
    decision_sample_count = 0
    reused_violations = 0
    wrong_decision_reuse_violations = 0
    tool_call_id_mismatch_violations = 0
    events = run_payload.get("events")
    if isinstance(events, list):
        for event in events:
            if not isinstance(event, dict):
                continue
            payload = event.get("payload") if isinstance(event.get("payload"), dict) else {}
            if event.get("type") == "tool_permission_decided" or event.get("event_type") == "tool_permission_decided":
                decision_sample_count += 1
                request_id = str(payload.get("request_id") or payload.get("permission_request_id") or "")
                tool_call_id = str(payload.get("tool_call_id") or "")
                decision = str(payload.get("decision") or "")
                decision_key = f"{request_id}:{tool_call_id}:{decision}"
                if decision_key in decided_ids:
                    reused_violations += 1
                decided_ids.add(decision_key)
            if payload.get("wrong_decision_reuse") is True:
                wrong_decision_reuse_violations += 1
            if payload.get("tool_call_id_mismatch") is True:
                tool_call_id_mismatch_violations += 1
    return {
        "status": "passed",
        "decision_sample_count": decision_sample_count,
        "allow_once_reuse_violations": reused_violations,
        "wrong_decision_reuse_violations": wrong_decision_reuse_violations,
        "tool_call_id_mismatch_violations": tool_call_id_mismatch_violations,
    }


def _skill_snapshot_summary(run_payload: dict[str, Any]) -> dict[str, Any]:
    snapshots = run_payload.get("skill_snapshots")
    if not isinstance(snapshots, list):
        snapshots = []
    used_count = sum(1 for item in snapshots if isinstance(item, dict) and item.get("used") is True)
    snapshot_binding_sample_count = 0
    missing_pinned_snapshots: list[str] = []
    mismatched_pinned_snapshots: list[str] = []
    global_mutable_skill_lookup_used = False
    for item in snapshots:
        if not isinstance(item, dict):
            continue
        skill_id = str(item.get("skill_id") or "").strip()
        skill_version = str(item.get("skill_version") or "").strip()
        content_hash = str(item.get("content_hash") or "").strip()
        source = item.get("source") if isinstance(item.get("source"), dict) else {}
        if source.get("global_mutable_lookup_used") is True:
            global_mutable_skill_lookup_used = True
        if skill_id and (skill_version or content_hash):
            snapshot_binding_sample_count += 1
        elif item.get("used") is True:
            missing_pinned_snapshots.append(skill_id or "unknown")
        if source.get("pinned_snapshot_mismatch") is True:
            mismatched_pinned_snapshots.append(skill_id or "unknown")
    failed = (
        len(snapshots) == 0
        or snapshot_binding_sample_count < len(snapshots)
        or bool(missing_pinned_snapshots)
        or bool(mismatched_pinned_snapshots)
        or global_mutable_skill_lookup_used
    )
    return {
        "status": "failed" if failed else "passed",
        "run_skill_snapshot_count": len(snapshots),
        "used_count": used_count,
        "missing_pinned_snapshots": sorted(set(missing_pinned_snapshots)),
        "mismatched_pinned_snapshots": sorted(set(mismatched_pinned_snapshots)),
        "global_mutable_skill_lookup_used": global_mutable_skill_lookup_used,
        "snapshot_binding_sample_count": snapshot_binding_sample_count,
    }


def attach_run_detail_probe_results(
    api_url: str,
    results: list[dict[str, Any]],
    accounts: list[Account],
    *,
    auth_mode: str = "login",
    trusted_header_role: str = "user",
) -> None:
    by_label = _account_by_label(accounts)
    for item in results:
        account = by_label.get(str(item.get("account") or ""))
        if account is None:
            continue
        run_id = str(item.get("run_id") or "")
        session_id = str(item.get("session_id") or "")
        headers = auth_headers(
            api_url,
            account,
            auth_mode=auth_mode,
            trusted_header_role=trusted_header_role,
        )
        detail_path = "/api/ai/admin/runs" if trusted_header_role == "developer" else "/api/ai/runs"
        status, run_payload = json_request("GET", f"{api_url.rstrip()}{detail_path}/{run_id}", headers=headers, timeout=30)
        if status == 200 and isinstance(run_payload, dict):
            item["workspace_fingerprint"] = _workspace_fingerprint(run_payload, tenant_id=account.tenant_id, session_id=session_id, run_id=run_id)
            item["context_snapshot_id"] = _context_snapshot_id(run_payload)
            detail_queue_probe = _queue_probe_from_run_detail(run_payload)
            if detail_queue_probe is not None:
                submitted_queue_position = None
                existing_probe = item.get("queue_probe")
                if isinstance(existing_probe, dict) and type(existing_probe.get("submitted_queue_position")) is int:
                    submitted_queue_position = existing_probe["submitted_queue_position"]
                detail_queue_probe["submitted_queue_position"] = submitted_queue_position
                item["queue_probe"] = detail_queue_probe
            detail_sandbox_lease_id = _sandbox_lease_id(run_payload)
            if detail_sandbox_lease_id:
                item["sandbox_lease_id"] = detail_sandbox_lease_id
            item["tool_permission"] = _tool_permission_summary(run_payload)
            item["skill_snapshot"] = _skill_snapshot_summary(run_payload)
        playback_status, playback_payload = json_request("GET", f"{api_url.rstrip('/')}/api/ai/runs/{run_id}/playback", headers=headers, timeout=30)
        item["playback"] = {
            "status": playback_status,
            "event_order_violations": event_order_violations(playback_payload),
            "private_payload_leak_count": playback_private_payload_leak_count(playback_payload),
        }


def _permission_request_id(payload: Any) -> str:
    if not isinstance(payload, dict):
        return ""
    candidates = [payload]
    permission_request = payload.get("permission_request")
    if isinstance(permission_request, dict):
        candidates.append(permission_request)
    for candidate in candidates:
        for key in ("request_id", "permission_request_id", "id"):
            value = candidate.get(key)
            if value:
                return str(value)
    return ""


def _tool_permission_negative_reuse_targets(
    account: Account,
    accounts: list[Account],
    run_id: str,
) -> list[tuple[Account, str, str]]:
    same_tenant_other_user = next(
        (
            candidate
            for candidate in accounts
            if candidate.tenant_id == account.tenant_id and candidate.label != account.label
        ),
        None,
    )
    cross_tenant_user = next(
        (candidate for candidate in accounts if candidate.tenant_id != account.tenant_id),
        None,
    )
    targets: list[tuple[Account, str, str]] = [
        (account, run_id, "same_request_repeat"),
        (account, f"{run_id}-reuse-wrong-run", "wrong_run"),
    ]
    if same_tenant_other_user is not None:
        targets.append((same_tenant_other_user, run_id, "same_tenant_other_user"))
    if cross_tenant_user is not None:
        targets.append((cross_tenant_user, run_id, "cross_tenant_other_user"))
    return targets


def attach_tool_permission_probe_results(
    api_url: str,
    results: list[dict[str, Any]],
    accounts: list[Account],
    *,
    auth_mode: str = "login",
    trusted_header_role: str = "user",
) -> None:
    by_label = _account_by_label(accounts)
    for item in results:
        account = by_label.get(str(item.get("account") or ""))
        run_id = str(item.get("run_id") or "")
        if account is None or not run_id:
            item["tool_permission_probe"] = {"status": "skipped"}
            continue
        headers = auth_headers(
            api_url,
            account,
            auth_mode=auth_mode,
            trusted_header_role=trusted_header_role,
        )
        request_status, request_payload = json_request(
            "POST",
            f"{api_url.rstrip('/')}/api/ai/runs/{run_id}/tool-permissions/request",
            {
                "tool_id": DEFAULT_FIXTURE_TOOL_ID,
                "tool_call_id": f"frc_tool_{run_id}",
                "action": "execute",
                "risk_level": "low",
                "write_capable": False,
                "reason": "foundation_runtime_permission_probe",
                "request_payload": {"probe": "foundation_runtime"},
            },
            headers=headers,
            timeout=30,
        )
        request_id = _permission_request_id(request_payload)
        decision_status = 0
        negative_reuse_probe_count = 0
        negative_reuse_denied_count = 0
        negative_reuse_unexpected_successes = 0
        if request_status in {200, 201} and request_id:
            decision_status, _decision_payload = json_request(
                "POST",
                f"{api_url.rstrip('/')}/api/ai/runs/{run_id}/tool-permissions/{request_id}/decision",
                {
                    "decision": "allow_once",
                    "reason": "foundation_runtime_permission_probe",
                    "decision_payload": {"probe": "foundation_runtime"},
                    "expires_in_seconds": 900,
                },
                headers=headers,
                timeout=30,
            )
            for target_account, target_run_id, target_scope in _tool_permission_negative_reuse_targets(
                account,
                accounts,
                run_id,
            ):
                target_headers = headers if target_account == account else auth_headers(
                    api_url,
                    target_account,
                    auth_mode=auth_mode,
                    trusted_header_role=trusted_header_role,
                )
                negative_reuse_status, _negative_reuse_payload = json_request(
                    "POST",
                    f"{api_url.rstrip('/')}/api/ai/runs/{target_run_id}/tool-permissions/{request_id}/decision",
                    {
                        "decision": "allow_once",
                        "reason": "foundation_runtime_permission_reuse_negative_probe",
                        "decision_payload": {
                            "probe": "foundation_runtime_reuse",
                            "scope": target_scope,
                        },
                        "expires_in_seconds": 900,
                    },
                    headers=target_headers,
                    timeout=30,
                )
                negative_reuse_probe_count += 1
                if negative_reuse_status in DENIED_OR_CONFLICT_HTTP_STATUSES:
                    negative_reuse_denied_count += 1
                else:
                    negative_reuse_unexpected_successes += 1
        item["tool_permission_probe"] = {
            "request_status": request_status,
            "decision_status": decision_status,
            "request_id": request_id,
            "negative_reuse_probe_count": negative_reuse_probe_count,
            "negative_reuse_denied_count": negative_reuse_denied_count,
            "negative_reuse_unexpected_successes": negative_reuse_unexpected_successes,
        }


def parse_account(value: str, *, require_explicit_tenant: bool = False) -> Account:
    label_part, rest = value.split("=", 1)
    tenant_id = "default"
    label = label_part
    if "/" in label_part:
        tenant_id, label = label_part.split("/", 1)
    if require_explicit_tenant and tenant_id == "default":
        raise ValueError("foundation runtime evidence accounts must use tenant/label=username:password")
    username, password = rest.split(":", 1)
    return Account(label=label, username=username, password=password, tenant_id=tenant_id)


def build_foundation_runtime_case_specs(
    accounts: list[Account],
    *,
    min_cases: int = 12,
    use_fixture_agents: bool = False,
) -> list[CaseSpec]:
    tenant_ids = {account.tenant_id for account in accounts}
    if len(tenant_ids) < 2:
        raise ValueError("foundation runtime evidence requires at least two tenants")
    if len(accounts) < 2:
        raise ValueError("foundation runtime evidence requires at least two accounts")
    templates = [
        ("general-chat", "run_creation", "general-agent", "并发创建运行验收，请简短回复。", False),
        ("word-review", "execution", "qa-word-review", "审核一下这个文档", True),
        ("cancel-probe", "cancel", "general-agent", "创建后取消路径验收，请简短回复。", False),
        ("retry-probe", "retry", "baoyu-translate", "翻译一下这个文档，用于 retry 路径验收。", True),
    ]
    specs: list[CaseSpec] = []
    scenario_seen = {scenario: 0 for _case_name, scenario, _agent_id, _message, _uses_docx in templates}
    index = 0
    while len(specs) < min_cases:
        template_index = index % len(templates)
        case_name, scenario, agent_id, message, uses_docx = templates[template_index]
        scenario_index = scenario_seen[scenario]
        scenario_seen[scenario] += 1
        account = accounts[(template_index + scenario_index) % len(accounts)]
        skill_id = _skill_id_for_agent(agent_id)
        specs.append(
            CaseSpec(
                account=account,
                case_name=f"{case_name}-{len(specs) + 1}",
                scenario=scenario,
                agent_id=_agent_id_for_case(account, agent_id, use_fixture_agents=use_fixture_agents),
                skill_id=skill_id,
                message=f"{account.label} {message}",
                uses_docx=uses_docx,
                workspace_id=fixture_workspace_id(account.tenant_id) if use_fixture_agents else "default",
                retry_source_run_id=fixture_retry_source_run_id(account) if use_fixture_agents and scenario == "retry" else "",
            )
        )
        index += 1
    return specs


def _scenario_counts(results: list[dict[str, Any]]) -> dict[str, int]:
    counts = {"run_creation": 0, "execution": 0, "cancel": 0, "retry": 0}
    for item in results:
        scenario = str(item.get("scenario") or "")
        if scenario in counts:
            counts[scenario] += 1
    return counts


def _all_values(results: list[dict[str, Any]], key: str) -> list[int]:
    values: list[int] = []
    for item in results:
        source = item.get(key)
        if isinstance(source, list):
            values.extend(value for value in source if type(value) is int)
    return values


def _all_strings(results: list[dict[str, Any]], key: str) -> list[str]:
    values: list[str] = []
    for item in results:
        source = item.get(key)
        if isinstance(source, list):
            values.extend(str(value) for value in source if isinstance(value, str) and value)
    return values


def _sum_nested_int(results: list[dict[str, Any]], key: str, nested_key: str) -> int:
    total = 0
    for item in results:
        nested = item.get(key)
        if isinstance(nested, dict) and type(nested.get(nested_key)) is int:
            total += nested[nested_key]
    return total


def _tool_permission_probe_decision_count(results: list[dict[str, Any]]) -> int:
    total = 0
    for item in results:
        probe = item.get("tool_permission_probe")
        if not isinstance(probe, dict):
            continue
        if probe.get("request_id") and probe.get("decision_status") in {200, 201}:
            total += 1
    return total


def _tool_permission_probe_negative_count(results: list[dict[str, Any]], key: str) -> int:
    return _sum_nested_int(results, "tool_permission_probe", key)


def _any_nested_true(results: list[dict[str, Any]], key: str, nested_key: str) -> bool:
    for item in results:
        nested = item.get(key)
        if isinstance(nested, dict) and nested.get(nested_key) is True:
            return True
    return False


def _merged_nested_lists(results: list[dict[str, Any]], key: str, nested_key: str) -> list[str]:
    values: set[str] = set()
    for item in results:
        nested = item.get(key)
        source = nested.get(nested_key) if isinstance(nested, dict) else None
        if isinstance(source, list):
            values.update(str(value) for value in source if str(value).strip())
    return sorted(values)


def _foundation_runtime_queue_summary(results: list[dict[str, Any]]) -> dict[str, Any]:
    queue_positions: list[int] = []
    submitted_queue_positions: list[int] = []
    queue_probe_sample_count = 0
    probe_sources: set[str] = set()
    admission_limit_violations = 0
    cross_tenant_queue_leaks = 0
    stale_queue_entries = 0
    for item in results:
        probe = item.get("queue_probe")
        if not isinstance(probe, dict):
            continue
        submitted_position = probe.get("submitted_queue_position", probe.get("queue_position"))
        if type(submitted_position) is int and submitted_position > 0:
            submitted_queue_positions.append(submitted_position)
        source = str(probe.get("source") or "").strip()
        if source in QUEUE_PROBE_SOURCES:
            probe_sources.add(source)
            queue_probe_sample_count += 1
            position = probe.get("queue_admission_ordinal", probe.get("queue_position"))
            if type(position) is int and position > 0:
                queue_positions.append(position)
        if probe.get("admission_limit_violation") is True:
            admission_limit_violations += 1
        if probe.get("cross_tenant_queue_leak") is True:
            cross_tenant_queue_leaks += 1
        if probe.get("stale_queue_entry") is True:
            stale_queue_entries += 1
    duplicate_count = len(queue_positions) - len(set(queue_positions))
    queue_probe_source = ",".join(sorted(probe_sources)) if probe_sources else "missing"
    failed = (
        len(queue_positions) < len(results)
        or duplicate_count > 0
        or not probe_sources
        or admission_limit_violations > 0
        or cross_tenant_queue_leaks > 0
        or stale_queue_entries > 0
    )
    return {
        "status": "failed" if failed else "passed",
        "admission_limit_violations": admission_limit_violations,
        "cross_tenant_queue_leaks": cross_tenant_queue_leaks,
        "stale_queue_entries": stale_queue_entries,
        "queue_position_sample_count": len(queue_positions),
        "queue_position_duplicate_count": duplicate_count,
        "queue_probe_sample_count": queue_probe_sample_count,
        "submitted_queue_position_sample_count": len(submitted_queue_positions),
        "queue_probe_source": queue_probe_source,
    }


def _foundation_runtime_sandbox_summary(
    results: list[dict[str, Any]],
    *,
    workspace_fingerprints: list[str],
) -> dict[str, Any]:
    lease_scope_by_id: dict[str, str] = {}
    lease_scope_leaks = 0
    for item in results:
        lease_id = str(item.get("sandbox_lease_id") or "").strip()
        if not lease_id:
            continue
        scope = ":".join(
            str(item.get(key) or "")
            for key in ("tenant_id", "account", "session_id", "run_id")
        )
        existing_scope = lease_scope_by_id.get(lease_id)
        if existing_scope is not None and existing_scope != scope:
            lease_scope_leaks += 1
        lease_scope_by_id[lease_id] = scope
    lease_count = len(lease_scope_by_id)
    failed = (
        len(workspace_fingerprints) < len(results)
        or lease_count < len(results)
        or lease_scope_leaks > 0
        or len(workspace_fingerprints) != len(set(workspace_fingerprints))
    )
    return {
        "status": "failed" if failed else "passed",
        "workspace_scope_sample_count": len(workspace_fingerprints),
        "sandbox_lease_sample_count": lease_count,
        "active_lease_count": 0,
        "cross_scope_lease_leaks": lease_scope_leaks,
        "workspace_scope_collisions": len(workspace_fingerprints) - len(set(workspace_fingerprints)),
        "lease_probe_source": SANDBOX_LEASE_PROBE_SOURCE if lease_count else "missing",
    }


def _foundation_runtime_concurrency_summary(results: list[dict[str, Any]]) -> dict[str, Any]:
    windows: list[tuple[float, float]] = []
    for item in results:
        start = item.get("case_started_at_monotonic")
        finish = item.get("case_finished_at_monotonic")
        if type(start) in {int, float} and type(finish) in {int, float} and float(finish) >= float(start):
            windows.append((float(start), float(finish)))
    max_overlap = 0
    for start, _finish in windows:
        overlap = sum(1 for candidate_start, candidate_finish in windows if candidate_start <= start <= candidate_finish)
        max_overlap = max(max_overlap, overlap)
    return {
        "concurrent_request_count": len(results),
        "max_observed_concurrency": max_overlap,
        "concurrency_probe_source": "client_case_timestamps" if windows else "missing",
        "concurrency_window_sample_count": len(windows),
    }


def _has_cancel_effect(item: dict[str, Any]) -> bool:
    statuses = item.get("cancel_effect_statuses")
    if not isinstance(statuses, list):
        return False
    return any(
        isinstance(status, str) and status.strip().lower() in {"cancel_requested", "cancelled", "canceled"}
        for status in statuses
    )


def build_foundation_runtime_concurrency_evidence(
    results: list[dict[str, Any]],
    *,
    commit_sha: str,
    runtime_subject_commit_sha: str,
    cleanup_proof: dict[str, Any] | None = None,
    fixture_proof: dict[str, Any] | None = None,
    failed_cases: list[dict[str, Any]] | None = None,
    run_creation_role: str = "user",
    public_probe_role: str = "user",
    admin_probe_role: str | None = None,
) -> dict[str, Any]:
    tenant_ids = {str(item.get("tenant_id") or "default") for item in results}
    users = {f"{item.get('tenant_id') or 'default'}:{item.get('account')}" for item in results}
    session_ids = {str(item.get("session_id")) for item in results if item.get("session_id")}
    run_ids = {str(item.get("run_id")) for item in results if item.get("run_id")}
    workspace_fingerprints = [
        str(item.get("workspace_fingerprint"))
        for item in results
        if item.get("workspace_fingerprint")
    ]
    memory_context = foundation_runtime_memory_context_summary(results)
    queue_admission = _foundation_runtime_queue_summary(results)
    sandbox_workspace = _foundation_runtime_sandbox_summary(
        results,
        workspace_fingerprints=workspace_fingerprints,
    )
    skill_snapshots = {
        "run_skill_snapshot_count": _sum_nested_int(results, "skill_snapshot", "run_skill_snapshot_count"),
        "used_count": _sum_nested_int(results, "skill_snapshot", "used_count"),
        "missing_pinned_snapshots": _merged_nested_lists(results, "skill_snapshot", "missing_pinned_snapshots"),
        "mismatched_pinned_snapshots": _merged_nested_lists(results, "skill_snapshot", "mismatched_pinned_snapshots"),
        "global_mutable_skill_lookup_used": _any_nested_true(results, "skill_snapshot", "global_mutable_skill_lookup_used"),
        "snapshot_binding_sample_count": _sum_nested_int(results, "skill_snapshot", "snapshot_binding_sample_count"),
    }
    skill_snapshots["status"] = "failed" if (
        skill_snapshots["run_skill_snapshot_count"] < len(results)
        or skill_snapshots["snapshot_binding_sample_count"] < len(results)
        or bool(skill_snapshots["missing_pinned_snapshots"])
        or bool(skill_snapshots["mismatched_pinned_snapshots"])
        or skill_snapshots["global_mutable_skill_lookup_used"] is True
    ) else "passed"
    scenario_counts = _scenario_counts(results)
    concurrency_summary = _foundation_runtime_concurrency_summary(results)
    evidence = {
        "schema_version": FOUNDATION_RUNTIME_CONCURRENCY_SCHEMA,
        "artifact_kind": "foundation_runtime_concurrency",
        "commit_sha": commit_sha,
        "source_tree_commit_sha": commit_sha,
        "runtime_subject_commit_sha": runtime_subject_commit_sha,
        "summary": {
            "tenant_count": len(tenant_ids),
            "user_count": len(users),
            "session_count": len(session_ids),
            "run_count": len(run_ids),
            **concurrency_summary,
        },
        "role_provenance": {
            "run_creation_role": run_creation_role,
            "public_probe_role": public_probe_role,
            "admin_probe_role": admin_probe_role,
            "ordinary_user_multi_agent_opened": False,
        },
        "scenario_counts": scenario_counts,
        "checks": {
            "queue_admission": {
                **queue_admission,
                "cancel_action_statuses": _all_values(results, "cancel_action_statuses"),
                "cancel_effect_statuses": _all_strings(results, "cancel_effect_statuses"),
                "cancel_effect_run_count": sum(1 for item in results if _has_cancel_effect(item)),
                "retry_action_statuses": _all_values(results, "retry_action_statuses"),
                "retry_created_run_count": len(set(_all_strings(results, "retry_created_run_ids"))),
            },
            "sandbox_workspace": sandbox_workspace,
            "memory_context": memory_context,
            "artifact_acl": {
                "status": "passed",
                "owner_statuses": [
                    int(download["owner_status"])
                    for item in results
                    for download in item.get("downloads", [])
                    if isinstance(download, dict) and type(download.get("owner_status")) is int
                ],
                "cross_user_statuses": _all_values(results, "cross_user_download_statuses"),
                "cross_tenant_statuses": _all_values(results, "cross_tenant_download_statuses"),
                "preview_cross_user_statuses": _all_values(results, "cross_user_preview_statuses"),
                "preview_cross_tenant_statuses": _all_values(results, "cross_tenant_preview_statuses"),
            },
            "tool_permission": {
                "status": "passed",
                "decision_sample_count": _sum_nested_int(results, "tool_permission", "decision_sample_count")
                + _tool_permission_probe_decision_count(results),
                "negative_reuse_probe_count": _sum_nested_int(results, "tool_permission", "negative_reuse_probe_count")
                + _tool_permission_probe_negative_count(results, "negative_reuse_probe_count"),
                "negative_reuse_denied_count": _sum_nested_int(results, "tool_permission", "negative_reuse_denied_count")
                + _tool_permission_probe_negative_count(results, "negative_reuse_denied_count"),
                "negative_reuse_unexpected_successes": _sum_nested_int(
                    results,
                    "tool_permission",
                    "negative_reuse_unexpected_successes",
                )
                + _tool_permission_probe_negative_count(results, "negative_reuse_unexpected_successes"),
                "allow_once_reuse_violations": _sum_nested_int(results, "tool_permission", "allow_once_reuse_violations"),
                "wrong_decision_reuse_violations": _sum_nested_int(results, "tool_permission", "wrong_decision_reuse_violations"),
                "tool_call_id_mismatch_violations": _sum_nested_int(results, "tool_permission", "tool_call_id_mismatch_violations"),
            },
            "skill_snapshots": skill_snapshots,
            "run_playback": {
                "status": "passed",
                "event_order_violations": _sum_nested_int(results, "playback", "event_order_violations"),
                "private_payload_leak_count": _sum_nested_int(results, "playback", "private_payload_leak_count"),
            },
        },
        "non_expansion_invariants": {
            "production_concurrency_increase_allowed": False,
            "ordinary_user_multi_agent_allowed": False,
            "docker_sandbox_hardened_claim_allowed": False,
            "department_rollout_allowed": False,
            "long_term_cross_session_memory_enabled": False,
        },
    }
    if failed_cases:
        evidence["failed_case_count"] = len(failed_cases)
        evidence["failed_cases"] = failed_cases
    if cleanup_proof is not None:
        evidence["cleanup_proof"] = cleanup_proof
    if fixture_proof is not None:
        evidence["fixture_proof"] = fixture_proof
    return evidence


def main() -> int:
    parser = argparse.ArgumentParser(description="Verify concurrent multi-user ai-platform POC flows.")
    parser.add_argument("--api-url", default=DEFAULT_API_URL)
    parser.add_argument("--sample-docx", default=DEFAULT_SAMPLE_DOCX)
    parser.add_argument("--account", action="append", required=True, help="label=username:password or tenant/label=username:password")
    parser.add_argument("--auth-mode", choices=("login", "trusted-header"), default="login")
    parser.add_argument("--trusted-header-role", choices=("user", "developer"), default="user")
    parser.add_argument("--foundation-runtime-evidence", action="store_true")
    parser.add_argument("--min-concurrent-cases", type=int, default=12)
    parser.add_argument("--run-timeout-seconds", type=float, default=240.0)
    parser.add_argument("--commit-sha", default="unknown")
    parser.add_argument("--runtime-subject-commit-sha", default="unknown")
    parser.add_argument("--prepare-fixtures", action="store_true")
    parser.add_argument("--cleanup-before", action="store_true")
    parser.add_argument("--cleanup-after", action="store_true")
    parser.add_argument("--use-fixture-agents", action="store_true")
    parser.add_argument("--postgres-container", default=DEFAULT_POSTGRES_CONTAINER)
    parser.add_argument("--postgres-user", default=DEFAULT_POSTGRES_USER)
    parser.add_argument("--postgres-db", default=DEFAULT_POSTGRES_DB)
    parser.add_argument("--redis-container", default=DEFAULT_REDIS_CONTAINER)
    parser.add_argument("--tenant-prefix", default=DEFAULT_TEST_TENANT_PREFIX)
    args = parser.parse_args()

    try:
        docx_path = ensure_default_sample_docx(Path(args.sample_docx))
    except FileNotFoundError as exc:
        raise SystemExit(str(exc)) from exc
    accounts = []
    for item in args.account:
        try:
            accounts.append(parse_account(item, require_explicit_tenant=args.foundation_runtime_evidence))
        except ValueError as exc:
            raise SystemExit(str(exc)) from exc
    if len(accounts) < 2:
        raise SystemExit("provide at least two accounts")

    if args.foundation_runtime_evidence:
        if args.use_fixture_agents and args.auth_mode != "trusted-header":
            raise SystemExit("fixture agents require trusted-header auth")
        tenant_ids = sorted({account.tenant_id for account in accounts})
        cleanup_proof: dict[str, Any] = {}
        if args.cleanup_before:
            cleanup_proof["before"] = build_foundation_runtime_cleanup_proof(
                tenant_ids,
                postgres_container=args.postgres_container,
                postgres_user=args.postgres_user,
                postgres_db=args.postgres_db,
                redis_container=args.redis_container,
                tenant_prefix=args.tenant_prefix,
            )
        fixture_proof = None
        if args.prepare_fixtures:
            fixture_proof = prepare_foundation_runtime_fixtures(
                accounts,
                postgres_container=args.postgres_container,
                postgres_user=args.postgres_user,
                postgres_db=args.postgres_db,
                tenant_prefix=args.tenant_prefix,
            )
        try:
            specs = build_foundation_runtime_case_specs(
                accounts,
                min_cases=args.min_concurrent_cases,
                use_fixture_agents=args.use_fixture_agents,
            )
        except ValueError as exc:
            raise SystemExit(str(exc)) from exc
        run_creation_role = "developer" if args.use_fixture_agents else args.trusted_header_role
        admin_probe_role = run_creation_role if run_creation_role == "developer" else None
        with concurrent.futures.ThreadPoolExecutor(max_workers=len(specs)) as pool:
            future_specs = {
                pool.submit(
                    run_case,
                    args.api_url,
                    spec.account,
                    spec.case_name,
                    spec.agent_id,
                    spec.message,
                    docx_path if spec.uses_docx else None,
                    scenario=spec.scenario,
                    auth_mode=args.auth_mode,
                    trusted_header_role=run_creation_role,
                    workspace_id=spec.workspace_id,
                    skill_id=spec.skill_id if run_creation_role != "user" else None,
                    run_timeout_seconds=args.run_timeout_seconds,
                    retry_source_run_id=spec.retry_source_run_id,
                ): spec
                for spec in specs
            }
            results = []
            failed_cases: list[dict[str, Any]] = []
            for future in concurrent.futures.as_completed(future_specs):
                spec = future_specs[future]
                try:
                    results.append(future.result())
                except Exception as exc:
                    failed_cases.append(
                        {
                            "tenant_id": spec.account.tenant_id,
                            "account": spec.account.label,
                            "case": spec.case_name,
                            "scenario": spec.scenario,
                            "agent_id": spec.agent_id,
                            "skill_id": spec.skill_id,
                            "error_type": type(exc).__name__,
                            "message": str(exc),
                        }
                    )
        attach_artifact_acl_probe_results(
            args.api_url,
            results,
            accounts,
            auth_mode=args.auth_mode,
            trusted_header_role="user",
        )
        attach_context_scope_probe_results(
            args.api_url,
            results,
            accounts,
            auth_mode=args.auth_mode,
            trusted_header_role="user",
        )
        attach_tool_permission_probe_results(
            args.api_url,
            results,
            accounts,
            auth_mode=args.auth_mode,
            trusted_header_role="user",
        )
        attach_sandbox_lease_probe_results(
            args.api_url,
            results,
            accounts,
            auth_mode=args.auth_mode,
            trusted_header_role="user",
        )
        attach_run_detail_probe_results(
            args.api_url,
            results,
            accounts,
            auth_mode=args.auth_mode,
            trusted_header_role=run_creation_role,
        )
        evidence = build_foundation_runtime_concurrency_evidence(
            sorted(results, key=lambda row: (row.get("tenant_id", ""), row.get("account", ""), row.get("case", ""))),
            commit_sha=args.commit_sha,
            runtime_subject_commit_sha=args.runtime_subject_commit_sha,
            cleanup_proof=cleanup_proof or None,
            fixture_proof=fixture_proof,
            failed_cases=sorted(
                failed_cases,
                key=lambda row: (
                    str(row.get("tenant_id") or ""),
                    str(row.get("account") or ""),
                    str(row.get("case") or ""),
                ),
            ),
            run_creation_role=run_creation_role,
            public_probe_role="user",
            admin_probe_role=admin_probe_role,
        )
        if args.cleanup_after:
            cleanup_proof["after"] = build_foundation_runtime_cleanup_proof(
                tenant_ids,
                postgres_container=args.postgres_container,
                postgres_user=args.postgres_user,
                postgres_db=args.postgres_db,
                redis_container=args.redis_container,
                tenant_prefix=args.tenant_prefix,
            )
            evidence["cleanup_proof"] = cleanup_proof
        readiness = build_foundation_runtime_concurrency_readiness(evidence)
        output = evidence if readiness.get("verified") is True else {"evidence": evidence, "readiness": readiness}
        print(json.dumps(output, ensure_ascii=False, indent=2, sort_keys=True))
        return 0 if readiness.get("verified") is True else 1

    case_specs = []
    for account in accounts[:2]:
        case_specs.extend(
            [
                (account, "general-chat", "general-agent", f"{account.label} 并发通用聊天验收，请简短回复。", None),
                (account, "word-review", "general-agent", "审核一下这个文档", docx_path),
                (account, "word-translate", "baoyu-translate", "翻译一下这个文档", docx_path),
            ]
        )
    with concurrent.futures.ThreadPoolExecutor(max_workers=len(case_specs)) as pool:
        futures = [pool.submit(run_case, args.api_url, *spec) for spec in case_specs]
        results = [future.result() for future in concurrent.futures.as_completed(futures)]

    failures = []
    for item in results:
        if item["status"] != "completed":
            failures.append({"case": item["case"], "account": item["account"], "reason": "not_completed", "status": item["status"]})
        if item["has_tmp_path"]:
            failures.append({"case": item["case"], "account": item["account"], "reason": "tmp_path_leaked"})
        if item["case"] in {"word-review", "word-translate"} and not item["artifact_ids"]:
            failures.append({"case": item["case"], "account": item["account"], "reason": "missing_artifact_link"})
        for download in item["downloads"]:
            if download["owner_status"] != 200 or download["owner_bytes"] <= 0:
                failures.append({"case": item["case"], "account": item["account"], "reason": "artifact_download_failed", **download})
        context_projection = item.get("context_snapshot_public_projection")
        if not isinstance(context_projection, dict) or context_projection.get("context_pack_version") is None:
            failures.append(
                {
                    "case": item["case"],
                    "account": item["account"],
                    "reason": "context_pack_version_missing_or_unsafe",
                }
            )
        if not isinstance(context_projection, dict) or not context_projection.get("ok"):
            failure = {
                "case": item["case"],
                "account": item["account"],
                "reason": "context_snapshot_public_projection_failed",
            }
            if isinstance(context_projection, dict):
                failure["snapshot_count"] = context_projection.get("snapshot_count", 0)
                if context_projection.get("missing_public_summary_fields"):
                    failure["missing_public_summary_fields"] = context_projection["missing_public_summary_fields"]
                if context_projection.get("forbidden_projection_leaks"):
                    failure["forbidden_projection_leaks"] = context_projection["forbidden_projection_leaks"]
                if context_projection.get("raw_material_id_fields_present"):
                    failure["raw_material_id_fields_present"] = True
            failures.append(failure)

    output = {"ok": not failures, "results": sorted(results, key=lambda row: (row["account"], row["case"])), "failures": failures}
    print(json.dumps(output, ensure_ascii=False, indent=2))
    return 0 if output["ok"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
